"""engine.py — provider-agnostic AI + Azure DevOps engine (no UI dependency).
Ported from the original QA tool scripts so the Flet UI can drive it directly.
Configure provider keys in AI_CONFIG below, or pass them at runtime via set_credentials().
"""
import os, re, json, base64, html as _html, requests, time

# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIG
# ═══════════════════════════════════════════════════════════════════════════════
# Azure DevOps organization — the {org} in https://dev.azure.com/{org}.
# DELIBERATELY has NO built-in default. It is supplied per-account at runtime from
# the user's own saved credentials (Setup → "Azure Organization", a required
# field) via set_credentials(org=…) / reset_session_credentials(org=…).
#
# This used to be hardcoded to a real customer's org name, which was wrong twice
# over: (1) a blank/missing cred silently fell back to it, so the app would issue
# requests against a STRANGER'S organization instead of saying "you haven't
# configured this" — the resulting 401/404 looked like a broken PAT rather than
# missing setup; (2) it shipped one customer's org name in a public repo.
# Blank is now the honest default, and _require_org() turns it into a clear,
# actionable error at the point of use. The env var is a headless/CI seed only
# (run_worker.py already passes AZURE_ORG explicitly for its own runs).
AZURE_ORG = (os.environ.get("AZURE_ORG") or "").strip()

AI_PROVIDER = "anthropic"   # overridden at runtime by the UI

# ── Output language for generated content + the Sprint Report ───────────────
# The registry is the SINGLE SOURCE OF TRUTH for every supported language: its
# dropdown label (native), the name handed to the AI prompt, and whether it
# lays out right-to-left. Add a row here to support a new language.
LANGUAGES = {
    "en": {"name": "English", "native": "English",    "rtl": False},
    "ar": {"name": "Arabic",  "native": "العربية",     "rtl": True},
    "fr": {"name": "French",  "native": "Français",    "rtl": False},
    "tr": {"name": "Turkish", "native": "Türkçe",      "rtl": False},
    "es": {"name": "Spanish", "native": "Español",     "rtl": False},
    "de": {"name": "German",  "native": "Deutsch",     "rtl": False},
    "nl": {"name": "Dutch",   "native": "Nederlands",  "rtl": False},
}

# Overridden at runtime by the UI via set_output_lang(). Holds a registry code.
OUTPUT_LANG = "ar"

def set_output_lang(lang):
    """Set the language for generated content. Accepts any LANGUAGES code
    (ar/en/fr/tr/es/de/nl); unknown values fall back to English. Legacy callers
    passing 'ar'/'en' or 'Arabic'/'English' keep working."""
    global OUTPUT_LANG
    code = str(lang or "").strip().lower()
    if code not in LANGUAGES:
        code = code[:2]
    OUTPUT_LANG = code if code in LANGUAGES else "en"
    return OUTPUT_LANG

def out_lang_name():
    """The language NAME to instruct the AI with (e.g. 'French')."""
    return LANGUAGES.get(OUTPUT_LANG, LANGUAGES["en"])["name"]

def out_is_rtl():
    """True when the OUTPUT language lays out right-to-left (Arabic today)."""
    return LANGUAGES.get(OUTPUT_LANG, LANGUAGES["en"])["rtl"]

AI_CONFIG = {
    "anthropic":    {"api_key": "your-anthropic-key-here", "model": "claude-sonnet-4-6", "vision": True},
    "openai":       {"api_key": "your-openai-key-here", "model": "gpt-4o", "vision": True},
    "gemini":       {"api_key": "your-gemini-key-here", "model": "gemini-1.5-pro", "vision": True},
    "azure_openai": {"api_key": "your-azure-openai-key-here", "endpoint": "https://YOUR-RESOURCE.openai.azure.com",
                     "deployment": "gpt-4o", "api_version": "2024-06-01", "vision": True},
    "ollama":       {"api_key": "", "base_url": "http://localhost:11434", "model": "llama3.1", "vision": False},
    "nvidia":       {"api_key": "nvapi-your-nvidia-key-here", "base_url": "https://integrate.api.nvidia.com/v1",
                     "model": "meta/llama-3.1-70b-instruct", "vision": False},
    # DeepSeek — OpenAI-compatible API (base_url https://api.deepseek.com).
    # New accounts get a one-time free token grant, then cheap pay-as-you-go.
    # "deepseek-chat" is the current alias for V4-Flash (non-thinking). NOTE: the
    # deepseek-chat / deepseek-reasoner aliases are scheduled for deprecation on
    # 2026-07-24 — after that switch model to "deepseek-v4-flash" (fast/cheap) or
    # "deepseek-v4-pro" (stronger reasoning). deepseek-chat is text-only.
    "deepseek":     {"api_key": "your-deepseek-key-here", "base_url": "https://api.deepseek.com",
                     "model": "deepseek-chat", "vision": False},
    # Qwen (Alibaba DashScope / Model Studio) — OpenAI-compatible. Default base_url
    # is the INTERNATIONAL (Singapore) endpoint, correct for accounts outside
    # mainland China (e.g. Egypt). For a Beijing-region key use
    # https://dashscope.aliyuncs.com/compatible-mode/v1 ; US: dashscope-us...
    # "qwen-plus" is a solid text default; for image input switch to "qwen-vl-max"
    # and set vision: True. New Model Studio accounts get limited trial credits.
    "qwen":         {"api_key": "your-qwen-key-here",
                     "base_url": "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
                     "model": "qwen-plus", "vision": False},
    # Manus — an AGENT API (not chat-completions). It speaks the OpenAI *Responses*
    # API at https://api.manus.im, authenticated by an "API_KEY" HEADER (the
    # openai client's api_key arg is just a placeholder). Tasks run ASYNCHRONOUSLY
    # and are billed in credits, so each call creates+polls a task. "models" are
    # agent profiles: manus-1.6 (general), -lite (cheap/fast), -max (deep). We use
    # task_mode "chat" (lightest) by default. See engine's `manus` call branch.
    "manus":        {"api_key": "your-manus-key-here", "base_url": "https://api.manus.im",
                     "model": "manus-1.6", "task_mode": "chat", "vision": True},
    # ── Free-tier providers (OpenAI-compatible — share the openai HTTP adapter) ──
    # Groq — LPU inference, very fast, no card. Keys: console.groq.com/keys.
    "groq":         {"api_key": "your-groq-key-here", "base_url": "https://api.groq.com/openai/v1",
                     "model": "llama-3.3-70b-versatile", "vision": False},
    # Cerebras — wafer-scale inference, no card, ~1M tokens/day. Keys: cloud.cerebras.ai.
    "cerebras":     {"api_key": "your-cerebras-key-here", "base_url": "https://api.cerebras.ai/v1",
                     "model": "llama-3.3-70b", "vision": False},
    # OpenRouter — one key, many models; the ":free" suffixed ids are free. Keys:
    # openrouter.ai/keys. base_url is OpenAI-compatible.
    "openrouter":   {"api_key": "your-openrouter-key-here", "base_url": "https://openrouter.ai/api/v1",
                     "model": "meta-llama/llama-3.3-70b-instruct:free", "vision": False},
    # Mistral — generous free "Experiment" tier (requires opting into data-training).
    # Keys: console.mistral.ai. OpenAI-compatible endpoint.
    "mistral":      {"api_key": "your-mistral-key-here", "base_url": "https://api.mistral.ai/v1",
                     "model": "mistral-large-latest", "vision": False},
    # MiniMax — OpenAI-compatible; built for coding/agentic + test-validated
    # loops (a good fit for test-case generation). Keys: platform.minimax.io.
    # base_url is the INTERNATIONAL endpoint. "MiniMax-M2" is the flagship;
    # newer point releases (MiniMax-M2.1 / -M2.5 / -M2.7) can be selected in the
    # model dropdown once a key is saved. Text-only. Time-limited free trial.
    "minimax":      {"api_key": "your-minimax-key-here", "base_url": "https://api.minimax.io/v1",
                     "model": "MiniMax-M2", "vision": False},
    # GLM (Zhipu AI / Z.AI) — OpenAI-compatible. Tops the open-weight
    # leaderboards. Keys: z.ai (free tier, no card). base_url is the
    # international Z.AI endpoint (mainland: https://open.bigmodel.cn/api/paas/v4).
    # "glm-4.5-flash" is the FREE, rate-limited model; switch to "glm-4.5" /
    # "glm-4.7" (or a -flash newer point release) for the stronger paid models.
    "glm":          {"api_key": "your-glm-key-here", "base_url": "https://api.z.ai/api/paas/v4",
                     "model": "glm-4.5-flash", "vision": False},
}

# OpenAI-compatible providers — all share one HTTP chat/model-list adapter. Add a
# provider here when its endpoint speaks the OpenAI /chat/completions + /models API.
OPENAI_COMPAT_PROVIDERS = ("openai", "nvidia", "deepseek", "qwen",
                           "groq", "cerebras", "openrouter", "mistral",
                           "minimax", "glm")

# Rough capability ranking for ordering the provider dropdown "most powerful
# first" (lower = stronger). Applied WITHIN each free/paid group in
# _provider_options, after the active provider. Deliberately subjective and
# time-sensitive — bump entries as models change. Anything not listed sorts
# last (99). NVIDIA/Groq/Cerebras/OpenRouter are model HOSTS, ranked on the
# strength of what they typically serve, not a model of their own.
POWER_RANK = {
    "anthropic": 0, "openai": 1, "gemini": 2, "azure_openai": 3,
    "deepseek": 4, "glm": 5, "qwen": 6, "minimax": 7, "mistral": 8,
    "nvidia": 9, "groq": 10, "cerebras": 11, "openrouter": 12,
    "manus": 13, "ollama": 14,
}

# ── AI usage / cost tracking ─────────────────────────────────────────────────
# USD per 1,000,000 tokens, {"in": ..., "out": ...}. This table is NOT
# authoritative — a provider's own invoice always wins — it only exists to put
# an approximate price next to the AI Usage report's EXACT token counts (those
# come straight from each provider's own response; see _norm_usage and the
# per-provider adapters above/below). Only models we're actually confident
# about are listed; everything else is left unpriced on purpose rather than
# guessed, so the report can honestly say "cost unknown" for a call instead of
# quietly showing a wrong number. Update this table when a provider changes
# published pricing (last reviewed 2026-07).
PRICING = {
    "anthropic": {
        "claude-opus-4-8":            {"in": 15.0,  "out": 75.0},
        "claude-sonnet-5":            {"in": 3.0,   "out": 15.0},
        "claude-sonnet-4-6":          {"in": 3.0,   "out": 15.0},
        "claude-haiku-4-5-20251001":  {"in": 0.8,   "out": 4.0},
    },
    "openai": {
        "gpt-4o":       {"in": 2.5,  "out": 10.0},
        "gpt-4o-mini":  {"in": 0.15, "out": 0.6},
    },
    "azure_openai": {
        "gpt-4o":       {"in": 2.5,  "out": 10.0},
        "gpt-4o-mini":  {"in": 0.15, "out": 0.6},
    },
    "gemini": {
        "gemini-1.5-pro":   {"in": 1.25,  "out": 5.0},
        "gemini-1.5-flash": {"in": 0.075, "out": 0.3},
    },
    "ollama": {},   # local — always free, handled as a special case in price_for()
}


def price_for(provider, model):
    """$/1M-token rate for one provider+model, or None if we don't have a
    confident published price for it (never guessed from a similar model)."""
    if provider == "ollama":
        return {"in": 0.0, "out": 0.0}
    return (PRICING.get(provider) or {}).get(model)


def _call_cost(provider, model, input_tokens, output_tokens):
    """Approximate USD cost for one call from EXACT token counts, or None if
    price_for() has no rate for this model — never a guessed/estimated price."""
    p = price_for(provider, model)
    if not p:
        return None
    try:
        return round((input_tokens or 0) / 1_000_000 * p["in"]
                     + (output_tokens or 0) / 1_000_000 * p["out"], 6)
    except Exception:
        return None


# Human-readable label for each `usage_tag` a call site can pass to
# ai_complete() — this is what turns an opaque tag like "generate_steps"
# into "Run · Steps" in the AI Usage report's Module column. Rows logged
# before this tagging existed (or any tag not listed here) fall back to
# "Other" — see _usage_module_label().
_USAGE_TAG_LABELS = {
    "ui_description":          "Run · UI description",
    "generate_steps":          "Run · Steps",
    "generate_titles":         "Run · Titles",
    "evaluate_existing_steps": "Run · Evaluate existing",
    "dedupe_ai_clusters":      "Dedup · AI clustering",
    "dedupe_titles_ai":        "Dedup · New-title check",
    "automation_compile":      "Automation · Compile",
    "automation_tiebreak":     "Automation · Tie-break",
    "automation_match_element": "Automation · Element match",
    "sprint_plan_complexity":  "Sprint Plan · Complexity",
}


def _usage_module_label(tag):
    if not tag:
        return "Other"
    return _USAGE_TAG_LABELS.get(tag, tag)


# Local, per-signed-in-user usage ledger — mirrors store.py's set_user()/
# CRED_FILE pattern (and main.py's _links_path()) so accounts sharing a device
# don't mix usage history, and the ledger works fully offline even when
# Supabase auth isn't configured/signed in. This is the source of truth for
# "my own usage"; the 'ai-usage' Edge Function upload (best-effort, see
# record_ai_usage below) is what lets an Admin see EVERYONE's usage.
import threading as _usage_threading   # local alias — the shared `_threading`
                                        # import lower in this file hasn't run
                                        # yet at this point in module load
_CURRENT_USAGE_USER = None
_USAGE_LOCK = _usage_threading.Lock()
_USAGE_MAX_LOCAL_RECORDS = 5000   # prune cap so the local ledger never grows unbounded


def set_current_user(user_id):
    """Point the local AI-usage ledger at this signed-in user's own file. Call
    with None on sign-out (reverts to the shared/local default file)."""
    global _CURRENT_USAGE_USER
    _CURRENT_USAGE_USER = user_id


def _usage_log_path():
    import platform_caps as _pc_dir
    d = _pc_dir.app_data_dir()   # writable on mobile too (see helper)
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass
    if _CURRENT_USAGE_USER:
        safe = re.sub(r"[^A-Za-z0-9._-]", "_", str(_CURRENT_USAGE_USER))[:80]
        return os.path.join(d, f"ai_usage_{safe}.json")
    return os.path.join(d, "ai_usage.json")


def _load_usage_records():
    p = _usage_log_path()
    if not os.path.exists(p):
        return []
    try:
        with open(p, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_usage_records(records):
    p = _usage_log_path()
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(records[-_USAGE_MAX_LOCAL_RECORDS:], f, ensure_ascii=False)
    except Exception:
        pass


def _record_ai_usage_sync(provider, model, usage, tag=None):
    """The actual usage-recording work (local ledger write + best-effort
    Supabase mirror) — runs on the single background worker thread started by
    record_ai_usage() below, never on the caller's own thread."""
    from datetime import datetime, timezone
    usage = usage or {}
    it, ot = usage.get("input_tokens"), usage.get("output_tokens")
    cost = _call_cost(provider, model, it, ot) if usage.get("exact") else None
    rec = {"ts": datetime.now(timezone.utc).isoformat(timespec="seconds"),
           "provider": provider, "model": model,
           "input_tokens": it, "output_tokens": ot,
           "exact": bool(usage.get("exact")), "cost_usd": cost, "tag": tag}
    try:
        with _USAGE_LOCK:
            recs = _load_usage_records()
            recs.append(rec)
            _save_usage_records(recs)
    except Exception:
        pass
    try:
        import auth_supabase as _auth
        _auth.log_ai_usage(provider, model, it or 0, ot or 0, tag)
    except Exception:
        pass


# ── background usage-logging worker ─────────────────────────────────────────
# Previously record_ai_usage() spawned a brand-new threading.Thread on EVERY
# AI call purely to log usage — fine for one call, wasteful for a long run
# with hundreds/thousands of calls (a lot of short-lived OS threads for a
# small, mostly-I/O task). Replaced with a single persistent daemon worker
# fed by a bounded queue: record_ai_usage() itself only enqueues (fast,
# never blocks, same fire-and-forget contract as before) and the one worker
# thread does the actual file/network work serially in the background.
_USAGE_QUEUE = None
_USAGE_WORKER_STARTED = False
_USAGE_WORKER_LOCK = _usage_threading.Lock()


def _usage_worker_loop():
    while True:
        item = _USAGE_QUEUE.get()
        try:
            provider, model, usage, tag = item
            _record_ai_usage_sync(provider, model, usage, tag)
        except Exception:
            pass


def _ensure_usage_worker():
    global _USAGE_QUEUE, _USAGE_WORKER_STARTED
    if _USAGE_WORKER_STARTED:
        return
    with _USAGE_WORKER_LOCK:
        if _USAGE_WORKER_STARTED:
            return
        import queue as _queue
        _USAGE_QUEUE = _queue.Queue(maxsize=2000)
        _usage_threading.Thread(target=_usage_worker_loop, daemon=True).start()
        _USAGE_WORKER_STARTED = True


def record_ai_usage(provider, model, usage, tag=None):
    """Record ONE call's usage — EXACT token counts as reported by the
    provider itself (never estimated), plus an approximate cost from PRICING
    if we have a rate for that model. Always written to the local per-user
    ledger (works fully offline), and best-effort mirrored to the shared
    'ai-usage' log (see auth_supabase.log_ai_usage) so an Admin can later pull
    a whole-org report. This function must NEVER raise and NEVER block its
    caller on network trouble — a logging failure must never break (or even
    visibly slow down) the AI call it's recording.

    Implementation note: this only ENQUEUES onto the shared background worker
    (see _ensure_usage_worker above) — it does not do the actual write itself,
    so it returns essentially immediately regardless of how many calls are
    in flight."""
    try:
        _ensure_usage_worker()
        _USAGE_QUEUE.put_nowait((provider, model, usage, tag))
    except Exception:
        # Queue full (extremely long run outpacing the writer) or any other
        # failure — drop the record rather than block or raise. Same
        # best-effort guarantee record_ai_usage has always made; a dropped
        # usage-log entry must never affect the AI call it was recording.
        pass


FEATURE_DESCRIPTION = ""   # optional global feature context for step generation

# Email
GMAIL_SENDER      = "wsstestteam2@gmail.com"
GMAIL_SENDER_NAME = "QA Studio"    # display name recipients see: From: "Name" <email>
GMAIL_APP_PASS = ""

# Runtime credentials (set by the UI)
AZURE_PAT = ""

def set_credentials(provider=None, api_key=None, pat=None, gmail=None,
                    org=None, gmail_sender=None, model=None, gmail_sender_name=None):
    global AI_PROVIDER, AZURE_PAT, GMAIL_APP_PASS, AZURE_ORG, GMAIL_SENDER, GMAIL_SENDER_NAME
    if provider:
        AI_PROVIDER = provider
    if api_key and AI_PROVIDER in AI_CONFIG:
        AI_CONFIG[AI_PROVIDER]["api_key"] = api_key
    if model and AI_PROVIDER in AI_CONFIG:
        # Azure routes by "deployment"; every other provider uses "model".
        key = "deployment" if AI_PROVIDER == "azure_openai" else "model"
        AI_CONFIG[AI_PROVIDER][key] = model.strip()
    if pat is not None:
        AZURE_PAT = pat
    if gmail is not None:
        GMAIL_APP_PASS = gmail
    if org:
        AZURE_ORG = org.strip()
    if gmail_sender:
        GMAIL_SENDER = gmail_sender.strip()
    if gmail_sender_name is not None:
        GMAIL_SENDER_NAME = gmail_sender_name.strip()


def reset_session_credentials(org="", pat=""):
    """Explicitly reset the per-account engine globals (AZURE_ORG, AZURE_PAT)
    to a specific account's own saved values (or blank) — called on every
    account switch (main.py's _switch_user_creds).

    Why this exists separately from set_credentials() above: set_credentials
    only ever ASSIGNS a field when the caller passes a truthy value (`if org:
    AZURE_ORG = ...`), by design — a Setup screen field left blank on Save
    should never blank out a value that was already there. That's the right
    behavior for "the user is editing a field", but it's the WRONG behavior
    for "a different account just signed in and this global must reflect
    THEIR value even if that value is empty" — using set_credentials for that
    case silently leaves the PREVIOUS account's org/PAT in place for an
    account that has none saved, which is exactly the class of cross-account
    leak documented in DEV_ROADMAP.md ("Setup — Azure Organization is now
    genuinely isolated per account"). Kept in one function (rather than each
    caller poking `E.AZURE_ORG = ...` directly, which is how that bug
    originally happened — some call sites did it, some didn't) so any FUTURE
    per-account global follows the same explicit-reset rule from one obvious
    place instead of needing to be independently rediscovered."""
    global AZURE_ORG, AZURE_PAT
    AZURE_ORG = (org or "").strip()
    AZURE_PAT = pat or ""


def current_model(provider=None):
    """Return the configured model id for a provider (or the active one)."""
    p = provider or AI_PROVIDER
    cfg = AI_CONFIG.get(p, {})
    return cfg.get("deployment") if p == "azure_openai" else cfg.get("model")


# ═══════════════════════════════════════════════════════════════════════════════
#  AI PROVIDER LAYER
# ═══════════════════════════════════════════════════════════════════════════════
class CreditBalanceError(Exception):
    pass

# ── Error classification ──────────────────────────────────────────────────────
# Categories returned by classify_ai_error(). The UI/log uses (category, message);
# the orchestrators use TRANSIENT_CATEGORIES to decide whether to retry.
TRANSIENT_CATEGORIES = {"rate_limit", "server", "overloaded", "timeout", "network"}

def _status_of(exc):
    """Best-effort HTTP status code from any provider SDK exception."""
    for attr in ("status_code", "status", "http_status", "code"):
        v = getattr(exc, attr, None)
        if isinstance(v, int):
            return v
        if isinstance(v, str) and v.isdigit():
            return int(v)
    resp = getattr(exc, "response", None)
    if resp is not None:
        sc = getattr(resp, "status_code", None)
        if isinstance(sc, int):
            return sc
    m = re.search(r"\b(4\d\d|5\d\d)\b", str(exc))
    return int(m.group(1)) if m else None

def _classified_error(cat, friendly):
    """RuntimeError carrying its already-known category as `.ai_category`.

    ai_complete() classifies the real provider exception once (typed SDK name,
    HTTP status, low-level message substrings like "connectionerror" /
    "getaddrinfo") and raises a plain RuntimeError with just the human-
    friendly text. If a caller further up the stack (e.g. run_steps' per-case
    except block) calls classify_ai_error() AGAIN on that RuntimeError, it has
    only the friendly text to work with — which doesn't contain any of those
    low-level substrings — so re-classifying it from scratch silently
    degraded every bubbled-up error to "unknown" regardless of its real
    category. Stamping `.ai_category` here lets classify_ai_error() short-
    circuit straight back to the original, correct category instead."""
    err = RuntimeError(friendly)
    err.ai_category = cat
    return err


def classify_ai_error(exc):
    """Map any provider exception to (category, friendly_message).

    Categories: auth, credit, rate_limit, bad_model, not_found, context_length,
    bad_request, content_filter, server, overloaded, network, timeout, unknown.
    Reads typed SDK exception names first, then HTTP status, then message text,
    so it works across the anthropic / openai / google SDKs.
    """
    prov = T_disp(AI_PROVIDER)
    cached_cat = getattr(exc, "ai_category", None)
    if cached_cat:
        return (cached_cat, str(exc))
    raw = str(exc or "")
    low = raw.lower()
    etype = type(exc).__name__.lower()
    status = _status_of(exc)

    # 0) credit / quota first (a 429 can mean either rate-limit OR out-of-quota)
    if _is_credit_error(raw):
        return ("credit", f"{prov}: account is out of credit/quota. Top up with the "
                          f"provider, or switch the AI Provider in Setup and Resume.")

    # 0.5) expired / invalid key — some providers (e.g. Gemini) return this as a
    # 400, so catch it by message BEFORE the generic bad_request branch below.
    if (("api key expired" in low) or ("api_key_invalid" in low)
            or ("api key not valid" in low) or ("api key is invalid" in low)
            or ("expired" in low and "key" in low) or ("renew" in low and "key" in low)):
        return ("auth", f"{prov}: API key expired/invalid. Renew or re-check the key "
                        f"in Setup and Save it, then Resume — or switch provider.")

    # 1) typed-exception names from the SDKs (most reliable)
    if "authentication" in etype or "permissiondenied" in etype:
        return ("auth", f"{prov}: API key rejected. Re-check the key in Setup, Save it, then Resume.")
    if "ratelimit" in etype:
        return ("rate_limit", f"{prov}: rate limited (429). Waiting before retry…")
    if "notfound" in etype:
        return ("bad_model", f"{prov}: model not found. Pick a valid model for this provider in Setup.")
    if "badrequest" in etype or "unprocessable" in etype or "invalidargument" in etype:
        if "context length" in low or "maximum context" in low or "too long" in low:
            return ("context_length", f"{prov}: input too long for this model's context window.")
        if "model" in low and ("not" in low or "invalid" in low or "does not exist" in low):
            return ("bad_model", f"{prov}: invalid model. Pick a valid model for this provider in Setup.")
        return ("bad_request", f"{prov}: request rejected — {re.sub(r'\s+', ' ', raw).strip()[:160]}")
    if "timeout" in etype or "timed out" in low:
        return ("timeout", f"{prov}: request timed out. Retrying…")
    if ("apiconnection" in etype or "connectionerror" in etype or "getaddrinfo" in low
            or "name or service not known" in low or "ssl" in low
            or "max retries" in low or "failed to establish" in low):
        return ("network", f"{prov}: cannot reach the provider — check your network/firewall.")
    if "overloaded" in etype or "overloaded" in low or status == 529:
        return ("overloaded", f"{prov}: provider overloaded. Retrying…")
    if "internalserver" in etype or "serviceunavailable" in etype:
        return ("server", f"{prov}: provider error ({status or '5xx'}). Retrying…")
    if "contentfilter" in etype or "content_filter" in low or ("blocked" in low and "safety" in low):
        return ("content_filter", f"{prov}: the response was blocked by a safety filter.")

    # 2) fall back to HTTP status code
    if status == 402:   # Payment Required — out of credit/quota (hard stop)
        return ("credit", f"{prov}: account is out of credit/quota (402). Top up with the "
                          f"provider, or switch the AI Provider in Setup and Resume.")
    if status == 401:
        return ("auth", f"{prov}: API key rejected (401). Re-check the key in Setup, Save it, then Resume.")
    if status == 403:
        return ("auth", f"{prov}: access denied (403). Check the key's permissions/region for this model.")
    if status == 404:
        return ("bad_model", f"{prov}: model/endpoint not found (404). Pick a valid model in Setup.")
    if status == 422:
        return ("bad_request", f"{prov}: request rejected (422) — check the model and parameters.")
    if status == 429:
        return ("rate_limit", f"{prov}: rate limited (429). Waiting before retry…")
    if status == 529:
        return ("overloaded", f"{prov}: provider overloaded (529). Retrying…")
    if status in (500, 502, 503, 504):
        return ("server", f"{prov}: provider error ({status}). Retrying…")

    # 3) message-text fallbacks
    if "invalid api key" in low or "incorrect api key" in low or "unauthorized" in low or "x-api-key" in low:
        return ("auth", f"{prov}: API key rejected. Re-check the key in Setup, Save it, then Resume.")
    if "model" in low and ("not found" in low or "does not exist" in low or "unknown model" in low):
        return ("bad_model", f"{prov}: model not found. Pick a valid model for this provider in Setup.")
    if "rate limit" in low or "429" in low:
        return ("rate_limit", f"{prov}: rate limited (429). Waiting before retry…")
    if "context length" in low or "maximum context" in low:
        return ("context_length", f"{prov}: input too long for this model's context window.")

    return ("unknown", (f"{prov}: {re.sub(r'\s+', ' ', raw).strip()[:180]}"
                        if raw else f"{prov}: unknown error."))

# Provider errors a user can fix by renewing a key, waiting, or switching provider
# — these PAUSE the run (so the user can act + Resume). Everything else (a bad
# JSON response, content filter, oversized context) falls back to raw steps.
# Errors worth PAUSING the run for so the user can switch provider / fix the key.
# Transient categories (rate_limit, server, overloaded, timeout) are deliberately
# EXCLUDED — ai_complete already retries those patiently, so pausing for them
# would just nag the user about something that clears on its own.
#
# "network" is the one exception: ai_complete DOES retry it too (budget=4,
# capped backoff ~20s total), but that budget assumes a brief blip. A real
# outage (wifi drops, VPN disconnects) outlives it, and every subsequent test
# case then burns through its own fresh 20s of retries and logs the exact same
# "cannot reach the provider" line — which just looks like the run is stuck
# repeating an error rather than actually failing. Once ai_complete's own
# budget is exhausted, treat it like auth/bad_model: stop the run with one
# clear message instead of nagging per test case.
_RECOVERABLE_AI_CATS = {"auth", "credit", "bad_model", "not_found", "network"}

def _is_recoverable_ai_error(exc):
    try:
        cat, _ = classify_ai_error(exc)
    except Exception:
        return False
    return cat in _RECOVERABLE_AI_CATS

def _ai_cfg():
    cfg = AI_CONFIG.get(AI_PROVIDER)
    if not cfg:
        raise RuntimeError(f"Unknown AI_PROVIDER '{AI_PROVIDER}'.")
    return cfg
def _is_credit_error(msg):
    """True when a provider says the ACCOUNT is out of credit/quota (a hard stop),
    as opposed to a transient per-minute rate limit (which clears on its own).
    Kept conservative so plain 'rate limit'/'too many requests' is NOT treated as
    out-of-credit — those still retry."""
    m = (msg or "").lower()
    return (
        "credit balance is too low" in m or "insufficient_quota" in m
        or "insufficient credit" in m or "insufficient credits" in m
        or "insufficient balance" in m or "out of credit" in m
        or "no credits" in m or "not enough credit" in m
        or "credit limit" in m or "credits have run out" in m
        or "credits exhausted" in m or "run out of credits" in m
        or "payment required" in m or "add credits" in m or "top up" in m
        or ("quota" in m and any(k in m for k in ("exceeded", "exhaust", "reached", "run out")))
        or ("billing" in m and "hard limit" in m))

def friendly_ai_error(msg):
    """Turn a raw provider error (often a long JSON 400/401) into one readable
    line for the activity log. Accepts an exception or a string."""
    try:
        exc = msg if isinstance(msg, BaseException) else Exception(str(msg or ""))
        _cat, friendly = classify_ai_error(exc)
        return friendly
    except Exception:
        return re.sub(r"\s+", " ", str(msg or "")).strip()[:200]

def T_disp(name):
    """Pretty provider name without importing the UI theme."""
    return {"anthropic": "Anthropic", "openai": "OpenAI", "gemini": "Gemini",
            "azure_openai": "Azure OpenAI", "ollama": "Ollama", "nvidia": "NVIDIA",
            "deepseek": "DeepSeek", "qwen": "Qwen", "manus": "Manus",
            "groq": "Groq", "cerebras": "Cerebras", "openrouter": "OpenRouter",
            "mistral": "Mistral", "minimax": "MiniMax",
            "glm": "GLM"}.get(name, str(name).title())

def active_providers():
    """Provider names that have a usable key."""
    out = []
    for name, cfg in AI_CONFIG.items():
        k = (cfg.get("api_key") or "").strip()
        if k and not k.startswith("your-") and "-here" not in k:
            out.append(name)
    return out

class EmptyAIResponse(Exception):
    """Raised when a provider returns no usable text (empty choices, None content,
    content-filter block, or truncated/blocked output)."""
    pass

def _extract_openai_text(resp):
    """Defensively pull text from an OpenAI-compatible chat completion."""
    choices = getattr(resp, "choices", None) or []
    if not choices:
        raise EmptyAIResponse("provider returned no choices")
    ch = choices[0]
    fr = getattr(ch, "finish_reason", None)
    msg = getattr(ch, "message", None)
    text = getattr(msg, "content", None) if msg is not None else None
    if text is None:
        refusal = getattr(msg, "refusal", None) if msg is not None else None
        if refusal:
            raise EmptyAIResponse(f"model refused: {refusal}")
        if fr == "content_filter":
            raise EmptyAIResponse("response blocked by content filter")
        raise EmptyAIResponse("empty content from provider")
    if fr == "length" and not str(text).strip():
        raise EmptyAIResponse("response truncated (max_tokens) with no content")
    return text

# ── HTTP adapter layer ────────────────────────────────────────────────────────
# Every provider is called over pure HTTP (requests) — no vendor SDKs — so
# connecting + generating never depend on an SDK's pinned Python/dependency
# versions (the cause of "provider won't connect" on end-user machines). One
# uniform transport for anthropic / openai-compatible (openai, nvidia, deepseek,
# qwen, ollama) / azure / gemini / manus. See PROVIDERS_REFACTOR_PLAN.md.


def _http_post_json(url, headers, payload, timeout):
    """POST JSON, return parsed dict. Pure-Python (requests). Raises RuntimeError
    with a short body on non-200."""
    r = requests.post(url, headers=headers, json=payload, timeout=(timeout or 120))
    if r.status_code != 200:
        try:
            body = (r.text or "")[:300]
        except Exception:
            body = ""
        raise RuntimeError(f"HTTP {r.status_code} from provider: {body}")
    return r.json()


def _extract_openai_text_json(data):
    """Same defensive extraction as _extract_openai_text, over a JSON dict."""
    choices = (data or {}).get("choices") or []
    if not choices:
        raise EmptyAIResponse("provider returned no choices")
    ch = choices[0] or {}
    fr = ch.get("finish_reason")
    msg = ch.get("message") or {}
    text = msg.get("content")
    if text is None:
        refusal = msg.get("refusal")
        if refusal:
            raise EmptyAIResponse(f"model refused: {refusal}")
        if fr == "content_filter":
            raise EmptyAIResponse("response blocked by content filter")
        raise EmptyAIResponse("empty content from provider")
    if fr == "length" and not str(text).strip():
        raise EmptyAIResponse("response truncated (max_tokens) with no content")
    return text


def _norm_usage(input_tokens=None, output_tokens=None, exact=False):
    """One normalized shape for exact-usage extraction across every provider
    adapter: {"input_tokens", "output_tokens", "total_tokens", "exact"}.
    exact=False means the provider's response didn't report usage at all —
    the caller must NOT estimate/guess a number in that case (see
    record_ai_usage / PRICING), just record the call as usage-unknown."""
    have = input_tokens is not None or output_tokens is not None
    return {"input_tokens": input_tokens, "output_tokens": output_tokens,
            "total_tokens": ((input_tokens or 0) + (output_tokens or 0)) if have else None,
            "exact": bool(exact and have)}


def _extract_openai_usage_json(data):
    """Exact token usage from an OpenAI-shaped response (openai, nvidia,
    deepseek, qwen, groq, cerebras, openrouter, mistral, azure_openai all
    share this 'usage': {prompt_tokens, completion_tokens} shape)."""
    u = (data or {}).get("usage") or {}
    if not u:
        return _norm_usage()
    return _norm_usage(u.get("prompt_tokens"), u.get("completion_tokens"), exact=True)


def _openai_compat_http(cfg, prompt_text, images, max_tokens, timeout, want_json=False):
    """Pure-HTTP chat completion for any OpenAI-compatible endpoint. Mirrors the
    openai SDK branch (text or text+images, optional JSON mode)."""
    base = (cfg.get("base_url") or "https://api.openai.com/v1").rstrip("/")
    if images:
        content = [{"type": "text", "text": prompt_text}]
        for im in images:
            content.append({"type": "image_url", "image_url": {
                "url": f"data:{im['media_type']};base64,{im['data']}"}})
    else:
        content = prompt_text
    payload = {"model": cfg["model"], "max_tokens": max_tokens,
               "messages": [{"role": "user", "content": content}]}
    if want_json:
        payload["response_format"] = {"type": "json_object"}
    headers = {"Content-Type": "application/json"}
    key = (cfg.get("api_key") or "").strip()
    if key:
        headers["Authorization"] = f"Bearer {key}"
    data = _http_post_json(base + "/chat/completions", headers, payload, timeout)
    return _extract_openai_text_json(data), _extract_openai_usage_json(data)


def _openai_compat_models_http(base_url, key, timeout=15):
    """Pure-HTTP GET /models for any OpenAI-compatible endpoint. Returns id list."""
    base = (base_url or "https://api.openai.com/v1").rstrip("/")
    headers = {"Content-Type": "application/json"}
    if (key or "").strip():
        headers["Authorization"] = f"Bearer {key.strip()}"
    r = requests.get(base + "/models", headers=headers, timeout=(timeout or 15))
    r.raise_for_status()
    return [m.get("id") for m in ((r.json() or {}).get("data") or []) if m.get("id")]


def _http_get_json(url, headers, timeout):
    """GET JSON, return parsed dict. Pure-Python (requests)."""
    r = requests.get(url, headers=headers, timeout=(timeout or 30))
    if r.status_code != 200:
        try:
            body = (r.text or "")[:300]
        except Exception:
            body = ""
        raise RuntimeError(f"HTTP {r.status_code} from provider: {body}")
    return r.json()


def _anthropic_http(cfg, prompt_text, images, max_tokens, timeout):
    """Pure-HTTP Anthropic Messages API (drops the `anthropic` SDK)."""
    content = []
    for im in images:
        content.append({"type": "image", "source": {"type": "base64",
            "media_type": im["media_type"], "data": im["data"]}})
    content.append({"type": "text", "text": prompt_text})
    headers = {"x-api-key": (cfg.get("api_key") or "").strip(),
               "anthropic-version": "2023-06-01", "content-type": "application/json"}
    payload = {"model": cfg["model"], "max_tokens": max_tokens,
               "messages": [{"role": "user", "content": content}]}
    data = _http_post_json("https://api.anthropic.com/v1/messages", headers, payload, timeout)
    blocks = (data or {}).get("content") or []
    out = "".join(b.get("text", "") for b in blocks if b.get("type") == "text").strip()
    if not out:
        raise EmptyAIResponse(f"empty response (stop_reason={(data or {}).get('stop_reason')})")
    u = (data or {}).get("usage") or {}
    usage = (_norm_usage(u.get("input_tokens"), u.get("output_tokens"), exact=True)
             if u else _norm_usage())
    return out, usage


def _anthropic_models_http(key, timeout=15):
    # Anthropic's /v1/models is cursor-paginated (has_more + last_id). Page through
    # so every available model is returned, not just the first page.
    headers = {"x-api-key": (key or "").strip(), "anthropic-version": "2023-06-01"}
    base = "https://api.anthropic.com/v1/models?limit=1000"
    ids, after = [], None
    for _ in range(20):  # safety cap on pages
        url = base + (f"&after_id={after}" if after else "")
        data = _http_get_json(url, headers, timeout)
        page = (data or {}).get("data") or []
        ids += [m.get("id") for m in page if m.get("id")]
        if not (data or {}).get("has_more"):
            break
        after = (data or {}).get("last_id") or (page[-1].get("id") if page else None)
        if not after:
            break
    return ids


def _azure_http(cfg, prompt_text, images, max_tokens, timeout):
    """Pure-HTTP Azure OpenAI chat completion (drops the openai SDK)."""
    if images:
        content = [{"type": "text", "text": prompt_text}]
        for im in images:
            content.append({"type": "image_url", "image_url": {
                "url": f"data:{im['media_type']};base64,{im['data']}"}})
    else:
        content = prompt_text
    endpoint = (cfg.get("endpoint") or "").rstrip("/")
    ver = cfg.get("api_version", "2024-06-01")
    url = f"{endpoint}/openai/deployments/{cfg['deployment']}/chat/completions?api-version={ver}"
    headers = {"api-key": (cfg.get("api_key") or "").strip(), "content-type": "application/json"}
    payload = {"messages": [{"role": "user", "content": content}], "max_tokens": max_tokens}
    data = _http_post_json(url, headers, payload, timeout)
    return _extract_openai_text_json(data), _extract_openai_usage_json(data)


def _azure_models_http(cfg, timeout=15):
    endpoint = (cfg.get("endpoint") or "").rstrip("/")
    ver = cfg.get("api_version", "2024-06-01")
    headers = {"api-key": (cfg.get("api_key") or "").strip()}
    data = _http_get_json(f"{endpoint}/openai/models?api-version={ver}", headers, timeout)
    return [m.get("id") for m in ((data or {}).get("data") or []) if m.get("id")]


def _gemini_http(cfg, prompt_text, images, max_tokens, timeout, want_json=False):
    """Pure-HTTP Gemini generateContent (drops google-generativeai/grpcio/protobuf)."""
    parts = [{"text": prompt_text}]
    for im in images:
        parts.append({"inline_data": {"mime_type": im["media_type"], "data": im["data"]}})
    gen_cfg = {"maxOutputTokens": max_tokens}
    if want_json:
        gen_cfg["responseMimeType"] = "application/json"
    key = (cfg.get("api_key") or "").strip()
    # Send the key in the x-goog-api-key HEADER, not the URL query string.
    # Keys in URLs are far more likely to be captured in proxy/server/access
    # logs than header auth (security review hardening).
    url = ("https://generativelanguage.googleapis.com/v1beta/models/"
           f"{cfg['model']}:generateContent")
    payload = {"contents": [{"parts": parts}], "generationConfig": gen_cfg}
    data = _http_post_json(url, {"content-type": "application/json",
                                 "x-goog-api-key": key}, payload, timeout)
    cands = (data or {}).get("candidates") or []
    if not cands:
        fb = ((data or {}).get("promptFeedback") or {}).get("blockReason")
        raise EmptyAIResponse(f"blocked by Gemini (block_reason={fb})")
    parts_out = ((cands[0].get("content") or {}).get("parts")) or []
    txt = "".join(p.get("text", "") for p in parts_out).strip()
    if not txt:
        raise EmptyAIResponse(f"empty response from Gemini (finish_reason={cands[0].get('finishReason')})")
    um = (data or {}).get("usageMetadata") or {}
    usage = (_norm_usage(um.get("promptTokenCount"), um.get("candidatesTokenCount"), exact=True)
             if um else _norm_usage())
    return txt, usage


def _gemini_models_http(key, timeout=15):
    # Gemini's ListModels is paginated (default pageSize ~50). Ask for a large
    # page and follow nextPageToken so the FULL catalogue is returned.
    key = (key or "").strip()
    base = "https://generativelanguage.googleapis.com/v1beta/models"
    _hdr = {"x-goog-api-key": key}   # key in header, not the URL (see _gemini call)
    ids, token = [], None
    for _ in range(20):  # safety cap on pages
        url = f"{base}?pageSize=1000" + (f"&pageToken={token}" if token else "")
        data = _http_get_json(url, _hdr, timeout)
        for m in ((data or {}).get("models") or []):
            if "generateContent" in (m.get("supportedGenerationMethods") or []):
                nm = m.get("name", "") or ""
                ids.append(nm.split("/", 1)[1] if nm.startswith("models/") else nm)
        token = (data or {}).get("nextPageToken")
        if not token:
            break
    return ids


def _manus_http(cfg, prompt_text, images, max_tokens, timeout):
    """Pure-HTTP Manus Responses API (create → poll → read). Drops the openai SDK."""
    base = (cfg.get("base_url") or "https://api.manus.im").rstrip("/")
    headers = {"API_KEY": (cfg.get("api_key") or "").strip(), "content-type": "application/json"}
    content = [{"type": "input_text", "text": prompt_text}]
    for im in images:
        content.append({"type": "input_image",
                        "image_url": f"data:{im['media_type']};base64,{im['data']}"})
    payload = {"model": cfg["model"], "input": [{"role": "user", "content": content}],
               "task_mode": cfg.get("task_mode") or "chat", "agent_profile": cfg["model"]}
    data = _http_post_json(base + "/responses", headers, payload, timeout)
    rid = (data or {}).get("id")
    status = (data or {}).get("status")
    deadline = time.time() + (timeout if timeout else 600)
    while status == "running" and time.time() < deadline:
        _interruptible_sleep(5)
        if _STOP_EVENT.is_set():
            break
        try:
            data = _http_get_json(f"{base}/responses/{rid}", headers, timeout)
        except Exception:
            break
        status = (data or {}).get("status")
    if status == "error":
        raise RuntimeError(f"Manus task failed (id {rid})")
    texts = []
    for msg in ((data or {}).get("output") or []):
        if msg.get("role") != "assistant":
            continue
        for part in (msg.get("content") or []):
            t = part.get("text")
            if t:
                texts.append(t)
    out = "\n".join(texts).strip()
    if not out:
        raise EmptyAIResponse(f"Manus returned no assistant text (status={status})")
    # Manus' agent-task Responses API doesn't report token usage — the call is
    # still counted (record_ai_usage logs it), just with tokens/cost unknown
    # rather than guessed.
    return out, _norm_usage()


def _ai_call_once(provider, cfg, prompt_text, images, max_tokens, timeout, want_json=False):
    """One provider call. Returns (text, usage) or raises (EmptyAIResponse / SDK
    error). `usage` is the EXACT token count the provider itself reported for
    THIS call (see _norm_usage) — never an estimate — with usage["exact"]=False
    when a provider's response doesn't include it at all.
    want_json=True asks the provider for strict JSON output where supported (used
    for the intent-compiler calls), so models like Gemini don't wrap the JSON in
    reasoning prose ('Wait, let's…') which then fails to parse."""
    if provider == "anthropic":
        return _anthropic_http(cfg, prompt_text, images, max_tokens, timeout)

    if provider in OPENAI_COMPAT_PROVIDERS:
        return _openai_compat_http(cfg, prompt_text, images, max_tokens, timeout, want_json)

    if provider == "azure_openai":
        return _azure_http(cfg, prompt_text, images, max_tokens, timeout)

    if provider == "gemini":
        return _gemini_http(cfg, prompt_text, images, max_tokens, timeout, want_json)

    if provider == "ollama":
        payload = {"model": cfg["model"],
                   "messages": [{"role": "user", "content": prompt_text}], "stream": False}
        r = requests.post(f"{cfg['base_url']}/api/chat", json=payload, timeout=timeout or 180)
        r.raise_for_status()
        data = r.json()
        txt = (data.get("message") or {}).get("content")
        if not (txt or "").strip():
            raise EmptyAIResponse("empty response from Ollama")
        # Ollama's /api/chat reports usage at the TOP level of the response
        # (not nested under "usage" like the OpenAI-shaped providers).
        usage = (_norm_usage(data.get("prompt_eval_count"), data.get("eval_count"), exact=True)
                 if ("prompt_eval_count" in data or "eval_count" in data) else _norm_usage())
        return txt, usage

    if provider == "manus":
        return _manus_http(cfg, prompt_text, images, max_tokens, timeout)

    raise RuntimeError(f"Unhandled provider '{provider}'")

# Module-level cooperative stop: run loops set this so long backoff sleeps inside
# ai_complete can bail out promptly when the user clicks Stop.
import threading as _threading
_STOP_EVENT = _threading.Event()

def request_stop():
    _STOP_EVENT.set()

def clear_stop():
    _STOP_EVENT.clear()

def _interruptible_sleep(seconds):
    """Sleep in small slices so a Stop request ends the wait quickly."""
    end = time.time() + max(0.0, seconds)
    while time.time() < end:
        if _STOP_EVENT.is_set():
            return
        time.sleep(min(0.25, end - time.time()))


class StopRequested(Exception):
    """Raised to unwind a run promptly when the user clicks Stop mid-generation."""
    pass


def _run_stopaware(fn, should_stop=None, on_slow=None, poll=0.3, slow_every=15):
    """Run a BLOCKING call fn() in a daemon thread and wait for it, but stay
    responsive: if should_stop() turns true we raise StopRequested immediately
    (the abandoned thread — e.g. a requests.post still waiting on a slow provider —
    finishes on its own timeout and is harmless, since generation has no side
    effects). on_slow(elapsed_seconds) is called about every `slow_every` seconds
    so the UI can show a 'still working' heartbeat instead of looking frozen."""
    should_stop = should_stop or (lambda: False)
    box = {}
    def _work():
        try: box["val"] = fn()
        except BaseException as e: box["err"] = e   # noqa: BLE001 — re-raised below
    t = _threading.Thread(target=_work, daemon=True)
    t.start()
    start = last_hb = time.time()
    while t.is_alive():
        if should_stop() or _STOP_EVENT.is_set():
            raise StopRequested()
        now = time.time()
        if on_slow and (now - last_hb) >= slow_every:
            try: on_slow(int(now - start))
            except Exception: pass
            last_hb = now
        t.join(timeout=poll)
    if "err" in box:
        raise box["err"]
    return box.get("val")

def _retry_after_seconds(exc):
    """Pull a Retry-After hint (seconds) from a provider exception, if any."""
    # OpenAI/Anthropic SDKs attach .response with headers; also check the message.
    try:
        resp = getattr(exc, "response", None)
        hdrs = getattr(resp, "headers", None) if resp is not None else None
        if hdrs:
            for k in ("retry-after", "Retry-After", "x-ratelimit-reset-requests",
                      "x-ratelimit-reset-tokens"):
                v = hdrs.get(k) if hasattr(hdrs, "get") else None
                if v:
                    m = re.search(r"[\d.]+", str(v))
                    if m:
                        return float(m.group(0))
    except Exception:
        pass
    # message text: "try again in 12s" / "retry after 5 seconds"
    try:
        m = re.search(r"(?:retry[- ]after|try again in)\D*([\d.]+)\s*(m|min|s|sec)?",
                      str(exc), re.I)
        if m:
            val = float(m.group(1))
            unit = (m.group(2) or "s").lower()
            return val * 60 if unit.startswith("m") else val
    except Exception:
        pass
    return None

# Shared (cross-thread) rate-limit cooldown. run_steps/run_titles/the dedup
# pass all run 2-worker pools against the SAME provider endpoint — before
# this gate, each worker discovered a rate limit independently and sat out
# its own private backoff, and the extra requests fired during the window
# count against the limit (often extending it). When any call gets
# rate-limited/overloaded, the wait it computed is published here; every
# other ai_complete() call waits out the remainder BEFORE issuing its
# request, instead of burning a request to learn what a sibling thread
# already knows. Combined with on_retry now being threaded through every
# Run call site, this is the fix for the "60s of total silence that was
# actually 3-4 invisible stacked retry cycles" failure mode.
_RL_GATE_LOCK = _threading.Lock()
_RL_COOLDOWN_UNTIL = 0.0


def ai_complete(prompt_text, images=None, max_tokens=4096, timeout=None,
                retries=3, on_retry=None, want_json=False, usage_out=None,
                usage_tag=None):
    """Call the active AI provider with defensive extraction + retry on transient
    errors (rate limit / 5xx / overloaded / timeout / network).

    Transient errors are retried patiently (rate-limit/overload get a larger
    budget and honor any Retry-After hint), so they almost never bubble up to the
    user. Raises CreditBalanceError for out-of-credit, or a RuntimeError carrying
    the friendly classified message for non-transient failures. `on_retry(msg)`
    (if given) is called before each retry so the UI can log "retrying…".

    Still returns a plain string on success, exactly as before — every existing
    call site keeps working unchanged. Two purely additive params for callers
    that care about usage: `usage_out` (a dict, filled in-place with this
    call's EXACT token usage before returning — see _norm_usage) and
    `usage_tag` (an optional short label like "generate_test_cases" recorded
    alongside the usage so a report can break totals down by feature, not
    just provider). Every successful call is ALSO recorded automatically
    (record_ai_usage, off-thread) regardless of whether a caller passes these.
    """
    global _RL_COOLDOWN_UNTIL
    images = images or []
    # Per-category retry budgets. Rate-limit / overloaded clear on their own, so
    # we wait them out generously instead of surfacing an error to the user.
    _BUDGET = {"rate_limit": 8, "overloaded": 8, "server": 5,
               "timeout": 4, "network": 4}
    attempt = 0
    last_friendly = None
    provider = None
    while True:
        # Provider/config are re-read EVERY attempt — switching providers in
        # Settings mid-run now rescues the in-flight call on its very next
        # retry, instead of the call grinding through the OLD provider's full
        # multi-minute retry budget first (reported live: user switched away
        # from a stalled provider at ~120s and the call kept retrying the old
        # one past 258s). A mid-call switch also resets the attempt counter —
        # the new provider gets a fresh budget, not the old one's exhausted
        # count.
        _prov_now = AI_PROVIDER
        if provider is not None and _prov_now != provider:
            if on_retry:
                try:
                    on_retry(f"provider switched to {T_disp(_prov_now)} — "
                             f"retrying there now…")
                except Exception:
                    pass
            attempt = 0
        provider = _prov_now
        cfg = _ai_cfg()
        attempt += 1
        # Respect a cooldown another thread just set (see _RL_COOLDOWN_UNTIL
        # above): don't fire a request that will almost certainly 429 and
        # start this call's own private backoff ladder on top.
        _cd = _RL_COOLDOWN_UNTIL - time.time()
        if _cd > 0:
            if on_retry:
                on_retry(f"{T_disp(provider)}: rate-limited — waiting "
                         f"{int(_cd) + 1}s (shared cooldown)…")
            _interruptible_sleep(min(_cd, 60))
        try:
            text, usage = _ai_call_once(provider, cfg, prompt_text, images,
                                        max_tokens, timeout, want_json)
            if usage_out is not None:
                try:
                    usage_out.update(usage or {})
                except Exception:
                    pass
            # Off-thread: recording (local ledger write + best-effort Supabase
            # upload) must never add latency to — or ever break — the actual
            # generation call that just succeeded. record_ai_usage() itself
            # now just enqueues onto a shared background worker (see its
            # docstring) instead of spawning a new thread per call, so this
            # call returns essentially immediately either way.
            try:
                record_ai_usage(provider, cfg.get("model"), usage, usage_tag)
            except Exception:
                pass
            return text
        except CreditBalanceError:
            raise
        except EmptyAIResponse as e:
            # empty/blocked: retry a couple of times (often transient), then give up
            cat, friendly = "empty", f"{T_disp(provider)}: {e}"
            last_friendly = friendly
            if attempt <= retries:
                _delay = min(2 * attempt, 8)
                # "Hit max_tokens with NO text" is not transient — it means the
                # model spent the whole budget before emitting its final answer
                # (with thinking models like Sonnet 5, reasoning tokens count
                # against the same cap, so a small cap can be consumed entirely
                # by thinking). Retrying the identical request only succeeds by
                # luck (thinking length varies) — double the cap instead so the
                # retry actually has room for both reasoning AND the answer.
                if "max_tokens" in str(e).lower() and max_tokens < 8192:
                    max_tokens = min(max_tokens * 2, 8192)
                    friendly += f" — raising max_tokens to {max_tokens}"
                if on_retry: on_retry(f"{friendly} — retrying ({attempt}/{retries})…")
                _interruptible_sleep(_delay); continue
            raise _classified_error(cat, friendly)
        except Exception as e:
            cat, friendly = classify_ai_error(e)
            # ANY error the classifier calls "credit" must raise CreditBalanceError
            # so the run PAUSES (via on_ai_error) instead of failing every
            # remaining case against an exhausted provider. The text-based
            # _is_credit_error misses an HTTP 402 whose raw string has no "out of
            # credit" phrase (OpenRouter and others return a bare 402 for
            # out-of-quota) — classify_ai_error catches it by STATUS, so trust it
            # too. Reported live: OpenRouter 402s logged per-case with no pause
            # while NVIDIA/Cerebras out-of-credit paused correctly.
            if cat == "credit" or _is_credit_error(str(e)):
                raise CreditBalanceError(friendly)
            last_friendly = friendly
            budget = max(retries, _BUDGET.get(cat, 0)) if cat in TRANSIENT_CATEGORIES else 0
            if budget and attempt <= budget:
                # Honor a server-provided Retry-After when present; otherwise back
                # off progressively. Rate-limit/overload wait longer (they clear).
                ra = _retry_after_seconds(e)
                if ra is not None:
                    _delay = min(max(ra, 1), 60)
                elif cat in ("rate_limit", "overloaded"):
                    _delay = min(5 + 5 * attempt, 45)      # 10,15,…cap 45s
                else:
                    _delay = min(2 * attempt, 20)
                if cat in ("rate_limit", "overloaded"):
                    # Publish the cooldown so concurrent workers wait it out
                    # too instead of independently rediscovering the limit.
                    with _RL_GATE_LOCK:
                        _RL_COOLDOWN_UNTIL = max(_RL_COOLDOWN_UNTIL,
                                                 time.time() + _delay)
                if on_retry:
                    on_retry(f"{friendly} — waiting {int(_delay)}s then retry "
                             f"({attempt}/{budget})…")
                    # Keep the retry note VISIBLE while waiting: at heartbeat-
                    # wired call sites the retry note shares its hb_id line
                    # with the "Still …ing — Ns so far" heartbeat, which fires
                    # every ~15s from the monitor thread and silently
                    # OVERWRITES it (confirmed live: 258s of backoff showed
                    # only the generic heartbeat, reading as a frozen call
                    # with no retries happening). Re-asserting the note every
                    # ~5s with a live countdown wins that overwrite race and
                    # doubles as progress feedback.
                    _end = time.time() + _delay
                    while True:
                        _left = _end - time.time()
                        if _left <= 0 or _STOP_EVENT.is_set():
                            break
                        _interruptible_sleep(min(5, _left))
                        _left = int(_end - time.time())
                        if _left > 0 and not _STOP_EVENT.is_set():
                            try:
                                on_retry(f"{friendly} — retrying in {_left}s "
                                         f"({attempt}/{budget})…")
                            except Exception:
                                pass
                else:
                    _interruptible_sleep(_delay)
                continue
            raise _classified_error(cat, friendly)


# Extra, VISIBLE backoff for a "network" failure at the orchestrator level
# (run_steps/run_titles), on top of ai_complete's own short internal retry
# budget (~20s). ai_complete's budget assumes a brief blip; a real outage —
# wifi drop, VPN reconnect — commonly clears within a couple of minutes, so
# it's worth waiting it out here before giving up on the whole run. Each wait
# is logged so it reads as "retrying" rather than "stuck".
_NETWORK_RETRY_WAITS = (20, 40, 80)   # seconds; ~140s of extra patience total

# One consistent activity-log line for every place a fatal provider error now
# PAUSES the run (instead of killing it) while waiting on the Run screen's
# Resume/Stop buttons — see run_steps' _gen_and_write wrapper, run_titles'
# _process_story_paused, and the empty-suite seeding path.
_PAUSED_ON_ERROR_MSG = ("Paused on provider error — switch the AI provider in "
                        "Setup, then click Resume; or Stop to end the run.")


def _paused_on_error_msg():
    """The pause line, NAMING the provider that actually failed.

    The old constant never said WHICH provider broke, so a log read back later
    (or by someone else) couldn't tell what to switch away from — and with the
    provider re-read per attempt, "Paused on provider error" could even refer
    to a different provider than the one currently selected in Setup. Resolved
    at call time, not import time, for exactly that reason."""
    try:
        prov = T_disp(AI_PROVIDER)
    except Exception:
        prov = ""
    if not prov:
        return _PAUSED_ON_ERROR_MSG
    return (f"Paused — {prov} failed. Switch the AI provider in Setup, then "
            f"click Resume; or Stop to end the run.")


def _call_with_network_retries(fn, cb, should_stop=None):
    """Call fn() (a zero-arg callable doing one AI generation call). If it
    fails with a 'network' category error, wait (with the backoff above) and
    retry the SAME call rather than surfacing the failure immediately — only
    once every wait has been used up (or Stop is clicked) does the final
    exception propagate to the caller, which then treats it like any other
    fatal error (stops the run with one clear message instead of hammering
    every remaining item). Any non-network error is raised immediately,
    unchanged — this only adds patience for the "provider unreachable" case."""
    should_stop = should_stop or (lambda: False)
    attempt = 0
    while True:
        try:
            return fn()
        except (CreditBalanceError, StopRequested):
            raise
        except Exception as e:
            cat, friendly = classify_ai_error(e)
            if cat != "network" or attempt >= len(_NETWORK_RETRY_WAITS):
                raise
            if should_stop() or _STOP_EVENT.is_set():
                raise
            w = _NETWORK_RETRY_WAITS[attempt]
            attempt += 1
            cb("log", {"msg": f"{friendly} — retrying in {w}s "
                              f"({attempt}/{len(_NETWORK_RETRY_WAITS)})…",
                       "tone": "warn", "ico": "⏳"})
            _interruptible_sleep(w)
            if should_stop() or _STOP_EVENT.is_set():
                raise


# ═══════════════════════════════════════════════════════════════════════════════
#  AZURE REST HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
from requests.adapters import HTTPAdapter

_AZ_SESSION = None


def _az_session():
    """Shared, connection-pooled requests.Session for Azure DevOps calls.

    A cold plan generation fires hundreds of small GETs — one per test suite, 16
    in parallel. A bare requests.get() opens a brand-new TCP+TLS connection every
    time; those handshakes are slow AND CPU-heavy, and the CPU work holds Python's
    GIL, which is what made the nav/UI stutter during a cold generate. Reusing
    pooled keep-alive connections removes the per-call handshake entirely: the same
    requests return the same results, but far faster and much lighter on the GIL —
    so we keep the worker count high without starving the UI thread. urllib3's
    connection pool is thread-safe, and we never mutate shared session state (auth
    is passed per request), so sharing it across the worker pool is safe."""
    global _AZ_SESSION
    if _AZ_SESSION is None:
        s = requests.Session()
        # Headroom so 16 concurrent workers (plus other fetches) never block on a
        # free connection. max_retries=0 keeps the existing raise-on-error behavior.
        ad = HTTPAdapter(pool_connections=16, pool_maxsize=32, max_retries=0)
        s.mount("https://", ad)
        s.mount("http://", ad)
        _AZ_SESSION = s
    return _AZ_SESSION


def _require_org():
    """Return the configured Azure organization, or raise a clear, actionable
    error if none is set.

    Every Azure URL in this module interpolates AZURE_ORG. With no default and
    no guard, an unconfigured org silently builds `https://dev.azure.com//_apis/…`
    — a malformed URL whose 404 reads as "check the project name spelling",
    sending the user to debug entirely the wrong thing. Failing loudly here,
    naming the exact Setup field to fill in, is the whole point of dropping the
    hardcoded fallback."""
    org = (AZURE_ORG or "").strip()
    if not org:
        raise RuntimeError(
            "No Azure DevOps organization is configured. Open Setup and fill in "
            "'Azure Organization' — it's the {org} part of https://dev.azure.com/{org}.")
    return org


def _azure_get(url, pat=None, timeout=12):
    # Single choke point for every Azure READ in this module (30 call sites), so
    # one check here covers them all without touching each URL build.
    _require_org()
    pat = pat or AZURE_PAT
    try:
        r = _az_session().get(url, auth=("", pat), timeout=timeout)
    except requests.exceptions.SSLError:
        raise RuntimeError("SSL error reaching Azure DevOps. Your network may block dev.azure.com.")
    except requests.exceptions.ConnectionError:
        raise RuntimeError("Cannot reach Azure DevOps (dev.azure.com). Check your network/firewall.")
    except requests.exceptions.Timeout:
        raise RuntimeError("Azure DevOps request timed out (12s). Network may be blocking it.")
    if r.status_code == 401:
        raise RuntimeError("Authentication failed (401). Check your PAT and its scopes.")
    if r.status_code == 403:
        raise RuntimeError("Access denied (403). Your PAT may lack Test Management permission.")
    if r.status_code == 404:
        raise RuntimeError("Not found (404). Check the project name spelling.")
    r.raise_for_status()
    return r.json()

def validate_pat(pat):
    """Returns (ok, message). Lightweight check that the PAT can reach the org."""
    try:
        _azure_get(f"https://dev.azure.com/{AZURE_ORG}/_apis/projects?api-version=7.0", pat)
        return True, "Valid"
    except Exception as e:
        return False, str(e)

def validate_api_key(timeout=20):
    """Cheap check that the configured AI provider key works.
    Returns (ok, category). category is one of: ok, credit, ratelimited, auth,
    bad_model, network, timeout, content_filter, server, overloaded,
    missing-package:<pkg>, or error:<message>. The UI maps these to friendly text.
    Uses a single direct call (no retry) so Connect is fast, with a short
    `timeout` (seconds) so a slow/unreachable provider can't hang Connect — the
    OpenAI SDK otherwise defaults to a 600s timeout.
    """
    cfg = _ai_cfg(); provider = AI_PROVIDER
    if provider == "manus":
        # Cheap key check: list tasks (no task creation / credits) via HTTP.
        try:
            base = (cfg.get("base_url") or "https://api.manus.im").rstrip("/")
            r = requests.get(base + "/v1/tasks?limit=1",
                             headers={"API_KEY": (cfg.get("api_key") or "").strip()},
                             timeout=timeout)
            if r.status_code in (200, 201):
                return True, "ok"
            if r.status_code in (401, 403):
                return False, "auth"
            if r.status_code == 429:
                return True, "ratelimited"
            if r.status_code == 402:
                return True, "credit"
            return False, f"error:HTTP {r.status_code}"
        except Exception as e:
            cat, friendly = classify_ai_error(e)
            if cat == "auth":
                return False, "auth"
            if cat == "credit":
                return True, "credit"
            if cat == "rate_limit":
                return True, "ratelimited"
            if cat in ("network", "timeout", "server", "overloaded"):
                return False, cat
            return False, "error:" + friendly
    if provider in OPENAI_COMPAT_PROVIDERS:
        # Validate with a cheap GET /models rather than a chat-completion "ping".
        # A ping cold-starts the SELECTED model, and large models (e.g. NVIDIA's
        # 400B MoE) can take longer than `timeout` to return the first token —
        # so Connect would time out even though the key is perfectly valid.
        # /models proves the key + endpoint fast and never loads a model.
        try:
            _openai_compat_models_http(cfg.get("base_url"),
                                       (cfg.get("api_key") or "").strip(), timeout)
            return True, "ok"          # 200 (any model list, even empty) → key works
        except Exception as e:
            cat, friendly = classify_ai_error(e)
            if cat == "auth":
                return False, "auth"
            if cat == "rate_limit":
                return True, "ratelimited"
            if cat == "credit":
                return True, "credit"
            if cat in ("network", "timeout", "server", "overloaded"):
                return False, cat
            return False, "error:" + friendly
    try:
        _ai_call_once(provider, cfg, "ping", [], 8, timeout)
        return True, "ok"
    except CreditBalanceError:
        return True, "credit"          # key valid, just out of credit
    except ModuleNotFoundError as e:
        missing = str(e).split("'")[-2] if "'" in str(e) else str(e)
        return False, f"missing-package:{missing}"
    except EmptyAIResponse:
        # got a (blocked/empty) response — that still means the key authenticated
        return True, "ok"
    except Exception as e:
        cat, friendly = classify_ai_error(e)
        if cat == "credit":
            return True, "credit"
        if cat == "rate_limit":
            return True, "ratelimited"  # key valid, just throttled
        if cat in ("auth",):
            return False, "auth"
        if cat in ("network", "timeout", "server", "overloaded", "content_filter"):
            return False, cat
        if cat == "bad_model":
            return False, "error:" + friendly
        return False, "error:" + friendly

# ── Model discovery ───────────────────────────────────────────────────────────
# Curated fallbacks shown when a live /models fetch fails or returns nothing.
# Chat/vision-capable text models only (no embeddings / audio / image-gen).
STATIC_MODELS = {
    "anthropic": ["claude-sonnet-4-6", "claude-opus-4-7", "claude-haiku-4-5",
                  "claude-3-7-sonnet-latest", "claude-3-5-sonnet-latest"],
    "openai":    ["gpt-4o", "gpt-4o-mini", "gpt-4.1", "gpt-4.1-mini", "o4-mini"],
    "gemini":    ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
    "nvidia":    ["meta/llama-3.1-70b-instruct", "meta/llama-3.1-8b-instruct",
                  "meta/llama-3.1-405b-instruct", "qwen/qwen2.5-72b-instruct",
                  "qwen/qwen3-235b-a22b", "deepseek-ai/deepseek-r1",
                  "nvidia/llama-3.1-nemotron-70b-instruct"],
    "deepseek":  ["deepseek-chat", "deepseek-reasoner"],
    "qwen":      ["qwen-plus", "qwen-max", "qwen-turbo", "qwen-vl-max", "qwen-vl-plus"],
    "azure_openai": ["gpt-4o", "gpt-4o-mini", "gpt-4.1"],
    "ollama":    ["llama3.1", "llama3.2", "qwen2.5", "mistral", "gemma2"],
    # Manus "models" are agent profiles, not chat models (no live /models list).
    "manus":     ["manus-1.6", "manus-1.6-lite", "manus-1.6-max"],
    "groq":      ["llama-3.3-70b-versatile", "llama-3.1-8b-instant",
                  "openai/gpt-oss-120b", "openai/gpt-oss-20b", "qwen/qwen3-32b"],
    "cerebras":  ["llama-3.3-70b", "llama-3.1-8b", "qwen-3-32b", "gpt-oss-120b"],
    "openrouter": ["meta-llama/llama-3.3-70b-instruct:free",
                   "deepseek/deepseek-chat-v3.1:free",
                   "google/gemini-2.0-flash-exp:free",
                   "qwen/qwen3-235b-a22b:free"],
    "mistral":   ["mistral-large-latest", "mistral-small-latest",
                  "open-mistral-nemo", "pixtral-large-latest"],
}

def _is_chat_model_id(provider, mid):
    """Keep chat/VLM models; hide only clearly non-chat families (embeddings,
    rerank, tts/audio, moderation/safety, OCR, and image/video GENERATORS). Uses
    PRECISE markers (e.g. 'diffusion', 'flux') instead of a bare 'image' so that
    vision/VLM CHAT models — like NVIDIA's qwen/qwen3.5-397b-a17b — are retrieved."""
    s = mid.lower()
    bad = ("embedding", "embed", "whisper", "tts", "audio", "speech",
           "moderation", "rerank", "guard", "safety", "transcribe", "ocr",
           "dall-e", "diffusion", "flux", "sdxl", "stable-diffusion",
           "text-to-image", "image-to-image", "code-search")
    if any(b in s for b in bad):
        return False
    if provider == "openai":
        # keep gpt-*/o*-family chat models
        return s.startswith(("gpt-", "o1", "o3", "o4", "chatgpt"))
    return True

def list_models(provider=None, api_key=None, base_url=None, timeout=15):
    """Return (models, source) where source is 'live' or 'static'.
    Fetches the provider's model catalogue with the SELECTED provider's key.
    Falls back to STATIC_MODELS on any error so the dropdown is never empty.
    """
    p = provider or AI_PROVIDER
    cfg = AI_CONFIG.get(p, {})
    key = (api_key or cfg.get("api_key") or "").strip()
    burl = base_url or cfg.get("base_url")
    static = STATIC_MODELS.get(p, [])

    # Manus has no live model list — the "models" are fixed agent profiles.
    if p == "manus":
        return (static, "static")

    def _ok(lst):
        # de-dupe, keep order, drop obvious non-chat ids, cap length
        seen, out = set(), []
        for m in lst:
            m = (m or "").strip()
            if not m or m in seen:
                continue
            if not _is_chat_model_id(p, m):
                continue
            seen.add(m); out.append(m)
        return out

    try:
        if p == "anthropic":
            ids = _ok(_anthropic_models_http(key, timeout))
            return (ids or static), ("live" if ids else "static")

        if p in OPENAI_COMPAT_PROVIDERS:
            ids = _ok(_openai_compat_models_http(burl, key, timeout))
            ids.sort()
            return (ids or static), ("live" if ids else "static")

        if p == "azure_openai":
            ids = _ok(_azure_models_http({**cfg, "api_key": key}, timeout))
            return (ids or static), ("live" if ids else "static")

        if p == "gemini":
            ids = _ok(_gemini_models_http(key, timeout))
            return (ids or static), ("live" if ids else "static")

        if p == "ollama":
            r = requests.get(f"{(burl or 'http://localhost:11434')}/api/tags", timeout=timeout)
            r.raise_for_status()
            ids = _ok([m.get("name") for m in (r.json().get("models") or [])])
            return (ids or static), ("live" if ids else "static")

    except Exception:
        return static, "static"
    return static, "static"


def fetch_projects(pat=None):
    data = _azure_get(f"https://dev.azure.com/{AZURE_ORG}/_apis/projects?api-version=7.0", pat)
    return sorted([p["name"] for p in data.get("value", [])])

def fetch_iterations(project, pat=None):
    """Each returned dict now also carries `start_date`/`finish_date`
    ("YYYY-MM-DD", or "" if that iteration has no dates configured in Azure
    DevOps) — the classification-node API already returns these under
    `attributes.startDate`/`finishDate` per node, just previously discarded
    here. Added so Task Manager can prorate its 170h/month workload
    benchmark against a SPRINT's own real date span, the same way it already
    does for an explicit Date range."""
    url = (f"https://dev.azure.com/{AZURE_ORG}/{project}"
           f"/_apis/wit/classificationnodes/iterations?$depth=10&api-version=7.0")
    data = _azure_get(url, pat)
    out = []
    def _walk(node, prefix):
        name = node.get("name", "")
        path = (prefix + "\\" + name) if prefix else name
        attrs = node.get("attributes") or {}
        out.append({"name": name, "path": path, "id": node.get("identifier", ""),
                    "start_date": (attrs.get("startDate") or "")[:10],
                    "finish_date": (attrs.get("finishDate") or "")[:10]})
        for ch in node.get("children", []) or []:
            _walk(ch, path)
    for child in data.get("children", []) or []:
        _walk(child, project)
    if not out:
        out.append({"name": project, "path": project, "id": data.get("identifier", ""),
                    "start_date": "", "finish_date": ""})
    return out

def _wiql_str(value):
    """Escape a value for interpolation into a WIQL string literal (single
    quotes doubled, same rule SQL-style query languages use). Centralizes the
    pattern that was previously copy-pasted at each WIQL-building call site —
    behavior is unchanged, this only removes the duplication."""
    return (value or "").replace("'", "''")


def fetch_stories_in_iteration(project, iteration_path, pat=None):
    pat = pat or AZURE_PAT
    safe = _wiql_str(iteration_path)
    url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/wiql?api-version=7.0"
    # Order by the sprint board's own stack-rank so the dropdown + generated plan
    # match the Taskboard/Backlog top-to-bottom order (not ascending id). The
    # rank field differs by process template — Scrum uses BacklogPriority, Agile/
    # CMMI use StackRank — so we order by BOTH: the process's UNused field is null
    # for every item, leaving the used one to decide (System.Id is the final tie
    # break). If a process lacks either field the WIQL 400s, so we fall back to
    # the old id ordering rather than failing the whole fetch.
    _base = ("SELECT [System.Id], [System.Title] FROM WorkItems "
             "WHERE [System.WorkItemType] = 'User Story' "
             f"AND [System.IterationPath] = '{safe}' ORDER BY ")
    _board_order = ("[Microsoft.VSTS.Common.StackRank], "
                    "[Microsoft.VSTS.Common.BacklogPriority], [System.Id]")
    r = requests.post(url, json={"query": _base + _board_order}, auth=("", pat), timeout=30)
    if r.status_code != 200:
        r = requests.post(url, json={"query": _base + "[System.Id]"},
                          auth=("", pat), timeout=30)
    if r.status_code != 200:
        raise RuntimeError(f"WIQL query failed (HTTP {r.status_code})")
    # WIQL returns workItems already in the ORDER BY (board) order — this id list
    # is the canonical order we must preserve.
    ids = [w["id"] for w in r.json().get("workItems", [])]
    if not ids:
        return []
    # The batched detail GETs don't guarantee response order (and never across
    # batches), so collect into a map and re-emit in the WIQL/board order.
    by_id = {}
    for i in range(0, len(ids), 200):
        batch = ids[i:i+200]
        burl = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
                f"?ids={','.join(map(str,batch))}&fields=System.Id,System.Title&api-version=7.0")
        br = requests.get(burl, auth=("", pat), timeout=30)
        if br.status_code == 200:
            for w in br.json().get("value", []):
                by_id[w["id"]] = {"id": w["id"],
                                  "title": w["fields"].get("System.Title", "")}
    return [by_id[i] for i in ids if i in by_id]

def board_order_rank(project, story_ids, pat=None):
    # {str(id): index} in Azure sprint-board order (StackRank/BacklogPriority,
    # then Id). Empty when not Azure / PAT-less, no numeric ids, or the WIQL
    # can't order (process lacks the fields) -- caller keeps its own order.
    pat = pat or AZURE_PAT
    ids = [str(i) for i in (story_ids or []) if str(i).isdigit()]
    if not ids or not pat or not AZURE_ORG:
        return {}
    url = ("https://dev.azure.com/" + AZURE_ORG + "/" + str(project) +
           "/_apis/wit/wiql?api-version=7.0")
    q = ("SELECT [System.Id] FROM WorkItems WHERE [System.Id] IN ("
         + ",".join(ids) + ") ORDER BY [Microsoft.VSTS.Common.StackRank], "
         "[Microsoft.VSTS.Common.BacklogPriority], [System.Id]")
    try:
        r = requests.post(url, json={"query": q}, auth=("", pat), timeout=30)
        if r.status_code != 200:
            return {}
        ordered = [str(w["id"]) for w in r.json().get("workItems", [])]
    except Exception:
        return {}
    return {sid: idx for idx, sid in enumerate(ordered)}


def sort_stories_by_board(project, stories, pat=None):
    # Sort story dicts (id + optional sprint) IN PLACE into Azure sprint-board
    # order: numbered sprints first (numeric), then each story by the board's
    # stack rank, then id. Non-Azure / unrankable stories fall back to a stable
    # (sprint, numeric-id) order. Returns the list.
    if not stories:
        return stories
    try:
        rank = board_order_rank(project, [s.get("id") for s in stories], pat)
    except Exception:
        rank = {}
    _big = len(rank)
    def _key(s):
        sp = re.search(r"\d+", str(s.get("sprint", "") or ""))
        sprint_key = (0, int(sp.group())) if sp else (1, 0)
        sid = str(s.get("id", ""))
        id_key = (0, int(sid)) if sid.isdigit() else (1, sid)
        return (sprint_key, rank.get(sid, _big), id_key)
    try:
        stories.sort(key=_key)
    except Exception:
        pass
    return stories


def create_plan_with_sprint_suites(project, name, iteration_path, cb=None, pat=None,
                                   story_ids=None):
    """Create a test plan, then add a requirement-based suite for every User Story
    in the chosen sprint (iteration_path). PAT-only — no AI calls.

    story_ids (optional): restrict the requirement suites to this SUBSET of the
    sprint's User Story ids (as chosen in the Create-plan dialog). None keeps the
    original behaviour — a suite for every story in the sprint. Order still
    follows the sprint board (fetch_stories_in_iteration).
    cb(event, payload) events:
        "plan"     -> {"plan_id": id}
        "stories"  -> {"total": N}
        "suite"    -> {"done": i, "total": N, "story_id": sid, "title": t, "ok": bool}
        "done"     -> {"plan_id": id, "story_ids": [...], "created": k, "skipped": s, "failed": f}
    Returns (plan_id, story_ids).
    """
    pat = pat or AZURE_PAT
    cb = cb or (lambda *a, **k: None)

    # 1) create the plan
    plan_id = create_test_plan(project, name, iteration_path, pat)
    cb("plan", {"plan_id": plan_id})

    # 2) User Stories in the sprint — filtered to the caller's selection if given.
    stories = fetch_stories_in_iteration(project, iteration_path, pat)
    if story_ids is not None:
        _want = {str(s) for s in story_ids}
        stories = [s for s in stories if str(s["id"]) in _want]
    total = len(stories)
    cb("stories", {"total": total})

    story_ids = [s["id"] for s in stories]
    if total == 0:
        cb("done", {"plan_id": plan_id, "story_ids": [], "created": 0, "skipped": 0, "failed": 0})
        return plan_id, []

    # 3) requirement suite per story
    root = _get_root_suite_id(project, plan_id, pat)
    created = skipped = failed = 0
    for i, s in enumerate(stories, 1):
        sid = s["id"]; title = s.get("title", "")
        ok = True
        try:
            res = create_requirement_suite(project, plan_id, sid, root, pat)
            if res is None:
                skipped += 1   # already existed
            else:
                created += 1
        except Exception:
            failed += 1; ok = False
        cb("suite", {"done": i, "total": total, "story_id": sid, "title": title, "ok": ok})

    cb("done", {"plan_id": plan_id, "story_ids": story_ids,
                "created": created, "skipped": skipped, "failed": failed})
    return plan_id, story_ids


def fetch_test_plans(project, pat=None):
    data = _azure_get(f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/testplan/plans?api-version=7.0", pat)
    return [{"id": p["id"], "name": p["name"]} for p in data.get("value", [])]

def create_test_plan(project, name, iteration_path, pat=None):
    pat = pat or AZURE_PAT
    url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/testplan/plans?api-version=7.0"
    body = {"name": name, "iteration": iteration_path}
    try:
        r = requests.post(url, json=body, auth=("", pat), timeout=30)
    except requests.exceptions.ConnectionError:
        raise RuntimeError("No internet connection or Azure DevOps is unreachable.")
    if r.status_code == 401:
        raise RuntimeError("Authentication failed (401). Check your PAT.")
    if r.status_code == 403:
        raise RuntimeError("Access denied (403). PAT needs Test Management (read & write).")
    if r.status_code == 400:
        raise RuntimeError(f"Invalid request (400). The iteration path may be wrong.\n{r.text[:160]}")
    r.raise_for_status()
    return r.json()["id"]

def _get_root_suite_id(project, plan_id, pat=None):
    data = _azure_get(f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/testplan/plans/{plan_id}?api-version=7.0", pat)
    return data.get("rootSuite", {}).get("id")


def sprint_report_data(project, iteration_path, pat=None):
    """Pull the User Stories AND Bugs in a sprint iteration with their states, for
    the Sprint (closure) Report. Returns:
        {"stories":[{id,title,state,parent,epic}], "bugs":[{id,title,state,tags}],
         "story_by_state":{state:count}, "bug_by_state":{state:count},
         "total_bugs":int, "regression_bugs":int, "sprint_bugs":int}
    'regression_bugs' = bugs tagged with 'Regression' (System.Tags); the rest are
    counted as sprint bugs. Best-effort: any failed call yields empty data."""
    from collections import Counter
    pat = pat or AZURE_PAT
    safe = _wiql_str(iteration_path)

    def _ids(wit):
        wiql = {"query": ("SELECT [System.Id] FROM WorkItems WHERE "
                          f"[System.WorkItemType] = '{wit}' AND "
                          f"[System.IterationPath] = '{safe}' ORDER BY [System.Id]")}
        url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/wiql?api-version=7.0"
        try:
            r = _az_session().post(url, json=wiql, auth=("", pat), timeout=30)
            return [w["id"] for w in r.json().get("workItems", [])] if r.status_code == 200 else []
        except Exception:
            return []

    def _items(ids, fields):
        out = []
        for i in range(0, len(ids), 200):
            b = ids[i:i + 200]
            url = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
                   f"?ids={','.join(map(str, b))}&fields={fields}&api-version=7.0")
            try:
                r = _az_session().get(url, auth=("", pat), timeout=30)
                if r.status_code == 200:
                    out += r.json().get("value", [])
            except Exception:
                pass
        return out

    stories = []
    for w in _items(_ids("User Story"),
                    "System.Id,System.Title,System.State,System.Parent"):
        f = w.get("fields", {})
        stories.append({"id": w["id"], "title": f.get("System.Title", ""),
                        "state": f.get("System.State", "Unknown"),
                        "parent": f.get("System.Parent")})

    # Resolve each story's Epic by walking the parent chain (Story → Feature →
    # Epic). We cache every work item we fetch and climb at most a few hops; if no
    # Epic is found we fall back to the nearest ancestor's title (usually the
    # Feature), and to "" when the story has no parent at all. Used to group the
    # report by epic (e.g. Admin Portal / Subscriber Portal).
    _wi = {}  # id -> {"type","title","parent"}

    def _fetch_wi(ids):
        ids = [i for i in ids if i and i not in _wi]
        if not ids:
            return
        for w in _items(ids, "System.Id,System.Title,System.WorkItemType,System.Parent"):
            f = w.get("fields", {})
            _wi[w["id"]] = {"type": f.get("System.WorkItemType", ""),
                            "title": f.get("System.Title", ""),
                            "parent": f.get("System.Parent")}

    try:
        pending = [s["parent"] for s in stories if s.get("parent")]
        for _hop in range(5):
            _fetch_wi(pending)
            pending = [v["parent"] for v in _wi.values()
                       if v.get("parent") and v["parent"] not in _wi]
            if not pending:
                break

        def _epic_of(pid):
            seen, cur, nearest = set(), pid, ""
            while cur and cur in _wi and cur not in seen:
                seen.add(cur)
                node = _wi[cur]
                if (node.get("type") or "").strip().lower() == "epic":
                    return node.get("title") or ""
                nearest = nearest or node.get("title") or ""
                cur = node.get("parent")
            return nearest

        for s in stories:
            s["epic"] = _epic_of(s.get("parent"))
    except Exception:
        for s in stories:
            s.setdefault("epic", "")

    bugs = []
    for w in _items(_ids("Bug"),
                    "System.Id,System.Title,System.State,System.Tags"):
        f = w.get("fields", {})
        bugs.append({"id": w["id"], "title": f.get("System.Title", ""),
                     "state": f.get("System.State", "Unknown"),
                     "tags": f.get("System.Tags", "") or ""})
    reg = sum(1 for b in bugs if "regression" in (b["tags"] or "").lower())
    return {"stories": stories, "bugs": bugs,
            "story_by_state": dict(Counter(s["state"] for s in stories)),
            "bug_by_state": dict(Counter(b["state"] for b in bugs)),
            "total_bugs": len(bugs), "regression_bugs": reg,
            "sprint_bugs": len(bugs) - reg}


def sprint_summary(project, plan_id, pat=None):
    """Public entry point — thin wrapper around _sprint_summary_impl().

    Reported bug: the Sprint Summary modal showed a bare, useless
    "Could not load summary: None" error with zero diagnostic value. That
    exact text comes from `str(ex)` being literally the string "None" —
    which happens when something deep in the pipeline raises, e.g.,
    `KeyError(None)` (a plain dict's `__str__`/`BaseException.__str__`
    renders a single `None` arg as the 4-character word "None", not an
    empty string). Nothing in this function's own body constructs an
    exception with a None argument, and every raw call it doesn't wrap
    itself already goes through `_azure_get` (which always raises with a
    real message) or is guarded by its own try/except — so the exact
    original raise site couldn't be pinned down without a live repro.
    Rather than leave that dead end in place, this wrapper guarantees NO
    caller can ever see a bare/None-ish message again: any exception from
    _sprint_summary_impl is re-raised with a real message, falling back to
    the exception's type name when the message itself is empty or "None"."""
    try:
        return _sprint_summary_impl(project, plan_id, pat)
    except Exception as ex:
        msg = str(ex).strip()
        if not msg or msg.lower() == "none":
            msg = f"{type(ex).__name__} — no further detail was available from Azure DevOps."
        raise RuntimeError(msg) from ex

def _sprint_summary_impl(project, plan_id, pat=None):
    """Build a status summary for the sprint behind a test plan.

    Reads the plan's iteration, finds every User Story in that iteration, tallies
    their states, and counts the test cases mapped to each story's suite.
    Returns a dict:
        {
          "plan_name": str, "iteration": str,
          "total_stories": int,
          "by_state": {state: count, ...},
          "stories": [{"id","title","state","test_cases","assigned_to"}, ...],
          "total_test_cases": int,
        }
    """
    pat = pat or AZURE_PAT
    # 1) plan → iteration path + name
    plan = _azure_get(f"https://dev.azure.com/{AZURE_ORG}/{project}"
                      f"/_apis/testplan/plans/{plan_id}?api-version=7.0", pat)
    plan_name = plan.get("name", str(plan_id))
    iteration = plan.get("iteration") or ""

    # 2) all User Stories in that iteration (id + title + state)
    #
    # "Assigned to Tester" here means the same custom picklist field the
    # Sprint Plan screen's "Assign to Testers" button writes (see
    # assign_testers() below) — a QA ownership field, distinct from
    # System.AssignedTo (the dev/PM assignee). Its reference name isn't
    # fixed (it's a Custom.* field that varies per project's process
    # template), so it has to be resolved by label first, same as
    # assign_testers() does.
    tester_field_ref = resolve_field_ref(project, "Assigned To Tester", pat)
    stories = []
    if iteration:
        safe = _wiql_str(iteration)
        wiql = {"query": ("SELECT [System.Id] FROM WorkItems "
                          "WHERE [System.WorkItemType] = 'User Story' "
                          f"AND [System.IterationPath] = '{safe}' ORDER BY [System.Id]")}
        url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/wiql?api-version=7.0"
        r = requests.post(url, json=wiql, auth=("", pat), timeout=30)
        ids = [w["id"] for w in r.json().get("workItems", [])] if r.status_code == 200 else []
        _fields = "System.Id,System.Title,System.State"
        if tester_field_ref:
            _fields += f",{tester_field_ref}"
        for i in range(0, len(ids), 200):
            batch = ids[i:i+200]
            burl = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
                    f"?ids={','.join(map(str,batch))}"
                    f"&fields={_fields}"
                    f"&api-version=7.0")
            br = requests.get(burl, auth=("", pat), timeout=30)
            if br.status_code == 200:
                for w in br.json().get("value", []):
                    f = w.get("fields", {})
                    # Assigned To Tester is USUALLY a plain picklist string
                    # (see tester_allowed_values()'s allowedValues) — but it
                    # turns out some projects' process templates define it as
                    # an IDENTITY field instead (an Azure AD user picker, same
                    # shape as System.AssignedTo: {"displayName":...,
                    # "uniqueName":...}), confirmed live — treating it as a
                    # bare string crashed with "'dict' object has no
                    # attribute 'strip'" the moment a story actually had one
                    # assigned. Handle both shapes rather than assuming one.
                    _tv = f.get(tester_field_ref) if tester_field_ref else None
                    if isinstance(_tv, dict):
                        tester_name = (_tv.get("displayName") or "").strip()
                    else:
                        tester_name = (str(_tv).strip() if _tv else "")
                    tester_name = tester_name or "Unassigned"
                    stories.append({"id": w["id"],
                                    "title": f.get("System.Title", ""),
                                    "state": f.get("System.State", "Unknown"),
                                    "assigned_to": tester_name})

    # 3) test-case counts per story (via that story's suite in the plan)
    story_ids = [s["id"] for s in stories]
    smap = {}
    if story_ids:
        try:
            smap = discover_suites_for_stories(project, plan_id, set(story_ids),
                                               create_missing=False)
        except Exception:
            smap = {}
    # Count test cases per story CONCURRENTLY — one suite fetch per story serially
    # made the sprint summary crawl on big sprints.
    import concurrent.futures as _cf
    def _count_story(s):
        suite_id = smap.get(s["id"])
        if not suite_id:
            return s["id"], 0
        try:
            return s["id"], len(fetch_test_cases_for_suite(project, plan_id, suite_id))
        except Exception:
            return s["id"], 0
    total_tc = 0
    if stories:
        with _cf.ThreadPoolExecutor(max_workers=min(16, len(stories))) as _ex:
            _counts = dict(_ex.map(_count_story, stories))
        for s in stories:
            s["test_cases"] = _counts.get(s["id"], 0)
            total_tc += s["test_cases"]

    # 4) tally states
    by_state = {}
    for s in stories:
        by_state[s["state"]] = by_state.get(s["state"], 0) + 1

    sort_stories_by_board(project, stories)
    return {
        "plan_name": plan_name,
        "iteration": iteration,
        "total_stories": len(stories),
        "by_state": by_state,
        "stories": stories,
        "total_test_cases": total_tc,
        "project": project,
        "org": AZURE_ORG,
    }

def create_requirement_suite(project, plan_id, story_id, root_suite_id=None, pat=None):
    pat = pat or AZURE_PAT
    if root_suite_id is None:
        root_suite_id = _get_root_suite_id(project, plan_id, pat)
    if not root_suite_id:
        raise RuntimeError("Could not find the plan's root suite.")
    url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/testplan/Plans/{plan_id}/suites?api-version=7.1"
    body = {"suiteType": "requirementTestSuite", "name": str(story_id),
            "requirementId": int(story_id), "parentSuite": {"id": int(root_suite_id)}}
    r = requests.post(url, json=body, auth=("", pat), timeout=30)
    if r.status_code in (200, 201):
        data = r.json()
        if isinstance(data.get("value"), list) and data["value"]:
            return data["value"][0].get("id")
        return data.get("id")
    if r.status_code == 400 and "already" in r.text.lower():
        return None
    raise RuntimeError(f"Create suite failed (HTTP {r.status_code})")


# ═══════════════════════════════════════════════════════════════════════════════
#  JSON parsing (robust for Qwen/DeepSeek quirks)
# ═══════════════════════════════════════════════════════════════════════════════
def _balanced_json_fragment(s):
    """From the first '{' or '[', return the balanced JSON fragment, ignoring
    brackets inside strings, and CLOSING a truncated tail (open string + open
    brackets) — so a response cut off at max_tokens is still recoverable."""
    starts = [i for i in (s.find("{"), s.find("[")) if i >= 0]
    if not starts:
        return None
    start = min(starts)
    pairs = {"}": "{", "]": "["}
    closer = {"{": "}", "[": "]"}
    stack, in_str, esc = [], False, False
    for i in range(start, len(s)):
        ch = s[i]
        if esc:
            esc = False; continue
        if ch == "\\":
            esc = True; continue
        if ch == '"':
            in_str = not in_str; continue
        if in_str:
            continue
        if ch in "{[":
            stack.append(ch)
        elif ch in "}]":
            if stack and stack[-1] == pairs[ch]:
                stack.pop()
                if not stack:
                    return s[start:i + 1]
            else:
                break
    if stack:   # truncated mid-structure → close what's still open
        frag = s[start:] + ('"' if in_str else "")
        return frag + "".join(closer[c] for c in reversed(stack))
    return None


def parse_json_robust(raw):
    if raw is None:
        raise ValueError("AI returned an empty response (None)")
    raw = str(raw)
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE).strip()
    if not raw:
        raise ValueError("AI returned an empty response")
    # strict=False tolerates literal newlines/tabs inside string values, which LLMs
    # emit constantly and which json.loads otherwise rejects with a control-char error.
    try: return json.loads(raw, strict=False)
    except json.JSONDecodeError: pass
    def _repair(s):
        s = re.sub(r"'([^'\n]*?)'(\s*[:,\]}])", r'"\1"\2', s)
        s = re.sub(r"([:,\[{]\s*)'([^'\n]*?)'", r'\1"\2"', s)
        out, in_str, esc = [], False, False
        for ch in s:
            if esc: out.append(ch); esc = False; continue
            if ch == "\\": out.append(ch); esc = True; continue
            if ch == '"': in_str = not in_str; out.append(ch); continue
            if in_str and ch == "\n": out.append("\\n"); continue
            if in_str and ch == "\t": out.append("\\t"); continue
            if in_str and ch == "\r": continue
            out.append(ch)
        return "".join(out)
    # largest balanced object/array (handles nesting + truncation), raw then repaired
    frag = _balanced_json_fragment(raw)
    for variant in ([frag, _repair(frag)] if frag else []):
        try: return json.loads(variant, strict=False)
        except Exception: continue
    # legacy greedy-array fallback
    m = re.search(r"\[.*\]", raw, re.DOTALL)
    candidate = m.group(0) if m else raw
    try: return json.loads(_repair(candidate), strict=False)
    except Exception: pass
    objs = re.findall(r"\{[^{}]+\}", candidate, re.DOTALL)
    out = []
    for o in objs:
        for variant in (o, _repair(o)):
            try: out.append(json.loads(variant, strict=False)); break
            except Exception: continue
    if out: return out
    raise ValueError(f"Cannot parse JSON:\n{raw[:300]}")


# ═══════════════════════════════════════════════════════════════════════════════
#  AZURE SDK CONNECTION (work items, test cases)
# ═══════════════════════════════════════════════════════════════════════════════
_wit_client = None
_test_client = None

def connect_azure_sdk(project):
    """Initialize the azure-devops SDK clients. Returns (wit_client, test_client)."""
    global _wit_client, _test_client
    from azure.devops.connection import Connection
    from msrest.authentication import BasicAuthentication
    # The SDK path (every work-item create/update/read) never goes through
    # _azure_get, so it needs its own guard — otherwise an unconfigured org
    # builds a base_url of "https://dev.azure.com/" and fails deep inside the SDK.
    org_url = f"https://dev.azure.com/{_require_org()}"
    creds = BasicAuthentication("", AZURE_PAT)
    conn = Connection(base_url=org_url, creds=creds)
    _wit_client  = conn.clients.get_work_item_tracking_client()
    _test_client = conn.clients.get_test_client()
    return _wit_client, _test_client


# ── Assign-to-tester (identity / picklist field on the sprint's user stories) ──
_FIELD_REF_CACHE = {}

def resolve_field_ref(project, label, pat=None):
    """Find a field's reference name by its display label (e.g. 'Assigned To
    Tester'), so callers don't need to know the internal Custom.* name."""
    key = (project, (label or "").strip().lower())
    if key in _FIELD_REF_CACHE:
        return _FIELD_REF_CACHE[key]
    url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/fields?api-version=7.0"
    ref = None
    try:
        data = _azure_get(url, pat)
        for f in data.get("value", []) or []:
            if (f.get("name", "") or "").strip().lower() == key[1]:
                ref = f.get("referenceName")
                break
    except Exception:
        ref = None
    _FIELD_REF_CACHE[key] = ref
    return ref

def tester_allowed_values(project, field_ref, wit_type="User Story", pat=None):
    """The field's allowed values (the Azure 'list') for a work item type, or []."""
    try:
        url = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitemtypes/"
               f"{wit_type}/fields/{field_ref}?api-version=7.0")
        data = _azure_get(url, pat)
        return list(data.get("allowedValues", []) or [])
    except Exception:
        return []

def _match_identity(name, allowed):
    """Map a resource name to one allowed value. Returns (value, error_or_None).
    Exact (case-insensitive) > startswith > contains; ambiguous → error."""
    n = (name or "").strip().lower()
    if not n:
        return None, "no name"
    for pred in (lambda a: a.strip().lower() == n,
                 lambda a: a.strip().lower().startswith(n),
                 lambda a: n in a.strip().lower()):
        hits = [a for a in allowed if pred(a)]
        if len(hits) == 1:
            return hits[0], None
        if len(hits) > 1:
            return None, f"'{name}' matches several testers ({', '.join(hits[:4])})"
    return None, f"no match for '{name}' in the Assigned To Tester list"

def assign_tester(project, work_item_id, value, field_ref):
    """Write the identity/picklist field on one work item. Raises on failure."""
    from azure.devops.v7_0.work_item_tracking.models import JsonPatchOperation
    if _wit_client is None:
        connect_azure_sdk(project)
    patch = [JsonPatchOperation(op="add", path=f"/fields/{field_ref}", value=value)]
    _wit_client.update_work_item(patch, id=int(work_item_id))

def assign_testers(project, assignments, field_label="Assigned To Tester", cb=None):
    """assignments = [{'id': story_id, 'name': resource_name}, …].
    Resolves the field by its label, matches each name to the field's allowed list
    (readable error when it doesn't match), and writes it.
    Returns {'ok': n, 'field': ref, 'errors': [str, …]}."""
    cb = cb or (lambda *a, **k: None)
    field_ref = resolve_field_ref(project, field_label)
    if not field_ref:
        return {"ok": 0, "field": None,
                "errors": [f"No field named '{field_label}' exists in this project."]}
    allowed = tester_allowed_values(project, field_ref)
    ok, errors = 0, []
    for a in assignments:
        sid, name = a.get("id"), (a.get("name") or "").strip()
        if not name:
            errors.append(f"Story {sid}: no assignee."); continue
        if allowed:
            value, err = _match_identity(name, allowed)
            if err:
                errors.append(f"Story {sid}: {err}."); continue
        else:
            value = name  # no enumerable list — let Azure resolve / reject it
        try:
            assign_tester(project, sid, value, field_ref)
            ok += 1
            cb(f"Assigned {value} → story {sid}", "ok")
        except Exception as e:
            msg = str(e)
            if any(t in msg for t in ("resolve", "TF401", "TF51", "not a valid")) \
               or "does not" in msg.lower():
                errors.append(f"Story {sid}: '{name}' isn't a valid Assigned To Tester.")
            else:
                errors.append(f"Story {sid}: {msg[:120]}")
    return {"ok": ok, "field": field_ref, "errors": errors}


def fetch_stories_in_plan(project, plan_id, pat=None):
    """User stories referenced by a test plan's requirement-based suites —
    independent of any sprint/iteration (so it works for plans that have no
    iteration set). Returns [{"id": int, "title": str}] sorted by id."""
    pat = pat or AZURE_PAT
    url = (f"https://dev.azure.com/{AZURE_ORG}/{project}"
           f"/_apis/testplan/Plans/{plan_id}/Suites?api-version=7.0&$expand=true")
    try:
        resp = requests.get(url, auth=("", pat), timeout=30)
    except requests.exceptions.ConnectionError:
        raise RuntimeError("No internet connection or Azure DevOps is unreachable.")
    if resp.status_code != 200:
        raise RuntimeError(f"Could not fetch suites (HTTP {resp.status_code})")
    ids = set()
    for suite in resp.json().get("value", []):
        rid = suite.get("requirementId")
        if rid:
            try:
                ids.add(int(rid))
            except Exception:
                pass
        else:  # QA-Studio-created suites may be named "<id>" or "<id>: title"
            try:
                ids.add(int(str(suite.get("name", "")).split(":")[0].strip()))
            except (ValueError, IndexError):
                pass
    ids = sorted(ids)
    titles, sprints = {}, {}
    for i in range(0, len(ids), 200):
        batch = ids[i:i + 200]
        burl = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
                f"?ids={','.join(map(str, batch))}"
                f"&fields=System.Id,System.Title,System.IterationPath&api-version=7.0")
        try:
            br = requests.get(burl, auth=("", pat), timeout=30)
            if br.status_code == 200:
                for w in br.json().get("value", []):
                    f = w.get("fields", {})
                    titles[int(w["id"])] = f.get("System.Title", "")
                    sprints[int(w["id"])] = (f.get("System.IterationPath", "") or "").split("\\")[-1]
        except Exception:
            pass
    return [{"id": sid, "title": titles.get(sid, ""), "sprint": sprints.get(sid, "")}
            for sid in ids]


_project_members_cache = {}   # project -> [{"name","email"}, ...], process-lifetime cache


def fetch_project_members(project, pat=None, force=False):
    """Display name + email for everyone on this Azure DevOps project's teams —
    powers the recipient picker on Setup / Sprint Plan / Regression Plan (pick
    real people instead of typing addresses from memory). Uses the Core API's
    teams + team-members endpoints, which work with the same PAT already used
    for everything else in the app (no extra scope needed). Cached per project
    for the life of the process; pass force=True to bypass the cache.
    """
    if not force and project in _project_members_cache:
        return _project_members_cache[project]
    pat = pat or AZURE_PAT
    try:
        teams = _azure_get(
            f"https://dev.azure.com/{AZURE_ORG}/_apis/projects/{project}/teams"
            f"?api-version=7.0", pat).get("value", [])
    except Exception as _ex:
        # Silent [] here is why "the recipient dropdown lists nobody" was
        # undiagnosable: the teams endpoint needs the PAT to carry
        # **Project and Team (Read)** scope, which the rest of the app never
        # exercises — so a PAT that works everywhere else still yields an empty
        # picker, with no error anywhere. Log the real reason; still fail-soft.
        try:
            import diag_log
            diag_log.log("engine.fetch_project_members.teams", _ex)
        except Exception:
            pass
        teams = []
    if not teams:
        try:
            import diag_log
            diag_log.log_warn(
                "engine.fetch_project_members",
                f"no teams returned for project {project!r} (org={AZURE_ORG!r}) — "
                f"the recipient picker will be empty. Most often the PAT is missing "
                f"the 'Project and Team (Read)' scope.")
        except Exception:
            pass

    import concurrent.futures as _cf

    def _team_members(team):
        tid = team.get("id")
        if not tid:
            return []
        try:
            return _azure_get(
                f"https://dev.azure.com/{AZURE_ORG}/_apis/projects/{project}/teams/"
                f"{tid}/members?api-version=7.0", pat).get("value", [])
        except Exception:
            return []

    out = {}
    if teams:
        with _cf.ThreadPoolExecutor(max_workers=min(8, len(teams))) as ex:
            for members in ex.map(_team_members, teams):
                for m in members:
                    ident = m.get("identity", {}) or {}
                    email = (ident.get("uniqueName") or "").strip()
                    name = (ident.get("displayName") or "").strip()
                    # Skip group identities (e.g. "[Project]\Team") which have no
                    # real "@" address to send mail to.
                    if email and "@" in email:
                        out[email.lower()] = {"name": name or email, "email": email}
    result = sorted(out.values(), key=lambda r: r["name"].lower())
    _project_members_cache[project] = result
    return result


# ═══════════════════════════════════════════════════════════════════════════════
#  TASK MANAGER — per-user task workload report (Original Estimate / Completed
#  Work, scoped to a sprint/iteration, matching the app's existing convention
#  of iteration-based scoping rather than raw date ranges — see Sprint Report /
#  Sprint Plan) + bulk "create a child Task under each selected story" tool.
# ═══════════════════════════════════════════════════════════════════════════════
def fetch_user_task_stats(project, iteration_path=None, assignee=None,
                          start_date=None, end_date=None, pat=None):
    """Every 'Task' work item assigned to `assignee` (an email — matched as a
    substring of the AssignedTo field, case-insensitive, since the plain
    (non-$expand) work-items endpoint returns AssignedTo as a "Display Name
    <email>" string rather than a structured identity object), scoped EITHER
    by `iteration_path` (a sprint, the original behavior) OR — when no sprint
    is given — by a [start_date, end_date] calendar range (both "YYYY-MM-DD"
    strings), matched against the task's [System.ChangedDate]. The Task
    Manager screen's Calculate now accepts either a sprint or a date range;
    this mirrors that at the query level rather than forcing iteration-only
    scoping. `iteration_path` takes precedence if both happen to be passed.
    Returns:
        {"iteration": str, "date_range": str, "assignee": str,
         "tasks": [{"id","title","state","parent_id","parent_title",
                    "original_estimate","completed_work"}, ...],
         "total_original_estimate": float, "total_completed_work": float,
         "count": int}
    Best-effort: any failed call yields an empty result rather than raising —
    same fail-soft convention as sprint_report_data."""
    pat = pat or AZURE_PAT
    if iteration_path:
        safe = _wiql_str(iteration_path)
        where = f"[System.IterationPath] = '{safe}'"
        date_range_label = ""
    else:
        from datetime import date as _date, timedelta as _timedelta
        # Half-open range [start, end+1) on ChangedDate so the END day is
        # covered in full rather than cut off at its midnight instant.
        try:
            end_next = (_date.fromisoformat(end_date) + _timedelta(days=1)).isoformat()
        except Exception:
            end_next = end_date
        where = (f"[System.ChangedDate] >= '{_wiql_str(start_date)}' AND "
                 f"[System.ChangedDate] < '{_wiql_str(end_next)}'")
        date_range_label = f"{start_date} – {end_date}"
    wiql = {"query": ("SELECT [System.Id] FROM WorkItems WHERE "
                      "[System.WorkItemType] = 'Task' AND "
                      f"{where} ORDER BY [System.Id]")}
    url = f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/wiql?api-version=7.0"
    try:
        r = _az_session().post(url, json=wiql, auth=("", pat), timeout=30)
        ids = [w["id"] for w in r.json().get("workItems", [])] if r.status_code == 200 else []
    except Exception:
        ids = []
    if not ids:
        return {"iteration": iteration_path or "", "date_range": date_range_label,
                "assignee": assignee, "tasks": [],
                "total_original_estimate": 0.0, "total_completed_work": 0.0, "count": 0}

    fields = ("System.Id,System.Title,System.State,System.AssignedTo,System.Parent,"
              "Microsoft.VSTS.Scheduling.OriginalEstimate,"
              "Microsoft.VSTS.Scheduling.CompletedWork")
    raw = []
    for i in range(0, len(ids), 200):
        b = ids[i:i + 200]
        burl = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
               f"?ids={','.join(map(str, b))}&fields={fields}&api-version=7.0")
        try:
            br = _az_session().get(burl, auth=("", pat), timeout=30)
            if br.status_code == 200:
                raw += br.json().get("value", [])
        except Exception:
            pass

    needle = (assignee or "").strip().lower()
    matched = []
    for w in raw:
        f = w.get("fields", {})
        assigned = str(f.get("System.AssignedTo", "") or "")
        if needle and needle not in assigned.lower():
            continue
        matched.append(w)

    # Resolve each matched task's parent story title (one extra batch fetch).
    parent_ids = sorted({f["fields"].get("System.Parent") for f in
                         [w for w in matched] if w.get("fields", {}).get("System.Parent")})
    parent_titles = {}
    for i in range(0, len(parent_ids), 200):
        b = parent_ids[i:i + 200]
        purl = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workitems"
               f"?ids={','.join(map(str, b))}&fields=System.Id,System.Title&api-version=7.0")
        try:
            pr = _az_session().get(purl, auth=("", pat), timeout=30)
            if pr.status_code == 200:
                for pw in pr.json().get("value", []):
                    parent_titles[pw["id"]] = pw.get("fields", {}).get("System.Title", "")
        except Exception:
            pass

    tasks, tot_est, tot_comp = [], 0.0, 0.0
    for w in matched:
        f = w.get("fields", {})
        est = float(f.get("Microsoft.VSTS.Scheduling.OriginalEstimate") or 0)
        comp = float(f.get("Microsoft.VSTS.Scheduling.CompletedWork") or 0)
        pid = f.get("System.Parent")
        tasks.append({"id": w["id"], "title": f.get("System.Title", ""),
                      "state": f.get("System.State", "Unknown"),
                      "parent_id": pid, "parent_title": parent_titles.get(pid, ""),
                      "original_estimate": est, "completed_work": comp})
        tot_est += est
        tot_comp += comp
    tasks.sort(key=lambda t: t["id"])
    return {"iteration": iteration_path or "", "date_range": date_range_label,
            "assignee": assignee, "tasks": tasks,
            "total_original_estimate": tot_est, "total_completed_work": tot_comp,
            "count": len(tasks)}


def create_child_tasks(project, items, pat=None, cb=None):
    """Create one 'Task' work item per entry, as a CHILD of the given story,
    assigned to the given person. items = [{"story_id": int, "title": str,
    "assigned_to": str, "due_date": "YYYY-MM-DD" (optional),
    "original_estimate": str/float (optional), "completed_work": str/float
    (optional)}, ...] (assigned_to is an email, matched the same way
    assign_tester's identity fields are — Azure resolves a bare email fine).
    Numeric fields are best-effort parsed — an unparsable/blank value is
    simply omitted rather than failing the whole item, since some projects'
    process rules only require Due Date, not the estimate fields (seen live:
    "TF401320: Rule Error for field Due Date. Error code: Required,
    InvalidEmpty" when Due Date was omitted entirely).
    Returns {"ok": int, "created": [{"story_id","task_id","title"}, ...],
    "errors": [str, ...]}. Best-effort per item: one failure doesn't stop the
    rest, same convention as assign_testers."""
    from azure.devops.v7_0.work_item_tracking.models import JsonPatchOperation
    pat = pat or AZURE_PAT
    cb = cb or (lambda *a, **k: None)
    if _wit_client is None:
        connect_azure_sdk(project)
    ok, created, errors = 0, [], []
    for it in items:
        sid = it.get("story_id")
        title = (it.get("title") or "").strip()
        assignee = (it.get("assigned_to") or "").strip()
        due = (it.get("due_date") or "").strip()
        if not sid or not title:
            errors.append(f"Story {sid}: missing title — skipped.")
            continue
        patch = [JsonPatchOperation(op="add", path="/fields/System.Title", value=title)]
        if assignee:
            patch.append(JsonPatchOperation(op="add", path="/fields/System.AssignedTo",
                                            value=assignee))
        if due:
            # Azure's work-item date fields want a full ISO datetime, not a
            # bare date — a plain "YYYY-MM-DD" is silently rejected by some
            # process templates' validators.
            patch.append(JsonPatchOperation(
                op="add", path="/fields/Microsoft.VSTS.Scheduling.DueDate",
                value=f"{due}T00:00:00Z"))
        for field_ref, key in (("Microsoft.VSTS.Scheduling.OriginalEstimate", "original_estimate"),
                              ("Microsoft.VSTS.Scheduling.CompletedWork", "completed_work")):
            raw = it.get(key)
            try:
                val = float(raw) if raw not in (None, "") else None
            except (TypeError, ValueError):
                val = None
            if val is not None:
                patch.append(JsonPatchOperation(op="add", path=f"/fields/{field_ref}", value=val))
        patch.append(JsonPatchOperation(op="add", path="/relations/-", value={
            "rel": "System.LinkTypes.Hierarchy-Reverse",
            "url": f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workItems/{sid}"}))
        try:
            wi = _wit_client.create_work_item(patch, project=project, type="Task")
            ok += 1
            created.append({"story_id": sid, "task_id": wi.id, "title": title})
            cb(f"Created task #{wi.id} \"{title}\" under story {sid}"
              + (f" → {assignee}" if assignee else ""), "ok")
        except Exception as e:
            msg = str(e)[:160]
            errors.append(f"Story {sid} (\"{title[:40]}\"): {msg}")
            cb(f"Story {sid}: failed — {msg}", "error")
    return {"ok": ok, "created": created, "errors": errors}


def discover_suites_for_stories(project, plan_id, story_ids, create_missing=True):
    """Match each story to a suite in the plan; auto-create requirement suites for
    any story without one (unless create_missing=False). Returns {story_id: suite_id}."""
    url = (f"https://dev.azure.com/{AZURE_ORG}/{project}"
           f"/_apis/testplan/Plans/{plan_id}/Suites?api-version=7.0&$expand=true")
    resp = requests.get(url, auth=("", AZURE_PAT), timeout=30)
    if resp.status_code != 200:
        raise RuntimeError(f"Could not fetch suites (HTTP {resp.status_code})")
    suites = resp.json().get("value", [])
    story_map = {}
    for suite in suites:
        suite_id = suite.get("id")
        req_id = suite.get("requirementId")
        if req_id and int(req_id) in story_ids:
            story_map[int(req_id)] = suite_id; continue
        name = suite.get("name", "")
        try:
            cand = int(name.split(":")[0].strip())
            if cand in story_ids:
                story_map[cand] = suite_id
        except (ValueError, IndexError):
            pass
    missing = [sid for sid in story_ids if sid not in story_map]
    if missing and create_missing:
        try: root_id = _get_root_suite_id(project, plan_id)
        except Exception: root_id = None
        for sid in missing:
            try:
                new_suite = create_requirement_suite(project, plan_id, sid, root_id)
                rr = requests.get(url, auth=("", AZURE_PAT), timeout=30)
                if rr.status_code == 200:
                    found = False
                    for s in rr.json().get("value", []):
                        if s.get("requirementId") and int(s["requirementId"]) == sid:
                            story_map[sid] = s.get("id"); found = True; break
                    if not found and new_suite:
                        story_map[sid] = new_suite
            except Exception:
                pass
    return story_map


def fetch_stories(story_ids):
    """Fetch work items (user stories) with title + acceptance criteria."""
    stories = []
    for sid in story_ids:
        try:
            wi = _wit_client.get_work_item(sid, expand="Relations")
            stories.append(wi)
        except Exception:
            pass
    return stories


# Matches an Azure DevOps work-item link as Azure's rich-text editor actually
# emits it when you insert a "link to work item" inside a field (the common
# way an AC field says "same as <link>" instead of restating the rules) —
# .../_workitems/edit/<id>, optionally with a trailing slash/query string —
# plus the legacy on-prem TFS #_a=edit&id=<id> form, as a fallback.
_AC_WI_LINK_RE = re.compile(r'_workitems/edit/(\d+)|[?&#]id=(\d+)', re.I)


def _resolve_ac_links(criteria, project, cb=None, cache=None, current_id=None):
    """Some stories' acceptance criteria is just a link to another story
    ("same as <link>") instead of restating the rules — left alone,
    generate_titles/generate_steps only ever see that raw link text, with
    nothing to actually test against. This scans the AC text for a linked
    work item id, fetches THAT story's own AcceptanceCriteria via the
    work-item API, and appends it as a clearly-labeled block so the AI has
    real requirements to generate from.

    `cache` should be one plain dict created ONCE per run and passed to
    every call (run_titles/run_steps both do this) — a shared "spec" story
    referenced by many stories in the same run is then only fetched once.
    Capped at the first 3 distinct linked ids found and does NOT recurse
    into whatever the linked story's own AC references — this is meant to
    be a bounded, cheap best-effort resolution, not a general graph walk.
    Failures (deleted item, no permission, wrong project) are logged as a
    warning and skipped rather than raising — a bad/stale link shouldn't be
    able to fail an otherwise-fine generation run."""
    if not criteria:
        return criteria
    cache = cache if cache is not None else {}
    cb = cb or (lambda *a, **k: None)
    seen, ids = set(), []
    for m in _AC_WI_LINK_RE.finditer(criteria):
        wid = m.group(1) or m.group(2)
        if not wid:
            continue
        wid = int(wid)
        if wid == current_id or wid in seen:
            continue
        seen.add(wid)
        ids.append(wid)
        if len(ids) >= 3:
            break
    if not ids:
        return criteria
    blocks = []
    for wid in ids:
        if wid not in cache:
            resolved = None
            try:
                wi = _wit_client.get_work_item(
                    wid, fields=["System.Title", "Microsoft.VSTS.Common.AcceptanceCriteria"])
                t = wi.fields.get("System.Title", "")
                ac = _strip_html(wi.fields.get("Microsoft.VSTS.Common.AcceptanceCriteria", ""))
                if ac:
                    resolved = f"[Referenced story #{wid} — {t}]\n{ac}"
                    cb("log", {"msg": f"Resolved acceptance-criteria link to story #{wid} "
                                       f"({t[:60]})", "tone": "dim"})
            except Exception as e:
                cb("log", {"msg": f"Could not resolve linked story #{wid} in acceptance "
                                   f"criteria: {str(e)[:80]}", "tone": "warn"})
            cache[wid] = resolved
        if cache[wid]:
            blocks.append(cache[wid])
    if not blocks:
        return criteria
    return criteria + "\n\n" + "\n\n".join(blocks)


def _downscale_image(raw_bytes, max_dim=1024):
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(raw_bytes))
        if img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB")
        w, h = img.size
        if max(w, h) > max_dim:
            scale = max_dim / float(max(w, h))
            img = img.resize((int(w*scale), int(h*scale)), Image.LANCZOS)
        out = io.BytesIO()
        img.save(out, format="JPEG", quality=80)
        return out.getvalue(), "image/jpeg"
    except Exception:
        return raw_bytes, None


def fetch_story_screenshots(story):
    shots = []
    for rel in (story.relations or []):
        if rel.rel != "AttachedFile":
            continue
        fname = rel.attributes.get("name", "").lower()
        ext = os.path.splitext(fname)[1]
        if ext not in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
            continue
        try:
            r = requests.get(rel.url, auth=("", AZURE_PAT), timeout=30)
            if r.status_code == 200:
                raw, mt = _downscale_image(r.content)
                media = mt or ("image/png" if ext == ".png" else "image/jpeg")
                shots.append({"b64": base64.b64encode(raw).decode(), "media_type": media, "name": fname})
        except Exception:
            pass
    return shots


def describe_story_ui(screenshots, story_title="", should_stop=None, on_slow=None,
                      on_retry=None):
    """`on_slow(elapsed_seconds)` — same convention as generate_titles/generate_steps
    — lets the caller log a heartbeat while this (vision, so often slower) call is
    still in flight, instead of the run going silent for however long a slow/
    rate-limited provider takes to answer. `should_stop` is accepted for API
    symmetry with those two but is NOT wired at its current call site (run_steps):
    this call happens mid-test-case, and the Stop button is documented/labelled
    as "Stop after current test case" — wiring should_stop here would let Stop
    abort mid-call instead, which is a behavior change this fix isn't making.
    Passing should_stop remains available for any future caller that wants it."""
    if not screenshots:
        return ""
    prompt = f"""
        أنت مهندس ضمان جودة خبير. لديك صورة/صور لواجهة المستخدم لميزة بعنوان: {story_title}
        صف الواجهة بدقة ووضوح باللغة العربية: العناصر الظاهرة (حقول، أزرار، قوائم، جداول،
        رسائل)، أسماؤها، وأي نصوص مهمة، وكيفية ترتيبها وتفاعلها.
        هذا الوصف سيُستخدم لكتابة خطوات اختبار، لذا ركّز على التفاصيل العملية القابلة للاختبار.
        أعد وصفاً نصياً فقط بدون أي صيغة JSON.
    """
    images = [{"media_type": sc["media_type"], "data": sc["b64"]} for sc in screenshots]
    try:
        return (_run_stopaware(
            lambda: ai_complete(prompt, images=images, max_tokens=1500,
                                usage_tag="ui_description", on_retry=on_retry),
            should_stop=should_stop, on_slow=on_slow) or "").strip()
    except StopRequested:
        raise
    except CreditBalanceError:
        raise
    except Exception:
        return ""


def fetch_test_cases_for_suite(project, plan_id, suite_id):
    # witFields=System.Id keeps the response to the bare work-item id per case
    # instead of the full test-case payload. The number of entries (what callers
    # count) is unchanged, but each response is a fraction of the size, so the
    # bulk per-suite counting during plan generation is dramatically faster.
    url = (f"https://dev.azure.com/{AZURE_ORG}/{project}/"
           f"_apis/testplan/Plans/{plan_id}/Suites/{suite_id}/TestCase"
           f"?witFields=System.Id&api-version=7.0")
    # Retry on throttling / transient server errors. Azure DevOps rate-limits bulk
    # counting (hundreds of suites) and returns HTTP 429 with a Retry-After header;
    # honoring it (instead of failing -> a silent count of 0, or hammering) keeps
    # the counts correct and avoids wasted requests.
    # Bounded retry: enough to ride out brief throttling, but capped so a heavily
    # rate-limited cold count (e.g. Regenerate over hundreds of suites) can't stall
    # for minutes. Worst case ≈ 3 attempts × ~6 s ≈ 18 s for a single stubborn suite.
    last = 0
    for _attempt in range(3):
        resp = _az_session().get(url, auth=("", AZURE_PAT), timeout=30)
        last = resp.status_code
        if last == 200:
            return resp.json().get("value", [])
        if last == 429:
            try:
                wait = float(resp.headers.get("Retry-After", "1"))
            except Exception:
                wait = 1.0
            time.sleep(min(max(wait, 0.5), 6))     # honor Retry-After, capped at 6 s
            continue
        if 500 <= last < 600:
            time.sleep(0.4 * (_attempt + 1))       # transient server error -> brief backoff
            continue
        break                                       # other 4xx -> not retryable
    raise RuntimeError(f"HTTP {last}")


def _strip_html(s):
    """Remove HTML tags/entities from an Azure step fragment → plain text."""
    if not s:
        return ""
    import html as _h
    s = re.sub(r"<br\s*/?>", " ", s, flags=re.I)
    s = re.sub(r"<[^>]+>", "", s)
    s = _h.unescape(s)
    return re.sub(r"\s+", " ", s).strip()


def parse_steps_xml(steps_xml):
    """Parse Azure's Microsoft.VSTS.TCM.Steps XML into a list of:
        {"index": int, "action": str, "expected": str}
    The format is <steps><step><parameterizedString>action</parameterizedString>
    <parameterizedString>expected</parameterizedString></step>...</steps>.
    """
    if not steps_xml or "<step " not in steps_xml:
        return []
    out = []
    try:
        import xml.etree.ElementTree as ET
        # wrap in case of stray entities; Azure XML is generally well-formed
        root = ET.fromstring(steps_xml)
        steps = root.findall(".//step")
        for i, st in enumerate(steps, 1):
            ps = st.findall("parameterizedString")
            action = _strip_html(ps[0].text if len(ps) >= 1 and ps[0].text else "")
            expected = _strip_html(ps[1].text if len(ps) >= 2 and ps[1].text else "")
            if action or expected:
                out.append({"index": i, "action": action, "expected": expected})
    except Exception:
        # Fallback: regex the parameterizedString pairs
        try:
            chunks = re.findall(r"<step\b.*?</step>", steps_xml, flags=re.S)
            for i, ch in enumerate(chunks, 1):
                ps = re.findall(r"<parameterizedString[^>]*>(.*?)</parameterizedString>", ch, flags=re.S)
                action = _strip_html(ps[0]) if len(ps) >= 1 else ""
                expected = _strip_html(ps[1]) if len(ps) >= 2 else ""
                if action or expected:
                    out.append({"index": i, "action": action, "expected": expected})
        except Exception:
            pass
    return out


def fetch_test_case_steps(tc_id):
    """Return parsed steps [{index,action,expected}] for a single test case id."""
    try:
        wi = _wit_client.get_work_item(tc_id, fields=["Microsoft.VSTS.TCM.Steps"])
        xml = (wi.fields or {}).get("Microsoft.VSTS.TCM.Steps", "") or ""
        return parse_steps_xml(xml)
    except Exception:
        return []


def fetch_test_case_detail(tc_id):
    """Title + parsed steps for one test case in a single work-item call. The suite
    listing is fetched with witFields=System.Id (fast bulk counting), so it carries
    NO title — the automation path uses this to get the real case name AND steps,
    which the classifier needs to place each case on the right page."""
    try:
        wi = _wit_client.get_work_item(
            tc_id, fields=["System.Title", "Microsoft.VSTS.TCM.Steps"])
        f = wi.fields or {}
        title = f.get("System.Title", "") or ""
        return title, parse_steps_xml(f.get("Microsoft.VSTS.TCM.Steps", "") or "")
    except Exception:
        return "", []


def fetch_test_case_title(tc_id):
    """Just the title for one test case. The suite listing is fetched with
    witFields=System.Id (fast counting) and carries no name, so the Run/Report
    paths use this to show real case titles instead of 'No Title'."""
    try:
        wi = _wit_client.get_work_item(tc_id, fields=["System.Title"])
        return (wi.fields or {}).get("System.Title", "") or ""
    except Exception:
        return ""


def fetch_existing_titles_for_suite(project, plan_id, suite_id):
    """Reliable {id, title} pairs for every test case already in a suite.

    fetch_test_cases_for_suite() is optimized for fast bulk counting
    (witFields=System.Id) and — as fetch_test_case_title's own docstring
    already notes — its workItem.name comes back BLANK because of that. The
    Run/Report screens already work around this by backfilling via
    fetch_test_case_title(id) per case; the duplicate-detection paths
    (dedupe_existing_suite, and the 'does this suite already have cases'
    checks in run_titles/run_steps) did NOT do this backfill, so they were
    silently comparing against blank titles — every suite looked empty and
    every dedup check trivially found 'no duplicates', no matter how many
    real test cases already existed in it. This is the actual reason
    duplicates kept accumulating regardless of how good the semantic/AI
    matching got: the matching never had real titles to work with.

    Backfills concurrently (same pattern as the Run/Report screens) so a
    large suite doesn't serialize one HTTP call per case."""
    raw = fetch_test_cases_for_suite(project, plan_id, suite_id)
    prelim = []
    for it in raw:
        wi = (it.get("workItem", {}) or {})
        tc_id = wi.get("id")
        if not tc_id:
            continue
        prelim.append({"id": int(tc_id), "title": (wi.get("name") or "").strip()})
    missing = [r for r in prelim if not r["title"]]
    if missing:
        import concurrent.futures as _cf
        with _cf.ThreadPoolExecutor(max_workers=min(16, len(missing))) as _ex:
            fetched = dict(_ex.map(lambda r: (r["id"], fetch_test_case_title(r["id"])), missing))
        for r in prelim:
            if not r["title"]:
                r["title"] = fetched.get(r["id"], "") or ""
    return prelim



import time

def _is_arabic_out():
    # Arabic SPECIFICALLY (not merely "non-English") — otherwise fr/tr/es/de/nl
    # would wrongly take the Arabic prompt + RTL layout. Arabic is the only
    # RTL output language today, so this doubles as the RTL gate.
    return OUTPUT_LANG == "ar"

def _coerce_step_list(data):
    """Flatten an AI JSON result into a list of step dicts, tolerating the same
    object-wrapping that JSON mode forces on OpenAI-compatible providers — e.g.
    {"steps":[{...}]} or {"الخطوات":[{...}]}, or a single {...} step object."""
    if isinstance(data, dict):
        for k in ("steps", "test_steps", "الخطوات", "خطوات", "cases", "items"):
            if isinstance(data.get(k), list):
                data = data[k]
                break
        else:
            lists = [v for v in data.values() if isinstance(v, list)]
            data = lists[0] if lists else [data]   # a lone step object → wrap it
    if not isinstance(data, list):
        return []
    return [it for it in data if isinstance(it, dict)]


def generate_steps(tc_title, acceptance_criteria, ui_description="", log=None,
                   should_stop=None, on_slow=None, on_retry=None):
    if _is_arabic_out():
        ui_block = f"\n        وصف واجهة المستخدم (مستخلص من الصور):\n        {ui_description}\n" if ui_description else ""
        text = f"""
        أنت مهندس ضمان جودة (QA) خبير. أنشئ خطوات اختبار تفصيلية لحالة الاختبار التالية.
        اكتب جميع الخطوات باللغة العربية فقط.
        أعد كائن JSON فقط يحتوي على مفتاح "steps" قيمته مصفوفة من الخطوات — بدون أي نص إضافي أو markdown.
        مهم: لا تستخدم علامات الاقتباس المزدوجة داخل قيم النصوص.

        قواعد صارمة لمنع الخطوات المكررة أو الزائدة:
        - كل خطوة = إجراء واحد فقط يقوم به المستخدم، مع نتيجته المتوقعة في حقل expected.
        - لا تنشئ خطوة منفصلة لمجرد إعادة وصف إجراء سابق أو نتيجته. مثال خاطئ يجب تجنّبه:
          أربع خطوات كلها تقول «تم النقر على أيقونة تغيير اللغة وظهرت القائمة» — هذا تكرار،
          اجعلها خطوة واحدة (النقر) ونتيجتها في expected (ظهور القائمة).
        - الشروط البيئية أو شروط الحالة (يوجد اتصال إنترنت، فتح المتصفح، المستخدم على صفحة
          تسجيل الدخول) تُكتب فقط في حقل precondition لأول خطوة مرتبطة، وليست خطوة إجراء مستقلة.
        - التحقق من نتيجة يوضع في حقل expected، لا كخطوة إجراء جديدة.
        - استخدم أقل عدد من الخطوات يغطي السيناريو بالكامل (عادة 2 إلى 6 خطوات).
        - لا تكرر نفس الإجراء عبر خطوات متعددة.

        قاعدة صارمة بخصوص النتيجة المتوقعة (expected) — الالتزام الحرفي بمعايير القبول:
        - لا تخترع نص رسالة أو سلوكاً محدداً غير مذكور صراحةً في "معايير القبول" أدناه.
          مثال خاطئ: كتابة "يظهر خطأ: يجب أن تكون كلمة المرور 8 أحرف على الأقل" بينما معايير
          القبول تذكر فقط حداً أدنى للطول دون تحديد نص رسالة لكل حالة انتهاك على حدة.
        - إذا كانت معايير القبول تحدد رسالة/نتيجة عامة واحدة تغطي كل إدخال غير صحيح (مثل رسالة
          خطأ عامة عند إدخال بيانات خاطئة)، استخدم تلك الرسالة العامة بالضبط كما وردت — لا تخترع
          رسالة أكثر تحديداً حتى لو بدت منطقية أو متوقعة من الناحية التقنية.
        - انسخ نص أي رسالة مذكورة في معايير القبول حرفياً (نفس الصياغة) بدلاً من إعادة صياغتها.

        قاعدة التأصيل — يجب أن ترتبط الحالة بمعايير القبول فعلياً:
        - أولاً، حدد الجملة في "معايير القبول" أدناه التي يختبرها عنوان حالة الاختبار هذا.
          انسخ تلك الجملة حرفياً في حقل "ac_quote".
        - كل قيمة في expected يجب أن تكون مبررة بتلك الجملة أو بجملة أخرى في معايير القبول.
          إذا لم تحدد معايير القبول نتيجة محددة لهذا السيناريو، اكتب النتيجة العامة التي
          تنص عليها المعايير فعلاً — لا تخترع نتيجة جديدة.
        - في action وprecondition، لا تذكر إلا شاشات وحقولاً وأزراراً واردة في معايير
          القبول أو وصف الميزة. لا تخترع عناصر واجهة غير مذكورة.

        قاعدة توقيت ظهور رسائل التحقق:
        - ما لم تنص معايير القبول صراحةً على تحقق فوري أثناء الكتابة أو عند مغادرة
          الحقل، افترض أن التحقق يحدث عند الضغط على زر الإرسال/الحفظ/الدخول
          المذكور في المعايير: أدخل القيمة غير الصحيحة، ثم اضغط الزر، وبعدها فقط
          توقّع ظهور الرسالة. لا تكتب أبداً خطوة تتوقع تحذيراً مباشرةً بعد الكتابة.
        - في سيناريو الإدخال الصحيح، النتيجة المتوقعة هي النتيجة الإيجابية التي
          تنص عليها المعايير (رسالة النجاح أو الانتقال للشاشة التالية) — لا تكتب
          أبداً نفياً بصيغة نص مثل «لا يظهر تحذير»: هذا وصف لسلوك وليس نصاً
          يظهر على الشاشة، ولا يمكن لأي فحص آلي أو يدوي العثور عليه.

        قاعدة بيئة التشغيل (حالات التوافق):
        - إذا كان عنوان الحالة يستهدف متصفحاً أو منصة محددة (Chrome أو Firefox أو
          Safari أو Edge أو iOS أو Android أو الجوال)، فاذكر تلك البيئة في حقل
          precondition للخطوة الأولى (فتح التطبيق على ذلك المتصفح أو الجهاز)، ثم
          نفّذ نفس التدفق الأساسي للقصة، واجعل expected يتضمن تطابق السلوك وسلامة
          العرض في تلك البيئة. لا تخترع سلوكاً خاصاً بالمنصة غير مذكور في معايير
          القبول.

        عنوان حالة الاختبار: {tc_title}
        معايير القبول: {acceptance_criteria}
        وصف الميزة: {FEATURE_DESCRIPTION}{ui_block}

        الصيغة:
        {{"ac_quote":"...","steps": [{{"precondition":"...","action":"...","expected":"..."}}]}}
    """
    else:
        ui_block = f"\n        UI description (extracted from screenshots):\n        {ui_description}\n" if ui_description else ""
        text = f"""
        You are an expert QA engineer. Generate detailed test steps for the following test case.
        Write ALL steps in {out_lang_name()} only.
        Return ONLY a JSON object with a "steps" key whose value is an array of step objects — no extra text or markdown.
        Important: do not use double quotes inside the string values.

        Strict rules to prevent repeated / redundant steps:
        - Each step = exactly ONE concrete user action, with its expected result in 'expected'.
        - Do NOT create a separate step that merely restates a previous action or its outcome.
          Bad example to avoid: four steps that all say "clicked the language icon and the menu
          appeared" — that is duplication; make it ONE step (the click) with the menu appearing
          in 'expected'.
        - Environmental / state preconditions (internet is available, browser opened, user is on
          the login page) go ONLY in the 'precondition' field of the first related step — never
          as their own action step.
        - Verifying an outcome goes in 'expected', not as a new action step.
        - Use the FEWEST steps that fully cover the scenario (usually 2-6).
        - Never repeat the same action across multiple steps.

        Strict rule for the 'expected' field — stay literal to the acceptance criteria:
        - Do NOT invent specific message text or behavior that isn't explicitly stated in the
          'Acceptance criteria' below. Bad example: writing "shows error: password must be at
          least 8 characters" when the acceptance criteria only states a minimum length
          constraint, without defining a distinct message text per violation.
        - If the acceptance criteria defines ONE general message/outcome that covers any invalid
          input (e.g. one generic error message for incorrect data), use that exact general
          message — do not invent a more specific one, even if it sounds technically plausible
          or expected.
        - Copy any message text given in the acceptance criteria verbatim (same wording) rather
          than paraphrasing it.

        Grounding rule — the test must trace to the acceptance criteria:
        - First, find the sentence in 'Acceptance criteria' below that this test case title
          verifies. Copy that sentence VERBATIM into the "ac_quote" field.
        - Every 'expected' value must be justified by that sentence or another sentence in the
          acceptance criteria. If the criteria do not state a specific outcome for this
          scenario, write the general outcome the criteria DO state — never invent one.
        - In 'action' and 'precondition', only mention screens, fields, and buttons that appear
          in the acceptance criteria or feature description. Do not invent UI elements.

        Validation-timing rule — WHEN the app shows validation messages:
        - Unless the acceptance criteria EXPLICITLY state live/inline validation
          (a message while typing, or on leaving the field), assume validation
          happens ON SUBMIT: enter the invalid value, then click the submit/
          save/login button the criteria mention, and only THEN expect the
          message. Never write a step that expects a warning immediately after
          typing.
        - For a VALID-input scenario, the expected result is the POSITIVE
          outcome the criteria state (the success message, or the next screen)
          — never a negative phrased as text such as "no warning appears": that
          describes behavior, it is not text rendered on screen, and no
          automated or manual check can ever find it.

        Runtime-environment rule (compatibility cases):
        - If the test case title targets a specific browser or platform (Chrome,
          Firefox, Safari, Edge, iOS, Android, mobile), state that environment in
          the FIRST step's 'precondition' (open the app on that browser/device),
          then perform the story's own main flow, and make 'expected' include the
          behavior and layout matching that environment. Do not invent
          platform-specific behavior the acceptance criteria never mention.

        Test case title: {tc_title}
        Acceptance criteria: {acceptance_criteria}
        Feature description: {FEATURE_DESCRIPTION}{ui_block}

        Format:
        {{"ac_quote":"...","steps": [{{"precondition":"...","action":"...","expected":"..."}}]}}
    """
    time.sleep(1)
    last_err = None
    for attempt in range(5):
        try:
            # _run_stopaware (not a bare ai_complete call) so a slow/rate-limited
            # provider gets a periodic on_slow(elapsed) heartbeat instead of the
            # run going silent for however long this attempt takes — same
            # mechanism generate_titles already uses. should_stop is intentionally
            # NOT wired at generate_steps' current call site (run_steps): see
            # describe_story_ui's docstring for why (Stop is "after current test
            # case", not mid-call) — passing should_stop stays available for any
            # future caller that wants it, with unchanged behavior for this one.
            data = parse_json_robust(_run_stopaware(
                lambda: ai_complete(text, max_tokens=4096, want_json=True,
                                    usage_tag="generate_steps", on_retry=on_retry),
                should_stop=should_stop, on_slow=on_slow))
            # Programmatic backstop for the grounding rule above: a token-
            # overlap check between the model's own ac_quote and the real
            # acceptance criteria. Doesn't block/retry (steps are still
            # returned either way) — just surfaces a case where the model
            # likely fabricated its grounding rather than quoting it, same
            # "distrust, don't silently drop" spirit as _check_ac_coverage.
            if isinstance(data, dict):
                quote = str(data.get("ac_quote", "")).strip()
                if quote and acceptance_criteria and log:
                    if not _rules_overlap(quote, acceptance_criteria):
                        log(f"Note: steps for \"{tc_title[:60]}\" cite acceptance-criteria "
                            f"text that doesn't match what's on the story — worth a spot check.",
                            "dim")
            return _coerce_step_list(data)
        except StopRequested:
            raise
        except CreditBalanceError:
            raise
        except Exception as e:
            last_err = e; es = str(e).lower()
            if "429" in es or "rate_limit" in es:
                w = 30*(attempt+1)
                if log: log(f"Rate limited — waiting {w}s (attempt {attempt+1}/5)…", "warn")
                time.sleep(w)
            elif any(k in es for k in ("500","502","503","out of memory","cuda","internal server","overloaded")):
                w = 10*(attempt+1)
                if log: log(f"Provider busy/GPU error — retrying in {w}s (attempt {attempt+1}/5)…", "warn")
                time.sleep(w)
            elif "empty response" in es or "cannot parse json" in es:
                if log: log(f"Bad/empty AI response — retrying (attempt {attempt+1}/5)…", "warn")
                time.sleep(3)
            else:
                raise
    raise RuntimeError(f"Failed after 5 attempts: {last_err}")


def evaluate_existing_steps(tc_title, criteria, existing_steps_xml, should_stop=None,
                            on_slow=None, on_retry=None):
    """Decide whether a test case's EXISTING steps already adequately cover
    its acceptance criteria, or need regenerating. `should_stop`/`on_slow`
    were added after this call was found to be the single biggest source of
    "the run looks frozen" reports for re-runs against a suite that already
    has most of its steps: it used to call ai_complete() directly with no
    stop-awareness or heartbeat at all — unlike generate_steps/describe_
    story_ui/the dedup AI calls, which all got this treatment earlier. A
    slow/rate-limited provider could silently stall for minutes with zero
    log output (confirmed live: 2:56 of total silence evaluating one case),
    and Stop couldn't interrupt it either. See run_steps' _gen_and_write for
    why this call moved off the main thread and into the worker pool too."""
    plain = re.sub(r"<[^>]+>", " ", existing_steps_xml or "")
    plain = _html.unescape(re.sub(r"\s+", " ", plain)).strip()[:4000]
    if _is_arabic_out():
        prompt = f"""
        أنت مهندس ضمان جودة خبير. لديك حالة اختبار بخطواتها الحالية، ومعايير القبول الخاصة بها.
        مهمتك: قرر هل الخطوات الحالية كافية وتغطي معايير القبول بشكل صحيح أم لا.
        اعتبر الخطوات غير كافية (adequate=false) إذا وُجد أي مما يلي:
        - خطوات مكررة تعيد وصف نفس الإجراء أو نتيجته أكثر من مرة.
        - شروط بيئية مكتوبة كخطوات إجراء مستقلة (اتصال إنترنت، فتح المتصفح، المستخدم على الصفحة).
        - نتيجة متوقعة مكتوبة كخطوات إجراء متعددة بدلاً من حقل النتيجة.
        عنوان حالة الاختبار: {tc_title}
        معايير القبول: {criteria}
        الخطوات الحالية: {plain}
        أعد فقط كائن JSON: {{"adequate": true/false, "reason": "سبب مختصر بالعربية"}}
    """
        fallback_reason = "تعذر فهم رد الذكاء الاصطناعي حول التقييم — تم الإبقاء على الخطوات الحالية كما هي احتياطياً"
    else:
        prompt = f"""
        You are an expert QA engineer. You have a test case with its current steps and its
        acceptance criteria. Your task: decide whether the current steps are adequate and
        correctly cover the acceptance criteria or not.
        Consider the steps INADEQUATE (adequate=false) if any of these are present:
        - repeated steps that restate the same action or its outcome more than once;
        - environmental preconditions written as their own action steps (internet available,
          browser opened, user is on the page);
        - an expected outcome written as several action steps instead of an 'expected' result.
        Test case title: {tc_title}
        Acceptance criteria: {criteria}
        Current steps: {plain}
        Return ONLY a JSON object: {{"adequate": true/false, "reason": "short reason in {out_lang_name()}"}}
    """
        fallback_reason = "Could not understand the AI's evaluation reply — left the existing steps unchanged as a precaution"
    # One retry before giving up on a parse failure — cheap (a bad JSON reply
    # is often a one-off formatting slip, not a systemic problem) and avoids
    # falling back on the very first hiccup.
    for attempt in range(2):
        raw = (_run_stopaware(lambda: ai_complete(prompt, max_tokens=1024,
                                                  usage_tag="evaluate_existing_steps",
                                                  on_retry=on_retry),
                              should_stop=should_stop, on_slow=on_slow) or "").strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        try:
            data = parse_json_robust(raw)
            if isinstance(data, list) and data: data = data[0]
            return {"adequate": bool(data.get("adequate", False)), "reason": str(data.get("reason", "")).strip()}
        except Exception:
            if attempt == 0:
                continue
            # Fail OPEN, not closed: a parse failure tells us nothing about
            # whether the existing steps are actually fine or not — it's a
            # formatting problem with the AI's reply, not a finding about
            # the test case. Treating that as "definitely inadequate" (the
            # old behavior) forced a regeneration on every parse hiccup,
            # which could itself land on a degenerate result (seen live: a
            # regenerated case with an empty action/expected step — see the
            # steps-count fix in _apply_result and the empty-step guard in
            # _gen_and_write) — silently replacing a possibly-fine test case
            # with a genuinely worse one, then flagging it "Needs your
            # review" with a reason that reads like a real coverage
            # judgment when it was actually just noise. Leaving the
            # existing steps untouched is the safer default: worst case is
            # a real problem goes unflagged until the next run, instead of
            # a fine test case getting silently degraded on this one.
            return {"adequate": True, "reason": fallback_reason}


def build_steps_xml(steps):
    n = len(steps)
    xml = f'<steps id="0" last="{n + 1}">'
    for i, s in enumerate(steps, 2):
        pre = s.get("precondition", "").strip()
        action = s.get("action", "").strip()
        exp = s.get("expected", "").strip()
        action_text = (f"الشرط المسبق: {pre}\nالإجراء: {action}" if pre else action) if _is_arabic_out() \
            else (f"Precondition: {pre}\nAction: {action}" if pre else action)
        action_html = "<DIV><P>" + _html.escape(action_text).replace("\n", "</P><P>") + "</P></DIV>" if action_text else "<DIV><P></P></DIV>"
        exp_html = "<DIV><P>" + _html.escape(exp).replace("\n", "</P><P>") + "</P></DIV>" if exp else "<DIV><P></P></DIV>"
        xml += (f'<step id="{i}" type="ValidateStep">'
                f'<parameterizedString isformatted="true">{_html.escape(action_html)}</parameterizedString>'
                f'<parameterizedString isformatted="true">{_html.escape(exp_html)}</parameterizedString>'
                f'<description/></step>')
    return xml + '</steps>'


def update_test_case_with_steps(tc_id, steps_xml, project, story_id=None):
    from azure.devops.v7_0.work_item_tracking.models import JsonPatchOperation
    patch = [JsonPatchOperation(op="add", path="/fields/Microsoft.VSTS.TCM.Steps", value=steps_xml)]
    if story_id:
        try:
            wi = _wit_client.get_work_item(tc_id, expand="Relations")
            rels = wi.relations or []
            linked = any(getattr(r, "rel", "") == "Microsoft.VSTS.Common.TestedBy-Reverse"
                         and str(getattr(r, "url", "")).rstrip("/").endswith(f"/{story_id}") for r in rels)
            if not linked:
                patch.append(JsonPatchOperation(op="add", path="/relations/-", value={
                    "rel": "Microsoft.VSTS.Common.TestedBy-Reverse",
                    "url": f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workItems/{story_id}"}))
        except Exception:
            pass
    _wit_client.update_work_item(patch, id=tc_id)


# ═══════════════════════════════════════════════════════════════════════════════
#  TITLES GENERATION
# ═══════════════════════════════════════════════════════════════════════════════
def _coerce_title_list(data):
    """Normalize an AI JSON result into a flat list of title strings.

    Providers differ: some return a bare array ["t1","t2"], but OpenAI-compatible
    endpoints in JSON mode (response_format=json_object) are REQUIRED to return an
    object, so a model wraps the array under some key — e.g. {"titles":[...]} or,
    worse, {"حالات الاختبار":[...]} (it invents a key). Iterating that dict would
    yield the KEY as a single bogus 'title'. This coerces every shape:
    array | {"titles":[...]}/{any:[...]} | [{"title": "..."}] → ["t1","t2",...]."""
    if data is None:
        return []
    if isinstance(data, dict):
        # prefer a recognizable key, else the first list value in the object
        for k in ("titles", "test_cases", "testcases", "cases", "items",
                  "results", "عناوين", "الحالات"):
            v = data.get(k)
            if isinstance(v, list):
                data = v
                break
        else:
            lists = [v for v in data.values() if isinstance(v, list)]
            data = lists[0] if lists else [v for v in data.values() if isinstance(v, str)]
    if not isinstance(data, list):
        return []
    out = []
    for item in data:
        if isinstance(item, dict):
            t = (item.get("title") or item.get("name") or item.get("عنوان")
                 or item.get("العنوان") or "")
        else:
            t = item
        t = str(t or "").strip()
        if t:
            out.append(t)
    return out


def _numbered_criteria(criteria):
    """Split a raw acceptance-criteria blob into AC1/AC2/... lines for the
    generation prompts' coverage-enumeration rule (see generate_titles'
    docstring note below). Splits on newlines and common bullet/numbering
    prefixes; falls back to treating the whole blob as a single AC1 if it
    doesn't look like a list. Returns (numbered_text, ac_ids) — ac_ids is
    the list of "AC1".."ACn" the model is expected to cover, used afterward
    to check which ones its output actually referenced.

    A `[Referenced story #id — Title]` line — the section header
    _resolve_ac_links prepends when a story's AC pointed at another story
    for its remaining requirements — is deliberately NOT given its own AC
    number: it's a label, not a testable requirement, and numbering it
    would waste a title slot / confuse the coverage check with a phantom
    "criterion" that has nothing to actually verify. It's kept in the
    numbered text as unnumbered context so the model still knows where the
    lines under it came from. This is exactly the mixed case where a
    story's own AC states some requirements directly and points to another
    story for the rest — both halves end up numbered here, just not the
    header line introducing the linked half."""
    raw = (criteria or "").strip()
    if not raw:
        return "", []
    lines = [re.sub(r"^[\s\-•*•]+|^\d+[.\)]\s*", "", ln).strip()
             for ln in re.split(r"[\r\n]+", raw)]
    lines = [ln for ln in lines if ln][:40]
    if not lines:
        lines = [raw]
    ac_ids = []
    numbered_parts = []
    for ln in lines:
        if re.match(r"^\[.*\]$", ln):
            numbered_parts.append(ln)
            continue
        aid = f"AC{len(ac_ids) + 1}"
        ac_ids.append(aid)
        numbered_parts.append(f"{aid}: {ln}")
    numbered = "\n".join(numbered_parts)
    return numbered, ac_ids


def _rules_overlap(quote, criteria):
    """Is `quote` plausibly lifted from `criteria` rather than fabricated?
    Primary check is a normalized substring match (the grounding rule asks
    for a verbatim copy); falls back to token-Jaccard overlap so minor
    whitespace/punctuation differences from the model don't produce false
    warnings. Shared by generate_steps' ac_quote backstop."""
    nq, nc = _norm_title(quote), _norm_title(criteria)
    if not nq:
        return True  # nothing to check against — don't warn on an empty quote
    if nq in nc:
        return True
    tq, tc_ = set(nq.split()), set(nc.split())
    if not tq:
        return True
    return (len(tq & tc_) / len(tq)) >= 0.6


def _check_ac_coverage(items, ac_ids, log=None, kind="titles"):
    """Programmatic backstop for the coverage-enumeration rule below: look at
    which 'ac' ids the model actually attached to its output and warn (does
    NOT retry/block — see generate_titles docstring) if any numbered
    criterion got zero titles. Purely a visibility aid for now; a targeted
    retry against just the missed ACs is a reasonable future extension but
    wasn't added here to keep this change proportionate to what was asked."""
    if not ac_ids or not log:
        return
    used = set()
    for it in items:
        if isinstance(it, dict):
            a = str(it.get("ac", "")).strip().upper()
            if a:
                used.add(a)
    missing = [a for a in ac_ids if a not in used]
    if missing and len(missing) < len(ac_ids):  # partial gap is the interesting case
        log(f"Note: AI {kind} didn't reference {', '.join(missing)} — "
            f"double-check those criteria got covered.", "dim")


def generate_titles(story_title, criteria, existing_titles=None, log=None,
                    should_stop=None, on_slow=None, on_retry=None):
    # Numbering the acceptance criteria as AC1/AC2/... and requiring the
    # model to tag each title with the AC it tests turns "cover the
    # requirements" (an abstract judgment weak/flash-tier models don't
    # reliably self-apply) into a mechanical enumerate-and-tag loop, which
    # they follow far more reliably — the same structural trick that fixed
    # the recurring dedup false-merge bug (see _ai_duplicate_clusters).
    # _check_ac_coverage below is the matching programmatic backstop: it
    # can't fix a gap, but it surfaces one instead of silently dropping it.
    ac_text, ac_ids = _numbered_criteria(criteria)
    if _is_arabic_out():
        ct = ac_text or "لا توجد معايير قبول. أنشئ عناوين عامة بناءً على العنوان."
        existing_block = ""
        if existing_titles:
            listed = "\n".join(f"- {t}" for t in existing_titles[:150])
            existing_block = f"""
            حالات الاختبار التالية موجودة بالفعل لهذه القصة — لا تكررها:
{listed}
            أنشئ فقط عناوين لسيناريوهات جديدة. إذا كانت جميعها مغطاة، أعد {{"titles": []}}.
        """
        coverage_block = ("""
        قواعد التغطية — كل معيار قبول يجب أن يُختبر:
        - معايير القبول أدناه مرقمة AC1 وAC2 وهكذا.
        - لكل معيار ACn، أنشئ عنواناً واحداً على الأقل يختبره (الحالة الصحيحة؛
          أضف حالات خاطئة/حدّية فقط إذا كان المعيار يذكر قيداً يمكن انتهاكه).
        - كل عنوان يجب أن يستند إلى معيار واحد بالضبط. لا تنشئ عناوين عن
          سلوكيات لم تُذكر في معايير القبول (لا عناوين عامة عن الأداء أو الأمان
          ما لم يذكرها معيار صراحة؛ عناوين التوافق تحكمها قواعد التوافق أدناه
          حصراً).
        - أرفق مع كل عنوان رقم المعيار الذي يختبره.
""" if ac_ids else "")
        prompt = f"""
        أنت مهندس QA خبير. أنشئ عناوين حالات اختبار لقصة المستخدم التالية.
        اكتب العناوين باللغة العربية فقط. لا تستخدم علامات اقتباس مزدوجة داخل النصوص.
        أعد كائن JSON فقط يحتوي على مفتاح "titles" قيمته مصفوفة كائنات {{"ac":"ACn","title":"..."}}.

        قواعد صارمة لمنع التكرار:
        - لا تنشئ عنوانين يختبران نفس السلوك بصياغة مختلفة. كل عنوان يجب أن يغطي
          سيناريو فريداً غير مغطى بأي عنوان آخر.
        - اعتبر العنوانين مكررين إذا كانا يتحققان من نفس القاعدة أو الشرط حتى لو
          اختلفت الكلمات. مثال على تكرار يجب تجنبه:
            • "التحقق من أن الحد الأدنى لحقل الإجابة بالعربي هو حرفان"
            • "التحقق من أن حقل الإجابة بالعربي لا يقبل أقل من 2 حرف"
          هذان عنوانان مكرران — اختر واحداً فقط.
        - لكل حقل أو قاعدة، أنشئ حالة اختبار واحدة فقط لكل سيناريو (صحيح / خاطئ /
          حدّي)، وليس عدة صياغات لنفس السيناريو.
        - ادمج الحالات المتشابهة في عنوان واحد واضح بدلاً من تكرارها.
        - قبل الإخراج، راجع القائمة واحذف أي عنوان يكرر معنى عنوان آخر.
{coverage_block}
        تغطية التوافق (مطلوبة دائماً — بالإضافة إلى ما سبق):
        - أضف بالضبط 3 عناوين إضافية بوسم "ac":"COMPAT":
          1) التحقق من أن التدفق الرئيسي للقصة يعمل بشكل متطابق على متصفحات
             Chrome وFirefox وSafari وEdge.
          2) التحقق من التدفق الرئيسي وسلامة العرض على iOS (متصفح Safari للجوال).
          3) التحقق من التدفق الرئيسي وسلامة العرض على Android (متصفح Chrome للجوال).
        - هذه العناوين تغيّر بيئة التشغيل فقط — التدفق المُختبَر هو نفسه التدفق
          الأساسي للقصة؛ لا تخترع سلوكاً جديداً غير مذكور في معايير القبول.
        - اذكر اسم الميزة داخل كل عنوان منها (مثال: التحقق من تسجيل الدخول على
          متصفحات Chrome وFirefox وSafari وEdge)، ولا تنشئ أي عنوان منها إذا كان
          مثيله موجوداً بالفعل ضمن الحالات الموجودة.

        العنوان: {story_title}
        معايير القبول: {ct}
        {existing_block}
        الصيغة: {{"titles": [{{"ac":"AC1","title":"عنوان 1"}},{{"ac":"AC2","title":"عنوان 2"}}]}}
    """
    else:
        ct = ac_text or "No acceptance criteria. Generate general titles based on the title."
        existing_block = ""
        if existing_titles:
            listed = "\n".join(f"- {t}" for t in existing_titles[:150])
            existing_block = f"""
            The following test cases already exist for this story — do NOT duplicate them:
{listed}
            Only generate titles for NEW scenarios. If all are covered, return {{"titles": []}}.
        """
        coverage_block = ("""
        Coverage rules — every criterion must be tested:
        - The acceptance criteria below are numbered AC1, AC2, ...
        - For EACH criterion ACn, write at least one title that tests it (the
          valid case; add invalid/boundary cases only if that criterion states
          a constraint that can be violated).
        - Every title MUST be based on exactly one criterion. Do NOT write
          titles about behaviors the criteria never mention (no generic
          performance or security titles unless a criterion explicitly
          mentions them; compatibility titles are governed EXCLUSIVELY by the
          compatibility rules below).
        - Tag each title with the criterion number it tests.
""" if ac_ids else "")
        prompt = f"""
        You are an expert QA engineer. Generate test case titles for the following user story.
        Write the titles in {out_lang_name()} only. Do not use double quotes inside the text.
        Return ONLY a JSON object with a "titles" key whose value is an array of
        {{"ac":"ACn","title":"..."}} objects.

        Strict rules to prevent duplication:
        - Do not create two titles that test the same behavior with different wording.
          Each title must cover a unique scenario not covered by any other title.
        - Treat two titles as duplicates if they verify the same rule or condition even if
          the words differ. Example of a duplicate to avoid:
            • "Verify the minimum length of the Arabic answer field is 2 characters"
            • "Verify the Arabic answer field does not accept fewer than 2 characters"
          These two are duplicates — pick only one.
        - For each field or rule, create only one test case per scenario (valid / invalid /
          boundary), not several phrasings of the same scenario.
        - Merge similar cases into one clear title instead of repeating them.
        - Before output, review the list and remove any title that repeats another's meaning.
{coverage_block}
        Compatibility coverage (ALWAYS required — in addition to the above):
        - Add EXACTLY 3 extra titles tagged "ac":"COMPAT":
          1) Verify the story's main flow works identically across the Chrome,
             Firefox, Safari and Edge browsers.
          2) Verify the main flow and layout render correctly on iOS (mobile Safari).
          3) Verify the main flow and layout render correctly on Android (mobile Chrome).
        - These titles change the RUNTIME ENVIRONMENT only — the flow under
          test is the story's own main flow; do not invent behavior the
          acceptance criteria never mention.
        - Name the feature inside each of these titles (e.g. Verify sign-in
          works identically across Chrome, Firefox, Safari and Edge), and skip
          any of them whose equivalent already exists among the existing cases.

        Title: {story_title}
        Acceptance criteria: {ct}
        {existing_block}
        Format: {{"titles": [{{"ac":"AC1","title":"title 1"}},{{"ac":"AC2","title":"title 2"}}]}}
    """
    should_stop = should_stop or (lambda: False)
    _interruptible_sleep(1)
    if should_stop() or _STOP_EVENT.is_set():
        raise StopRequested()
    last_err = None
    for attempt in range(5):
        if should_stop() or _STOP_EVENT.is_set():
            raise StopRequested()
        try:
            # Run the (possibly slow) provider call stop-aware: Stop unwinds it at
            # once instead of waiting out the request timeout, and on_slow lets the
            # UI log a heartbeat while a large model is still generating.
            data = parse_json_robust(_run_stopaware(
                lambda: ai_complete(prompt, max_tokens=4096, want_json=True,
                                    usage_tag="generate_titles", on_retry=on_retry),
                should_stop=should_stop, on_slow=on_slow))
            raw_items = data.get("titles") if isinstance(data, dict) else data
            if isinstance(raw_items, list):
                _check_ac_coverage(raw_items, ac_ids, log=log, kind="titles")
            return _coerce_title_list(data)
        except StopRequested:
            raise
        except CreditBalanceError:
            raise
        except Exception as e:
            last_err = e; es = str(e).lower()
            if "429" in es or "rate_limit" in es:
                _interruptible_sleep(30*(attempt+1))
            elif any(k in es for k in ("500","502","503","cuda","out of memory","overloaded")):
                _interruptible_sleep(10*(attempt+1))
            elif "empty response" in es or "cannot parse json" in es:
                _interruptible_sleep(3)
            else:
                raise
    raise RuntimeError(f"Failed after 5 attempts: {last_err}")


def create_test_case(project, plan_id, suite_id, title, story_id):
    """Create a test case work item and add it to the suite. Returns tc_id."""
    from azure.devops.v7_0.work_item_tracking.models import JsonPatchOperation
    patch = [JsonPatchOperation(op="add", path="/fields/System.Title", value=title)]
    if story_id:
        patch.append(JsonPatchOperation(op="add", path="/relations/-", value={
            "rel": "Microsoft.VSTS.Common.TestedBy-Reverse",
            "url": f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/wit/workItems/{story_id}"}))
    wi = _wit_client.create_work_item(patch, project=project, type="Test Case")
    tc_id = wi.id
    try:
        _test_client.add_test_cases_to_suite(project=project, plan_id=plan_id,
                                             suite_id=suite_id, test_case_ids=str(tc_id))
    except Exception:
        pass
    return tc_id


def _norm_title(t):
    t = re.sub(r"[^\w\s\u0600-\u06FF]", "", str(t))
    return re.sub(r"\s+", " ", t).strip().lower()


def _classify_delete_error(status_code, raw):
    """Turn an Azure DevOps delete failure (HTTP status + response text, or a
    raw exception string) into a plain-English reason — WITHOUT throwing away
    the server's own error text. Previously any status code that wasn't
    401/403/404/5xx (e.g. 400 Bad Request) fell straight to a bare
    "HTTP {code}", discarding `raw` entirely — which is exactly why a 400
    showed no detail at all. Now every path keeps a truncated snippet of
    what Azure DevOps actually said, appended to the short label."""
    raw_l = (raw or "").lower()
    snippet = re.sub(r"\s+", " ", (raw or "")).strip()[:220]

    if status_code == 403 or "permission" in raw_l or "not authorized" in raw_l or "access is denied" in raw_l:
        base = ("no permission to delete work items in this project — ask an "
                "Azure DevOps admin to grant 'Delete and restore work items'")
    elif status_code == 401 or "unauthorized" in raw_l:
        base = "sign-in rejected — check the PAT in Setup"
    elif status_code == 404 or "does not exist" in raw_l or "has been deleted" in raw_l:
        base = "already deleted"
    elif "vs403" in raw_l or ("state" in raw_l and "transition" in raw_l):
        base = "blocked by a work item rule in this project"
    elif status_code and status_code >= 500:
        base = f"Azure DevOps server error (HTTP {status_code})"
    elif status_code:
        base = f"HTTP {status_code}"
    else:
        base = "unknown error"

    if snippet and snippet.lower() not in base.lower():
        return f"{base} — {snippet}"
    return base or "unknown error"


def delete_test_case(project, plan_id, suite_id, tc_id):
    """Remove a duplicate test case from the SUITE so it stops showing up in
    runs/reports. Does NOT hard-delete the underlying work item.

    Two things learned the hard way, in order:
      1. Test Case work items reject the generic Work Item Tracking delete API
         (_wit_client.delete_work_item / _apis/wit/workitems/{id}) outright —
         HTTP 400 "You cannot delete or restore test work items using this
         API. Use Test Management REST API to delete test artifacts."
      2. Switching to the Test Management API's own delete_test_case (the
         "correct" endpoint per that message) still fails for most accounts —
         hard-deleting a work item needs the elevated 'Delete and restore work
         items' permission, which is separate from normal suite-editing rights
         and most PATs don't have it (confirmed live: AccessDeniedException).

    Removing the test case from the suite only needs suite-edit permission
    (routine, already granted), and it achieves the actual goal: the duplicate
    no longer appears anywhere the suite is used. The work item itself is left
    alone rather than hard-deleted. Returns (ok, reason) — reason is "" on
    success, otherwise a short, readable diagnosis."""
    try:
        _test_client.remove_test_cases_from_suite_url(
            project=project, plan_id=plan_id, suite_id=suite_id, test_case_ids=str(tc_id))
        return True, ""
    except Exception as e1:
        try:
            url = (f"https://dev.azure.com/{AZURE_ORG}/{project}/_apis/test/Plans/"
                   f"{plan_id}/Suites/{suite_id}/testcases/{tc_id}?api-version=7.0")
            r = requests.delete(url, auth=("", AZURE_PAT), timeout=30)
            if r.status_code in (200, 204):
                return True, ""
            return False, _classify_delete_error(r.status_code, r.text)
        except Exception as e2:
            # e1 is from the SDK call (usually carries the real Azure DevOps
            # error text, e.g. a permission message); e2 is just "the REST
            # fallback also failed" (often a bare network/timeout error) —
            # prefer e1's message when there is one.
            return False, _classify_delete_error(None, str(e1) or str(e2))


def dedupe_existing_suite(project, plan_id, suite_id, cb=None, do_delete=True,
                          should_stop=None, story_id=None, story_title=None):
    """Find duplicate test cases ALREADY in a suite and remove the less complete
    one of each duplicate group, keeping the most accurate (most steps, then
    oldest id). Duplicates are matched two ways: first the cheap semantic-key
    check ('لا يقبل أقل من 2 حرف' vs 'لا يقبل أقل من حرفين' — shared words/
    hand-mapped synonyms), then an AI mop-up pass over whatever's still
    unclustered, to catch true synonyms/paraphrase the static check can't
    ('requests submitted for the branch' vs 'actions taken for the branch').

    This runs ONCE PER STORY (run_steps calls it per suite in a loop), so its
    own "Cleaned up…" summary line fires once per story too — with no story
    context that line just looks like the same message repeating for no
    reason in the activity log. `story_id`/`story_title`, when given, are
    prefixed onto the three possible summary outcomes below so each one is
    identifiable on its own.

    Returns {"removed": [ {id,title,kept_id} ], "kept": [...], "groups": n}.
    """
    story_tag = f"Story {story_id} · {story_title} — " if story_id else ""
    # Direction follows the STORY TITLE's own script (mirrors run_steps' own
    # "Story {sid} → suite {suite_id} · {title}" line) — not just "a tag is
    # present" — so an English title doesn't get forced right-to-left.
    _story_tag_ar = bool(story_tag) and any('؀' <= c <= 'ۿ' for c in (story_title or ""))
    cb = cb or (lambda *a, **k: None)
    should_stop = should_stop or (lambda: False)
    try:
        # fetch_existing_titles_for_suite (not the raw fetch_test_cases_for_suite)
        # -- the raw suite listing's workItem.name always comes back blank (see
        # its own docstring), which used to make every duplicate check here
        # compare against empty titles and silently find nothing to do, no
        # matter how many real test cases were already in the suite.
        cases = fetch_existing_titles_for_suite(project, plan_id, suite_id)
    except Exception as e:
        cb("log", {"msg": f"Could not read suite for dedup: {str(e)[:80]}", "tone": "warn"})
        return {"removed": [], "groups": 0}

    # Build records: {id, title, step_count}. The per-case step fetch is done
    # CONCURRENTLY (was serial — slow on big suites).
    import concurrent.futures as _cf

    def _rec(c):
        tc_id = c.get("id")
        title = c.get("title", "")
        if not tc_id:
            return None
        try:
            sc = len(fetch_test_case_steps(tc_id))
        except Exception:
            sc = 0
        return {"id": int(tc_id), "title": title, "steps": sc,
                "key": _semantic_key(title), "norm": _norm_title(title)}

    if cases:
        with _cf.ThreadPoolExecutor(max_workers=min(16, len(cases))) as _ex:
            recs = [r for r in _ex.map(_rec, cases) if r]
    else:
        recs = []

    # Group by semantic key (and exact-norm), then within each group decide keeper
    groups = {}
    for r in recs:
        # find an existing group whose key is a near-duplicate
        placed = False
        for gk in list(groups.keys()):
            if r["norm"] and any(r["norm"] == x["norm"] for x in groups[gk]):
                groups[gk].append(r); placed = True; break
            if _is_near_duplicate(r["key"], {gk}):
                groups[gk].append(r); placed = True; break
        if not placed:
            groups[r["key"]] = [r]

    # AI mop-up: the semantic-key grouping above only catches duplicates that
    # share literal words or a hand-mapped synonym. Ask the AI to look at
    # whatever's still standing alone (unclustered) and catch the rest — same
    # reasoning as _dedupe_titles_ai, applied to cases already sitting in the
    # suite. Best-effort and cheap: skipped entirely if Stop was clicked or
    # there's nothing ambiguous left to check.
    if not should_stop():
        singles = [(gk, members[0]) for gk, members in groups.items() if len(members) == 1]
        if len(singles) >= 2:
            items = [{"id": r["id"], "title": r["title"]} for _gk, r in singles]
            ai_groups = _ai_duplicate_clusters(
                items, should_stop=should_stop, story_title=story_title,
                on_slow=lambda s: cb("log", {
                    "msg": f"Still checking for duplicates with {T_disp(AI_PROVIDER)} "
                           f"({current_model() or 'model'}) — {s}s so far. Large models "
                           f"on free tiers can be slow; the run is not frozen.",
                    "tone": "dim", "ico": "⏳", "hb_id": f"dedupe:{suite_id}"}),
                on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                              "hb_id": f"dedupe:{suite_id}"}))
            if ai_groups:
                # Union-find over `singles` indices: the AI can legitimately
                # return overlapping/chained groups for a tight cluster of
                # near-identical titles (e.g. [A,B] and [B,C] instead of one
                # [A,B,C]) — merging pair-by-pair either crashes on a key
                # that's already been merged away, or silently leaves the
                # cluster split in two. Union-find collapses any chain into
                # one connected cluster regardless of how the AI grouped it,
                # and each surviving key is only ever touched once.
                parent = list(range(len(singles)))
                def _find(i):
                    while parent[i] != i:
                        parent[i] = parent[parent[i]]
                        i = parent[i]
                    return i
                def _union(a, b):
                    ra, rb = _find(a), _find(b)
                    if ra != rb:
                        parent[rb] = ra
                for idxs in ai_groups:
                    valid = [i for i in idxs if 0 <= i < len(singles)]
                    for i in valid[1:]:
                        _union(valid[0], i)
                clusters = {}
                for i in range(len(singles)):
                    clusters.setdefault(_find(i), []).append(i)
                for idx_list in clusters.values():
                    if len(idx_list) < 2:
                        continue
                    gks = [singles[i][0] for i in idx_list]
                    target_gk = gks[0]
                    for gk in gks[1:]:
                        if gk in groups and gk != target_gk:
                            groups[target_gk].extend(groups.pop(gk))

    removed = []       # successfully deleted duplicates
    kept_dupes = []    # duplicates we could NOT delete (left in the suite)
    dup_groups = 0
    # If the very first delete fails because of missing permissions, every
    # other delete this run will fail for the exact same reason — retrying 18
    # more times just wastes API calls and floods the log with the same line.
    # Stop attempting deletes (but keep detecting/reporting groups) once that
    # happens, and say so once instead of repeating it per case.
    perm_blocked_reason = None
    # IDs kept as the "most complete" survivor of a duplicate group THIS RUN.
    # A keeper's step count can legitimately exceed what its own narrow title
    # implies — it's often the winner precisely because a deleted duplicate's
    # scope got folded into it (e.g. an "email+password entry" case kept over
    # an "email field validation" case, absorbing the latter's steps in the
    # process). If the very same run then runs evaluate_existing_steps against
    # that keeper using ONLY its own title/AC — with no memory of what was
    # just merged into it — a keeper can be judged "inadequate" for its title
    # and shrunk back down, silently erasing the coverage that was only ever
    # preserved inside the now-deleted duplicates. run_steps checks this set
    # to skip evaluation of a same-run keeper rather than risk that.
    keeper_ids_this_run = []
    for gk, members in groups.items():
        if len(members) < 2:
            continue
        dup_groups += 1
        # keeper = most steps (more complete/accurate), tie-break = smallest id (oldest)
        members.sort(key=lambda m: (-m["steps"], m["id"]))
        keeper = members[0]
        keeper_ids_this_run.append(keeper["id"])
        n_dupes = len(members) - 1
        cb("log", {"msg": f"{len(members)} test cases test the same thing — "
                          f"keeping #{keeper['id']} ({keeper['steps']} step"
                          + ("s" if keeper['steps'] != 1 else "") + ", the most complete)",
                   "tone": "info", "ico": "≡", "ar": True,
                   "id": keeper["id"], "detail": keeper["title"]})
        for victim in members[1:]:
            if not do_delete or perm_blocked_reason:
                reason_note = f" — {perm_blocked_reason}" if perm_blocked_reason else ""
                cb("log", {"msg": f"{victim['title']}", "tone": "warn", "ar": True,
                           "id": victim["id"],
                           "detail": f"kept (not removed{reason_note}) · duplicate of #{keeper['id']}"})
                kept_dupes.append({"id": victim["id"], "title": victim["title"], "kept_id": keeper["id"]})
                continue
            ok, reason = delete_test_case(project, plan_id, suite_id, victim["id"])
            if ok:
                # log the OLD (removed) test — its id + title — and the id we kept.
                # tone "warn" (amber) + an explicit trash icon — a removal is a
                # DIFFERENT kind of outcome than a newly created/updated test case,
                # so it must not read as the same green ✓ "success" those use.
                cb("log", {"msg": f"{victim['title']}", "tone": "warn", "ar": True,
                           "ico": "🗑", "id": victim["id"],
                           "detail": f"removed #{victim['id']} (duplicate) · kept "
                                     f"#{keeper['id']} — {keeper['title']} instead"})
                removed.append({"id": victim["id"], "title": victim["title"], "kept_id": keeper["id"]})
            else:
                # delete failed → the duplicate is STILL there; never count it as removed
                cb("log", {"msg": f"{victim['title']}", "tone": "err", "ar": True,
                           "id": victim["id"],
                           "detail": f"could not remove — {reason} · duplicate of #{keeper['id']}"})
                kept_dupes.append({"id": victim["id"], "title": victim["title"], "kept_id": keeper["id"]})
                if "permission" in reason.lower():
                    perm_blocked_reason = reason
    if dup_groups:
        n_dupe_cases = sum(len(m) - 1 for m in groups.values() if len(m) >= 2)
        if removed:
            tail = (f"; {len(kept_dupes)} couldn't be removed" if kept_dupes else "")
            cb("log", {"msg": story_tag + f"Cleaned up {len(removed)} duplicate test case"
                              + ("s" if len(removed) != 1 else "")
                              + f" ({dup_groups} set" + ("s" if dup_groups != 1 else "")
                              + " of test cases that covered the same thing)"
                              + tail,
                       "tone": "warn" if kept_dupes else "ok", "ar": _story_tag_ar,
                       "replace_wip": f"dedupe:{suite_id}"})
        else:
            extra = f" ({perm_blocked_reason})" if perm_blocked_reason else ""
            cb("log", {"msg": story_tag + f"Found {n_dupe_cases} duplicate test case"
                              + ("s" if n_dupe_cases != 1 else "")
                              + f" across {dup_groups} set" + ("s" if dup_groups != 1 else "")
                              + f", but couldn't remove any{extra}.",
                       "tone": "err", "ar": _story_tag_ar,
                       "replace_wip": f"dedupe:{suite_id}"})
    else:
        # All three summary branches carry "replace_wip": f"dedupe:{suite_id}" —
        # clears the AI-mop-up heartbeat above (if it ever fired) once the pass
        # actually resolves. Without this, a heartbeat that fired even once is
        # never removed — nothing else here logs a matching completion line the
        # UI could hitch replace_wip onto — so it sits in the activity log
        # forever, frozen at its last elapsed value (seen live: "Still checking
        # for duplicates… — 15s so far" displayed well after the suite's dedup
        # pass had already logged its own result).
        cb("log", {"msg": story_tag + "No duplicate test cases found in the suite.",
                   "tone": "dim", "ar": _story_tag_ar,
                   "replace_wip": f"dedupe:{suite_id}"})
    return {"removed": removed, "kept": kept_dupes, "groups": dup_groups,
            "keeper_ids": keeper_ids_this_run}


def dedupe_case_list(cases, log=None, should_stop=None, story_title=None):
    """Skip-only duplicate detection for a list of in-memory test cases,
    used by the automation flow to avoid generating a script twice for the
    same scenario. Unlike dedupe_existing_suite this makes NO Azure DevOps
    calls and deletes nothing — duplicates are simply left out of the
    returned list. Same two-layer matching: cheap semantic-key grouping
    first, then an AI mop-up pass over whatever's still unclustered, to
    catch true synonyms/paraphrase the static check can't — including the
    structural target/rule extraction + programmatic backstop added to
    _ai_duplicate_clusters, since this function calls it directly and gets
    both automatically.

    `cases` — list of {"id": int, "title": str, "steps": [...]}.
    `story_title`, when given, is passed through to _ai_duplicate_clusters
    for the same grounding reason dedupe_existing_suite passes it — without
    it the model only sees bare titles and can more easily conflate "same
    feature area" with "same scenario".
    Returns the deduplicated list (most-complete case kept per duplicate
    group; non-duplicate cases pass through untouched).
    """
    log = log or (lambda *a, **k: None)
    should_stop = should_stop or (lambda: False)
    if not cases:
        return cases
    _dedupe_start = time.time()

    recs = []
    for c in cases:
        title = c.get("title", "") or ""
        recs.append({"id": c.get("id"), "title": title,
                     "steps": len(c.get("steps") or []),
                     "key": _semantic_key(title), "norm": _norm_title(title),
                     "_orig": c})

    groups = {}
    order = []  # first-seen order of group keys, so kept cases stay grouped
    for r in recs:
        placed = False
        for gk in list(groups.keys()):
            if r["norm"] and any(r["norm"] == x["norm"] for x in groups[gk]):
                groups[gk].append(r); placed = True; break
            if _is_near_duplicate(r["key"], {gk}):
                groups[gk].append(r); placed = True; break
        if not placed:
            groups[r["key"]] = [r]
            order.append(r["key"])

    if not should_stop():
        singles = [(gk, members[0]) for gk, members in groups.items() if len(members) == 1]
        if len(singles) >= 2:
            items = [{"id": r["id"], "title": r["title"]} for _gk, r in singles]
            # The "— Ns so far…" shape is collapsed IN PLACE by _auto_logmsg's
            # keyed retry/heartbeat collapse (main.py) — one line per story,
            # updated with the latest elapsed value, keyed by the message base
            # (the story-title snippet makes two concurrently-deduping
            # stories' heartbeats distinct keys instead of fighting over one).
            ai_groups = _ai_duplicate_clusters(
                items, should_stop=should_stop, story_title=story_title,
                on_slow=lambda s: log("  still checking for duplicates (%s) — %ds so far…"
                                      % ((story_title or "story")[:30], s), "dim"),
                on_retry=lambda m: log(f"  {m}", "warn"))
            if ai_groups:
                # Union-find — see dedupe_existing_suite for why: the AI can
                # return overlapping/chained groups for a tight cluster of
                # near-identical titles, and pairwise merging either crashes
                # or leaves the cluster split in two.
                parent = list(range(len(singles)))
                def _find(i):
                    while parent[i] != i:
                        parent[i] = parent[parent[i]]
                        i = parent[i]
                    return i
                def _union(a, b):
                    ra, rb = _find(a), _find(b)
                    if ra != rb:
                        parent[rb] = ra
                for idxs in ai_groups:
                    valid = [i for i in idxs if 0 <= i < len(singles)]
                    for i in valid[1:]:
                        _union(valid[0], i)
                clusters = {}
                for i in range(len(singles)):
                    clusters.setdefault(_find(i), []).append(i)
                for idx_list in clusters.values():
                    if len(idx_list) < 2:
                        continue
                    gks = [singles[i][0] for i in idx_list]
                    target_gk = gks[0]
                    for gk in gks[1:]:
                        if gk in groups and gk != target_gk:
                            groups[target_gk].extend(groups.pop(gk))

    kept = []
    n_skipped = 0
    dup_groups = 0
    for gk in order:
        members = groups.get(gk)
        if not members:
            continue
        if len(members) < 2:
            kept.append(members[0]["_orig"])
            continue
        dup_groups += 1
        members.sort(key=lambda m: (-m["steps"], m["id"]))
        keeper = members[0]
        kept.append(keeper["_orig"])
        for victim in members[1:]:
            n_skipped += 1
            log(f"  skipping duplicate: {victim['title'][:70]} "
                f"(same as #{keeper['id']})", "dim")
    if dup_groups:
        log(f"Skipped {n_skipped} duplicate test case" + ("s" if n_skipped != 1 else "")
            + f" before sequencing ({dup_groups} set" + ("s" if dup_groups != 1 else "")
            + " of test cases that covered the same thing) — generating "
            + "automation for the most complete case in each set only. "
            + f"⏱ {_fmt_mmss(time.time() - _dedupe_start)}", "warn")
    return kept


# Arabic filler/stop words that don't change a test's meaning — removed before
# comparing two titles for semantic equivalence.
_AR_STOP = {
    "التحقق", "من", "أن", "ان", "هو", "هي", "عند", "في", "على", "إلى", "الى",
    "مع", "عن", "لا", "هذا", "هذه", "يتم", "يجب", "كان", "تكون", "حقل", "الحقل",
    "لحقل", "قيمة", "القيمة", "رسالة", "الرسالة", "زر", "الزر", "صفحة", "الصفحة",
    "إمكانية", "امكانية", "ظهور", "وجود", "بشكل", "صحيح", "الحد",
}
# Synonym groups → canonical token, so "الحد الأدنى ... حرفان" and
# "لا يقبل أقل من 2 حرف" map onto shared concept tokens.
_AR_SYN = {
    # length-amount tokens all collapse to "minlen" concept
    "حرفان": "minlen", "حرفين": "minlen", "حرف": "minlen", "أحرف": "minlen",
    "احرف": "minlen", "2": "minlen", "٢": "minlen",
    # "minimum" and "does not accept less than" express the SAME rule → one token
    "الأدنى": "minrule", "الادنى": "minrule", "أدنى": "minrule", "ادنى": "minrule",
    "أقل": "minrule", "اقل": "minrule",
    "الأقصى": "maxrule", "الاقصى": "maxrule", "أقصى": "maxrule", "اقصى": "maxrule",
    "أكثر": "maxrule", "اكثر": "maxrule",
    "يقبل": "accept", "قبول": "accept", "تقبل": "accept",
    "الإجابة": "answer", "الاجابة": "answer", "إجابة": "answer", "اجابة": "answer",
    "السؤال": "question", "سؤال": "question",
    "العربي": "ar", "بالعربي": "ar", "العربية": "ar", "بالعربية": "ar", "عربي": "ar",
    "الإنجليزي": "en", "الانجليزي": "en", "بالإنجليزي": "en",
    "بالانجليزي": "en", "الإنجليزية": "en", "الانجليزية": "en", "بالإنجليزية": "en",
    "إجباري": "required", "اجباري": "required", "الزامي": "required",
    "إلزامي": "required", "مطلوب": "required",
    "فارغ": "empty", "فارغاً": "empty", "فارغا": "empty", "تركه": "empty",
    # field-name tokens
    "الاسم": "name", "الإسم": "name", "اسم": "name", "إسم": "name",
    "البريد": "email", "الايميل": "email", "الإيميل": "email", "الالكتروني": "email",
    "الإلكتروني": "email", "الهاتف": "phone", "الجوال": "phone", "الموبايل": "phone",
    "رقم": "number", "الرقم": "number",
    # number words → digits so "حرفين" ~ "2 حرف", "ثلاثة" ~ "3"
    "حرفين": "minlen", "حرفان": "minlen",
    "واحد": "1", "واحده": "1", "واحدة": "1", "اثنين": "2", "اثنان": "2",
    "ثلاثة": "3", "ثلاثه": "3", "اربعة": "4", "أربعة": "4", "اربعه": "4",
    "خمسة": "5", "خمسه": "5", "ستة": "6", "سته": "6",
    "٣": "3", "٤": "4", "٥": "5", "٦": "6", "٧": "7", "٨": "8", "٩": "9", "١": "1",
}

def _semantic_key(title):
    """Reduce an Arabic title to a frozenset of meaning tokens, so two titles
    that test the same rule with different wording collapse to the same key."""
    norm = _norm_title(title)
    toks = []
    for w in norm.split():
        w = _AR_SYN.get(w, w)
        if w in _AR_STOP:
            continue
        toks.append(w)
    return frozenset(toks)

def _is_near_duplicate(key, seen_keys, threshold=0.8):
    """True if `key` overlaps an already-seen key by >= threshold (Jaccard)."""
    if not key:
        return False
    for k in seen_keys:
        if not k:
            continue
        inter = len(key & k)
        union = len(key | k)
        if union and inter / union >= threshold:
            return True
        # also treat full subset of a short key as duplicate
        if key <= k or k <= key:
            if min(len(key), len(k)) >= 2:
                return True
    return False


def _dedupe_ai_debug_path():
    import platform_caps as _pc_dir
    d = _pc_dir.app_data_dir()   # writable on mobile too (see helper)
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass
    return os.path.join(d, "dedupe_ai_debug.log")


def _log_dedupe_ai_call(kind, story_title, sent_items, raw_text, result):
    """Best-effort diagnostic log for the two AI duplicate-detection passes
    (_ai_duplicate_clusters over an existing suite, _dedupe_titles_ai over
    freshly-proposed titles) — never raises, never blocks, and never affects
    what the caller does with `result`.

    Root-causing a false-positive merge (the AI grouping two genuinely
    unrelated test cases as 'duplicates') used to be pure guesswork after the
    fact: the model's raw JSON was parsed and immediately discarded, so there
    was no way to tell whether a bad group came from a single wrong judgment
    in one AI response, or from the union-find chaining two separate,
    individually-plausible pairings into a bigger cluster the AI never
    actually endorsed as a whole (see dedupe_existing_suite's union-find
    comment). This appends one JSON line per call that returned at least one
    group, capturing exactly what was sent, what the model said verbatim, and
    what was parsed out of it — capped at the last ~300 lines so the file
    never grows unbounded."""
    if not result:
        return
    try:
        entry = {
            "ts": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "kind": kind,                 # "existing_suite" | "new_titles"
            "story_title": story_title,
            "sent_items": sent_items,     # exactly what was numbered in the prompt
            "raw_response": raw_text,     # the model's response, verbatim
            "result": result,             # what was parsed out of it
        }
        p = _dedupe_ai_debug_path()
        lines = []
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    lines = f.readlines()
            except Exception:
                lines = []
        lines.append(json.dumps(entry, ensure_ascii=False) + "\n")
        with open(p, "w", encoding="utf-8") as f:
            f.writelines(lines[-300:])
    except Exception:
        pass


def _ai_duplicate_clusters(items, should_stop=None, story_title=None, on_slow=None,
                           on_retry=None):
    """items: [{"id":..,"title":..}, ...] — test cases the cheap semantic-key
    check found NO overlap for (each still alone in its own group). Asks the
    AI which of these titles describe the exact same scenario as another one
    in the list, despite different wording/synonyms it can't be expected to
    know ahead of time via a hand-maintained table.

    `story_title`, when given, is included as context so the AI grounds its
    judgment in what THIS story actually covers — without it, the model has
    only bare titles to go on and tends to treat "same feature area" (e.g.
    every check that happens to mention the same screen) as if it meant
    "same scenario", which is exactly how unrelated test cases end up
    wrongly clustered and deleted as "duplicates".

    Returns a list of groups, each a list of INDICES into `items` (index 0 =
    first item, etc.) that should be treated as one duplicate cluster.
    Best-effort: returns [] on any failure, timeout, or Stop.

    `on_slow(elapsed_seconds)` — same convention as generate_titles/generate_steps:
    called periodically while waiting on a slow/retrying provider, so a caller
    can log a heartbeat instead of this looking frozen (a free-tier model's
    internal retry-with-backoff can legitimately run several minutes)."""
    if len(items) < 2:
        return []
    should_stop = should_stop or (lambda: False)
    try:
        if should_stop() or _STOP_EVENT.is_set():
            return []
        ar = _is_arabic_out()
        shown = items[:150]
        listed = "\n".join(f"{i+1}. {it['title']}" for i, it in enumerate(shown))
        story_ctx_ar = (f"كل هذه الحالات تنتمي إلى القصة: \"{story_title}\" — قارن بينها "
                        f"ضمن نطاق هذه القصة فقط، ولا تفترض أن التشابه في الموضوع يعني "
                        f"نفس السيناريو.\n\n") if story_title else ""
        story_ctx_en = (f"All of these test cases belong to the story: \"{story_title}\" — "
                        f"compare them within that story's scope only, and do not assume "
                        f"that sharing a topic means they share a scenario.\n\n") if story_title else ""
        if ar:
            prompt = f"""
            أنت مراجع ضمان جودة. مهمتك الوحيدة هي إيجاد عناوين حالات الاختبار المكررة
            فعلياً — أي نفس الفحص مكتوباً مرتين بصياغة مختلفة. اتبع الإجراء أدناه
            بالترتيب وبدقة، ولا تتخطَّ أي خطوة.

            التعريف
            يُعتبر عنوانان مكررين فقط إذا تطابق الشرطان التاليان معاً:
            (أ) نفس الهدف — نفس الحقل أو الزر أو العنصر أو الإجراء؛ و
            (ب) نفس القاعدة — نفس الشرط أو السلوك المتوقع أو قاعدة التحقق التي
                يتم فحصها بالضبط.
            إذا تطابق (أ) واختلف (ب) بأي شكل — حتى بكلمة واحدة مثل «رقم» مقابل
            «حرف كبير» مقابل «رمز» مقابل «8 أحرف» — فهما ليسا مكررين.

            العناوين التي تشترك في نفس الميزة أو الشاشة أو الحقل أو قالب الجملة ليست
            مكررة ما لم تكن القاعدة المفحوصة نفسها متطابقة. تشابه الصياغة لا يثبت
            شيئاً. الحقل الواحد كثيراً ما يخضع لعدة قواعد تحقق منفصلة (إلزامي، حد
            أدنى للطول، حد أقصى للطول، احتواء رقم، احتواء حرف كبير، احتواء رمز، منع
            المسافات، ...). كل قاعدة من هذه القواعد حالة اختبار مستقلة يجب الإبقاء
            عليها كلها.

            الإجراء (إلزامي)
            الخطوة 1 — لكل عنوان مرقم، استخرج قيمتين قصيرتين:
              "target": الحقل أو العنصر أو الإجراء محل الاختبار
              "rule":   الشرط المحدد الذي يُفحص، منقولاً من نص العنوان حرفياً قدر
                        الإمكان
            الخطوة 2 — قارن العناوين مثنى مثنى. اعتبر عنوانين مكررين فقط إذا تطابق
              "target" وكان "rule" يعني نفس الشيء تماماً. إذا ذكرت القاعدتان متطلبين
              مختلفين أو حدين مختلفين أو نوعين مختلفين من المحارف فهما اختباران
              مختلفان — لا تجمعهما أبداً.
            الخطوة 3 — كوّن المجموعات فقط من الأزواج التي تأكدت في الخطوة 2. عند أي
              شك في أي زوج لا تجمعه؛ تفويت تكرار لا يسبب ضرراً، أما الدمج الخاطئ
              فيحذف حالة اختبار حقيقية نهائياً.

            صيغة المخرجات
            أعد كائن JSON واحداً فقط دون أي نص آخر، بهذين المفتاحين وبهذا الترتيب:
            {{
              "analysis": [{{"n": الرقم, "target": "...", "rule": "..."}}, ...],
              "groups": [[أرقام العناوين المكررة فعلاً], ...]
            }}
            يجب أن يحتوي "analysis" على عنصر واحد لكل عنوان في القائمة. أدرج في
            "groups" فقط المجموعات التي تضم رقمين أو أكثر. إن لم يوجد أي تكرار:
            "groups": [].

            مثال — ادرسه جيداً؛ فهو يوضح الخطأ الأكثر شيوعاً:
            العناوين:
            1. التحقق أن كلمة المرور تحتوي على رمز واحد على الأقل
            2. التحقق من أن كلمة المرور تحتوي على حرف كبير واحد على الأقل
            3. التحقق من أن كلمة المرور تحتوي على رقم واحد على الأقل
            4. التأكد من أن كلمة المرور يجب أن تتضمن رقماً واحداً على الأقل
            الناتج الصحيح:
            {{
              "analysis": [
                {{"n": 1, "target": "حقل كلمة المرور", "rule": "رمز واحد على الأقل"}},
                {{"n": 2, "target": "حقل كلمة المرور", "rule": "حرف كبير واحد على الأقل"}},
                {{"n": 3, "target": "حقل كلمة المرور", "rule": "رقم واحد على الأقل"}},
                {{"n": 4, "target": "حقل كلمة المرور", "rule": "رقم واحد على الأقل"}}
              ],
              "groups": [[3, 4]]
            }}
            العناوين 1 و2 و3 تتشارك نفس الحقل ونفس قالب الجملة تقريباً، لكن «رمز»
            و«حرف كبير» و«رقم» ثلاث قواعد مختلفة ← ثلاثة اختبارات مختلفة ← لا تُجمع
            أبداً. أما العنوانان 3 و4 فيفحصان نفس القاعدة (احتواء رقم) بصياغتين
            مختلفتين ← مكرران.

            أمثلة سريعة إضافية:
            - «زر الدخول معطل عند ترك كلمة المرور فارغة» مقابل «التحقق من بقاء زر
              الدخول معطلاً عند فراغ حقل كلمة المرور» ← مكرر (نفس الهدف ونفس
              القاعدة بصياغة مختلفة).
            - «التحقق من أن حقل المستخدمين إلزامي» مقابل «التحقق من أن الحد الأقصى
              لطول حقل الاسم 100 حرف» ← ليس مكرراً (هدف مختلف وقاعدة مختلفة).
            - المتصفح/نظام التشغيل/الجهاز جزء من القاعدة: نفس التدفق على Chrome
              مقابل نفس التدفق على Safari أو iOS أو Android ← بيئتا تشغيل
              مختلفتان ← ليسا مكررين أبداً؛ أدرج البيئة ضمن "rule" عند استخراجها.

            {story_ctx_ar}عناوين حالات الاختبار المرقمة الموجودة في المجموعة:
            {listed}

            نفّذ الآن الخطوة 1 ثم الخطوة 2 ثم الخطوة 3، وأعد كائن JSON فقط.
            """
        else:
            prompt = f"""
            You are a QA reviewer. Your ONLY job is to find test-case titles that are true
            duplicates — the SAME check written twice in different words. Follow the
            procedure below exactly, in order. Do not skip any step.

            DEFINITION
            Two titles are duplicates ONLY if BOTH of the following are identical:
            (A) the same target — the same field, button, element, or action; AND
            (B) the same rule — the exact same condition, expected behavior, or
                validation rule being checked.
            If (A) matches but (B) differs in ANY way — even by a single word such as
            "digit" vs "uppercase letter" vs "symbol" vs "8 characters" — they are NOT
            duplicates.

            Titles that share the same feature, screen, field, or sentence template are
            NOT duplicates unless the checked rule itself is identical. Similar wording
            proves nothing. One single field often has MANY separate validation rules
            (required, minimum length, maximum length, must contain a digit, must contain
            an uppercase letter, must contain a symbol, no spaces, ...). Each of those
            rules is a separate test case and every one of them must be kept.

            PROCEDURE (mandatory)
            Step 1 — For EVERY numbered title, extract two short values:
              "target": the field/element/action under test
              "rule":   the specific condition being checked, copied as literally as
                        possible from the title
            Step 2 — Compare titles pairwise. Two titles are duplicates ONLY if their
              "target" is the same AND their "rule" means exactly the same thing. If the
              two rules name different requirements, limits, quantities, or character
              types, they are DIFFERENT tests — never group them.
            Step 3 — Build groups only from pairs confirmed in Step 2. If you are unsure
              about any pair, DO NOT group it. Missing a duplicate is harmless; a wrong
              merge permanently deletes a real test case.

            OUTPUT FORMAT
            Return ONE JSON object and nothing else, with exactly these two keys in this
            order:
            {{
              "analysis": [{{"n": <number>, "target": "...", "rule": "..."}}, ...],
              "groups": [[numbers that are true duplicates], ...]
            }}
            "analysis" must contain one entry per input title. "groups" must contain only
            groups of 2 or more numbers. If there are no duplicates: "groups": [].

            EXAMPLE — study it carefully; it shows the single most common mistake:
            Input titles:
            1. Verify the password contains at least one symbol
            2. Verify the password contains at least one uppercase letter
            3. Verify the password contains at least one digit
            4. Confirm password must include at least 1 number
            Correct output:
            {{
              "analysis": [
                {{"n": 1, "target": "password field", "rule": "at least one symbol"}},
                {{"n": 2, "target": "password field", "rule": "at least one uppercase letter"}},
                {{"n": 3, "target": "password field", "rule": "at least one digit"}},
                {{"n": 4, "target": "password field", "rule": "at least one digit"}}
              ],
              "groups": [[3, 4]]
            }}
            Titles 1, 2 and 3 share the same field and almost the same sentence, but
            "symbol", "uppercase letter" and "digit" are three DIFFERENT rules, so they
            are three different tests and must NOT be grouped. Titles 3 and 4 check the
            same rule (contains a digit) in different words, so they ARE duplicates.

            Other quick examples:
            - "Login button is disabled when password is empty" vs "Verify the login
              button stays disabled when the password field is blank" → duplicates
              (same target, same rule, different wording).
            - "Verify the users field is required" vs "Verify the name field max length
              is 100" → NOT duplicates (different target AND different rule).
            - The browser/OS/device is PART of the rule: the same flow on Chrome vs the
              same flow on Safari, iOS or Android → different runtime environments →
              NEVER duplicates; include the environment in the extracted "rule".

            {story_ctx_en}Numbered test-case titles already in the suite:
            {listed}

            Now perform Step 1, Step 2, and Step 3, and return only the JSON object.
            """
        raw = _run_stopaware(lambda: ai_complete(prompt, max_tokens=1536, want_json=True,
                                                 usage_tag="dedupe_ai_clusters",
                                                 on_retry=on_retry),
                             should_stop=should_stop, on_slow=on_slow)
        data = parse_json_robust(raw)
        raw_groups = data.get("groups") if isinstance(data, dict) else None
        if not raw_groups:
            return []
        # Programmatic backstop: cross-check the model's OWN structured
        # "analysis" (its {n, target, rule} extraction, forced by the
        # prompt's Step 1) before trusting a group in "groups" — a weak
        # model can still merge on lexical similarity between TITLES despite
        # the prompt's instructions, but if its own analysis names a
        # DIFFERENT "rule" for two members of the group it just returned
        # (e.g. "at least one digit" vs "at least one uppercase letter"),
        # that is a direct self-contradiction and the group is dropped
        # rather than trusted. If the model omitted/malformed the analysis
        # for a group's numbers, this falls back to trusting the group as
        # before, so a formatting slip doesn't silently disable detection
        # entirely — it's a backstop against the specific "grouped despite
        # naming different rules" failure, not a hard requirement.
        analysis = data.get("analysis") if isinstance(data, dict) else None
        rule_by_n = {}
        if isinstance(analysis, list):
            for entry in analysis:
                try:
                    rule_by_n[int(entry.get("n"))] = str(entry.get("rule", "")).strip()
                except Exception:
                    continue
        def _rules_match(r1, r2):
            a, b = _norm_title(r1 or ""), _norm_title(r2 or "")
            if not a or not b or a == b:
                return True
            ta, tb = set(a.split()), set(b.split())
            if not ta or not tb:
                return True
            union = len(ta | tb)
            return union > 0 and len(ta & tb) / union >= 0.6
        out = []
        for g in (raw_groups or []):
            nums, idxs = [], set()
            for n in (g or []):
                try:
                    ni = int(n)
                except Exception:
                    continue
                nums.append(ni)
                idxs.add(ni - 1)
            idxs = sorted(i for i in idxs if 0 <= i < len(shown))
            if len(idxs) < 2:
                continue
            rules = [rule_by_n[n] for n in nums if n in rule_by_n]
            if len(rules) >= 2 and not all(_rules_match(rules[0], r) for r in rules[1:]):
                continue
            out.append(idxs)
        _log_dedupe_ai_call(
            "existing_suite", story_title, shown, raw,
            [[shown[i] for i in g] for g in out])
        return out
    except Exception:
        return []


def _dedupe_titles_ai(candidate_titles, existing_titles, story_title, log=None,
                      should_stop=None, on_slow=None, on_retry=None):
    """Second-pass duplicate check, run AFTER the cheap word-overlap filter
    (_is_near_duplicate). That filter only catches duplicates that share enough
    literal words/hand-mapped synonyms (_AR_SYN) — it has no way to know that,
    say, "requests submitted for the branch" and "actions taken for the branch"
    assert the same thing, since those are different roots, not spelling
    variants. A hand-curated synonym table can never cover every domain a
    story might be about, so instead this asks the AI provider itself — which
    already understands synonyms/paraphrase in any domain and language — to
    flag any remaining candidate that duplicates the MEANING of an existing
    test case or an earlier candidate. Deliberately runs only on whatever
    survives the free filter first, keeping the prompt (and cost) small.

    Best-effort: any failure (timeout, bad JSON, provider error) just returns
    candidate_titles unchanged — this is a quality refinement, never a reason
    to block generation.

    `on_slow(elapsed_seconds)` — same convention as generate_titles/generate_steps,
    see _ai_duplicate_clusters for why this matters."""
    if not candidate_titles:
        return candidate_titles
    should_stop = should_stop or (lambda: False)
    try:
        if should_stop() or _STOP_EVENT.is_set():
            raise StopRequested()
        ar = _is_arabic_out()
        ex_list = (existing_titles or [])[:150]
        ex_block = "\n".join(f"{i+1}. {t}" for i, t in enumerate(ex_list)) or ("لا يوجد" if ar else "(none)")
        cand_block = "\n".join(f"{i+1}. {t}" for i, t in enumerate(candidate_titles))
        if ar:
            prompt = f"""
            أنت مراجع ضمان جودة. قصة المستخدم: {story_title}

            حالات اختبار "موجودة بالفعل" لهذه القصة:
            {ex_block}

            حالات اختبار "جديدة مقترحة" (مرقمة):
            {cand_block}

            اتبع الإجراء التالي بالترتيب لتحديد أي حالة جديدة تكرر حالة موجودة أو
            حالة جديدة أخرى سابقة لها رقم أصغر.

            التعريف: يُعتبر عنوانان مكررين فقط إذا تطابق الهدف (نفس الحقل/العنصر/
            الإجراء) والقاعدة (نفس الشرط أو السلوك المتوقع بالضبط) معاً. الحقل
            الواحد كثيراً ما يخضع لعدة قواعد تحقق منفصلة (حد أدنى للطول، احتواء
            حرف كبير، احتواء رقم، احتواء رمز، ...) — كل قاعدة حالة اختبار مستقلة،
            حتى لو تشابهت صياغة العناوين بشدة.

            الخطوة 1 — لكل حالة "جديدة مقترحة"، استخرج "target" (الهدف) و"rule"
            (القاعدة المحددة، منقولة حرفياً قدر الإمكان) من نص عنوانها.
            الخطوة 2 — قارن كل حالة جديدة بكل الحالات الموجودة بالفعل وبالحالات
            الجديدة الأخرى. اعتبرها مكررة فقط إذا تطابق "target" و"rule" تماماً.
            إذا اختلفت القاعدة بأي شكل — حتى بكلمة واحدة مثل «رقم» مقابل «حرف
            كبير» — فهما اختباران مختلفان، لا تُدرجها كتكرار.
            الخطوة 3 — عند أي شك لا تُدرج الرقم كتكرار؛ تفويت تكرار غير ضار، أما
            حذف حالة فريدة بالخطأ فيفقد تغطية اختبار حقيقية.

            المتصفح/نظام التشغيل/الجهاز جزء من القاعدة: نفس التدفق على Chrome
            مقابل نفس التدفق على Safari أو iOS أو Android ← بيئتا تشغيل مختلفتان
            ← ليسا مكررين أبداً؛ أدرج البيئة ضمن "rule" عند استخراجها.

            مثال يوضح الخطأ الأكثر شيوعاً: "التحقق أن كلمة المرور تحتوي على رمز
            واحد على الأقل" و"التحقق أن كلمة المرور تحتوي على حرف كبير واحد على
            الأقل" و"التحقق أن كلمة المرور تحتوي على رقم واحد على الأقل" تتشارك
            نفس الحقل وقالب الجملة لكن «رمز» و«حرف كبير» و«رقم» ثلاث قواعد
            مختلفة ← ثلاث حالات مختلفة، لا تُجمع أبداً كمكررة لبعضها.

            أعد فقط كائن JSON بهذين المفتاحين:
            {{"analysis": [{{"n": الرقم, "target": "...", "rule": "..."}}, ...],
              "duplicate_numbers": [الأرقام المكررة الواجب حذفها]}}
            يجب أن يحتوي "analysis" على عنصر واحد لكل حالة جديدة مقترحة.
            إن لم يوجد أي تكرار أعد "duplicate_numbers": [].
            """
        else:
            prompt = f"""
            You are a QA reviewer. User story: {story_title}

            Test cases that "already exist" for this story:
            {ex_block}

            "Newly proposed" test cases (numbered):
            {cand_block}

            Follow this procedure, in order, to find which new cases duplicate an
            existing case or an earlier-numbered new case.

            DEFINITION: two titles are duplicates ONLY if BOTH the target (same
            field/element/action) AND the rule (the exact same condition or
            expected behavior) match. One field often has several separate
            validation rules (minimum length, requires an uppercase letter,
            requires a digit, requires a symbol, ...) — each rule is its own
            test case, even when the titles look very similar.

            Step 1 — For every "newly proposed" title, extract "target" and
            "rule" (copied as literally as possible from the title).
            Step 2 — Compare each new title against every existing title and every
            other new title. Call it a duplicate ONLY if target AND rule both
            match exactly. If the rule differs in ANY way — even one word like
            "digit" vs "uppercase letter" — they are different tests; do not
            flag it.
            Step 3 — When in doubt, do NOT flag it as a duplicate. Missing a
            duplicate is harmless; wrongly dropping a unique case loses real
            test coverage.

            The browser/OS/device is PART of the rule: the same flow on Chrome
            vs the same flow on Safari, iOS or Android → different runtime
            environments → NEVER duplicates; include the environment in the
            extracted "rule".

            Example of the most common mistake: "Verify the password contains at
            least one symbol", "...at least one uppercase letter", and "...at
            least one digit" share the same field and almost the same sentence,
            but "symbol", "uppercase letter", and "digit" are three DIFFERENT
            rules → three different tests, never group them as duplicates of
            each other.

            Return ONLY a JSON object with exactly these two keys:
            {{"analysis": [{{"n": <number>, "target": "...", "rule": "..."}}, ...],
              "duplicate_numbers": [numbers to remove]}}
            "analysis" must contain one entry per newly-proposed title.
            If there are no duplicates, return "duplicate_numbers": [].
            """
        raw = _run_stopaware(lambda: ai_complete(prompt, max_tokens=1536, want_json=True,
                                                 usage_tag="dedupe_titles_ai",
                                                 on_retry=on_retry),
                             should_stop=should_stop, on_slow=on_slow)
        data = parse_json_robust(raw)
        nums = data.get("duplicate_numbers") if isinstance(data, dict) else None
        if not nums:
            return candidate_titles
        drop = set()
        for n in nums:
            try:
                drop.add(int(n) - 1)
            except Exception:
                pass
        if not drop:
            return candidate_titles
        # NOTE: no programmatic backstop here, unlike _ai_duplicate_clusters.
        # That one can cross-check a group's members against EACH OTHER's
        # extracted "rule" because every member of a group is a numbered item
        # in the SAME list the model analyzed. Here, "analysis" only covers
        # the newly-proposed candidates — a dropped candidate could equally
        # have been matched against an EXISTING title (which has no extracted
        # rule to compare against) or another candidate, and the flat
        # duplicate_numbers list doesn't say which. Any veto rule built on
        # that ambiguity risks silently disabling itself (a candidate's own
        # rule simply won't resemble anything when it matched an existing
        # title, which is the common/expected case) — worse than no check at
        # all, since it would look like protection without providing any.
        # The structural prompt fix above (forced target/rule extraction,
        # step-by-step comparison, explicit password-rule example) is the
        # actual defense here; stakes are also lower than the existing-suite
        # cleanup — a wrongly-dropped candidate here just never gets created,
        # nothing real is deleted.
        if not drop:
            return candidate_titles
        kept = [t for i, t in enumerate(candidate_titles) if i not in drop]
        _log_dedupe_ai_call(
            "new_titles", story_title,
            {"existing": ex_list, "candidates": candidate_titles}, raw,
            [candidate_titles[i] for i in sorted(drop)])
        if log and len(kept) != len(candidate_titles):
            log(f"AI review caught {len(candidate_titles) - len(kept)} additional "
                f"duplicate title(s) the quick filter missed", "dim")
        return kept
    except StopRequested:
        raise
    except Exception:
        return candidate_titles


# ═══════════════════════════════════════════════════════════════════════════════
#  ORCHESTRATORS — driven by the UI via a callback
#  cb(event_type, payload) where event_type in:
#    'log'    payload={'msg','tone','id','indent'}
#    'stat'   payload={'total','stories_done','total_stories','done','skipped','errors'}
#    'progress' payload={'pct','label'}
#    'story'  payload={'id','title'}
#    'done'   payload={'summary', 'action_items', ...}
# ═══════════════════════════════════════════════════════════════════════════════
class StopRequested(Exception):
    pass

class _TrackerOps:
    """The tracker-facing operations run_titles/run_steps depend on.

    Both generators interleave AI work with 14 direct calls to Azure-specific
    functions (suite discovery, story fetch, dedupe, case create, step write).
    That coupling is what stopped generation running on any other tracker.

    Rather than rewrite two ~300-line functions, those calls now go through one
    injectable object. `ops=None` binds every attribute to the SAME module-level
    function called before, so the default path is pure indirection with ZERO
    behaviour change — provable by inspection, asserted by verify_ops.py.
    A non-Azure backend supplies its own (see backend_setup.generation_ops).
    """

    __slots__ = ("connect", "discover_suites", "fetch_stories", "dedupe_suite",
                 "existing_titles", "create_case", "cases_for_suite",
                 "case_title", "write_steps", "case_detail")

    def __init__(self, **kw):
        for name in self.__slots__:
            setattr(self, name, kw[name])


def _default_ops():
    """Azure ops — literally the existing module-level functions."""
    return _TrackerOps(
        connect=connect_azure_sdk,
        discover_suites=discover_suites_for_stories,
        fetch_stories=fetch_stories,
        dedupe_suite=dedupe_existing_suite,
        existing_titles=fetch_existing_titles_for_suite,
        create_case=create_test_case,
        cases_for_suite=fetch_test_cases_for_suite,
        case_title=fetch_test_case_title,
        write_steps=update_test_case_with_steps,
        case_detail=fetch_test_case_detail,   # (title, parsed steps) for Automation
    )


def run_titles(project, plan_id, story_ids, cb, should_stop=lambda: False,
               on_ai_error=None, gate=None, ops=None):
    _o = ops or _default_ops()
    wit, test = _o.connect(project)
    cb("log", {"msg": "Discovering suites for stories…", "tone": "dim"})
    story_suite_map = _o.discover_suites(project, plan_id, set(story_ids))
    # DO NOT re-add a "one story per suite" guard here (tried in #120, reverted
    # in #125 after it broke every non-Azure run). `story_suite_map` is NOT
    # story→suite one-to-one: backend_setup.generation_ops.discover_suites
    # deliberately keys the SAME story under up to four ALIASES — sid (int),
    # str(sid), story.ref.id and story.ref.key — all pointing at one suite, so
    # that `story_suite_map.get(sid)` matches regardless of int-vs-str. A guard
    # that treats a repeated suite id as "a duplicate story" therefore deletes
    # that story's own aliases, including the key the caller looks up: seen live
    # on Azure→TestRail as "Suite 45 already covered by story 101048 — skipping
    # duplicate story 101048" (the SAME id) followed by "No suite for story
    # 101048 — skipped", i.e. a run that processed 0 stories.
    # The hybrid double-count this was meant to fix is already handled where it
    # actually belongs — the case-level idempotence guard further down, which
    # dedupes by test-case work-item id and is alias-agnostic.
    stories = _o.fetch_stories(story_ids)

    total_created = 0; errors = 0; stories_done = 0
    total_stories = len(stories)
    _titles_start = time.time()
    per_story_stats = {}   # {sid: {"id","title","total","ok","skipped","err","suite"}}
    cb("stat", {"total": 0, "stories_done": 0, "total_stories": total_stories,
                "done": 0, "skipped": 0, "errors": 0})

    # Each story's own pipeline (dedup its suite → generate titles → AI-dedupe
    # those titles → create the test cases) is independent of every other
    # story's — different suite, different AI calls — so a couple of stories
    # run concurrently instead of strictly one at a time, same small bounded-
    # pool rationale as run_steps' Steps-generation/dedup-check pools (still
    # hits the same rate-limited free-tier endpoint either way, this overlaps
    # LATENCY across a couple of stories rather than multiplying request
    # volume). cb is shadowed with a lock so concurrent stories logging
    # through it stays fully serialized, same guarantee as before.
    _cb_lock = _threading.Lock()
    _real_cb = cb
    def cb(kind, payload):
        with _cb_lock:
            _real_cb(kind, payload)

    _fatal = {"hit": False}   # first fatal (credit/auth/bad_model/not_found/network) wins
    _ac_link_cache = {}       # shared across all stories this run — see _resolve_ac_links

    def _process_story(story):
        """Runs on a worker thread: dedup this story's suite, generate +
        AI-dedupe titles, create the resulting test cases. Returns a result
        dict describing exactly what happened; does NOT touch total_created/
        errors/stories_done/per_story_stats directly — the main thread does
        all of that bookkeeping from the result, same separation as
        run_steps' _gen_and_write/_apply_result split."""
        sid = story.id
        title = story.fields.get("System.Title", "No Title")
        criteria = story.fields.get("Microsoft.VSTS.Common.AcceptanceCriteria", "")
        criteria = _resolve_ac_links(criteria, project, cb=cb, cache=_ac_link_cache, current_id=sid)
        cb("story", {"id": sid, "title": title})
        suite_id = story_suite_map.get(sid)
        if not suite_id:
            cb("log", {"msg": f"No suite for story {sid} — skipped", "tone": "warn"})
            return {"sid": sid, "title": title, "suite_id": None}

        ps_total = 0; ps_ok = 0; ps_skipped = 0; ps_err = 0; ps_secs = 0.0

        # Remove pre-existing duplicate test cases in this suite first, keeping the
        # most complete one of each group (catches dupes from prior runs / manual entry).
        try:
            _o.dedupe_suite(project, plan_id, suite_id, cb=cb, do_delete=True,
                                  should_stop=should_stop, story_id=sid, story_title=title)
        except Exception as de:
            cb("log", {"msg": f"Dedup skipped: {str(de)[:80]}", "tone": "warn"})
        # existing titles (fetch_existing_titles_for_suite backfills real titles —
        # the raw suite listing's workItem.name is always blank, which used to
        # make this look empty even when the suite was already full)
        existing_titles = []
        try:
            for it in _o.existing_titles(project, plan_id, suite_id):
                nm = (it.get("title") or "").strip()
                if nm: existing_titles.append(nm)
        except Exception:
            pass
        if existing_titles:
            # Prefixed with the story tag for the same reason dedupe_existing_suite's
            # summary lines are (above) — with 2 stories now running concurrently,
            # this generic one-liner is indistinguishable from another story's
            # identical line without it.
            _title_is_ar = any('؀' <= c <= 'ۿ' for c in title)
            cb("log", {"msg": f"Story {sid} · {title} — Suite already has "
                              f"{len(existing_titles)} test case(s) — only new added",
                       "tone": "warn", "ar": _title_is_ar})
        # Move the headline status off "Starting…" while the (often slow) title
        # generation runs — otherwise, since the per-case progress event only
        # fires once the FIRST case starts, a long free-tier title call leaves the
        # status stuck on "Starting…" even though it's working. Label only (no pct)
        # so the % and bar are untouched; the per-case progress below overrides it.
        cb("progress", {"label": "Generating test-case titles…"})
        try:
            titles = _call_with_network_retries(lambda: generate_titles(
                title, criteria, existing_titles,
                log=lambda m, t="warn": cb("log", {"msg": m, "tone": t}),
                should_stop=should_stop,
                on_slow=lambda s: (cb("progress",
                                      {"label": f"Generating test-case titles… {s}s"}),
                                   cb("log", {
                    "msg": f"Still generating titles with {T_disp(AI_PROVIDER)} "
                           f"({current_model() or 'model'}) — {s}s so far. Large models "
                           f"on free tiers can be slow; click Stop to cancel.",
                    "tone": "dim", "ico": "⏳", "hb_id": f"titles:{sid}"})),
                on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                              "hb_id": f"titles:{sid}"}),
            ), cb, should_stop=should_stop)
        except StopRequested:
            # Caught explicitly (not the generic except below) — StopRequested
            # carries no message text, so classify_ai_error() would otherwise
            # misreport this expected Stop-triggered cancellation as a fake
            # "{provider}: unknown error." (same bug fixed in run_steps).
            return {"sid": sid, "title": title, "suite_id": suite_id, "stopped": True,
                    "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}
        except CreditBalanceError:
            return {"sid": sid, "title": title, "suite_id": suite_id, "fatal": "credit",
                    "fatal_summary": "Stopped — out of AI credits",
                    "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}
        except Exception as e:
            cat, friendly = classify_ai_error(e)
            # Same reasoning as run_steps: a config error or a real network
            # outage hits every remaining story identically, so stop once
            # with one clear message instead of repeating the same failure
            # per story.
            if cat in ("auth", "bad_model", "not_found", "network"):
                return {"sid": sid, "title": title, "suite_id": suite_id, "fatal": cat,
                        "fatal_summary": f"Stopped — {friendly}",
                        "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}
            cb("log", {"msg": f"AI error: {e}", "tone": "err"})
            ps_err += 1
            return {"sid": sid, "title": title, "suite_id": suite_id, "ai_error": True,
                    "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}

        # Clears the "Still generating titles…" heartbeat above once the call
        # actually resolves. Without this, a heartbeat that fired even once is
        # never removed — the only log lines below this point fire
        # CONDITIONALLY (only if something was actually dropped/skipped), so a
        # clean run with nothing to skip would leave the heartbeat as a
        # permanent ghost, frozen at its last elapsed value.
        cb("log", {"msg": f"Generated {len(titles)} candidate title(s)", "tone": "dim",
                   "replace_wip": f"titles:{sid}"})

        existing_norm = {_norm_title(t) for t in existing_titles}
        seen_keys = {_semantic_key(t) for t in existing_titles}
        unique = []
        dropped_dupes = 0
        for t in titles:
            nk = _norm_title(t)
            sk = _semantic_key(t)
            if nk in existing_norm or _is_near_duplicate(sk, seen_keys):
                dropped_dupes += 1
                continue
            unique.append(t)
            existing_norm.add(nk)
            seen_keys.add(sk)
        if dropped_dupes:
            ps_skipped += dropped_dupes
            cb("log", {"msg": f"Skipped {dropped_dupes} duplicate/near-duplicate title"
                              + ("s" if dropped_dupes > 1 else ""), "tone": "dim", "ico": "⏭"})
        # Second pass: the word-overlap filter above only catches duplicates that
        # share literal words/hand-mapped synonyms. Ask the AI to catch the rest
        # (true synonyms/paraphrases it understands but the static filter can't).
        _before_ai = len(unique)
        try:
            unique = _dedupe_titles_ai(
                unique, existing_titles, title,
                log=lambda m, t="dim": cb("log", {"msg": m, "tone": t}),
                should_stop=should_stop,
                on_slow=lambda s: cb("log", {
                    "msg": f"Still checking for duplicate titles with {T_disp(AI_PROVIDER)} "
                           f"({current_model() or 'model'}) — {s}s so far. Large models "
                           f"on free tiers can be slow; the run is not frozen.",
                    "tone": "dim", "ico": "⏳", "hb_id": f"dedupetitles:{sid}"}),
                on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                              "hb_id": f"dedupetitles:{sid}"}))
        except StopRequested:
            return {"sid": sid, "title": title, "suite_id": suite_id, "stopped": True,
                    "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}
        # Clears the "Still checking for duplicate titles…" heartbeat above —
        # same reasoning as the generate_titles cleanup a few lines up.
        # _dedupe_titles_ai's own internal log() call (the "AI review caught…"
        # line) fires conditionally AND doesn't carry a replace_wip anyway
        # (its log callback signature is just (msg, tone)), so it can't be
        # relied on to clean this up either.
        cb("log", {"msg": "Title dedup check complete", "tone": "dim",
                   "replace_wip": f"dedupetitles:{sid}"})
        if len(unique) != _before_ai:
            ps_skipped += (_before_ai - len(unique))
        for tc_title in unique:
            if should_stop(): break
            _tc_start = time.time()
            try:
                tc_id = _o.create_case(project, plan_id, suite_id, tc_title, sid)
                ps_ok += 1
                _elapsed = time.time() - _tc_start
                ps_secs += _elapsed
                cb("log", {"msg": tc_title, "tone": "ok", "id": tc_id, "ar": True,
                           "secs": round(_elapsed, 1), "detail": f"⏱ {_fmt_mmss(_elapsed)}"})
            except Exception as e:
                ps_err += 1
                ps_secs += (time.time() - _tc_start)
                cb("log", {"msg": f"{tc_title} — {e}", "tone": "err", "ar": True})

        # total = ok + skipped + err by construction (not incremented separately
        # during the creation loop above) — report.py's per-story percentage ring
        # assumes `total` is the full candidate count so `processed
        # (ok+skipped+err) / total` naturally lands at 100%. The old formula only
        # counted post-dedup attempted titles into `total`, excluding whatever
        # dedup had already filtered out — any story with a nonzero skipped count
        # (title-level dedup catching duplicates, the normal/expected case for a
        # suite that already has content) showed >100% in the report (seen live:
        # 200%/166% on stories with 11 and 2 skipped titles respectively). This
        # exact "total set upfront to the full candidate count" invariant is
        # already how run_steps' story_prog avoids the same issue.
        ps_total = ps_ok + ps_skipped + ps_err
        return {"sid": sid, "title": title, "suite_id": suite_id,
                "total": ps_total, "ok": ps_ok, "skipped": ps_skipped, "err": ps_err, "secs": ps_secs}

    _TITLES_WORKERS = 2
    import concurrent.futures as _cf_rt
    _futs = {}   # future -> story

    def _apply_story_result(fut, story):
        """Main-thread-only: applies one completed story's bookkeeping/log
        lines. Always called from _drain_titles, which runs on the main
        thread and always finishes (including the "Story X completed" line
        below) BEFORE the dispatch loop submits the next story into a freed
        slot. That ordering is what guarantees a story's completion is never
        logged AFTER the next story that took its slot has already logged
        its own start — without it, the ThreadPoolExecutor auto-reuses a
        freed worker thread for the next queued task the instant it's idle,
        which can race ahead of this thread's own completion logging (seen
        live: "Story 101049 · Profile" logging before "Story 101048
        completed", even though 101048 genuinely finished first — same
        thread, so execution order was correct, only the LOG order raced)."""
        try:
            res = fut.result()
        except Exception as e:
            cb("log", {"msg": f"Story {story.id} — unexpected error: {str(e)[:120]}", "tone": "err"})
            return
        suite_id = res.get("suite_id")
        if suite_id is None:
            return   # "no suite" — already logged inside _process_story, no stats entry
        sid = res["sid"]; title = res["title"]
        ps = per_story_stats.setdefault(sid, {"id": sid, "title": title, "total": 0,
                                              "ok": 0, "skipped": 0, "err": 0, "suite": suite_id})
        ps["total"] += res.get("total", 0)
        ps["ok"] += res.get("ok", 0)
        ps["skipped"] += res.get("skipped", 0)
        ps["err"] += res.get("err", 0)
        ps["secs"] = ps.get("secs", 0.0) + res.get("secs", 0.0)
        nonlocal total_created, errors, stories_done
        total_created += res.get("ok", 0)
        errors += res.get("err", 0)
        if res.get("fatal"):
            if not _fatal["hit"]:
                _fatal.update(hit=True, summary=res["fatal_summary"], reason=res["fatal"])
            return
        if res.get("stopped") or res.get("ai_error"):
            return   # matches the original sequential code's `continue` — not "completed"
        stories_done += 1
        cb("log", {"msg": f"Story {sid} completed", "tone": "ok", "ico": "└"})
        cb("stat", {"total": total_created, "stories_done": stories_done,
                    "total_stories": total_stories, "done": total_created,
                    "skipped": 0, "errors": errors})

    def _process_story_paused(story):
        """Pause-and-ask wrapper around _process_story — the run_titles twin of
        run_steps' _gen_and_write wrapper (see its docstring). A story's FATAL
        provider error (credit/auth/bad_model/not_found/network) blocks on
        on_ai_error; Resume retries the WHOLE story pipeline — safe, because
        the only fatal-capable AI call (generate_titles) runs BEFORE any title
        is created in Azure, and the dedup passes are already non-fatal — and
        Stop lets the fatal result flow to the drain loop unchanged."""
        while True:
            res = _process_story(story)
            if not res.get("fatal") or on_ai_error is None:
                return res
            _friendly = ((res.get("fatal_summary") or "").replace("Stopped — ", "", 1)
                         or "AI provider error")
            cb("log", {"msg": _friendly, "tone": "err"})
            cb("log", {"msg": _paused_on_error_msg(), "tone": "warn", "ico": "⏸"})
            if on_ai_error(_friendly) == "retry":
                cb("log", {"msg": f"Retrying story {res.get('sid')} with "
                                  f"{T_disp(AI_PROVIDER)}…", "tone": "dim"})
                continue
            return res

    def _drain_titles(wait_for_all):
        """Harvest completed futures and apply their results on the main
        thread — same rationale (and the same timeout=0-can't-really-wait
        pitfall this avoids) as run_steps' own _drain. Called BEFORE every
        new submission in the dispatch loop below, so a freed slot's
        completion is always logged before that slot's next story starts."""
        if not _futs:
            return
        if wait_for_all:
            done_set, _ = _cf_rt.wait(list(_futs), return_when=_cf_rt.ALL_COMPLETED)
        elif len(_futs) >= _TITLES_WORKERS:
            done_set, _ = _cf_rt.wait(list(_futs), return_when=_cf_rt.FIRST_COMPLETED)
        else:
            done_set, _ = _cf_rt.wait(list(_futs), timeout=0, return_when=_cf_rt.ALL_COMPLETED)
        for fut in done_set:
            story = _futs.pop(fut)
            _apply_story_result(fut, story)

    with _cf_rt.ThreadPoolExecutor(max_workers=_TITLES_WORKERS) as _ex:
        for story in stories:
            if should_stop() or _fatal["hit"]:
                break
            if gate and not gate():   # manual Pause point (False = Stop clicked)
                break
            _drain_titles(wait_for_all=False)
            fut = _ex.submit(_process_story_paused, story)
            _futs[fut] = story
        _drain_titles(wait_for_all=True)

    # round per-story seconds
    for v in per_story_stats.values():
        v["secs"] = round(v.get("secs", 0.0), 1)
    _total_secs = round(time.time() - _titles_start, 1)
    if _fatal["hit"]:
        cb("done", {"summary": _fatal["summary"], "reason": _fatal["reason"],
                    "created": total_created, "errors": errors,
                    "stories_done": stories_done, "total_stories": total_stories,
                    "per_story": list(per_story_stats.values()),
                    "total_secs": _total_secs})
        return
    cb("done", {"summary": f"{total_created} created · {errors} failed",
                "created": total_created, "errors": errors,
                "stories_done": stories_done, "total_stories": total_stories,
                "per_story": list(per_story_stats.values()),
                "total_secs": _total_secs})


def run_steps(project, plan_id, story_ids, cb, should_stop=lambda: False,
              existing_mode="skip", dedupe_existing=True, on_ai_error=None,
              gate=None, ops=None):
    """existing_mode: 'skip' or 'evaluate'. dedupe_existing: remove pre-existing
    duplicate test cases in each suite before processing."""
    _o = ops or _default_ops()
    # The step-generation loop below runs a small bounded pool of workers so
    # multiple cases' generate_steps() calls (the dominant per-case cost) can
    # be in flight at once instead of strictly one-at-a-time. Every existing
    # log line / heartbeat / progress update in this function calls cb() —
    # previously always from this one function running on a single thread, so
    # cb() was implicitly never called concurrently. Shadowing cb here with a
    # lock-guarded wrapper preserves that same guarantee (cb is still only
    # ever invoked one call at a time, fully serialized) even though the work
    # THAT LEADS UP TO some of those calls may now run on worker threads —
    # every single cb(...) call site below this line (unchanged, pre-existing
    # ones included) gets this for free without needing to touch each one.
    _cb_lock = _threading.Lock()
    _real_cb = cb
    def cb(kind, payload):
        with _cb_lock:
            _real_cb(kind, payload)

    wit, test = _o.connect(project)
    # Timer starts HERE, before suite discovery/removal/seeding — those used to
    # run before _run_start was set further down, so the pre-existing-duplicate
    # removal pass (which can take a while, especially with its AI mop-up) was
    # invisible in the run's reported "Time" stat even though it's real elapsed
    # time the user waited through.
    _run_start = time.time()
    cb("log", {"msg": "Discovering suites for stories…", "tone": "dim"})
    story_suite_map = _o.discover_suites(project, plan_id, set(story_ids))
    # DO NOT re-add a "one story per suite" guard here (tried in #120, reverted
    # in #125 after it broke every non-Azure run). `story_suite_map` is NOT
    # story→suite one-to-one: backend_setup.generation_ops.discover_suites
    # deliberately keys the SAME story under up to four ALIASES — sid (int),
    # str(sid), story.ref.id and story.ref.key — all pointing at one suite, so
    # that `story_suite_map.get(sid)` matches regardless of int-vs-str. A guard
    # that treats a repeated suite id as "a duplicate story" therefore deletes
    # that story's own aliases, including the key the caller looks up: seen live
    # on Azure→TestRail as "Suite 45 already covered by story 101048 — skipping
    # duplicate story 101048" (the SAME id) followed by "No suite for story
    # 101048 — skipped", i.e. a run that processed 0 stories.
    # The hybrid double-count this was meant to fix is already handled where it
    # actually belongs — the case-level idempotence guard further down, which
    # dedupes by test-case work-item id and is alias-agnostic.
    stories = _o.fetch_stories(story_ids)
    story_ctx = {}
    _ac_link_cache = {}   # shared across all stories this run — see _resolve_ac_links
    for s in stories:
        _criteria = s.fields.get("Microsoft.VSTS.Common.AcceptanceCriteria", "")
        _criteria = _resolve_ac_links(_criteria, project, cb=cb, cache=_ac_link_cache, current_id=s.id)
        story_ctx[s.id] = {
            "title": s.fields.get("System.Title", "No Title"),
            "criteria": _criteria,
            "screenshots": fetch_story_screenshots(s),
            # Guards the once-per-story ui_desc lazy-fill below (_gen_and_write)
            # now that several of a story's cases can reach that fill point on
            # DIFFERENT worker threads at once — pre-created here (single-
            # threaded at this point) so there's no race creating the lock
            # itself; _gen_and_write also lazily setdefault()s it as a
            # belt-and-suspenders fallback.
            "ui_desc_lock": _threading.Lock(),
        }
    # Log each story → suite mapping up front
    for sid in story_ids:
        suite_id = story_suite_map.get(sid)
        title = story_ctx.get(sid, {}).get("title", "")
        if suite_id:
            # "ar" must reflect whether the TITLE is actually Arabic, not just
            # "a title exists" — bool(title) forced every story line (English
            # titles included, e.g. "Change Language"/"Profile"/"Date") into
            # right-aligned RTL in the log and the run email report.
            _title_is_ar = any('؀' <= c <= 'ۿ' for c in title)
            cb("log", {"msg": f"Story {sid} → suite {suite_id} · {title}",
                       "tone": "story", "ico": "▸", "ar": _title_is_ar})
        else:
            cb("log", {"msg": f"Story {sid} — no suite found/created, skipped",
                       "tone": "warn", "ico": "⚠"})
    # ── Remove pre-existing duplicate test cases in each suite ──
    # Catches duplicates already sitting in Azure (from prior runs or manual
    # entry), keeping the most complete one of each group and deleting the rest.
    # Track which test case IDs were kept as a same-run dedup "winner": a keeper
    # can absorb scope from the duplicates it survives (that's often WHY it has
    # more steps and wins), so the evaluate-existing-steps pass below must not
    # judge it against its own narrow title alone — that would shrink it back
    # down and permanently lose coverage that only existed in the now-deleted
    # duplicates.
    _dedup_keeper_ids_this_run = set()
    import concurrent.futures as _cf_rs
    if dedupe_existing:
        cb("log", {"msg": "Checking suites for duplicate test cases…", "tone": "dim"})
        # Each suite's dedup pass is independent (its own suite, its own AI
        # mop-up call), so run a couple concurrently instead of strictly one
        # at a time — same bounded-pool rationale as the Steps-generation
        # worker pool further below (kept small on purpose: this still hits
        # the same rate-limited free-tier endpoint, just a couple of calls in
        # flight instead of one, not meant to multiply request volume). cb is
        # already lock-guarded from this point on (see _cb_lock above), so
        # concurrent dedupe_existing_suite calls logging through it is safe.
        _DEDUP_WORKERS = 2
        _dedup_jobs = [(sid, suite_id) for sid, suite_id in story_suite_map.items() if suite_id]
        with _cf_rs.ThreadPoolExecutor(max_workers=_DEDUP_WORKERS) as _dedup_ex:
            _dedup_futs = {}
            for sid, suite_id in _dedup_jobs:
                if should_stop():
                    break
                fut = _dedup_ex.submit(
                    _o.dedupe_suite, project, plan_id, suite_id, cb=cb,
                    do_delete=True, should_stop=should_stop, story_id=sid,
                    story_title=story_ctx.get(sid, {}).get("title", ""))
                _dedup_futs[fut] = (sid, suite_id)
            for fut in _cf_rs.as_completed(_dedup_futs):
                sid, suite_id = _dedup_futs[fut]
                try:
                    _dd_res = fut.result()
                    _dedup_keeper_ids_this_run.update(_dd_res.get("keeper_ids", []) or [])
                except Exception as de:
                    cb("log", {"msg": f"Story {sid} — dedup skipped for suite {suite_id}: "
                                      f"{str(de)[:80]}", "tone": "warn"})
    # ── Seed titles for empty suites ──
    # If a story's suite has NO test cases yet, generate titles first (same dedup
    # as the titles tool) and create the test cases, so the steps run can proceed
    # on a fresh plan. Suites that already have test cases are left untouched —
    # those go through the normal Skip/Evaluate path below.
    seeded_total = 0
    for sid, suite_id in story_suite_map.items():
        if should_stop(): break
        if not suite_id:
            continue
        # Current test cases in this suite — presence only (no titles needed
        # for this decision), so the fast raw listing is fine here. The bug
        # this replaces: this used to filter on workItem.name, which is
        # always blank on the raw listing (see fetch_existing_titles_for_suite's
        # docstring) — so every ALREADY-POPULATED suite looked empty and got
        # seeded with a fresh batch of titles on top of what was already there.
        try:
            has_existing = bool(_o.cases_for_suite(project, plan_id, suite_id))
        except Exception:
            has_existing = False
        if has_existing:
            continue  # suite already populated → handled by Skip/Evaluate later

        ctx = story_ctx.get(sid, {})
        s_title = ctx.get("title", "")
        s_criteria = ctx.get("criteria", "")
        cb("log", {"msg": f"Suite {suite_id} is empty — generating test case titles…",
                   "tone": "dim", "ico": "▸"})
        # Sentinel loop instead of a bare try: a FATAL provider error here can
        # now PAUSE via on_ai_error (switch provider + Resume retries this
        # seeding call; Stop keeps the old end-the-run behavior) instead of
        # unconditionally killing the whole run.
        _titles = None
        while _titles is None:
            try:
                _titles = _call_with_network_retries(lambda: generate_titles(
                    s_title, s_criteria, [],
                    log=lambda m, t="warn": cb("log", {"msg": m, "tone": t}),
                    should_stop=should_stop,
                    on_slow=lambda s: cb("log", {
                        "msg": f"Still generating titles with {T_disp(AI_PROVIDER)} "
                               f"({current_model() or 'model'}) — {s}s so far. Large models "
                               f"on free tiers can be slow; click Stop to cancel.",
                        "tone": "dim", "ico": "⏳", "hb_id": f"seedtitles:{sid}"}),
                    on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                                  "hb_id": f"seedtitles:{sid}"}),
                ), cb, should_stop=should_stop) or []
            except StopRequested:
                _titles = "STOP"
            except CreditBalanceError:
                if on_ai_error:
                    cb("log", {"msg": "Out of AI credits.", "tone": "err"})
                    cb("log", {"msg": _paused_on_error_msg(), "tone": "warn", "ico": "⏸"})
                    if on_ai_error("Out of AI credits — add credits or "
                                   "switch to another provider") == "retry":
                        continue
                cb("done", {"summary": "Stopped — out of AI credits", "reason": "credit"}); return
            except Exception as e:
                cat, friendly = classify_ai_error(e)
                if cat in ("auth", "bad_model", "not_found", "network"):
                    if on_ai_error:
                        cb("log", {"msg": friendly, "tone": "err"})
                        cb("log", {"msg": _paused_on_error_msg(), "tone": "warn", "ico": "⏸"})
                        if on_ai_error(friendly) == "retry":
                            continue
                    cb("log", {"msg": friendly, "tone": "err"})
                    cb("done", {"summary": f"Stopped — {friendly}", "reason": cat}); return
                cb("log", {"msg": f"Title generation failed for story {sid}: {e}", "tone": "err"})
                _titles = "SKIP"
        if _titles == "STOP":
            break
        if _titles == "SKIP":
            continue
        titles = _titles

        # Clears the "Still generating titles…" heartbeat above once the call
        # actually resolves — see the identical fix/comment in run_titles's
        # _process_story for why this is needed (nothing below fires
        # unconditionally, so a heartbeat that fired even once would
        # otherwise be a permanent ghost).
        cb("log", {"msg": f"Generated {len(titles)} candidate title(s)", "tone": "dim",
                   "replace_wip": f"seedtitles:{sid}"})

        # de-duplicate (exact + semantic), same as the titles tool
        seen_norm = set(); seen_keys = set(); unique = []; dropped = 0
        for t in titles:
            nk = _norm_title(t); sk = _semantic_key(t)
            if nk in seen_norm or _is_near_duplicate(sk, seen_keys):
                dropped += 1; continue
            unique.append(t); seen_norm.add(nk); seen_keys.add(sk)
        if dropped:
            cb("log", {"msg": f"Skipped {dropped} duplicate/near-duplicate title"
                              + ("s" if dropped > 1 else ""), "tone": "dim", "ico": "⏭"})
        # Second pass: catch synonym/paraphrase duplicates the word-overlap
        # filter above can't see (same reasoning as the titles-tool path).
        _before_ai2 = len(unique)
        try:
            unique = _dedupe_titles_ai(
                unique, [], s_title,
                log=lambda m, t="dim": cb("log", {"msg": m, "tone": t}),
                should_stop=should_stop,
                on_slow=lambda s: cb("log", {
                    "msg": f"Still checking for duplicate titles with {T_disp(AI_PROVIDER)} "
                           f"({current_model() or 'model'}) — {s}s so far. Large models "
                           f"on free tiers can be slow; the run is not frozen.",
                    "tone": "dim", "ico": "⏳", "hb_id": f"dedupetitles:{sid}"}),
                on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                              "hb_id": f"dedupetitles:{sid}"}))
        except StopRequested:
            break
        # Clears the "Still checking for duplicate titles…" heartbeat above —
        # same reasoning as the generate_titles cleanup a few lines up.
        cb("log", {"msg": "Title dedup check complete", "tone": "dim",
                   "replace_wip": f"dedupetitles:{sid}"})
        if len(unique) != _before_ai2:
            dropped += (_before_ai2 - len(unique))
        for tc_title in unique:
            if should_stop(): break
            try:
                new_id = _o.create_case(project, plan_id, suite_id, tc_title, sid)
                seeded_total += 1
                cb("log", {"msg": tc_title + " — test case created", "tone": "ok",
                           "id": new_id, "ar": True})
            except Exception as e:
                cb("log", {"msg": f"{tc_title} — {e}", "tone": "err", "ar": True})
    if seeded_total:
        cb("log", {"msg": f"Created {seeded_total} new test case"
                          + ("s" if seeded_total > 1 else "")
                          + " — now generating steps…", "tone": "ok", "ico": "✓"})

    # Build flat test-case list (now includes any freshly-seeded cases)
    suite_test_cases = []
    for sid, suite_id in story_suite_map.items():
        try:
            for tc in _o.cases_for_suite(project, plan_id, suite_id):
                suite_test_cases.append((tc, sid, suite_id))
        except Exception:
            pass

    # IDEMPOTENCE GUARD (roadmap #120): a given test case must be processed at
    # most once. On the Azure→TestRail hybrid a single story could get its cases
    # enumerated twice (the run reading both the Azure suite and the TestRail
    # write-target suite for the same story), which doubled the count (e.g.
    # "68" = 34×2), showed the suite twice in the Run rail, and listed the story
    # twice in the report. Collapse by the case's unique work-item id, keeping the
    # first occurrence. A correct Azure run has no duplicates here, so this is a
    # no-op there — it never drops a distinct case or story.
    if suite_test_cases:
        _seen_tc, _deduped = set(), []
        for _entry in suite_test_cases:
            _tc = _entry[0]
            _cid = (_tc.get("workItem", {}) or {}).get("id") if isinstance(_tc, dict) else None
            if _cid is not None:
                if _cid in _seen_tc:
                    continue
                _seen_tc.add(_cid)
            _deduped.append(_entry)
        if len(_deduped) != len(suite_test_cases):
            cb("log", {"msg": f"Skipped {len(suite_test_cases) - len(_deduped)} duplicate "
                              f"test-case enumeration(s) (same case seen twice).",
                       "tone": "warn", "ico": "⚠"})
        suite_test_cases = _deduped

    # Titles: the suite listing carries only the work-item id (witFields=System.Id),
    # so fetch each case's title CONCURRENTLY → real titles in the log/report instead
    # of "No Title".
    _title_map = {}
    _tc_ids = [i for i in (tc.get("workItem", {}).get("id")
                           for tc, _, _ in suite_test_cases) if i]
    if _tc_ids:
        try:
            import concurrent.futures as _cf
            with _cf.ThreadPoolExecutor(max_workers=min(16, len(_tc_ids))) as _ex:
                for _tid, _t in _ex.map(lambda i: (i, _o.case_title(i)), _tc_ids):
                    _title_map[_tid] = _t
        except Exception:
            pass

    total = len(suite_test_cases)
    # Count stories that actually have test cases to process
    _stories_with_tc = set(sid for _, sid, _ in suite_test_cases)
    total_stories = len(_stories_with_tc) if _stories_with_tc else len(story_suite_map)
    done = ok = err = skipped = stories_done = 0
    action_items = []
    skipped_items = []
    from collections import Counter as _C
    ok_by_story = _C(); skip_by_story = _C(); err_by_story = _C()
    time_by_story = {}        # {sid: cumulative seconds}
    # _run_start now set at the top of the function (see comment there) so the
    # dedupe-existing/seeding passes above count toward the total elapsed time.
    from collections import Counter
    remaining = Counter(sid for _, sid, _ in suite_test_cases)
    story_total = Counter(sid for _, sid, _ in suite_test_cases)
    # Per-story progress: {sid: {"total","done","ok","skipped","err","title","suite"}}
    story_prog = {}
    for sid in story_total:
        title = story_ctx.get(sid, {}).get("title", "")
        suite = story_suite_map.get(sid)
        story_prog[sid] = {"total": story_total[sid], "done": 0, "ok": 0,
                           "skipped": 0, "err": 0, "title": title, "suite": suite}
        cb("story_progress", {"id": sid, **story_prog[sid]})

    cb("stat", {"total": total, "stories_done": 0, "total_stories": total_stories,
                "done": 0, "skipped": 0, "errors": 0, "created": seeded_total})

    # ── bounded concurrency for the expensive part (evaluate + generate_steps + write) ──
    # Only the CHEAP, non-AI checks stay sequential on the main thread: the
    # existing-steps fetch, the instant "already has steps, skipped" (skip
    # mode) short-circuit, and the dedup-keeper short-circuit. Every case that
    # isn't short-circuited there — a brand-new case, OR an existing one that
    # needs the evaluate_existing_steps AI decision — is handed to a small
    # worker pool where _gen_and_write does the whole rest of the pipeline:
    # the evaluate decision (when applicable), the once-per-story UI-
    # description fetch, and generate_steps() + the Azure write. This used to
    # split evaluate_existing_steps out as its OWN sequential main-thread
    # phase before any case ever reached this pool — for a story whose cases
    # mostly already have steps, that left the pool idle while cases were
    # evaluated one at a time (see _gen_and_write's docstring for the full
    # story). _STEPS_WORKERS is deliberately small: each in-flight call still
    # goes through the same ai_complete() retry/backoff machinery as before,
    # so this isn't meant to hammer a rate-limited/free-tier provider much
    # harder than the sequential path already could — it's meant to overlap
    # LATENCY (waiting on the network) across a few cases, not multiply
    # request volume. (_cf_rs already imported above, for the dedup-check
    # worker pool.)
    _STEPS_WORKERS = 2
    _executor = _cf_rs.ThreadPoolExecutor(max_workers=_STEPS_WORKERS)
    _inflight = {}   # future -> None
    _fatal = {"hit": False}   # first fatal (credit/auth/bad_model/not_found/network) wins

    def _finish_case(story_id):
        """The per-story progress/stat/log tail that used to run unconditionally
        at the bottom of every loop iteration — factored out so both the
        synchronous skip/evaluate paths AND async-completed generate results
        can call it the same way, with identical behavior to before."""
        nonlocal stories_done
        sp = story_prog.get(story_id)
        if sp is not None:
            sp["done"] = sp["total"] - (remaining[story_id] - 1)
            sp["ok"] = ok_by_story.get(story_id, 0)
            sp["skipped"] = skip_by_story.get(story_id, 0)
            sp["err"] = err_by_story.get(story_id, 0)
            cb("story_progress", {"id": story_id, **sp})
        remaining[story_id] -= 1
        if remaining[story_id] == 0:
            stories_done += 1
            cb("log", {"msg": f"Story {story_id} completed · all test cases processed",
                       "tone": "ok", "ico": "└"})
        pct = int(done / total * 100) if total else 0
        cb("stat", {"total": total, "stories_done": stories_done, "total_stories": total_stories,
                    "done": ok, "skipped": skipped, "errors": err, "created": seeded_total})
        cb("progress", {"pct": pct, "label": f"{pct}% · {done} of {total}"})

    def _gen_attempt(tc_id, tc_title, criteria, ctx, story_id, has_existing,
                     existing_xml, _tc_start, seq=""):
        """Runs on a worker thread. Owns the WHOLE post-dedup pipeline for one
        case: the evaluate_existing_steps AI decision when applicable, the
        once-per-story UI-description fetch (lazily, guarded by ctx's own
        lock), and generate_steps() + the Azure write.

        This used to be split in two: evaluate_existing_steps ran
        SEQUENTIALLY on the main thread for every case BEFORE any dispatch to
        this worker pool, and only generate+write was pooled. For a story
        whose cases mostly already have steps (the common case re-running
        against an established suite), that left the 2-worker pool sitting
        idle while cases were evaluated one at a time on the main thread —
        and evaluate_existing_steps had no should_stop/on_slow awareness at
        all, so a slow/rate-limited provider could silently stall the WHOLE
        run for minutes with nothing in the log (confirmed live: 2:56 of
        total silence evaluating a single case) and no way for Stop to
        interrupt it. Folding the decision into this worker fixes both: up to
        _STEPS_WORKERS cases' evaluate calls now genuinely run concurrently,
        and evaluate_existing_steps gets the same heartbeat/should_stop
        handling generate_steps already had.

        Returns a result dict describing exactly what happened; does NOT
        touch any shared counter — _apply_result (main thread only) does all
        bookkeeping. cb() itself IS called from here — safe because cb was
        shadowed with a lock at the top of run_steps."""
        inadequate_reason = ""
        if has_existing and existing_mode == "evaluate":
            try:
                if should_stop():
                    return {"ok": False, "stopped": True, "tc_id": tc_id, "tc_title": tc_title,
                            "story_id": story_id, "tc_start": _tc_start,
                            "stopped_wip_id": f"eval:{tc_id}"}
                cb("log", {"msg": tc_title + " — checking if existing steps are adequate…",
                           "tone": "dim", "id": tc_id, "seq": seq, "ar": True, "wip": True,
                           "wip_id": f"eval:{tc_id}"})
                verdict = evaluate_existing_steps(
                    tc_title, criteria, existing_xml, should_stop=should_stop,
                    on_slow=lambda s: cb("log", {
                        "msg": f"Still checking existing steps with {T_disp(AI_PROVIDER)} "
                               f"({current_model() or 'model'}) — {s}s so far. Large "
                               f"models on free tiers can be slow; the run is not frozen.",
                        "tone": "dim", "ico": "⏳", "id": tc_id, "indent": True,
                        "hb_id": f"eval:{tc_id}"}),
                    on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                                  "id": tc_id, "indent": True,
                                                  "hb_id": f"eval:{tc_id}"}))
            except StopRequested:
                # Must be caught explicitly, same reasoning as the generate-
                # phase StopRequested handling below: it carries no message
                # text, so if it fell through to the generic `except
                # Exception` below it would be misreported as "تعذر التقييم"
                # (evaluation failed) instead of an honest stop — or, worse,
                # if it escaped this function entirely uncaught (as it did
                # before this refactor, when evaluate_existing_steps ran on
                # the main thread with no try/except StopRequested around
                # it at all), it would propagate all the way out of
                # run_steps and surface in main.py's generic exception
                # handler as "Run failed: " (empty message) — a real bug
                # reported live, root-caused to exactly this gap.
                return {"ok": False, "stopped": True, "tc_id": tc_id, "tc_title": tc_title,
                        "story_id": story_id, "tc_start": _tc_start,
                        "stopped_wip_id": f"eval:{tc_id}"}
            except CreditBalanceError:
                return {"ok": False, "credit": True, "tc_id": tc_id, "tc_title": tc_title,
                        "story_id": story_id, "tc_start": _tc_start}
            except Exception:
                verdict = {"adequate": False, "reason": "تعذر التقييم"}
            if verdict.get("adequate"):
                return {"ok": True, "skip_adequate": True, "tc_id": tc_id, "tc_title": tc_title,
                        "story_id": story_id,
                        "reason": verdict.get("reason", "Existing steps adequate"),
                        "tc_start": _tc_start}
            inadequate_reason = verdict.get("reason", "")

        # UI description cached once per story, lazily, guarded by ctx's own
        # lock — several of this story's cases can reach this at once now
        # that evaluate_existing_steps also runs on a worker (previously this
        # ran on the main thread, one case at a time, so a plain "'ui_desc'
        # not in ctx" check was already safe with no lock needed). Double-
        # checked locking: skip the lock entirely once ui_desc is set.
        if "ui_desc" not in ctx:
            _uidesc_lock = ctx.setdefault("ui_desc_lock", _threading.Lock())
            with _uidesc_lock:
                if "ui_desc" not in ctx:
                    if ctx.get("screenshots"):
                        cb("log", {"msg": f"read {len(ctx['screenshots'])} screenshot(s) once — UI described",
                                   "tone": "dim", "ico": "👁", "indent": True})
                    try:
                        ctx["ui_desc"] = _call_with_network_retries(
                            lambda: describe_story_ui(
                                ctx.get("screenshots"), ctx.get("title", ""),
                                on_slow=lambda s: cb("log", {
                                    "msg": f"Still describing the UI with {T_disp(AI_PROVIDER)} "
                                           f"({current_model() or 'model'}) — {s}s so far. "
                                           f"Vision calls on free-tier providers can be slow.",
                                    "tone": "dim", "ico": "⏳", "indent": True,
                                    "hb_id": f"uidesc:{story_id}"}),
                                on_retry=lambda m: cb("log", {"msg": m, "tone": "warn",
                                                              "ico": "⏳", "indent": True,
                                                              "hb_id": f"uidesc:{story_id}"}),
                            ), cb, should_stop=should_stop)
                        # Clears the heartbeat line above once the call actually
                        # resolves — without this a heartbeat that fired even
                        # once is never removed (nothing else logs a "this call
                        # finished" line to hitch replace_wip onto), so it sits
                        # in the activity log forever frozen at its last value.
                        cb("log", {"msg": "UI description ready", "tone": "dim", "ico": "👁",
                                   "indent": True, "replace_wip": f"uidesc:{story_id}"})
                    except StopRequested:
                        # THE bug reported live: this call previously had no
                        # except StopRequested clause at all (it only caught
                        # CreditBalanceError), so clicking Stop while this was
                        # in flight propagated all the way out of run_steps
                        # uncaught, landing in main.py's generic exception
                        # handler as "Run failed: " (StopRequested's message
                        # text is empty). Fixed by handling it explicitly here,
                        # same as every other AI call site in this run.
                        return {"ok": False, "stopped": True, "tc_id": tc_id, "tc_title": tc_title,
                                "story_id": story_id, "tc_start": _tc_start,
                                "stopped_wip_id": f"uidesc:{story_id}"}
                    except CreditBalanceError:
                        return {"ok": False, "credit": True, "tc_id": tc_id, "tc_title": tc_title,
                                "story_id": story_id, "tc_start": _tc_start}
        ui_desc = ctx.get("ui_desc", "")

        cb("log", {"msg": tc_title + " — generating…", "tone": "info",
                   "id": tc_id, "seq": seq, "ar": True, "wip": True, "wip_id": tc_id,
                   "replace_wip": f"eval:{tc_id}"})
        try:
            steps = _call_with_network_retries(
                lambda: generate_steps(
                    tc_title, criteria, ui_desc,
                    log=lambda m, t="warn": cb("log", {"msg": m, "tone": t}),
                    on_slow=lambda s: cb("log", {
                        "msg": f"Still generating steps with {T_disp(AI_PROVIDER)} "
                               f"({current_model() or 'model'}) — {s}s so far. Large "
                               f"models on free tiers can be slow; the run is not frozen.",
                        "tone": "dim", "ico": "⏳", "id": tc_id, "indent": True,
                        "hb_id": tc_id}),
                    on_retry=lambda m: cb("log", {"msg": m, "tone": "warn", "ico": "⏳",
                                                  "id": tc_id, "indent": True,
                                                  "hb_id": tc_id}),
                ), cb, should_stop=should_stop)
            # Drop steps with neither action nor expected text — seen live: a
            # technically-valid JSON reply (e.g. {"steps":[{}]}) with empty
            # string values sails straight through parse_json_robust/
            # _coerce_step_list as an apparently successful "1 steps" result,
            # gets written to Azure DevOps as a genuinely blank step row, and
            # only shows up as "the test case has no steps" when someone
            # actually opens it — nothing before this point ever checked the
            # step's own content, only that the list was non-empty. If EVERY
            # step turns out empty this way, that's not a usable result at
            # all — raise so it's counted/logged as a real failure instead of
            # silently writing nothing useful.
            steps = [s for s in steps if (s.get("action", "").strip() or s.get("expected", "").strip())]
            if not steps:
                raise RuntimeError("AI returned steps with no action/expected text")
            _o.write_steps(tc_id, build_steps_xml(steps), project, story_id)
            return {"ok": True, "tc_id": tc_id, "tc_title": tc_title, "story_id": story_id,
                    "steps": steps, "inadequate_reason": inadequate_reason,
                    "tc_start": _tc_start}
        except StopRequested:
            # Stop was clicked while THIS case was mid-generation. Must be
            # caught before the generic `except Exception` below — StopRequested
            # carries no message text (str(StopRequested()) == ""), so if it
            # fell through to classify_ai_error() like a real provider error,
            # every case in flight at the moment of Stop would misleadingly
            # log as "{provider}: unknown error." (classify_ai_error's generic
            # fallback for an exception with empty text) — a real bug seen live:
            # clicking Stop while 2 cases were generating produced two fake
            # "NVIDIA: unknown error." lines that looked like provider failures
            # but were just this cancellation being misclassified.
            return {"ok": False, "stopped": True, "tc_id": tc_id, "tc_title": tc_title,
                    "story_id": story_id, "tc_start": _tc_start, "stopped_wip_id": tc_id}
        except CreditBalanceError:
            return {"ok": False, "credit": True, "tc_id": tc_id, "tc_title": tc_title,
                    "story_id": story_id, "tc_start": _tc_start}
        except Exception as e:
            return {"ok": False, "error": e, "tc_id": tc_id, "tc_title": tc_title,
                    "story_id": story_id, "tc_start": _tc_start}

    def _gen_and_write(*a):
        """Pause-and-ask wrapper around _gen_attempt. A FATAL provider error
        (out of credits / auth / bad model / not found / network) used to STOP
        the whole run outright ("Out of AI credits — run stopped", reported
        live with a run at 0/21). With `on_ai_error` wired (main.py's
        _run_on_ai_error), the run PAUSES instead: the callback blocks on a
        Condition showing a dialog where the user can switch provider and
        Resume → this case retries from scratch on the (re-read-per-attempt)
        provider; Stop → the fatal result flows to _apply_result and ends the
        run exactly as before. Ordinary per-case errors keep flowing through
        untouched — the run already continues past those. Concurrent workers
        hitting errors while paused wait on the same gate and reuse the same
        decision (same pattern as Automation's _auto_on_ai_error)."""
        while True:
            res = _gen_attempt(*a)
            if on_ai_error is None or res.get("stopped"):
                return res
            if res.get("credit"):
                friendly = "Out of AI credits — add credits or switch to another provider"
            elif res.get("error") is not None:
                _cat, friendly = classify_ai_error(res["error"])
                if _cat not in ("auth", "bad_model", "not_found", "network"):
                    return res
            else:
                return res
            cb("log", {"msg": friendly, "tone": "err", "id": res.get("tc_id")})
            cb("log", {"msg": _paused_on_error_msg(), "tone": "warn", "ico": "⏸"})
            if on_ai_error(friendly) == "retry":
                cb("log", {"msg": (res.get("tc_title") or "") +
                                  f" — retrying with {T_disp(AI_PROVIDER)}…",
                           "tone": "dim", "id": res.get("tc_id"), "ar": True})
                continue
            return res

    def _apply_result(res):
        """Main-thread-only: identical bookkeeping to what the old inline
        generate_steps() success/error handling did, now applied to a
        completed future's result instead of an immediate return value."""
        nonlocal ok, done, err, skipped
        story_id = res["story_id"]; tc_id = res["tc_id"]; tc_title = res["tc_title"]
        if res.get("credit"):
            if not _fatal["hit"]:
                _fatal.update(hit=True, summary="Stopped — out of AI credits", reason="credit")
            return
        if res.get("stopped"):
            # Honest "stopped", not a fake provider error — see the
            # StopRequested notes in _gen_and_write. Not counted toward
            # ok/err/done: this case simply didn't finish, same as if the
            # loop above had never dispatched it in the first place.
            # stopped_wip_id targets whichever wip/heartbeat line was actually
            # showing when the stop happened (eval:{tc_id}, uidesc:{story_id},
            # or plain tc_id for the generate phase) — falls back to tc_id so
            # this still behaves correctly for any older/other caller.
            cb("log", {"msg": tc_title + " — stopped (Stop was clicked while this "
                              "was generating)", "tone": "dim", "id": tc_id,
                       "ar": True, "replace_wip": res.get("stopped_wip_id", tc_id)})
            return
        if res.get("skip_adequate"):
            # existing steps judged adequate — the AI-decision analogue of the
            # "already has steps, skipped" fast path above, just resolved via
            # a (potentially slow) AI call on a worker thread instead of a
            # cheap sync check. Mirrors the original inline log/stats exactly.
            skipped += 1; done += 1; skip_by_story[story_id] += 1
            _el = round(time.time() - res["tc_start"], 1)
            skipped_items.append({"id": tc_id, "title": tc_title,
                                  "reason": res.get("reason", "Existing steps adequate"),
                                  "secs": _el})
            cb("log", {"msg": tc_title + " — existing steps adequate", "tone": "ok",
                       "id": tc_id, "ar": True, "secs": _el,
                       "detail": f"existing steps adequate · ⏱ {_fmt_mmss(_el)}",
                       "replace_wip": f"eval:{tc_id}"})
            time_by_story[story_id] = time_by_story.get(story_id, 0.0) + (time.time() - res["tc_start"])
            _finish_case(story_id)
            return
        if not res["ok"]:
            e = res["error"]
            cat, friendly = classify_ai_error(e)
            if cat in ("auth", "bad_model", "not_found", "network"):
                cb("log", {"msg": tc_title + f" — {friendly}", "tone": "err",
                           "id": tc_id, "ar": True, "replace_wip": tc_id})
                if not _fatal["hit"]:
                    _fatal.update(hit=True, summary=f"Stopped — {friendly}", reason=cat)
                return
            err += 1; done += 1; err_by_story[story_id] += 1
            cb("log", {"msg": tc_title + f" — {friendly}", "tone": "err", "id": tc_id,
                       "ar": True, "replace_wip": tc_id})
            time_by_story[story_id] = time_by_story.get(story_id, 0.0) + (time.time() - res["tc_start"])
            _finish_case(story_id)
            return
        steps = res["steps"]
        ok += 1; done += 1; ok_by_story[story_id] += 1
        # pre/action/expected each count steps with actual non-empty content
        # in that field — NOT just len(steps) repeated three times (the old
        # behavior). That made "action"/"expected" meaningless as a sanity
        # check: a step with genuinely blank action/expected text (seen
        # live — the AI's JSON parsed fine but with empty string values)
        # still showed as "1 steps · pre 0 · action 1 · expected 1", looking
        # completely normal in the log while the case being written to Azure
        # DevOps actually had no real step content — exactly what surfaced
        # later as "the test case has no steps." pre already counted real
        # content correctly; action/expected now match that same pattern,
        # so a genuinely broken result is visible here instead of hidden.
        npre = sum(1 for s in steps if s.get("precondition", "").strip())
        nact = sum(1 for s in steps if s.get("action", "").strip())
        nexp = sum(1 for s in steps if s.get("expected", "").strip())
        _elapsed = time.time() - res["tc_start"]
        cb("log", {"msg": tc_title, "tone": "ok", "id": tc_id, "ar": True,
                   "replace_wip": tc_id,
                   "secs": round(_elapsed, 1),
                   "detail": f"{len(steps)} steps · pre {npre} · action {nact} · "
                             f"expected {nexp} · ⏱ {_fmt_mmss(_elapsed)}"})
        if res.get("inadequate_reason"):
            action_items.append({"id": tc_id, "title": tc_title,
                                 "reason": res["inadequate_reason"],
                                 "secs": round(_elapsed, 1)})
        time_by_story[story_id] = time_by_story.get(story_id, 0.0) + _elapsed
        _finish_case(story_id)

    def _drain(wait_for_all):
        """Process completed futures. wait_for_all=False only harvests
        whatever's already done UNLESS the pool is already at capacity, in
        which case it genuinely BLOCKS until at least one slot frees up —
        that's what actually keeps the in-flight window bounded to
        _STEPS_WORKERS. wait_for_all=True blocks until every submitted
        future is done (used once at the very end).

        BUG FIX: this used to pass timeout=0 even in the "pool full" branch,
        which made the wait() call return IMMEDIATELY regardless of
        return_when — concurrent.futures.wait() can't honor FIRST_COMPLETED
        with a zero timeout, there's no time budget left to actually wait in.
        Net effect: the in-flight window was never really bounded — the scan
        loop kept dispatching new work as fast as it could decide it, so many
        more than _STEPS_WORKERS cases could pile up in the executor's own
        internal queue and then fire their 'generating…' lines back-to-back
        in a burst as workers cycled through the backlog (visible in the app
        as several simultaneous un-replaced 'generating…' lines instead of
        the intended 2)."""
        if not _inflight:
            return
        if wait_for_all:
            done_set, _ = _cf_rs.wait(list(_inflight), return_when=_cf_rs.ALL_COMPLETED)
        elif len(_inflight) >= _STEPS_WORKERS:
            # Pool at capacity: block for real until a slot frees up (no
            # timeout — this is the actual backpressure mechanism).
            done_set, _ = _cf_rs.wait(list(_inflight), return_when=_cf_rs.FIRST_COMPLETED)
        else:
            # Pool has room: just a non-blocking opportunistic harvest of
            # anything already finished — never delays dispatching the next
            # case.
            done_set, _ = _cf_rs.wait(list(_inflight), timeout=0,
                                      return_when=_cf_rs.ALL_COMPLETED)
        for fut in done_set:
            _inflight.pop(fut, None)
            try:
                _apply_result(fut.result())
            except Exception:
                pass   # a worker itself should never raise (all paths return a dict) —
                       # belt-and-suspenders so a bug here can't crash the whole run

    _user_stopped = False
    # Dispatch-order case counter — stamped as "seq" onto every per-case log
    # payload so the Run log shows "#k/N" progress on each case line, same as
    # the Automation log's "Sequencing case k/N" / "#k/N" convention. Rendered
    # by main.py's _render_one_log as its own LTR mono chip (like the [id]
    # chip), so it can't bidi-scramble against an Arabic title.
    _case_seq = 0
    try:
        for tc, story_id, suite_id in suite_test_cases:
            if _fatal["hit"]:
                break
            if should_stop():
                # Deliberately does NOT wait for in-flight generate_steps
                # calls below (see the shutdown(wait=...) call in `finally`)
                # — matching this codebase's existing "Stop should feel
                # instant, not wait out an in-flight AI call" precedent
                # (see DEV_ROADMAP.md's Automation-Stop fix). Any case already
                # submitted to the pool still finishes and still writes its
                # steps to Azure in the background (that work can't be
                # un-started), it just won't be counted in this run's summary.
                _user_stopped = True
                break
            if gate and not gate():   # manual Pause point (False = Stop clicked)
                _user_stopped = True
                break
            _tc_start = time.time()
            _case_seq += 1
            _seq = "%d/%d" % (_case_seq, total)
            wi = tc.get("workItem", {})
            tc_id = wi.get("id")
            tc_title = wi.get("name") or _title_map.get(tc_id) or "No Title"
            ctx = story_ctx.get(story_id, {})
            criteria = ctx.get("criteria", "")

            # Live progress the instant this case starts (so the bar leaves "Starting…"
            # and the active story flips to "Running" immediately, not only when it ends).
            _start_pct = int(done / total * 100) if total else 0
            cb("progress", {"pct": _start_pct, "label": f"{_start_pct}% · {done} of {total}"})
            _sp = story_prog.get(story_id)
            if _sp is not None and _sp.get("done", 0) == 0:
                # mark the story active (done stays 0 but emit so the card shows Running)
                cb("story_progress", {"id": story_id, **_sp, "_active": True})

            # existing steps?
            existing_xml = ""
            try:
                ex = wit.get_work_item(tc_id, fields=["Microsoft.VSTS.TCM.Steps"])
                existing_xml = (ex.fields or {}).get("Microsoft.VSTS.TCM.Steps", "") or ""
            except Exception:
                pass
            has_existing = bool(existing_xml and "<step " in existing_xml)

            if has_existing and existing_mode == "skip":
                skipped += 1; done += 1; skip_by_story[story_id] += 1
                _el = round(time.time() - _tc_start, 1)
                skipped_items.append({"id": tc_id, "title": tc_title,
                                      "reason": "Already had steps", "secs": _el})
                cb("log", {"msg": tc_title + " — already has steps, skipped", "tone": "skip",
                           "id": tc_id, "seq": _seq, "ico": "⏭", "ar": True, "secs": _el,
                           "detail": f"skipped · ⏱ {_fmt_mmss(_el)}"})
                time_by_story[story_id] = time_by_story.get(story_id, 0.0) + (time.time() - _tc_start)
                _finish_case(story_id)
                continue

            if has_existing and tc_id in _dedup_keeper_ids_this_run:
                # This case was just kept moments ago as the most-complete
                # survivor of a duplicate group in THIS run — its step count
                # may already exceed what its own title implies because it
                # absorbed scope from the (now permanently deleted) duplicates
                # it beat out. Judging it against its own narrow title here
                # would risk shrinking it back down and losing that coverage
                # for good, since the siblings that used to carry it are gone.
                # Leave it untouched this run regardless of existing_mode.
                # Fast/sync, no AI call — stays on the main thread, same as
                # the "already has steps, skipped" branch above.
                skipped += 1; done += 1; skip_by_story[story_id] += 1
                _el = round(time.time() - _tc_start, 1)
                skipped_items.append({"id": tc_id, "title": tc_title,
                                      "reason": "Kept as most-complete duplicate survivor this run",
                                      "secs": _el})
                cb("log", {"msg": tc_title + " — kept as the most complete of a duplicate "
                                  "set moments ago, left as-is this run to avoid losing "
                                  "the coverage that made it the keeper",
                           "tone": "ok", "id": tc_id, "seq": _seq, "ar": True, "secs": _el,
                           "detail": f"dedup keeper — left as-is · ⏱ {_fmt_mmss(_el)}"})
                time_by_story[story_id] = time_by_story.get(story_id, 0.0) + (time.time() - _tc_start)
                _finish_case(story_id)
                continue

            # Everything else — a brand-new case, OR an existing one that
            # needs the (potentially slow) evaluate_existing_steps AI
            # decision — now dispatches straight to the worker pool.
            # evaluate_existing_steps used to run HERE, synchronously, one
            # case at a time, before ANY case reached the pool — for a story
            # whose cases mostly already have steps (the common case
            # re-running against an established suite), that left both
            # workers idle while cases were evaluated one after another, each
            # a slow AI call with zero should_stop/heartbeat awareness
            # (confirmed live: a single evaluate call silently stalling the
            # whole run for 2:56, and clicking Stop during one of these
            # calls used to escape run_steps entirely uncaught and surface
            # in main.py as a fake "Run failed: " — see _gen_and_write's own
            # docstring). The evaluate decision, the once-per-story
            # UI-description fetch, and generate_steps()+write all now run
            # inside _gen_and_write on a worker thread instead.
            _drain(wait_for_all=False)
            # The lambda folds this case's "seq" into WHATEVER result dict
            # _gen_and_write returns (there are ~10 return-shape variants) so
            # _apply_result can stamp it onto every completion line without
            # each return literal needing its own edit. Captured via default
            # arg — a plain closure would late-bind the loop's _seq.
            fut = _executor.submit(
                lambda *a, _s=_seq: {**_gen_and_write(*a), "seq": _s},
                tc_id, tc_title, criteria, ctx, story_id, has_existing,
                existing_xml, _tc_start, _seq)
            _inflight[fut] = None

        # Drain whatever's still in flight before finishing up — ALWAYS, including
        # on a user Stop.
        #
        # This used to be skipped when the user stopped (`if not _user_stopped`)
        # so that Stop would "feel instant". That was wrong on three counts, all
        # reported live:
        #   1. It contradicts the button's own label. "Stop after current test
        #      case" promises the in-flight case finishes. Abandoning it is what
        #      the OTHER label ("Stop now") would mean.
        #   2. It left the UI permanently lying. Un-harvested futures never reach
        #      _apply_result, so the "generating…" line each case logged (wip_id =
        #      tc_id) never receives its replace_wip — the spinners spin forever
        #      while the run header says "Stopped".
        #   3. The abandoned work did NOT stop. Those cases kept running on pool
        #      threads and kept writing their steps to the tracker, invisibly,
        #      after the run had reported itself finished.
        # Draining costs the time of the in-flight AI calls (bounded by the pool
        # size), and in exchange every case resolves honestly: it either completes
        # normally or returns its StopRequested result, and BOTH paths log a line
        # carrying replace_wip — which is what actually clears the spinner.
        if _user_stopped and _inflight:
            # Say so. The drain below waits for real AI calls to return, which
            # can take a while — and this codebase has a long history of silent
            # waits reading as "the app froze" (see the heartbeat entries in
            # DEV_ROADMAP). The count makes the wait legible and bounded.
            _n = len(_inflight)
            cb("log", {"msg": f"Stop requested — finishing {_n} test case"
                              f"{'' if _n == 1 else 's'} already in progress…",
                       "tone": "dim", "ico": "⏳"})
        _drain(wait_for_all=True)
    finally:
        # wait=True for the same reason: shutting the pool down without waiting
        # is what let those threads keep writing after the run was "over".
        _executor.shutdown(wait=True)

    if _fatal["hit"]:
        cb("done", {"summary": _fatal["summary"], "reason": _fatal["reason"],
                    "action_items": action_items}); return

    per_story = []
    for sid, sp in story_prog.items():
        per_story.append({"id": sid, "title": sp["title"], "suite": sp["suite"],
                          "total": sp["total"], "ok": ok_by_story.get(sid, 0),
                          "skipped": skip_by_story.get(sid, 0), "err": err_by_story.get(sid, 0),
                          "secs": round(time_by_story.get(sid, 0.0), 1)})
    _total_secs = round(time.time() - _run_start, 1)
    cb("done", {"summary": f"{ok} updated · {skipped} skipped · {err} failed",
                "updated": ok, "skipped": skipped, "errors": err,
                "created": seeded_total,
                "stories_done": stories_done, "total_stories": total_stories,
                "action_items": action_items, "skipped_items": skipped_items,
                "per_story": per_story, "total_secs": _total_secs})


def validate_stories_in_plan(project, plan_id, story_ids):
    """Read-only check: returns (found, missing) story-id lists for the plan.
    A story is 'found' if it maps to a requirement suite already in the plan."""
    smap = discover_suites_for_stories(project, plan_id, set(story_ids), create_missing=False)
    found = [sid for sid in story_ids if sid in smap]
    missing = [sid for sid in story_ids if sid not in smap]
    return found, missing


def count_test_cases(project, plan_id, story_ids):
    """Real number of existing test cases across the given stories (parallel).
    Used for the live estimate so it shows a true count, not a guess."""
    import concurrent.futures as _cf
    try:
        smap = discover_suites_for_stories(project, plan_id, set(story_ids),
                                           create_missing=False)
    except Exception:
        return 0
    suites = [sid for sid in smap.values() if sid]
    if not suites:
        return 0
    def _one(suite_id):
        try:
            return len(fetch_test_cases_for_suite(project, plan_id, suite_id))
        except Exception:
            return 0
    with _cf.ThreadPoolExecutor(max_workers=min(16, len(suites))) as _ex:
        return sum(_ex.map(_one, suites))


def count_existing_steps(project, plan_id, story_ids):
    """Count test cases that already have steps (for the existing-steps modal)."""
    import concurrent.futures as _cf
    wit, test = connect_azure_sdk(project)
    smap = discover_suites_for_stories(project, plan_id, set(story_ids), create_missing=False)
    # Fetch each suite's test cases CONCURRENTLY (was a serial loop per suite).
    suites = [s for s in smap.values() if s]
    ids = []
    if suites:
        def _fetch_ids(suite_id):
            out = []
            try:
                for tc in fetch_test_cases_for_suite(project, plan_id, suite_id):
                    wid = tc.get("workItem", {}).get("id")
                    if wid:
                        out.append(wid)
            except Exception:
                pass
            return out
        with _cf.ThreadPoolExecutor(max_workers=min(16, len(suites))) as _ex:
            for _lst in _ex.map(_fetch_ids, suites):
                ids.extend(_lst)
    have = 0
    for i in range(0, len(ids), 200):
        try:
            for w in wit.get_work_items(ids[i:i+200], fields=["Microsoft.VSTS.TCM.Steps"]):
                sf = (w.fields or {}).get("Microsoft.VSTS.TCM.Steps", "") or ""
                if sf and "<step " in sf: have += 1
        except Exception:
            pass
    return have, len(ids)


# ═══════════════════════════════════════════════════════════════════════════════
#  EMAIL REPORT
# ═══════════════════════════════════════════════════════════════════════════════
# ── just-in-time shared-sender credential refresh ──────────────────────────
# The Gmail sender credential (App Password / sender / sender name) can be set
# org-wide by an admin AFTER a user is already signed in. The app fetches it at
# sign-in and caches it in GMAIL_APP_PASS, so a long-signed-in session would
# otherwise keep sending with a stale/empty value — the exact "admin set it but
# the member's report never sent" bug. main.py registers a best-effort refresher
# via set_sender_refresher(); every send site calls ensure_sender_creds() right
# before it checks/uses the credential, so the latest server value is always
# picked up just-in-time. It is debounced so a screen's own pre-check plus
# send_report()'s backstop don't double-fetch within one send.
_SENDER_REFRESHER = None
_sender_refresh_ts = 0.0
_SENDER_REFRESH_MIN_GAP = 8.0   # seconds


def set_sender_refresher(fn):
    """Register (or clear, with None) a best-effort callable that refreshes the
    shared sender creds. main.py wires this to a synchronous org-settings fetch."""
    global _SENDER_REFRESHER
    _SENDER_REFRESHER = fn


def ensure_sender_creds():
    """Best-effort: pull the latest org-shared sender creds just-in-time, then
    report whether a Gmail App Password is now available. NEVER raises; on any
    failure (offline, not signed in, no refresher registered) it falls back to
    whatever GMAIL_APP_PASS already holds, so it can only ever upgrade the cached
    value, never break a send. Debounced so repeated calls a few seconds apart
    (a screen pre-check followed by send_report) only fetch once."""
    global _sender_refresh_ts
    fn = _SENDER_REFRESHER
    if fn is not None:
        import time as _t
        now = _t.time()
        if now - _sender_refresh_ts >= _SENDER_REFRESH_MIN_GAP:
            _sender_refresh_ts = now
            try:
                fn()
            except Exception:
                pass
    return bool(GMAIL_APP_PASS)


def send_report(to_addrs, subject, html_body, attachments=None, refresh_sender=True):
    """Send an HTML email via Gmail SMTP, with optional file attachments.
    Returns (ok, error_msg).

    ``refresh_sender`` is normally True so report sends pick up the latest
    organization setting. Setup's explicit sender test sets it to False after
    placing the just-saved values in memory, preventing a background refresh
    for a SuperAdmin's own organization from replacing the selected org's
    settings mid-test.
    """
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.image import MIMEImage
    from email.mime.application import MIMEApplication
    # Backstop: refresh the shared sender creds just before sending, in case this
    # site reached us without its own pre-check having done so.
    if refresh_sender:
        ensure_sender_creds()
    if not GMAIL_APP_PASS or not to_addrs:
        return False, "No Gmail password or recipients configured."
    # multipart/related so the HTML can reference the logo as cid:qastudio-logo
    from email.utils import formataddr
    msg = MIMEMultipart("related")
    msg["Subject"] = subject
    # Show a friendly display name to the recipient: From: "QA Studio" <addr>.
    # A blank name falls back to the bare address.
    msg["From"] = formataddr(((GMAIL_SENDER_NAME or "").strip(), GMAIL_SENDER))
    msg["To"] = ", ".join(to_addrs) if isinstance(to_addrs, list) else to_addrs
    _alt = MIMEMultipart("alternative")
    _alt.attach(MIMEText(html_body, "html"))
    msg.attach(_alt)
    # file attachments (e.g. the regression plan as Word/Excel/PDF)
    for _path in (attachments or []):
        try:
            with open(_path, "rb") as _af:
                _part = MIMEApplication(_af.read())
            _part.add_header("Content-Disposition", "attachment",
                             filename=os.path.basename(_path))
            msg.attach(_part)
        except Exception:
            pass
    # inline brand logo (safe no-op if the file is missing)
    try:
        _lp = _logo_path()
        if _lp:
            with open(_lp, "rb") as _f:
                _img = MIMEImage(_f.read())
            _img.add_header("Content-ID", f"<{LOGO_CID}>")
            _img.add_header("Content-Disposition", "inline",
                            filename=os.path.basename(_lp))
            msg.attach(_img)
    except Exception:
        pass
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=20) as s:
            s.login(GMAIL_SENDER, GMAIL_APP_PASS)
            s.send_message(msg)
        return True, ""
    except smtplib.SMTPAuthenticationError:
        return False, ("Gmail rejected the App Password. Generate a new 16-character "
                       "App Password (2-Step Verification must be on) and save it.")
    except smtplib.SMTPRecipientsRefused:
        return False, "Recipient address was refused — check the report email addresses."
    except smtplib.SMTPSenderRefused:
        return False, f"Sender {GMAIL_SENDER} was refused by Gmail."
    except smtplib.SMTPException as e:
        m = str(e).lower()
        if "username and password not accepted" in m or "badcredentials" in m or "535" in m:
            return False, ("Gmail rejected the App Password. Generate a new 16-character "
                           "App Password (2-Step Verification must be on) and save it.")
        return False, f"Email failed to send ({type(e).__name__})."
    except Exception as e:
        m = str(e).lower()
        if "timed out" in m or "timeout" in m or "connection" in m or "ssl" in m:
            return False, "Could not reach Gmail SMTP — check your network/firewall (port 465)."
        return False, "Email failed to send."


def _fmt_secs(s):
    """Human-friendly duration: 45s, 1m 20s, 2m."""
    try:
        s = float(s)
    except Exception:
        return ""
    if s < 60:
        return f"{s:.0f}s"
    m = int(s // 60); sec = int(round(s - m * 60))
    return f"{m}m {sec}s" if sec else f"{m}m"


def _fmt_mmss(s):
    """Duration as mm:ss (e.g. 0:45, 1:57, 12:03)."""
    try:
        s = float(s)
    except Exception:
        return ""
    m = int(s // 60); sec = int(round(s - m * 60))
    if sec == 60:
        m += 1; sec = 0
    return f"{m}:{sec:02d}"


# ---- brand logo, embedded inline in emails via Content-ID ----
LOGO_CID = "qastudio-logo"

def _logo_path():
    """Path to the inline email logo (transparent Q mark), next to this file.
    Prefers qa-logo-email.png — a plain, no-glow crop of the mark made
    specifically for email (the login/sidebar app.png and qa-logo.png now
    carry a soft blurred glow baked into the image, which suits an in-app
    dark surface but isn't what we want floating on a plain email background),
    falling back to the glow versions if that file isn't present."""
    try:
        here = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        return ""
    for name in ("qa-logo-email.png", "qa-logo.png", "app.png"):
        p = os.path.join(here, name)
        if os.path.exists(p):
            return p
    return ""

def _logo_tag(size=42):
    """<img> referencing the CID-embedded logo; degrades to alt text if blocked."""
    return (f"<img src='cid:{LOGO_CID}' width='{size}' height='{size}' alt='QA Studio' "
            f"style='display:block;border:0;outline:none;text-decoration:none' />")


def build_report_email(tool, summary, stats, action_items=None, skipped_items=None,
                       per_story=None, plan_url=None, total_secs=None, log_lines=None,
                       org=None, project=None, url_map=None):
    """Restrained, email-safe (table + inline-style) HTML run report.
    Renders consistently across Outlook / Gmail / Apple Mail; web fonts fall
    back to system fonts. Drives off the same data the in-app report uses."""
    import datetime as _dt

    # ---- palette ----
    # VIOLET* matches the app's actual brand accent (theme.py's VIOLET/VIOLET_INK/
    # VIOLET_SOFT — a cyan/teal-blue, not literal indigo/violet despite the
    # variable names carried over from an earlier palette). This used to be a
    # separate, unrelated royal-blue (#3A57D6) hardcoded just for the email, so
    # the report's accent color didn't match the rest of the app at all.
    PAPER="#E9E8EE"; CARD="#FFFFFF"; TINT="#FAFAFC"
    INK="#1B1A22"; INK2="#6B6975"; INK3="#9C9AA6"
    LINE="#E8E7EE"; LINE2="#F1F0F5"
    VIOLET="#0E9CC0"; VIOLET_INK="#0B6E86"; VIOLET_SOFT="#D6F4FB"
    GREEN="#1F8A52"; GREEN_SOFT="#E7F4ED"
    RED="#D6414A"; RED_SOFT="#FBEAEC"
    AMBER="#AB780C"; AMBER_SOFT="#F7EFD8"
    UI='"Segoe UI",Roboto,Helvetica,Arial,sans-serif'
    MONO='"SFMono-Regular",Consolas,Menlo,monospace'
    AR='"Segoe UI","Tahoma",Arial,sans-serif'

    def _wi_url(item_id):
        # Backend-aware first: `url_map` carries the correct per-item link
        # (Jira issue / Xray or Zephyr test / TestRail case / Azure work item),
        # resolved through the backend seam by the caller. The dev.azure.com
        # form is a last-resort fallback — correct ONLY for Azure, so it must
        # not fire when a url_map entry exists (that was the "email links point
        # at Azure on a Jira connection" bug).
        if item_id and url_map:
            u = url_map.get(str(item_id))
            if u:
                return u
        if not (org and project and item_id):
            return ""
        return f"https://dev.azure.com/{org}/{project}/_workitems/edit/{item_id}"

    def _intval(v):
        try:
            return int(str(v).split('/')[0].strip())
        except Exception:
            return 0

    def _is_ar(s):
        return any('\u0600' <= c <= '\u06ff' for c in str(s))

    stats = stats or {}
    is_steps = "step" in str(tool).lower()
    review_n = len(action_items or [])
    failed_n = _intval(stats.get("Failed", 0))
    stopped = "stop" in str(summary).lower()

    # ---- status + accent ----
    if failed_n > 0:
        pill_fg, pill_bg, pill_txt = RED, RED_SOFT, "Completed with errors"; accent = RED
    elif stopped:
        pill_fg, pill_bg, pill_txt = AMBER, AMBER_SOFT, "Stopped early"; accent = AMBER
    else:
        pill_fg, pill_bg, pill_txt = GREEN, GREEN_SOFT, "Completed"; accent = VIOLET
    check_ic = "&#10003;" if pill_fg != AMBER else "&#9632;"
    status_pill = (f"<span style='display:inline-block;background:{pill_bg};color:{pill_fg};"
                   f"font-size:11px;font-weight:700;letter-spacing:.4px;padding:5px 12px;"
                   f"border-radius:20px'>{check_ic}&nbsp; {pill_txt.upper()}</span>")

    # ---- headline ----
    if is_steps:
        n = _intval(stats.get("Updated", 0))
        headline = (f"<b style='color:{VIOLET_INK}'>{n} test case" + ("s" if n != 1 else "") + "</b> updated") if n else "No test cases updated"
    else:
        n = _intval(stats.get("Created", 0))
        headline = (f"<b style='color:{VIOLET_INK}'>{n} title" + ("s" if n != 1 else "") + "</b> created") if n else "No new titles created"
    sub = _html.escape(str(summary or ""))

    hero = (f"{status_pill}"
            f"<div style='font-size:25px;font-weight:700;letter-spacing:-.5px;color:{INK};"
            f"line-height:1.15;margin:14px 0 0'>{headline}</div>"
            f"<div style='font-size:13px;color:{INK2};font-weight:600;margin-top:8px'>{sub}</div>")

    # ---- masthead ----
    today = _dt.date.today().strftime("%d %b %Y")
    kind = "Test Case Steps" if is_steps else "Test Case Titles"
    masthead = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        f"<td width='40' valign='middle' style='padding-right:13px'>{_logo_tag(40)}</td>"
        f"<td valign='middle'>"
        f"<div style='font-size:15px;font-weight:700;color:{INK};letter-spacing:-.2px'>QA Studio</div>"
        f"<div style='font-size:12px;font-weight:700;color:{VIOLET_INK};margin-top:2px'>{kind} &middot; Run report</div>"
        f"</td>"
        f"<td valign='middle' align='right' style='font-family:{MONO};font-size:11px;"
        f"color:{INK3};font-weight:700'>{today}</td>"
        f"</tr></table>")

    # ---- metric strip ----
    items = list(stats.items())
    if is_steps:
        merged = []
        for k, v in items:
            merged.append((k, v))
            if k == "Updated":
                merged.append(("Review", review_n))
        items = merged

    def _mcolor(k, v):
        iv = _intval(v)
        if k in ("Updated", "Created") and iv > 0: return GREEN
        if k == "Review" and iv > 0: return AMBER
        if k == "Failed" and iv > 0: return RED
        if k in ("Time", "Stories"): return INK
        return INK3
    mcells = ""
    for i, (k, v) in enumerate(items):
        col = _mcolor(k, v)
        bl = "" if i == 0 else f"border-left:1px solid {LINE2};"
        vsize = "18px" if k == "Time" else "24px"
        mcells += (f"<td width='1' style='{bl}padding:13px 6px 14px;text-align:center;vertical-align:top'>"
                   f"<div style='font-size:9.5px;font-weight:700;letter-spacing:1px;color:{INK3};"
                   f"text-transform:uppercase'>{_html.escape(str(k))}</div>"
                   f"<div style='font-family:{MONO};font-size:{vsize};font-weight:700;color:{col};"
                   f"margin-top:6px;line-height:1'>{_html.escape(str(v))}</div></td>")
    metrics = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
               f"style='border:1px solid {LINE};border-radius:12px;table-layout:fixed'>"
               f"<tr>{mcells}</tr></table>")

    # ---- cta ----
    cta_row = ""
    if plan_url:
        safe_url = _html.escape(str(plan_url), quote=True)
        cta_row = (f"<tr><td style='padding:20px 32px 0'>"
                   f"<a href='{safe_url}' style='display:inline-block;background:{VIOLET};color:#fff;"
                   f"text-decoration:none;font-size:13px;font-weight:700;padding:12px 22px;"
                   f"border-radius:11px'>Open test plan in Azure DevOps &rarr;</a></td></tr>")

    # ---- section heading helper ----
    def _sec_head(dot, title, count, desc=""):
        d = (f"<div style='font-size:12.5px;color:{INK2};font-weight:600;margin-top:7px;"
             f"line-height:1.5'>{desc}</div>") if desc else ""
        return (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
                f"<td valign='middle' style='padding-right:10px'>"
                f"<span style='display:inline-block;width:9px;height:9px;border-radius:50%;"
                f"background:{dot}'></span></td>"
                f"<td valign='middle' style='font-size:14.5px;font-weight:700;color:{INK};"
                f"letter-spacing:-.2px'>{title}</td>"
                f"<td valign='middle' style='padding-left:9px'><span style='font-family:{MONO};"
                f"font-size:11px;font-weight:700;color:{INK2};background:{LINE2};border-radius:20px;"
                f"padding:3px 9px'>{count}</span></td>"
                f"</tr></table>{d}")

    # ---- case card (review / skipped) ----
    def _case_card(a, rail, tag_fg, tag_bg, label, bg=CARD):
        title = _html.escape(str(a.get("title", "")))
        reason = _html.escape(str(a.get("reason", "")))
        item_id = _html.escape(str(a.get("id", "")))
        rtl = "direction:rtl;text-align:right;" if _is_ar(a.get("title", "")) else ""
        url = _wi_url(a.get("id"))
        open_link = (f"<a href='{_html.escape(url, quote=True)}' style='color:{VIOLET_INK};"
                     f"font-size:11.5px;font-weight:700;text-decoration:none'>Open &rarr;</a>") if url else ""
        return (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                f"style='margin-top:11px;background:{bg};border:1px solid {LINE};"
                f"border-left:3px solid {rail};border-radius:11px'><tr>"
                f"<td style='padding:13px 16px'>"
                f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
                f"<td valign='middle'><span style='display:inline-block;background:{tag_bg};color:{tag_fg};"
                f"font-size:10px;font-weight:700;letter-spacing:.5px;text-transform:uppercase;"
                f"padding:3px 9px;border-radius:6px'>{label}</span> "
                f"<span style='font-family:{MONO};font-size:11.5px;font-weight:700;color:{INK3}'>#{item_id}</span></td>"
                f"<td valign='middle' align='right'>{open_link}</td>"
                f"</tr></table>"
                f"<div style='font-family:{AR};font-size:14px;font-weight:600;color:{INK};"
                f"margin-top:11px;line-height:1.5;{rtl}'>{title}</div>"
                + (f"<div style='font-family:{AR};font-size:12.5px;color:{INK2};margin-top:6px;"
                   f"line-height:1.7;{rtl}'>{reason}</div>" if reason else "")
                + f"</td></tr></table>")

    sections = ""

    # per-story (top of the report)
    if per_story:
        rows = ""
        _ps_shown = per_story[:80]
        for sp in _ps_shown:
            sid = _html.escape(str(sp.get("id", "")))
            title = _html.escape(str(sp.get("title", "")))
            total = int(sp.get("total", 0) or 0)
            ok = int(sp.get("ok", 0) or 0); sk = int(sp.get("skipped", 0) or 0); er = int(sp.get("err", 0) or 0)
            secs = sp.get("secs", None)
            rtl = "direction:rtl;text-align:right;" if _is_ar(sp.get("title", "")) else ""
            chips = ""
            if ok: chips += (f"<span style='display:inline-block;background:{GREEN_SOFT};color:{GREEN};"
                             f"font-family:{MONO};font-size:11px;font-weight:700;padding:2px 9px;"
                             f"border-radius:7px;margin-left:5px'>&#10003; {ok}</span>")
            if sk: chips += (f"<span style='display:inline-block;background:{LINE2};color:{INK2};"
                             f"font-family:{MONO};font-size:11px;font-weight:700;padding:2px 9px;"
                             f"border-radius:7px;margin-left:5px'>{sk} skip</span>")
            if er: chips += (f"<span style='display:inline-block;background:{RED_SOFT};color:{RED};"
                             f"font-family:{MONO};font-size:11px;font-weight:700;padding:2px 9px;"
                             f"border-radius:7px;margin-left:5px'>&#10005; {er}</span>")
            tsub = f" &middot; {_fmt_mmss(secs)}" if secs not in (None, "", 0) else ""
            su = _wi_url(sp.get("id"))
            tlink = (f"<a href='{_html.escape(su, quote=True)}' style='color:{INK};"
                     f"text-decoration:none'>{title}</a>" if su else title)
            rows += (f"<tr><td style='padding:13px 0;border-top:1px solid {LINE2};vertical-align:top'>"
                     f"<div style='font-size:14px;font-weight:700;color:{INK};{rtl}'>{tlink}</div>"
                     f"<div style='font-family:{MONO};font-size:11px;font-weight:600;color:{INK3};"
                     f"margin-top:3px'>#{sid} &middot; {total} test case" + ("s" if total != 1 else "") + tsub + "</div></td>"
                     f"<td align='right' style='padding:13px 0;border-top:1px solid {LINE2};"
                     f"vertical-align:top;white-space:nowrap'>{chips}</td></tr>")
        _ps_more = (f"<div style='font-size:12px;color:{INK3};margin-top:10px'>&hellip; and "
                    f"{len(per_story)-80} more &middot; see the attached export / open the full "
                    f"report in QA Studio</div>") if len(per_story) > 80 else ""
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head(VIOLET, 'Per-story breakdown', len(per_story))}"
                     f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                     f"style='margin-top:6px'>{rows}</table>{_ps_more}</td></tr>")

    # needs review
    if action_items:
        _ai_shown = action_items[:40]
        cards = "".join(_case_card(a, AMBER, AMBER, AMBER_SOFT, "Review") for a in _ai_shown)
        _ai_more = (f"<div style='font-size:12px;color:{INK3};margin-top:10px'>&hellip; and "
                    f"{len(action_items)-40} more &middot; see the attached export / open the "
                    f"full report in QA Studio</div>") if len(action_items) > 40 else ""
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head(AMBER, 'Needs your review', len(action_items), 'Steps that no longer match the story&rsquo;s acceptance criteria were regenerated &mdash; confirm them before the next run.')}"
                     f"{cards}{_ai_more}</td></tr>")

    # skipped
    if skipped_items:
        shown = skipped_items[:40]
        cards = "".join(_case_card(a, "#CBC9D4", INK2, LINE2, "Skipped", bg=TINT) for a in shown)
        more = (f"<div style='font-size:12px;color:{INK3};margin-top:10px'>&hellip; and {len(skipped_items)-40} more</div>") if len(skipped_items) > 40 else ""
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head('#CBC9D4', 'Skipped', len(skipped_items), 'Existing steps were judged adequate and left untouched. No action needed.')}"
                     f"{cards}{more}</td></tr>")

    # activity log (full, scrollable where supported)
    if log_lines:
        tone_color = {"ok": GREEN, "err": RED, "warn": AMBER, "skip": INK3,
                      "review": AMBER, "story": VIOLET, "dim": INK3, "info": VIOLET_INK}
        default_ico = {"ok": "&#10003;", "err": "&#10005;", "warn": "&#9888;", "skip": "&#9197;",
                       "review": "&#9888;", "story": "&#9656;", "dim": "&middot;", "info": "&bull;"}
        rows = ""
        shown = log_lines[:120]
        for ln in shown:
            tone = ln.get("tone", "dim")
            col = tone_color.get(tone, INK)
            raw_ico = ln.get("ico")
            ico = _html.escape(str(raw_ico)) if raw_ico else default_ico.get(tone, "&middot;")
            msg = _html.escape(str(ln.get("msg", "")))
            item_id = _html.escape(str(ln.get("id", "")))
            detail = _html.escape(str(ln.get("detail", "")))
            # Align each line by the TITLE's own script, not the legacy "ar"
            # flag (hardcoded True in the run engine from the Arabic-only era),
            # so French/German/etc. titles align LTR in the email.
            is_ar = _is_ar(ln.get("msg", ""))
            is_story = (tone == "story")
            indent = "padding-left:30px;" if ln.get("indent") else ""
            u = _wi_url(ln.get("id"))
            tdir = "rtl" if is_ar else "ltr"; talign = "right" if is_ar else "left"
            fam = AR if is_ar else UI
            id_chip = (f"<span style='font-family:{MONO};font-size:10.5px;font-weight:700;"
                       f"color:{INK3};background:#EEEDF3;border-radius:5px;padding:1px 6px'>{item_id}</span> ") if item_id else ""
            title_color = col if is_story else INK
            mlink = (f"<a href='{_html.escape(u, quote=True)}' style='color:{title_color};"
                     f"text-decoration:none'>{msg}</a>" if u else msg)
            detail_html = (f"<div style='font-family:{MONO};font-size:10.5px;font-weight:600;"
                           f"color:{INK3};margin-top:3px'>{detail}</div>") if detail else ""
            bg = "#F4F3F8" if tone in ("info", "story") and not is_ar else CARD
            rows += (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                     f"style='border-top:1px solid {LINE2};{indent}background:{bg}'><tr>"
                     f"<td width='18' valign='top' style='padding:8px 0 8px 0;color:{col};"
                     f"font-family:{MONO};font-size:13px;font-weight:700;text-align:center'>{ico}</td>"
                     f"<td valign='top' style='padding:8px 0 8px 9px;direction:{tdir};text-align:{talign}'>"
                     f"<div style='font-family:{fam};font-size:12.5px;font-weight:600;color:{title_color};"
                     f"line-height:1.5'>{id_chip}{mlink}</div>{detail_html}</td></tr></table>")
        more = (f"<div style='padding:11px 15px;border-top:1px solid {LINE};background:#F4F3F8;"
                f"text-align:center;font-size:11.5px;font-weight:600;color:{INK2}'>"
                f"&hellip; and {len(log_lines)-120} more lines &middot; open the full trace in QA Studio</div>") if len(log_lines) > 120 else ""
        if is_steps:
            legend = (f"<span style='color:{GREEN};font-weight:700'>&#9632;</span> updated &nbsp; "
                      f"<span style='color:{AMBER};font-weight:700'>&#9632;</span> review &nbsp; "
                      f"<span style='color:{INK3};font-weight:700'>&#9632;</span> kept")
        else:
            legend = (f"<span style='color:{GREEN};font-weight:700'>&#9632;</span> created &nbsp; "
                      f"<span style='color:{AMBER};font-weight:700'>&#9632;</span> removed &nbsp; "
                      f"<span style='color:{INK3};font-weight:700'>&#9632;</span> skipped")
        toolbar = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                   f"style='background:#F4F3F8;border-bottom:1px solid {LINE}'><tr>"
                   f"<td style='padding:9px 15px;font-family:{MONO};font-size:10.5px;font-weight:700;color:{INK2}'>"
                   f"{len(log_lines)} lines &middot; full trace</td>"
                   f"<td align='right' style='padding:9px 15px;font-family:{UI};font-size:10px;"
                   f"font-weight:700;color:{INK3}'>{legend}</td></tr></table>")
        # No inner-scroll container: mobile/desktop Outlook can't scroll a nested
        # div, which would clip the log. Render the lines expanded (capped above)
        # so every client shows the full trace.
        log_box = (f"<div style='border:1px solid {LINE};border-radius:12px;overflow:hidden;background:{TINT}'>"
                   f"{toolbar}"
                   f"{rows}"
                   f"{more}</div>")
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head(INK, 'Run activity log', str(len(log_lines)) + ' lines')}"
                     f"<div style='margin-top:14px'>{log_box}</div></td></tr>")

    footer = (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
              f"<td valign='middle' style='padding-right:9px'>{_logo_tag(24)}</td>"
              f"<td valign='middle' style='font-size:11.5px;font-weight:600;color:{INK3}'>"
              f"Generated by QA Studio &middot; Azure DevOps + AI</td></tr></table>"
              + (f"<div style='font-family:{MONO};font-size:11px;color:{INK3};margin-top:8px;line-height:1.6'>"
                 f"Org: {_html.escape(str(org))} &middot; Project: {_html.escape(str(project))}</div>" if (org and project) else ""))

    return f"""<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>
</head>
<body style='margin:0;padding:0;background:{PAPER};-webkit-text-size-adjust:100%'>
<center style='width:100%;background:{PAPER}'>
<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='background:{PAPER}'><tr>
<td align='center' style='padding:26px 12px 48px'>
<table role='presentation' width='640' cellpadding='0' cellspacing='0' style='width:640px;max-width:640px;background:{CARD};border:1px solid #DEDDE6;border-radius:16px;overflow:hidden;font-family:{UI};color:{INK}'>
  <tr><td style='height:3px;line-height:3px;font-size:0;background:{accent}'>&nbsp;</td></tr>
  <tr><td style='padding:24px 32px 0'>{masthead}</td></tr>
  <tr><td style='padding:18px 32px 4px'>{hero}</td></tr>
  <tr><td style='padding:18px 32px 0'>{metrics}</td></tr>
  {cta_row}
  {sections}
  <tr><td style='padding:20px 32px 26px;border-top:1px solid {LINE};background:{TINT}'>{footer}</td></tr>
</table>
</td></tr></table></center></body></html>"""

def build_automation_report_email(target, summary, stats, project_dir=None, git_url=None,
                                  git_branch=None, log_lines=None, org=None, project=None,
                                  failed=False, stopped=False):
    """Restrained, email-safe HTML summary of an Automation-screen run — same
    masthead/hero/metric-strip/log/footer chrome as build_report_email, so all
    three report emails (run report, sprint summary, automation) read as one
    consistent brand rather than three different designs."""
    import datetime as _dt

    PAPER="#E9E8EE"; CARD="#FFFFFF"; TINT="#FAFAFC"
    INK="#1B1A22"; INK2="#6B6975"; INK3="#9C9AA6"
    LINE="#E8E7EE"; LINE2="#F1F0F5"
    VIOLET="#0E9CC0"; VIOLET_INK="#0B6E86"; VIOLET_SOFT="#D6F4FB"
    GREEN="#1F8A52"; GREEN_SOFT="#E7F4ED"
    RED="#D6414A"; RED_SOFT="#FBEAEC"
    AMBER="#AB780C"; AMBER_SOFT="#F7EFD8"
    UI='"Segoe UI",Roboto,Helvetica,Arial,sans-serif'
    MONO='"SFMono-Regular",Consolas,Menlo,monospace'
    AR='"Segoe UI","Tahoma",Arial,sans-serif'

    def _is_ar(s):
        return any('؀' <= c <= 'ۿ' for c in str(s))

    stats = stats or {}
    tgt_label = {"selenium": "Selenium (Java · TestNG)", "playwright": "Playwright",
                 "cypress": "Cypress"}.get(str(target).lower(), str(target or "").title())

    if failed:
        pill_fg, pill_bg, pill_txt = RED, RED_SOFT, "Failed"; accent = RED
    elif stopped:
        pill_fg, pill_bg, pill_txt = AMBER, AMBER_SOFT, "Stopped early"; accent = AMBER
    else:
        pill_fg, pill_bg, pill_txt = GREEN, GREEN_SOFT, "Completed"; accent = VIOLET
    check_ic = "&#10003;" if pill_fg != AMBER else "&#9632;"
    status_pill = (f"<span style='display:inline-block;background:{pill_bg};color:{pill_fg};"
                   f"font-size:11px;font-weight:700;letter-spacing:.4px;padding:5px 12px;"
                   f"border-radius:20px'>{check_ic}&nbsp; {pill_txt.upper()}</span>")

    # Same numeric-headline convention as build_report_email ("<b>N test cases</b>
    # updated") rather than restating the target framework — the framework is
    # already in the masthead/metric strip, so the headline stays a result count.
    try:
        _n = int(stats.get("Stories", 0) or 0)
    except Exception:
        _n = 0
    if _n:
        headline = f"<b style='color:{VIOLET_INK}'>{_n} stor" + ("y" if _n == 1 else "ies") + "</b> automated"
    else:
        headline = "No stories automated"
    hero = (f"{status_pill}"
            f"<div style='font-size:25px;font-weight:700;letter-spacing:-.5px;color:{INK};"
            f"line-height:1.15;margin:14px 0 0'>{headline}</div>"
            f"<div style='font-size:13px;color:{INK2};font-weight:600;margin-top:8px'>"
            f"{_html.escape(str(summary or ''))}</div>")

    today = _dt.date.today().strftime("%d %b %Y")
    masthead = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        f"<td width='40' valign='middle' style='padding-right:13px'>{_logo_tag(40)}</td>"
        f"<td valign='middle'>"
        f"<div style='font-size:15px;font-weight:700;color:{INK};letter-spacing:-.2px'>QA Studio</div>"
        f"<div style='font-size:12px;font-weight:700;color:{VIOLET_INK};margin-top:2px'>"
        f"{tgt_label} &middot; Automation report</div></td>"
        f"<td valign='middle' align='right' style='font-family:{MONO};font-size:11px;"
        f"color:{INK3};font-weight:700'>{today}</td>"
        f"</tr></table>")

    def _mcolor(k, v):
        try: iv = int(str(v).split('/')[0].strip())
        except Exception: iv = 0
        if k in ("Stories", "Test cases") and iv > 0: return GREEN
        if k in ("Self-healed", "Skipped") and iv > 0: return AMBER
        if k == "Time": return INK
        return INK3
    mcells = ""
    items = list(stats.items())
    for i, (k, v) in enumerate(items):
        col = _mcolor(k, v)
        bl = "" if i == 0 else f"border-left:1px solid {LINE2};"
        vsize = "18px" if k == "Time" else "24px"
        mcells += (f"<td width='1' style='{bl}padding:13px 6px 14px;text-align:center;vertical-align:top'>"
                   f"<div style='font-size:9.5px;font-weight:700;letter-spacing:1px;color:{INK3};"
                   f"text-transform:uppercase'>{_html.escape(str(k))}</div>"
                   f"<div style='font-family:{MONO};font-size:{vsize};font-weight:700;color:{col};"
                   f"margin-top:6px;line-height:1'>{_html.escape(str(v))}</div></td>")
    metrics = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
               f"style='border:1px solid {LINE};border-radius:12px;table-layout:fixed'>"
               f"<tr>{mcells}</tr></table>")

    cta_row = ""
    if git_url:
        safe_url = _html.escape(str(git_url), quote=True)
        cta_row = (f"<tr><td style='padding:20px 32px 0'>"
                   f"<a href='{safe_url}' style='display:inline-block;background:{VIOLET};color:#fff;"
                   f"text-decoration:none;font-size:13px;font-weight:700;padding:12px 22px;"
                   f"border-radius:11px'>Open Git repository &rarr;</a></td></tr>")

    def _sec_head(dot, title, count):
        return (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
                f"<td valign='middle' style='padding-right:10px'>"
                f"<span style='display:inline-block;width:9px;height:9px;border-radius:50%;"
                f"background:{dot}'></span></td>"
                f"<td valign='middle' style='font-size:14.5px;font-weight:700;color:{INK};"
                f"letter-spacing:-.2px'>{title}</td>"
                f"<td valign='middle' style='padding-left:9px'><span style='font-family:{MONO};"
                f"font-size:11px;font-weight:700;color:{INK2};background:{LINE2};border-radius:20px;"
                f"padding:3px 9px'>{count}</span></td>"
                f"</tr></table>")

    sections = ""

    # destination (project folder + git branch/target) — quick facts, no card noise
    facts = []
    if project_dir:
        facts.append(("Output folder", str(project_dir)))
    if git_url:
        facts.append(("Repository", str(git_url) + (f" @ {git_branch}" if git_branch else "")))
    if facts:
        rows = "".join(
            f"<tr><td style='padding:9px 0;border-top:1px solid {LINE2};font-size:11.5px;"
            f"font-weight:700;color:{INK3};width:130px;vertical-align:top'>{_html.escape(k)}</td>"
            f"<td style='padding:9px 0;border-top:1px solid {LINE2};font-family:{MONO};"
            f"font-size:12px;font-weight:600;color:{INK};word-break:break-all'>{_html.escape(v)}</td></tr>"
            for k, v in facts)
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head(VIOLET, 'Destination', len(facts))}"
                     f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                     f"style='margin-top:6px'>{rows}</table></td></tr>")

    # activity log (msg/tone only — the Automation log has no id/detail/ico fields)
    if log_lines:
        tone_color = {"ok": GREEN, "err": RED, "warn": AMBER, "story": VIOLET,
                      "dim": INK3, "info": VIOLET_INK}
        default_ico = {"ok": "&#10003;", "err": "&#10005;", "warn": "&#9888;",
                       "story": "&#9656;", "dim": "&middot;", "info": "&bull;"}
        rows = ""
        shown = log_lines[:150]
        for ln in shown:
            tone = ln.get("tone", "dim")
            col = tone_color.get(tone, INK)
            raw_ico = ln.get("ico")
            ico = _html.escape(str(raw_ico)) if raw_ico else default_ico.get(tone, "&middot;")
            raw_msg = ln.get("msg", "")
            msg = _html.escape(str(raw_msg))
            # Arabic story/test-case titles must render RTL + right-aligned, same
            # as the Run report's log — otherwise they render left-to-right and
            # left-aligned like English lines, which reads wrong for Arabic text.
            is_ar = _is_ar(raw_msg)
            tdir = "rtl" if is_ar else "ltr"; talign = "right" if is_ar else "left"
            fam = AR if is_ar else UI
            bg = "#F4F3F8" if tone in ("info", "story") and not is_ar else CARD
            rows += (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                     f"style='border-top:1px solid {LINE2};background:{bg}'><tr>"
                     f"<td width='18' valign='top' style='padding:8px 0 8px 0;color:{col};"
                     f"font-family:{MONO};font-size:13px;font-weight:700;text-align:center'>{ico}</td>"
                     f"<td valign='top' style='padding:8px 0 8px 9px;direction:{tdir};text-align:{talign}'>"
                     f"<div style='font-family:{fam};font-size:12.5px;font-weight:600;color:{INK};"
                     f"line-height:1.5'>{msg}</div></td></tr></table>")
        more = (f"<div style='padding:11px 15px;border-top:1px solid {LINE};background:#F4F3F8;"
                f"text-align:center;font-size:11.5px;font-weight:600;color:{INK2}'>"
                f"&hellip; and {len(log_lines)-150} more lines &middot; open the full trace in "
                f"QA Studio</div>") if len(log_lines) > 150 else ""
        toolbar = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                   f"style='background:#F4F3F8;border-bottom:1px solid {LINE}'><tr>"
                   f"<td style='padding:9px 15px;font-family:{MONO};font-size:10.5px;font-weight:700;"
                   f"color:{INK2}'>{len(log_lines)} lines &middot; full trace</td></tr></table>")
        log_box = (f"<div style='border:1px solid {LINE};border-radius:12px;overflow:hidden;"
                   f"background:{TINT}'>{toolbar}{rows}{more}</div>")
        sections += (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                     f"{_sec_head(INK, 'Automation activity log', str(len(log_lines)) + ' lines')}"
                     f"<div style='margin-top:14px'>{log_box}</div></td></tr>")

    footer = (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
              f"<td valign='middle' style='padding-right:9px'>{_logo_tag(24)}</td>"
              f"<td valign='middle' style='font-size:11.5px;font-weight:600;color:{INK3}'>"
              f"Generated by QA Studio &middot; Azure DevOps + AI</td></tr></table>"
              + (f"<div style='font-family:{MONO};font-size:11px;color:{INK3};margin-top:8px;line-height:1.6'>"
                 f"Org: {_html.escape(str(org))} &middot; Project: {_html.escape(str(project))}</div>" if (org and project) else ""))

    return f"""<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>
</head>
<body style='margin:0;padding:0;background:{PAPER};-webkit-text-size-adjust:100%'>
<center style='width:100%;background:{PAPER}'>
<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='background:{PAPER}'><tr>
<td align='center' style='padding:26px 12px 48px'>
<table role='presentation' width='640' cellpadding='0' cellspacing='0' style='width:640px;max-width:640px;background:{CARD};border:1px solid #DEDDE6;border-radius:16px;overflow:hidden;font-family:{UI};color:{INK}'>
  <tr><td style='height:3px;line-height:3px;font-size:0;background:{accent}'>&nbsp;</td></tr>
  <tr><td style='padding:24px 32px 0'>{masthead}</td></tr>
  <tr><td style='padding:18px 32px 4px'>{hero}</td></tr>
  <tr><td style='padding:18px 32px 0'>{metrics}</td></tr>
  {cta_row}
  {sections}
  <tr><td style='padding:20px 32px 26px;border-top:1px solid {LINE};background:{TINT}'>{footer}</td></tr>
</table>
</td></tr></table></center></body></html>"""

def build_sprint_summary_email(data):
    """Restrained, email-safe (table + inline-style) Sprint Summary email.
    Matches the run-report design; logo is embedded inline via Content-ID."""
    import datetime as _dt

    # Same VIOLET* fix as build_report_email above — matches the app's actual
    # cyan/teal accent (theme.py) instead of an unrelated hardcoded royal-blue.
    PAPER="#E9E8EE"; CARD="#FFFFFF"; TINT="#FAFAFC"
    INK="#1B1A22"; INK2="#6B6975"; INK3="#9C9AA6"
    LINE="#E8E7EE"; LINE2="#F1F0F5"
    VIOLET="#0E9CC0"; VIOLET_INK="#0B6E86"; VIOLET_SOFT="#D6F4FB"
    GREEN="#1F8A52"; GREEN_SOFT="#E7F4ED"
    RED="#D6414A"; AMBER="#AB780C"; AMBER_SOFT="#F7EFD8"
    UI='"Segoe UI",Roboto,Helvetica,Arial,sans-serif'
    MONO='"SFMono-Regular",Consolas,Menlo,monospace'
    AR='"Segoe UI","Tahoma",Arial,sans-serif'

    plan_name = _html.escape(str(data.get("plan_name", "")))
    iteration = _html.escape(str(data.get("iteration", "") or "—"))
    total_stories = data.get("total_stories", 0)
    total_tc = data.get("total_test_cases", 0)
    by_state = data.get("by_state", {})
    stories = data.get("stories", [])
    _proj = data.get("project", "")
    _org = data.get("org", AZURE_ORG)

    def _is_ar(s):
        return any('\u0600' <= c <= '\u06ff' for c in str(s))

    def _state_colors(state):
        s = (state or "").lower()
        if s in ("done", "closed", "completed", "resolved"): return (GREEN, GREEN_SOFT)
        if s in ("active", "in progress", "committed", "doing"): return (VIOLET_INK, VIOLET_SOFT)
        if s in ("new", "to do", "proposed", "open"): return (AMBER, AMBER_SOFT)
        return (INK2, LINE2)

    # masthead
    today = _dt.date.today().strftime("%d %b %Y")
    masthead = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        f"<td width='40' valign='middle' style='padding-right:13px'>{_logo_tag(40)}</td>"
        f"<td valign='middle'>"
        f"<div style='font-size:15px;font-weight:700;color:{INK};letter-spacing:-.2px'>QA Studio</div>"
        f"<div style='font-size:12px;font-weight:700;color:{VIOLET_INK};margin-top:2px'>Sprint Summary &middot; Report</div>"
        f"</td>"
        f"<td valign='middle' align='right' style='font-family:{MONO};font-size:11px;color:{INK3};font-weight:700'>{today}</td>"
        f"</tr></table>")

    # hero
    hero = (f"<span style='display:inline-block;background:{VIOLET_SOFT};color:{VIOLET_INK};"
            f"font-size:11px;font-weight:700;letter-spacing:.4px;padding:5px 12px;border-radius:20px'>SPRINT SNAPSHOT</span>"
            f"<div style='font-size:23px;font-weight:700;letter-spacing:-.5px;color:{INK};line-height:1.2;margin:14px 0 0'>{plan_name}</div>"
            f"<div style='font-family:{MONO};font-size:12px;color:{INK2};font-weight:600;margin-top:6px'>{iteration}</div>")

    # metric strip
    metrics_data = [("Stories", total_stories, VIOLET_INK), ("Test Cases", total_tc, GREEN), ("Statuses", len(by_state), INK)]
    mcells = ""
    for i,(k,v,col) in enumerate(metrics_data):
        bl = "" if i == 0 else f"border-left:1px solid {LINE2};"
        mcells += (f"<td width='1' style='{bl}padding:14px 8px 15px;text-align:center'>"
                   f"<div style='font-size:9.5px;font-weight:800;letter-spacing:1px;color:{INK3};text-transform:uppercase'>{_html.escape(str(k))}</div>"
                   f"<div style='font-family:{MONO};font-size:25px;font-weight:700;color:{col};margin-top:6px;line-height:1'>{v}</div></td>")
    metrics = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
               f"style='border:1px solid {LINE};border-radius:12px;table-layout:fixed'><tr>{mcells}</tr></table>")

    def _sec_head(dot, title, count):
        return (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
                f"<td valign='middle' style='padding-right:10px'><span style='display:inline-block;width:9px;height:9px;border-radius:50%;background:{dot}'></span></td>"
                f"<td valign='middle' style='font-size:14.5px;font-weight:800;color:{INK};letter-spacing:-.2px'>{title}</td>"
                f"<td valign='middle' style='padding-left:9px'><span style='font-family:{MONO};font-size:11px;font-weight:800;color:{INK2};background:{LINE2};border-radius:20px;padding:3px 9px'>{count}</span></td>"
                f"</tr></table>")

    # status breakdown — wrapping chips
    chips = ""
    for st, cnt in sorted(by_state.items(), key=lambda x: -x[1]):
        fg, bg = _state_colors(st)
        chips += (f"<div style='display:inline-block;vertical-align:top;background:{bg};border-radius:11px;"
                  f"padding:13px 10px;text-align:center;min-width:92px;margin:0 8px 8px 0;box-sizing:border-box'>"
                  f"<div style='font-family:{MONO};font-size:22px;font-weight:700;color:{fg}'>{cnt}</div>"
                  f"<div style='font-size:11px;color:{INK2};font-weight:700;margin-top:3px;line-height:1.3'>{_html.escape(str(st))}</div></div>")
    status_block = (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                    f"{_sec_head(VIOLET, 'Status breakdown', len(by_state))}"
                    f"<div style='font-size:0;margin-top:14px'>{chips}</div></td></tr>") if by_state else ""

    # stories — capped like every other email report's row list (see
    # build_report_email's per_story/skipped_items, build_ai_usage_email's
    # rows): a big sprint can run into hundreds of stories, and listing every
    # one makes the email huge (some inboxes clip or flag oversized mail) and
    # unreadable on a phone. The in-app Sprint Report/Plan screens already
    # show the full list; the email is a summary, not the source of truth.
    _stories_shown = stories[:150]
    rows = ""
    for s in _stories_shown:
        title = _html.escape(str(s.get("title", "")))
        sid = _html.escape(str(s.get("id", "")))
        state = str(s.get("state", ""))
        tc = int(s.get("test_cases", 0) or 0)
        # Same "QA Assigned" field the in-app Sprint Summary modal shows
        # (modals.py's assigned_cell) — was missing here, so the email never
        # reflected who each story is assigned to.
        assigned_name = _html.escape(str(s.get("assigned_to") or "Unassigned"))
        assigned_col = (INK3 if assigned_name == "Unassigned" else INK2)
        fg, bg = _state_colors(state)
        rtl = "direction:rtl;text-align:right;" if _is_ar(title) else ""
        wi = (f"https://dev.azure.com/{_org}/{_proj}/_workitems/edit/{s.get('id','')}" if _proj and s.get("id") else "")
        tlink = (f"<a href='{_html.escape(wi, quote=True)}' style='color:{INK};text-decoration:none'>{title}</a>" if wi else title)
        idlink = (f"<a href='{_html.escape(wi, quote=True)}' style='color:{VIOLET_INK};text-decoration:none'>#{sid} &rarr;</a>" if wi else f"#{sid}")
        rows += (f"<tr><td style='padding:13px 0;border-top:1px solid {LINE2};vertical-align:middle'>"
                 f"<div style='font-size:13.5px;font-weight:700;color:{INK};{rtl}'>{tlink}</div>"
                 f"<div style='font-family:{MONO};font-size:11px;font-weight:600;color:{INK3};margin-top:3px'>{idlink}</div></td>"
                 f"<td align='right' width='100' style='padding:13px 0;border-top:1px solid {LINE2};vertical-align:middle;white-space:nowrap'>"
                 f"<span style='font-size:11.5px;font-weight:600;color:{assigned_col}'>{assigned_name}</span></td>"
                 f"<td align='right' width='55' style='padding:13px 0;border-top:1px solid {LINE2};vertical-align:middle;white-space:nowrap'>"
                 f"<span style='font-family:{MONO};font-size:11px;font-weight:700;color:{INK2}'>{tc} TC</span></td>"
                 f"<td align='right' width='110' style='padding:13px 0;border-top:1px solid {LINE2};vertical-align:middle;white-space:nowrap'>"
                 f"<span style='display:inline-block;background:{bg};color:{fg};font-size:11px;font-weight:800;padding:3px 10px;border-radius:20px'>{_html.escape(state)}</span></td>"
                 f"</tr>")
    _stories_more = (f"<div style='font-size:12px;color:{INK3};margin-top:10px'>&hellip; and "
                     f"{len(stories)-150} more &middot; see the full list in the Sprint Report/"
                     f"Plan screen</div>") if len(stories) > 150 else ""
    # Full 4-column header (STORY / ASSIGNED / TC / STATUS), matching the
    # in-app Sprint Summary modal's table_header (modals.py) — previously
    # only the "Assigned" cell had a label, leaving the rest of the row
    # looking headerless. TC and status also split back into their own
    # columns (were combined in one cell) so each gets its own label.
    # Column widths (real HTML width= attributes, not just CSS) MUST match
    # the data rows' <td width=...> above exactly — Outlook's Word rendering
    # engine ignores CSS max-width entirely and won't infer a column's width
    # from an empty sibling cell, so mismatched/missing widths let whichever
    # cell has content balloon and swallow the row instead of lining up.
    def _hdr(t):
        return (f"<span style='font-size:9.5px;font-weight:800;letter-spacing:.6px;"
                f"color:{INK3};text-transform:uppercase'>{t}</span>")
    story_header = (f"<tr>"
                     f"<td style='padding-bottom:6px'>{_hdr('Story')}</td>"
                     f"<td align='right' width='100' style='padding-bottom:6px'>{_hdr('Assigned')}</td>"
                     f"<td align='right' width='55' style='padding-bottom:6px'>{_hdr('TC')}</td>"
                     f"<td align='right' width='110' style='padding-bottom:6px'>{_hdr('Status')}</td>"
                     f"</tr>") if stories else ""
    story_block = (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                   f"{_sec_head(INK, 'Stories', len(stories))}"
                   f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='margin-top:10px'>{story_header}{rows}</table>{_stories_more}</td></tr>") if stories else ""

    footer = (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
              f"<td valign='middle' style='padding-right:9px'>{_logo_tag(24)}</td>"
              f"<td valign='middle' style='font-size:11.5px;font-weight:600;color:{INK3}'>Generated by QA Studio &middot; Azure DevOps + AI</td></tr></table>"
              + (f"<div style='font-family:{MONO};font-size:11px;color:{INK3};margin-top:8px'>Org: {_html.escape(str(_org))} &middot; Project: {_html.escape(str(_proj))}</div>" if _proj else ""))

    return f"""<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>
</head>
<body style='margin:0;padding:0;background:{PAPER};-webkit-text-size-adjust:100%'>
<center style='width:100%;background:{PAPER}'>
<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='background:{PAPER}'><tr>
<td align='center' style='padding:26px 12px 48px'>
<table role='presentation' width='640' cellpadding='0' cellspacing='0' style='width:640px;max-width:640px;background:{CARD};border:1px solid #DEDDE6;border-radius:16px;overflow:hidden;font-family:{UI};color:{INK}'>
  <tr><td style='height:3px;line-height:3px;font-size:0;background:{VIOLET}'>&nbsp;</td></tr>
  <tr><td style='padding:24px 32px 0'>{masthead}</td></tr>
  <tr><td style='padding:18px 32px 4px'>{hero}</td></tr>
  <tr><td style='padding:18px 32px 0'>{metrics}</td></tr>
  {status_block}
  {story_block}
  <tr><td style='padding:20px 32px 26px;border-top:1px solid {LINE};background:{TINT}'>{footer}</td></tr>
</table>
</td></tr></table></center></body></html>"""


# ═══════════════════════════════════════════════════════════════════════════════
#  AI USAGE REPORT (per-user, or whole-org for Admins — see
#  auth_supabase.admin_get_ai_usage)
# ═══════════════════════════════════════════════════════════════════════════════

def usage_report_all_users(start_date=None, end_date=None):
    """Build an AI usage report for an optional [start_date, end_date]
    ('YYYY-MM-DD' strings, inclusive; None = no bound). Despite the name
    (kept for backward compatibility — this is the one function every caller
    already uses), the SCOPE is decided server-side by the caller's role, not
    by this function: a non-Admin gets a report of only their OWN usage; an
    Admin gets one across every signed-in user. Returns (ok,
    report_or_message) — ok is False with a friendly message for every
    failure mode (not configured, not signed in, offline, function not
    deployed), never a raise.

    report = {
      "start", "end", "generated",
      "rows": [{"date","user","provider","model","module","calls",
                "input_tokens","output_tokens","cost_usd"}, ...],  # grouped +
                                                                    # summed
      "totals": {"calls","input_tokens","output_tokens","cost_usd"},
      "unpriced_calls": N,   # calls whose model has no rate in PRICING —
                             # excluded from totals["cost_usd"], never guessed
      "truncated": bool,     # the server-side row cap was hit (see ai-usage)
    }
    Token counts are EXACT (each row is built from what the provider itself
    reported for that call — see _norm_usage). Cost is an ESTIMATE from the
    local PRICING table, computed here (not server-side) so a price update
    never needs redeploying the Edge Function.
    """
    import datetime as _dt
    import auth_supabase as auth
    ok, res = auth.admin_get_ai_usage(start_date, end_date)
    if not ok:
        return False, res
    raw_rows = res if isinstance(res, list) else []

    from collections import defaultdict
    # Grouped by tag too (the raw usage_tag a call site passed to ai_complete,
    # e.g. "generate_steps") — NOT the friendly label, so two tags that happen
    # to map to the same label can never collide. Historical rows logged
    # before call sites were tagged have tag=None, which buckets separately
    # as "Other" via _usage_module_label() below — they never silently merge
    # into a real module's numbers.
    buckets = defaultdict(lambda: {"calls": 0, "input_tokens": 0, "output_tokens": 0,
                                   "first_time": None, "last_time": None})
    for r in raw_rows:
        ts = str(r.get("created_at") or "")
        date = ts[:10] or "unknown"
        # "HH:MM" slice of the UTC timestamp (created_at is always UTC — see
        # record_ai_usage) — kept alongside the day-level grouping below rather
        # than used AS the grouping key, so rows stay one-per-day/user/provider/
        # model/module (summed) instead of exploding into one row per exact
        # timestamp. Tracked as a min/max range so a day's row can show WHEN
        # during the day the usage happened, not just which day.
        time_str = ts[11:16] if len(ts) >= 16 else None
        user = r.get("user_email") or "(unknown)"
        provider = r.get("provider") or "unknown"
        model = r.get("model") or "unknown"
        tag = r.get("tag") or None
        b = buckets[(date, user, provider, model, tag)]
        b["calls"] += 1
        b["input_tokens"] += int(r.get("input_tokens") or 0)
        b["output_tokens"] += int(r.get("output_tokens") or 0)
        if time_str:
            if b["first_time"] is None or time_str < b["first_time"]:
                b["first_time"] = time_str
            if b["last_time"] is None or time_str > b["last_time"]:
                b["last_time"] = time_str

    rows = []
    totals = {"calls": 0, "input_tokens": 0, "output_tokens": 0, "cost_usd": 0.0}
    unpriced_calls = 0
    # key=... coerces the tag slot to "" for the sort only (buckets themselves
    # still key on the real None) — plain tuple sort would otherwise raise
    # comparing None against a str the moment two rows share every other
    # field but one has a tag and the other doesn't (an untagged historical
    # row next to a freshly-tagged one on the same date/user/provider/model).
    for (date, user, provider, model, tag), b in sorted(
            buckets.items(), key=lambda kv: kv[0][:4] + (kv[0][4] or "",)):
        cost = _call_cost(provider, model, b["input_tokens"], b["output_tokens"])
        # "date" stays a bare YYYY-MM-DD (used for sorting above and by any
        # machine consumer of the JSON export); "date_range" is the display
        # string every human-facing renderer (screen table, xlsx/docx/pdf,
        # email) should use instead — same day, plus the UTC time span the
        # calls in this row actually happened within.
        if b["first_time"] and b["last_time"]:
            date_range = (f"{date} ({b['first_time']} UTC)" if b["first_time"] == b["last_time"]
                          else f"{date} ({b['first_time']}–{b['last_time']} UTC)")
        else:
            date_range = date
        rows.append({"date": date, "date_range": date_range, "user": user,
                     "provider": provider, "model": model,
                     "module": _usage_module_label(tag),
                     "calls": b["calls"], "input_tokens": b["input_tokens"],
                     "output_tokens": b["output_tokens"], "cost_usd": cost})
        totals["calls"] += b["calls"]
        totals["input_tokens"] += b["input_tokens"]
        totals["output_tokens"] += b["output_tokens"]
        if cost is None:
            unpriced_calls += b["calls"]
        else:
            totals["cost_usd"] += cost
    totals["cost_usd"] = round(totals["cost_usd"], 4)

    report = {"start": start_date, "end": end_date,
             "generated": _dt.datetime.now(_dt.timezone.utc).isoformat(timespec="seconds"),
             "rows": rows, "totals": totals, "unpriced_calls": unpriced_calls,
             "row_count_raw": len(raw_rows)}
    return True, report


def _usage_out_dir():
    import platform_caps as _pc
    d = os.path.join(_pc.export_base_dir(), "QA Studio", "AI Usage Reports")
    os.makedirs(d, exist_ok=True)
    return d


def _usage_stamp(report):
    from datetime import datetime as _dtc
    rng = ""
    if report.get("start") or report.get("end"):
        rng = f"_{report.get('start') or 'start'}_to_{report.get('end') or 'now'}"
    return f"AIUsage{rng}_{_dtc.now():%Y%m%d-%H%M}"


def export_usage_json(report):
    p = os.path.join(_usage_out_dir(), _usage_stamp(report) + ".json")
    with open(p, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    return p


def export_usage_xlsx(report):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook()
    ws = wb.active
    ws.title = "AI Usage"
    head = Font(bold=True, color="FFFFFF", name="Segoe UI")
    fill = PatternFill("solid", fgColor="3A57D6")
    thin = Border(*[Side(style="thin", color="E6E8F1")] * 4)
    cols = ["Date", "User", "Provider", "Model", "Module", "Calls", "Input Tokens",
            "Output Tokens", "Cost (USD)"]
    for c, name in enumerate(cols, 1):
        cell = ws.cell(1, c, name)
        cell.font = head; cell.fill = fill; cell.border = thin
        cell.alignment = Alignment(horizontal="center")
    r = 2
    for row in report["rows"]:
        vals = [row.get("date_range", row["date"]), row["user"], row["provider"], row["model"],
                row.get("module", "Other"), row["calls"],
                row["input_tokens"], row["output_tokens"],
                (round(row["cost_usd"], 4) if row["cost_usd"] is not None else "—")]
        for c, v in enumerate(vals, 1):
            ws.cell(r, c, v).border = thin
        r += 1
    t = report["totals"]
    ws.cell(r, 5, "TOTAL").font = Font(bold=True)
    for c, v in ((6, t["calls"]), (7, t["input_tokens"]), (8, t["output_tokens"]),
                 (9, t["cost_usd"])):
        cell = ws.cell(r, c, v)
        cell.font = Font(bold=True)
    if report.get("unpriced_calls"):
        r += 2
        ws.cell(r, 1, f"{report['unpriced_calls']} call(s) use a model with no "
                       "published price and are excluded from the cost total above.")
    for c, w in zip("ABCDEFGHI", (26, 28, 12, 22, 22, 8, 13, 13, 12)):
        ws.column_dimensions[c].width = w
    p = os.path.join(_usage_out_dir(), _usage_stamp(report) + ".xlsx")
    wb.save(p)
    return p


def export_usage_docx(report):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    try:
        ns = doc.styles["Normal"]
        ns.font.name = "Segoe UI"; ns.font.size = Pt(10)
    except Exception:
        pass
    h = doc.add_heading("AI Usage Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = sub.add_run(f"{report.get('start') or 'earliest'} → {report.get('end') or 'latest'}")
    rn.font.size = Pt(11); rn.font.color.rgb = RGBColor(0x6A, 0x4D, 0xFF)

    heads = ["Date", "User", "Provider", "Model", "Module", "Calls", "Input", "Output", "Cost (USD)"]
    tbl = doc.add_table(rows=1, cols=len(heads))
    try:
        tbl.style = "Medium Shading 1 Accent 1"
    except Exception:
        pass
    for i, hd in enumerate(heads):
        tbl.rows[0].cells[i].text = hd
    for row in report["rows"]:
        c = tbl.add_row().cells
        vals = [row.get("date_range", row["date"]), row["user"], row["provider"], row["model"],
                row.get("module", "Other"), str(row["calls"]),
                str(row["input_tokens"]), str(row["output_tokens"]),
                (f'{row["cost_usd"]:.4f}' if row["cost_usd"] is not None else "—")]
        for i, v in enumerate(vals):
            c[i].text = v
    t = report["totals"]
    c = tbl.add_row().cells
    c[4].text = "TOTAL"
    c[5].text = str(t["calls"]); c[6].text = str(t["input_tokens"])
    c[7].text = str(t["output_tokens"]); c[8].text = f'{t["cost_usd"]:.4f}'
    for cell in c:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    if report.get("unpriced_calls"):
        doc.add_paragraph(f"{report['unpriced_calls']} call(s) use a model with no "
                          "published price and are excluded from the cost total above.")
    p = os.path.join(_usage_out_dir(), _usage_stamp(report) + ".docx")
    doc.save(p)
    return p


def export_usage_pdf(report):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    p = os.path.join(_usage_out_dir(), _usage_stamp(report) + ".pdf")
    doc = SimpleDocTemplate(p, pagesize=landscape(A4), topMargin=14 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    elems = [Paragraph("AI Usage Report", styles["Title"]),
             Paragraph(f"{report.get('start') or 'earliest'} &rarr; {report.get('end') or 'latest'}",
                       styles["Normal"]),
             Spacer(1, 6 * mm)]
    data = [["Date", "User", "Provider", "Model", "Module", "Calls", "Input", "Output", "Cost (USD)"]]
    for row in report["rows"]:
        data.append([row.get("date_range", row["date"]), row["user"], row["provider"], row["model"],
                    row.get("module", "Other"),
                    str(row["calls"]), str(row["input_tokens"]), str(row["output_tokens"]),
                    (f'{row["cost_usd"]:.4f}' if row["cost_usd"] is not None else "—")])
    t = report["totals"]
    data.append(["", "", "", "", "TOTAL", str(t["calls"]), str(t["input_tokens"]),
                str(t["output_tokens"]), f'{t["cost_usd"]:.4f}'])
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3A57D6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E6E8F1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#F7F7FB")]),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#ECEBF5")),
        ("ALIGN", (5, 1), (-1, -1), "RIGHT"),
    ]))
    elems.append(tbl)
    if report.get("unpriced_calls"):
        elems.append(Spacer(1, 4 * mm))
        elems.append(Paragraph(f"{report['unpriced_calls']} call(s) use a model with no "
                               "published price and are excluded from the cost total above.",
                               styles["Normal"]))
    doc.build(elems)
    return p


def build_ai_usage_email(report):
    """Restrained, email-safe AI Usage report — same visual language as
    build_sprint_summary_email (masthead/hero/metric strip/table/footer)."""
    PAPER = "#E9E8EE"; CARD = "#FFFFFF"; TINT = "#FAFAFC"
    INK = "#1B1A22"; INK2 = "#6B6975"; INK3 = "#9C9AA6"
    LINE = "#E8E7EE"; LINE2 = "#F1F0F5"
    VIOLET = "#0E9CC0"; VIOLET_INK = "#0B6E86"; VIOLET_SOFT = "#D6F4FB"
    GREEN = "#1F8A52"; AMBER = "#AB780C"
    UI = '"Segoe UI",Roboto,Helvetica,Arial,sans-serif'
    MONO = '"SFMono-Regular",Consolas,Menlo,monospace'

    import datetime as _dt
    today = _dt.date.today().strftime("%d %b %Y")
    rng = f"{report.get('start') or 'earliest'} &rarr; {report.get('end') or 'latest'}"
    t = report.get("totals", {})
    rows = report.get("rows", [])

    masthead = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        f"<td width='40' valign='middle' style='padding-right:13px'>{_logo_tag(40)}</td>"
        f"<td valign='middle'>"
        f"<div style='font-size:15px;font-weight:700;color:{INK};letter-spacing:-.2px'>QA Studio</div>"
        f"<div style='font-size:12px;font-weight:700;color:{VIOLET_INK};margin-top:2px'>AI Usage &middot; Report</div>"
        f"</td>"
        f"<td valign='middle' align='right' style='font-family:{MONO};font-size:11px;color:{INK3};font-weight:700'>{today}</td>"
        f"</tr></table>")

    hero = (f"<span style='display:inline-block;background:{VIOLET_SOFT};color:{VIOLET_INK};"
            f"font-size:11px;font-weight:700;letter-spacing:.4px;padding:5px 12px;border-radius:20px'>WHOLE-ORG USAGE</span>"
            f"<div style='font-family:{MONO};font-size:12px;color:{INK2};font-weight:600;margin-top:10px'>{rng}</div>")

    metrics_data = [("Calls", t.get("calls", 0), VIOLET_INK),
                    ("Input Tokens", t.get("input_tokens", 0), INK),
                    ("Output Tokens", t.get("output_tokens", 0), INK),
                    ("Est. Cost", f'${t.get("cost_usd", 0):.2f}', GREEN)]
    mcells = ""
    for i, (k, v, col) in enumerate(metrics_data):
        bl = "" if i == 0 else f"border-left:1px solid {LINE2};"
        mcells += (f"<td width='1' style='{bl}padding:14px 8px 15px;text-align:center'>"
                  f"<div style='font-size:9.5px;font-weight:800;letter-spacing:1px;color:{INK3};text-transform:uppercase'>{_html.escape(str(k))}</div>"
                  f"<div style='font-family:{MONO};font-size:22px;font-weight:700;color:{col};margin-top:6px;line-height:1'>{v}</div></td>")
    metrics = (f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
              f"style='border:1px solid {LINE};border-radius:12px;table-layout:fixed'><tr>{mcells}</tr></table>")

    # flat rows table, one line per date/user/provider/model — capped so the
    # email stays a reasonable size; the full data is in the attached/exported file.
    shown = rows[:200]
    trows = ""
    for row in shown:
        cost = f'${row["cost_usd"]:.4f}' if row["cost_usd"] is not None else "&mdash;"
        trows += (f"<tr><td style='padding:8px 0;border-top:1px solid {LINE2};font-family:{MONO};"
                  f"font-size:11px;color:{INK2}'>{_html.escape(row.get('date_range', row['date']))}</td>"
                  f"<td style='padding:8px 10px;border-top:1px solid {LINE2};font-size:11.5px;font-weight:600;color:{INK}'>{_html.escape(row['user'])}</td>"
                  f"<td style='padding:8px 10px;border-top:1px solid {LINE2};font-size:11px;color:{INK2}'>{_html.escape(row['provider'])}</td>"
                  f"<td style='padding:8px 10px;border-top:1px solid {LINE2};font-family:{MONO};font-size:10.5px;color:{INK2}'>{_html.escape(row['model'])}</td>"
                  f"<td style='padding:8px 10px;border-top:1px solid {LINE2};font-size:11px;color:{INK2}'>{_html.escape(row.get('module') or 'Other')}</td>"
                  f"<td align='right' style='padding:8px 10px;border-top:1px solid {LINE2};font-family:{MONO};font-size:11px;color:{INK}'>{row['calls']}</td>"
                  f"<td align='right' style='padding:8px 10px;border-top:1px solid {LINE2};font-family:{MONO};font-size:11px;color:{INK}'>{row['input_tokens']}</td>"
                  f"<td align='right' style='padding:8px 10px;border-top:1px solid {LINE2};font-family:{MONO};font-size:11px;color:{INK}'>{row['output_tokens']}</td>"
                  f"<td align='right' style='padding:8px 0;border-top:1px solid {LINE2};font-family:{MONO};font-size:11px;font-weight:700;color:{GREEN}'>{cost}</td></tr>")
    more = (f"<tr><td colspan='9' style='padding:10px 0;border-top:1px solid {LINE2};text-align:center;"
           f"font-size:11px;color:{INK3}'>&hellip; and {len(rows) - 200} more rows &middot; see the attached export</td></tr>"
           ) if len(rows) > 200 else ""
    theads = "".join(f"<th align='{'left' if i < 5 else 'right'}' style='padding:0 10px 8px 0;font-size:10px;"
                     f"font-weight:800;letter-spacing:.4px;color:{INK3};text-transform:uppercase'>{h}</th>"
                     for i, h in enumerate(["Date", "User", "Provider", "Model", "Module", "Calls", "In", "Out", "Cost"]))
    table_block = (f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE}'>"
                  f"<div style='font-size:14.5px;font-weight:800;color:{INK};letter-spacing:-.2px'>Usage by date &middot; user &middot; model</div>"
                  f"<div style='margin-top:12px;overflow-x:auto;-webkit-overflow-scrolling:touch'>"
                  f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='min-width:520px'>"
                  f"<tr>{theads}</tr>{trows}{more}</table></div></td></tr>") if rows else (
        f"<tr><td style='padding:26px 32px;border-top:1px solid {LINE};text-align:center;"
        f"color:{INK3};font-size:12.5px'>No AI usage recorded in this range.</td></tr>")

    unpriced_note = (
        f"<tr><td style='padding:0 32px 18px'><div style='font-size:11px;color:{AMBER};"
        f"background:#FBF3DC;border-radius:8px;padding:9px 12px'>{report['unpriced_calls']} call(s) "
        f"use a model with no published price and are excluded from the cost total above.</div></td></tr>"
    ) if report.get("unpriced_calls") else ""

    footer = (f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
             f"<td valign='middle' style='padding-right:9px'>{_logo_tag(24)}</td>"
             f"<td valign='middle' style='font-size:11.5px;font-weight:600;color:{INK3}'>"
             f"Generated by QA Studio &middot; token counts are exact (from each provider's own "
             f"response); cost is an estimate from a locally maintained price table.</td></tr></table>")

    return f"""<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>
</head>
<body style='margin:0;padding:0;background:{PAPER};-webkit-text-size-adjust:100%'>
<center style='width:100%;background:{PAPER}'>
<table role='presentation' width='100%' cellpadding='0' cellspacing='0' style='background:{PAPER}'><tr>
<td align='center' style='padding:26px 12px 48px'>
<table role='presentation' width='680' cellpadding='0' cellspacing='0' style='width:680px;max-width:680px;background:{CARD};border:1px solid #DEDDE6;border-radius:16px;overflow:hidden;font-family:{UI};color:{INK}'>
  <tr><td style='height:3px;line-height:3px;font-size:0;background:{VIOLET}'>&nbsp;</td></tr>
  <tr><td style='padding:24px 32px 0'>{masthead}</td></tr>
  <tr><td style='padding:18px 32px 4px'>{hero}</td></tr>
  <tr><td style='padding:18px 32px 0'>{metrics}</td></tr>
  {table_block}
  {unpriced_note}
  <tr><td style='padding:20px 32px 26px;border-top:1px solid {LINE};background:{TINT}'>{footer}</td></tr>
</table>
</td></tr></table></center></body></html>"""


# ═══════════════════════════════════════════════════════════════════════════════
#  SELF-HEALING AUTOMATION GENERATION (Selenium + Java + TestNG)
# ═══════════════════════════════════════════════════════════════════════════════
#  Flow (no live browser at generation time):
#   1. compile_test_case()          — AI turns noisy steps into atomic INTENTS.
#   2. validate_and_sequence_suite()— classify + order cases (logged-out → login → app).
#   3. build_selfhealing_project()  — emit a Maven/TestNG project whose Healer
#                                     resolves each element at RUNTIME (seed locator,
#                                     else AI-picked from the live DOM), then caches it.
#   4. push_to_git()                — commits + pushes the project to a Git repo.
# ═══════════════════════════════════════════════════════════════════════════════

def _safe_class_name(text, fallback="Story"):
    """Turn an arbitrary string into a valid Java class identifier."""
    # transliterate-ish: keep ascii letters/digits, capitalize words
    t = re.sub(r"[^0-9A-Za-z]+", " ", str(text)).strip()
    if not t:
        return fallback
    parts = [p for p in t.split(" ") if p]
    name = "".join(p[:1].upper() + p[1:] for p in parts)
    if name and name[0].isdigit():
        name = fallback + name
    return name or fallback


_HARVEST_JS = r"""
function robustCss(el){
  if(el.id) return '#'+CSS.escape(el.id);
  if(el.name) return el.tagName.toLowerCase()+'[name="'+el.name+'"]';
  let path=[], e=el;
  while(e && e.nodeType===1 && path.length<5){
    let sel=e.tagName.toLowerCase();
    if(e.className && typeof e.className==='string'){
      let c=e.className.trim().split(/\s+/).filter(Boolean).slice(0,2);
      if(c.length) sel+='.'+c.map(x=>CSS.escape(x)).join('.');
    }
    let p=e.parentNode, idx=1, sib=e;
    while(sib=sib.previousElementSibling){ if(sib.tagName===e.tagName) idx++; }
    sel+=':nth-of-type('+idx+')';
    path.unshift(sel);
    e=e.parentNode;
    if(e && e.id){ path.unshift('#'+CSS.escape(e.id)); break; }
  }
  return path.join(' > ');
}
function xpathOf(el){
  if(el.id) return '//*[@id="'+el.id+'"]';
  let parts=[], e=el;
  while(e && e.nodeType===1){
    let idx=1, sib=e;
    while(sib=sib.previousElementSibling){ if(sib.tagName===e.tagName) idx++; }
    parts.unshift(e.tagName.toLowerCase()+'['+idx+']');
    e=e.parentNode;
  }
  return '/'+parts.join('/');
}
const sel='input,button,a,select,textarea,[role=button],[role=link],[role=tab],[role=menuitem],[role=option],[role=checkbox],[role=switch],[contenteditable=true]';
function anameOf(el){
  // best-effort accessible name: aria-label, associated <label>, title, alt, text
  let n = el.getAttribute('aria-label') || '';
  if(!n && el.id){
    try{ var lab=document.querySelector('label[for="'+CSS.escape(el.id)+'"]');
         if(lab) n=(lab.innerText||'').trim(); }catch(e){}
  }
  if(!n){ var pl=el.closest('label'); if(pl) n=(pl.innerText||'').trim(); }
  if(!n) n = el.getAttribute('title') || el.getAttribute('alt') || '';
  return (n||'').trim().slice(0,80);
}
const els=[...document.querySelectorAll(sel)];
return els.slice(0,250).map((el,i)=>({
  idx: i,
  tag: el.tagName.toLowerCase(),
  type: el.getAttribute('type')||'',
  role: el.getAttribute('role')||'',
  id: el.id||'',
  name: el.getAttribute('name')||'',
  testid: el.getAttribute('data-testid')||el.getAttribute('data-test')||el.getAttribute('data-cy')||'',
  text: (el.innerText||el.value||'').trim().slice(0,60),
  placeholder: el.getAttribute('placeholder')||'',
  aria: el.getAttribute('aria-label')||'',
  aname: anameOf(el),
  cls: ((typeof el.className==='string'?el.className:'')+' '+
        (el.querySelector('i,svg,[class*=icon i]')?
          (el.querySelector('i,svg,[class*=icon i]').getAttribute('class')||''):'')).trim().slice(0,120),
  svgicon: (function(){
    var s=el.getAttribute('data-svgicon')||el.getAttribute('data-svg-icon')||
          el.getAttribute('ng-reflect-svg-icon')||el.getAttribute('data-icon')||'';
    if(!s){var d=el.querySelector('[data-svgicon],[data-svg-icon],[ng-reflect-svg-icon],[data-icon]');
           if(d) s=d.getAttribute('data-svgicon')||d.getAttribute('data-svg-icon')||
                   d.getAttribute('ng-reflect-svg-icon')||d.getAttribute('data-icon')||'';}
    return (s||'').trim().slice(0,40);
  })(),
  disabled: !!(el.disabled||el.getAttribute('aria-disabled')==='true'),
  visible: !!(el.offsetWidth||el.offsetHeight||el.getClientRects().length),
  css: robustCss(el),
  xpath: xpathOf(el)
}));
"""


# Error / validation message nodes — these are usually spans/divs/[role=alert]
# and are invisible to the interactive harvest above. Captured separately so the
# DOM-diff assertion binder (and negative-login error capture) can find them.
_ERROR_HARVEST_JS = r"""
function robustCss(el){
  if(el.id) return '#'+CSS.escape(el.id);
  if(el.name) return el.tagName.toLowerCase()+'[name="'+el.name+'"]';
  let path=[], e=el;
  while(e && e.nodeType===1 && path.length<5){
    let sel=e.tagName.toLowerCase();
    if(e.className && typeof e.className==='string'){
      let c=e.className.trim().split(/\s+/).filter(Boolean).slice(0,2);
      if(c.length) sel+='.'+c.map(x=>CSS.escape(x)).join('.');
    }
    let p=e.parentNode, idx=1, sib=e;
    while(sib=sib.previousElementSibling){ if(sib.tagName===e.tagName) idx++; }
    sel+=':nth-of-type('+idx+')';
    path.unshift(sel);
    e=e.parentNode;
    if(e && e.id){ path.unshift('#'+CSS.escape(e.id)); break; }
  }
  return path.join(' > ');
}
function xpathOf(el){
  if(el.id) return '//*[@id="'+el.id+'"]';
  let parts=[], e=el;
  while(e && e.nodeType===1){
    let idx=1, sib=e;
    while(sib=sib.previousElementSibling){ if(sib.tagName===e.tagName) idx++; }
    parts.unshift(e.tagName.toLowerCase()+'['+idx+']');
    e=e.parentNode;
  }
  return '/'+parts.join('/');
}
const sel="[role=alert],[role=status],[aria-live],.alert,.alert-error,.alert-danger,"+
  ".error,.has-error,.invalid-feedback,.help-block,.field-error,.form-error,.toast,"+
  ".kc-feedback-text,.pf-c-form__helper-text,#input-error,.message,.notification,"+
  "[id*=error i],[class*=error i],[class*=invalid i],[class*=feedback i],[class*=danger i],[class*=toast i]";
const out=[]; const seen=new Set();
[...document.querySelectorAll(sel)].forEach(el=>{
  const txt=(el.innerText||el.textContent||'').trim();
  if(!txt) return;
  if(txt.length>220) return;
  const key=robustCss(el);
  if(seen.has(key)) return; seen.add(key);
  out.push({
    tag: el.tagName.toLowerCase(), type:'', role: el.getAttribute('role')||'',
    id: el.id||'', name: el.getAttribute('name')||'', testid:'',
    text: txt.slice(0,120), placeholder:'',
    aria: el.getAttribute('aria-label')||'', aname:'',
    disabled:false,
    visible: !!(el.offsetWidth||el.offsetHeight||el.getClientRects().length),
    css: key, xpath: xpathOf(el), is_error: true
  });
});
return out.slice(0,80);
"""


# ═══════════════════════════════════════════════════════════════════════════════
#  INTENT-DRIVEN EXPLORER  (compile → deterministic execute → AI tie-break)
#  Replaces the old "ask the AI to pick 1 of 120 elements per step" approach,
#  which produced repeated clicks (assertions/restated steps treated as actions)
#  and false guesses (preconditions treated as actions). Now:
#    1. compile_test_case()  — LLM turns messy steps into typed intents ONCE per
#       case (precondition / action / assertion), with page-language keywords.
#    2. _rank_candidates()   — deterministic locator binding against the live DOM.
#    3. AI is used only to break ties among a short candidate list, never to
#       invent locators. Assertions bind by DOM-diff (what newly appeared).
# ═══════════════════════════════════════════════════════════════════════════════
import unicodedata as _ud

_AR_DIACRITICS = "".join(chr(c) for c in list(range(0x0610, 0x061B)) +
                         list(range(0x064B, 0x0660)) + [0x0670, 0x0640])  # +tatweel

def _norm(s):
    """Normalize text for language-agnostic matching: lowercase, strip Arabic
    diacritics/tatweel, unify alef/ya/ta-marbuta, collapse whitespace."""
    s = (s or "").strip().lower()
    s = "".join(ch for ch in s if ch not in _AR_DIACRITICS)
    s = (s.replace("\u0623", "\u0627").replace("\u0625", "\u0627").replace("\u0622", "\u0627")
           .replace("\u0649", "\u064a").replace("\u0629", "\u0647"))
    s = _ud.normalize("NFKC", s)
    return re.sub(r"\s+", " ", s).strip()


# Framework-generated ids that change between renders/sessions and MUST NOT be
# captured as saved locators (the generated Java would break next run):
# PrimeNG (pn_id_*), Angular CDK/Material (cdk-*, mat-*), React useId (:r0:),
# GUIDs, and PrimeNG panel/header patterns like pn_id_18_0_header.
_VOLATILE_ID = re.compile(
    r"(^pn_id|^cdk-|^mat-|^ui-id-|^:r[0-9a-z]+:|"
    r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}|_[0-9]+_[0-9]+)", re.I)


def _xq(s):
    """Quote a string for an XPath literal, handling embedded quotes via concat()."""
    s = s or ""
    if '"' not in s:
        return '"' + s + '"'
    if "'" not in s:
        return "'" + s + "'"
    return "concat(" + ", '\"', ".join('"%s"' % p for p in s.split('"')) + ")"


_LOGIN_CTX_KWS = ("login", "log in", "logon", "sign in", "signin", "authenticat",
                  "تسجيل الدخول", "تسجيل دخول", "الدخول", "كلمة المرور", "كلمة مرور",
                  "اسم المستخدم", "اسم مستخدم")
_NEG_LOGIN_KWS = ("invalid", "wrong", "incorrect", "fail", "empty", "blank", "without",
                  "locked", "lockout", "bad credentials", "required",
                  "خاطئ", "خاطئة", "غير صحيح", "غير صحيحة", "بيانات غير", "فارغ",
                  "بدون", "فشل", "خطأ", "مطلوب")
_PRESENCE_KWS = ("التحقق من وجود", "من وجود", "وجود", "موجود", "ظهور",
                 "verify the existence", "existence of", "presence of", "exists",
                 "is displayed", "is visible", "is present")
_EMPTY_FIELD_KWS = ("empty", "blank", "without", "leave it", "leave the", "do not enter",
                    "don't enter", "no password", "no username", "missing",
                    "فارغ", "بدون", "اترك", "دون إدخال", "لا تدخل")


def _tc_blob(tc):
    blob = (tc.get("title", "") or "")
    for s in (tc.get("steps") or []):
        blob += " " + (s.get("action", "") or "") + " " + \
                (s.get("expected", "") or "") + " " + (s.get("precondition", "") or "")
    return blob.lower()


def _is_negative_login_tc(tc):
    """A negative/validation LOGIN case — must run on a fresh login page so the
    bad submit surfaces real error-state locators."""
    low = _tc_blob(tc)
    return (any(k in low for k in _LOGIN_CTX_KWS) and
            any(k.lower() in low for k in _NEG_LOGIN_KWS))


def _classify_case(tc):
    """Return 'negative_login' | 'presence' | 'interaction'."""
    if _is_negative_login_tc(tc):
        return "negative_login"
    # presence is a property of the case's INTENT (its title), not of an "appears"
    # word that may show up in any interaction case's expected result.
    title = _norm(tc.get("title", ""))
    if any(_norm(k) in title for k in _PRESENCE_KWS):
        return "presence"
    return "interaction"


# Words that say a case belongs on the LOGGED-OUT login page (where the Keycloak
# language DROPDOWN lives) vs. the authenticated app (where a single language
# TOGGLE button lives). Inferred from the case's own title/steps — no hard-coded
# per-story rules — exactly as the AI-authored case intends.
_LOGIN_PAGE_KWS = ("صفحة تسجيل الدخول", "صفحه تسجيل الدخول", "تسجيل الدخول",
                   "قبل تسجيل الدخول", "شاشة الدخول", "login page", "sign-in page",
                   "sign in page", "on login", "locale", "kc_locale")
_DROPDOWN_KWS = ("قائمة منسدلة", "قائمه منسدله", "منسدلة", "منسدله", "القائمة المنسدلة",
                 "الاختيار بين", "قائمة اللغات", "dropdown", "drop-down", "drop down")


def _infer_page_context(tc, case_type="interaction"):
    """Decide where this case runs: 'login' (logged-out, language dropdown) or
    'app' (post-login, single language toggle). Read from the case text. Ambiguous
    cases default to 'app' (the common case) and the choice is logged so a wrong
    inference is visible in the activity feed."""
    if case_type == "negative_login":
        return "login"
    # Login keywords in the STEPS are almost always a "log in first" PRECONDITION,
    # not the thing under test — so infer page context from the TITLE (the case's
    # real intent). Otherwise every app case with a login precondition would be
    # dragged onto the login page (→ everything "logged-out", 0 app cases).
    title = _norm(tc.get("title", ""))
    if any(_norm(k) in title for k in _LOGIN_PAGE_KWS):
        return "login"
    blob = _norm(_tc_blob(tc))
    # a language case described as a DROPDOWN / choose-between is the login page;
    # the in-app control is a single toggle with no dropdown.
    lang = any(_norm(k) in blob for k in ("اللغة", "لغة", "language", "locale"))
    if lang and any(_norm(k) in blob for k in _DROPDOWN_KWS):
        return "login"
    return "app"


def _wants_empty_field(text):
    t = (text or "").lower()
    return any(k in t for k in _EMPTY_FIELD_KWS)


# Programmatic backstop for compile_test_case's verb field — a weak model
# will often pattern-match a step's OWN wording into "verb" instead of the
# 6 allowed enum values (e.g. "press", "enter", or an Arabic verb), and will
# sometimes phrase an assertion step ("verify X appears") as if it were an
# action. Both get corrected here rather than trusted from the prompt alone
# — same philosophy as the dedup prompts' structured-output backstop:
# constrain in the prompt, THEN verify/repair in Python, since a weak model
# following instructions is a probability, not a guarantee.
_VERB_SYNONYMS = {
    "press": "click", "tap": "click", "اضغط": "click", "انقر": "click",
    "enter": "type", "write": "type", "fill": "type", "input": "type",
    "اكتب": "type", "أدخل": "type", "ادخل": "type",
    "open": "navigate", "goto": "navigate", "go": "navigate",
    "انتقل": "navigate", "افتح": "navigate",
    "choose": "select", "pick": "select", "اختر": "select",
}
_ASSERTION_VERB_WORDS = {"verify", "check", "assert", "expect", "confirm",
                         "تأكد", "تحقق", "يظهر", "تظهر"}
_VALID_ACTION_VERBS = ("navigate", "click", "type", "select", "hover", "wait")


def compile_test_case(tc, story=None, log=None, case_type="interaction",
                      case_label="", meta_out=None):
    """STAGE 1 — turn a test case's raw steps into a normalized, deduplicated list
    of typed INTENTS. The LLM reads the (often messy, Arabic) steps and returns
    JSON; it never sees locators, so it cannot hallucinate them.

    case_type ('presence'|'interaction'|'negative_login') shapes the output: a
    presence case must NOT be walked as a long interaction.

    Each intent:
      {"role":"precondition"|"action"|"assertion",
       "verb":"navigate|click|type|select|hover|wait",   # action only
       "target":"<human description>",
       "keywords":["visible text / aria tokens in the PAGE language", ...],
       "kind":"button|link|input|select|checkbox|menuitem|text|any",
       "value":"<text to type/select, '' for empty-field cases>",
       "check":"visible|hidden|text_contains|url_contains|enabled|disabled|count",
       "expected":"<expected value for the check>",
       "from_steps":[1-based original step indices this intent came from]}

    Returns a list of intents, or [] on failure (caller falls back to raw steps).
    """
    log = log or (lambda *a, **k: None)
    steps = tc.get("steps") or []
    raw = []
    for i, s in enumerate(steps, 1):
        raw.append({"n": i, "precondition": (s.get("precondition", "") or "").strip(),
                    "action": (s.get("action", "") or "").strip(),
                    "expected": (s.get("expected", "") or "").strip()})
    lang = "Arabic" if _is_arabic_out() else "English"
    # Concrete per-case-type templates instead of a soft description — a weak
    # model follows "typically N actions + M assertions" far more reliably
    # than an abstract instruction like "don't over-split" (see the deleted
    # SIZE rule below, which just told the model to self-correct — something
    # small models don't reliably do while generating token-by-token).
    shape = {
        "presence": ("PRESENCE/visibility case. Typical shape: 0-1 navigate action, "
                     "then exactly 1 assertion that the element is visible. No "
                     "click/select/type — this is a look-only check."),
        "negative_login": ("NEGATIVE-LOGIN case. Typical shape: 1-2 type actions "
                           "(the invalid/empty credentials, value=\"\" for an empty "
                           "field), 1 click action (submit), then exactly 1 assertion "
                           "that the error/validation message is visible."),
        "interaction": ("INTERACTION case. Typical shape: one action per distinct "
                        "UI operation the user performs, then 1-2 assertions TOTAL "
                        "for the outcome — not one assertion per step."),
    }.get(case_type, "One action per distinct UI operation, then 1-2 assertions total.")
    prompt = (
        "ROLE\n"
        "You convert ONE UI test case into an ordered list of atomic INTENTS that a "
        "Selenium walker executes step by step.\n\n"
        "INPUT\n"
        "The raw steps may be Arabic or English and are usually noisy: preconditions "
        "written as steps, the same action restated across several steps, or an action "
        "merged with its expected result. Read intent, don't transcribe literally. Each "
        "raw step below carries its own \"n\" — use those exact values in from_steps, "
        "never invent your own numbering.\n\n"
        f"CASE TYPE: {case_type} — {shape}\n\n"
        "OUTPUT\n"
        'Return ONLY a JSON object shaped {"intents": [ <intent>, ... ]} — no markdown, '
        "no commentary. Each <intent> object has exactly these fields:\n"
        '  "role"      : "precondition" | "action" | "assertion"\n'
        '  "verb"      : for role=action one of [navigate,click,type,select,hover,wait]; else ""\n'
        '  "target"    : short human name of the element or goal\n'
        '  "keywords"  : 1-6 short tokens — the EXACT visible text/aria-label/placeholder '
        f'as it appears on the page, in {lang} FIRST, then an English guess when unsure; '
        "each token ≤3 words (never a whole sentence copied from the step); "
        "for icon-only buttons add icon tokens (globe,language,lang,flag,world,translate)\n"
        '  "kind"      : expected element type (button,input,link,menuitem,text,...)\n'
        '  "value"     : text to type ("" for an empty-field check); for select = the option\'s visible text\n'
        '  "check"     : optional assertion kind, else ""\n'
        '  "expected"  : the expected-outcome text, if any, else ""\n'
        '  "from_steps": array of the original step "n" value(s) this intent came from\n\n'
        "RULES\n"
        "1) ROLES\n"
        "   • precondition = environmental/state setup with NO UI action (internet up, "
        "browser open, user already on page X). Never invent a click for these.\n"
        "   • action = exactly ONE real UI operation. Collapse repeated/restated steps "
        "for the SAME operation into a single action; never emit the same action twice in a row.\n"
        "   • assertion = a verification of an outcome. Collapse hard: at most ONE per "
        "distinct observable outcome; most cases need 1-2 assertions TOTAL, not one per "
        "step. An assertion is never a click.\n"
        "2) VERB WORDS THAT MEAN ASSERTION, NOT ACTION — if a step's wording is "
        "verify/check/confirm/expect/تأكد/تحقق/يظهر/تظهر (\"appears\"/\"shows\"), that step "
        "is role=assertion with verb=\"\", NEVER an action. Verb synonyms for real "
        "actions: press/tap/اضغط/انقر → click; enter/write/fill/اكتب/أدخل → type; "
        "open page/انتقل → navigate; choose/اختر → select.\n"
        "3) CUSTOM DROPDOWN (PrimeNG/Material, not a native <select>) = TWO actions: "
        "click the trigger to open it, then click the option. For the option set "
        'kind="menuitem" and value = its visible text (e.g. English / العربية).\n'
        '4) EMPTY-FIELD validation (leave a field blank) = a type action with value="".\n'
        '5) PAGE: also decide WHERE this case runs. "login" ONLY if it tests the '
        'sign-in page itself (its own fields/buttons/language dropdown) while logged '
        'OUT. If a precondition says the user is already logged in, or the steps use '
        'in-app elements (profile avatar, header menus, edit/save icons), it is '
        '"app" — EVEN IF the title mentions passwords or logout: menu items listed '
        'in a title (e.g. "تغيير كلمة المرور - تسجيل الخروج") are things shown INSIDE '
        'the app, not the login form.\n\n'
        f"TEST CASE TITLE: {tc.get('title','')}\n"
        f"ACCEPTANCE CRITERIA: {((story or {}).get('criteria') or '')[:800]}\n"
        f"RAW STEPS (JSON): {json.dumps(raw, ensure_ascii=False)[:5000]}\n\n"
        "EXAMPLE — study it carefully; it shows precondition handling, collapsing a "
        "restated step into one action, and merging into one assertion:\n"
        'Steps: {"n":1,"action":"الانترنت متاح"} {"n":2,"action":"افتح صفحة تسجيل الدخول"} '
        '{"n":3,"action":"اضغط زر الدخول"} {"n":4,"action":"اضغط على زر الدخول مرة أخرى"} '
        '{"n":5,"expected":"تظهر رسالة: البريد مطلوب"}\n'
        'Output: {"page":"login","intents": [\n'
        '  {"role":"precondition","verb":"","target":"internet available","keywords":[],'
        '"kind":"","value":"","check":"","expected":"","from_steps":[1]},\n'
        '  {"role":"action","verb":"navigate","target":"login page","keywords":['
        '"تسجيل الدخول","login"],"kind":"link","value":"","check":"","expected":"",'
        '"from_steps":[2]},\n'
        '  {"role":"action","verb":"click","target":"login button","keywords":['
        '"الدخول","login"],"kind":"button","value":"","check":"","expected":"",'
        '"from_steps":[3,4]},\n'
        '  {"role":"assertion","verb":"","target":"required-email message","keywords":['
        '"البريد مطلوب","required"],"kind":"text","value":"","check":"text_visible",'
        '"expected":"البريد مطلوب","from_steps":[5]}\n'
        "]}\n"
        "Note steps 3+4 collapsed into ONE click action (from_steps:[3,4]), and the "
        "precondition has verb=\"\" with no invented click.\n\n"
        'Return ONLY {"page":"login|app","intents":[...]}. Every intent has all 9 fields. No markdown, no explanation.'
    )
    try:
        # Retry lines are indented 2 spaces (aligns them with the "  ⏱ …"
        # case lines around them instead of sitting flush-left, visually
        # orphaned) and prefixed with case_label ("#k/N · ") — with 2
        # concurrent compile workers, an unlabeled "request timed out —
        # retrying…" line can't be tied back to WHICH case is retrying.
        out = parse_json_robust(ai_complete(prompt, max_tokens=4096, timeout=90,
                                            on_retry=lambda m: log(
                                                f"  {case_label}{m}", "dim"),
                                            want_json=True,
                                            usage_tag="automation_compile"))
        # JSON mode forces an object on OpenAI-compatible providers, so unwrap the
        # intents array from whatever key the model used (it may invent one).
        if isinstance(out, dict):
            # The model's page-context judgment (rule 5 above) — reported via
            # meta_out (same additive out-param pattern as ai_complete's
            # usage_out) so the caller can override the title-keyword
            # heuristic, which misreads titles that merely ENUMERATE menu
            # items ("… تغيير كلمة المرور - تسجيل الخروج") as login cases.
            _pg = str(out.get("page") or "").strip().lower()
            if meta_out is not None and _pg in ("login", "app"):
                meta_out["page"] = _pg
            out = (out.get("intents") or out.get("items") or out.get("steps")
                   or next((v for v in out.values() if isinstance(v, list)), None)
                   or [])
        if not isinstance(out, list) or not out:
            return []
        _n_steps = len(raw)
        clean = []
        for it in out:
            if not isinstance(it, dict):
                continue
            role = (it.get("role") or "action").strip().lower()
            if role not in ("precondition", "action", "assertion"):
                role = "action"
            verb_raw = (it.get("verb") or "").strip().lower()
            # A verify/check-worded "verb" means this is really an assertion,
            # regardless of what role the model assigned it — this is the
            # single most common weak-model mistake for this prompt (see
            # compile_test_case's own docstring/comments above).
            if verb_raw in _ASSERTION_VERB_WORDS:
                role = "assertion"
            if role == "action":
                verb = _VERB_SYNONYMS.get(verb_raw, verb_raw)
                if verb not in _VALID_ACTION_VERBS:
                    # Unmapped/unknown verb text — fall back by element kind
                    # rather than dropping the intent or guessing wildly.
                    _kind_hint = (it.get("kind") or "").strip().lower()
                    verb = "type" if _kind_hint in ("input", "textarea") else "click"
            else:
                verb = ""
            fs = it.get("from_steps") or []
            if isinstance(fs, int):
                fs = [fs]
            # Clamp from_steps to the actual valid range — the prompt now
            # tells the model to reuse each raw step's own "n", but a
            # hallucinated/off-by-one reference is still possible; silently
            # dropping out-of-range values is safer than keeping a reference
            # to a step that doesn't exist.
            from_steps = sorted({int(x) for x in fs if str(x).strip().lstrip("-").isdigit()
                                 and 1 <= int(x) <= _n_steps})
            clean.append({
                "role": role,
                "verb": verb,
                "target": (it.get("target") or "").strip(),
                "keywords": [str(k) for k in (it.get("keywords") or []) if str(k).strip()][:6],
                "kind": (it.get("kind") or "any").strip().lower(),
                "value": str(it.get("value") or ""),
                "check": (it.get("check") or "").strip().lower(),
                "expected": str(it.get("expected") or ""),
                "from_steps": from_steps,
            })
        # Safety net: collapse consecutive actions that repeat the SAME verb
        # + target (the LLM sometimes still emits a restated step as its own
        # action despite the prompt's example demonstrating this exact
        # collapse), and consecutive assertions that check the same thing.
        # Merge their from_steps so each original step still receives an
        # assert_locator / the right action mapping.
        collapsed = []
        for it in clean:
            if (it["role"] == "assertion" and collapsed and
                    collapsed[-1]["role"] == "assertion" and
                    _norm(collapsed[-1]["target"]) == _norm(it["target"]) and
                    _norm(collapsed[-1]["expected"]) == _norm(it["expected"])):
                collapsed[-1]["from_steps"] = sorted(set(collapsed[-1]["from_steps"] + it["from_steps"]))
                continue
            if (it["role"] == "action" and collapsed and
                    collapsed[-1]["role"] == "action" and
                    collapsed[-1]["verb"] == it["verb"] and
                    _norm(collapsed[-1]["target"]) == _norm(it["target"])):
                collapsed[-1]["from_steps"] = sorted(set(collapsed[-1]["from_steps"] + it["from_steps"]))
                continue
            collapsed.append(it)
        return collapsed
    except CreditBalanceError:
        raise
    except Exception as e:
        # Recoverable provider errors (expired/invalid key, rate limit, outage)
        # propagate so the run can PAUSE and let the user fix it + Resume.
        if _is_recoverable_ai_error(e):
            raise
        # Genuine non-recoverable issues (e.g. a malformed JSON response) fall
        # back to raw steps with a single clean line — no giant JSON dump.
        log(f"    compile failed ({friendly_ai_error(e)}) — using raw steps", "warn")
        return []


def _intents_from_raw_steps(tc):
    """Fallback when the compiler is unavailable: derive simple intents from the
    raw steps so the walk never regresses below the old behavior."""
    intents = []
    for i, s in enumerate(tc.get("steps") or [], 1):
        action = (s.get("action", "") or "").strip()
        exp = (s.get("expected", "") or "").strip()
        disp = action
        for pfx in ("الشرط المسبق:", "الإجراء:", "Precondition:", "Action:"):
            disp = disp.replace(pfx, " ")
        disp = disp.strip()
        if not disp and not exp:
            intents.append({"role": "precondition", "verb": "", "target": "",
                            "keywords": [], "kind": "any", "value": "", "check": "",
                            "expected": "", "from_steps": [i]})
            continue
        if disp:
            low = _norm(disp)
            verb = ("type" if any(k in low for k in ("type", "enter", "ادخل", "أدخل", "اكتب", "كتابة"))
                    else "select" if any(k in low for k in ("select", "اختر", "اختيار"))
                    else "click")
            intents.append({"role": "action", "verb": verb, "target": disp,
                            "keywords": [w for w in re.split(r"[\s,.:؛،]+", disp) if len(w) > 2][:6],
                            "kind": "input" if verb == "type" else "any",
                            "value": "" if _wants_empty_field(disp) else "",
                            "check": "", "expected": "", "from_steps": [i]})
        if exp:
            intents.append({"role": "assertion", "verb": "", "target": exp,
                            "keywords": [w for w in re.split(r"[\s,.:؛،]+", exp) if len(w) > 2][:6],
                            "kind": "any", "value": "", "check": "visible",
                            "expected": "", "from_steps": [i]})
    return intents


def _el_haystack(el):
    return _norm(" ".join(str(el.get(k, "")) for k in
                          ("text", "aname", "aria", "placeholder", "name", "id",
                           "role", "testid", "type", "cls", "svgicon")))


def _kind_matches(kind, el):
    if not kind or kind == "any":
        return False
    tag = (el.get("tag") or "").lower(); typ = (el.get("type") or "").lower()
    role = (el.get("role") or "").lower(); cls = _norm(el.get("cls", ""))
    menu_cls = any(t in cls for t in ("menu-item", "menuitem", "dropdown-item",
                                      "dropdownitem", "list-item", "option"))
    m = {"button": tag == "button" or typ in ("button", "submit") or role == "button",
         "link": tag == "a" or role == "link",
         "input": tag in ("input", "textarea") and typ not in ("button", "submit", "checkbox"),
         "select": tag == "select" or role in ("combobox", "listbox"),
         "checkbox": typ == "checkbox" or role in ("checkbox", "switch"),
         "menuitem": role in ("menuitem", "option", "tab") or menu_cls}
    return bool(m.get(kind, False))


def _rank_candidates(intent, elements):
    """STAGE 2 — deterministic scoring of live elements against an intent.
    Returns a list of (score, element) sorted high→low. No LLM involved."""
    kws = [_norm(k) for k in (intent.get("keywords") or []) if _norm(k)]
    tgt = _norm(intent.get("target", ""))
    kind = intent.get("kind", "any")
    verb = intent.get("verb", "")
    ranked = []
    for el in elements:
        if not el.get("visible", True):
            continue
        hay = _el_haystack(el)
        score = 0.0
        for k in kws:
            if k and k in hay:
                score += 2.0
                if hay == k or (el.get("text") and _norm(el["text"]) == k):
                    score += 1.0           # exact label match
        # token overlap with the target description
        for tok in (t for t in tgt.split(" ") if len(t) > 2):
            if tok in hay:
                score += 0.5
        if _kind_matches(kind, el):
            score += 1.0
        if verb == "type" and el.get("tag") in ("input", "textarea"):
            score += 0.5
        if score > 0:
            ranked.append((score, el))
    ranked.sort(key=lambda t: t[0], reverse=True)
    return ranked


def _tiebreak_with_ai(intent, shortlist, cb):
    """The ONLY place the LLM picks an element — and only among a short list of
    real candidates (never the full DOM), already pre-filtered/scored by
    _rank_candidates. Returns the chosen element or None."""
    brief = [{"idx": e["idx"], "tag": e.get("tag"), "type": e.get("type"),
              "text": e.get("text"), "aname": e.get("aname"), "aria": e.get("aria"),
              "placeholder": e.get("placeholder"), "id": e.get("id")}
             for e in shortlist]
    prompt = (
        "Pick the ONE element that best matches the intent. Reply ONLY JSON.\n"
        f"INTENT: {json.dumps({k: intent.get(k) for k in ('role','verb','target','keywords','kind')}, ensure_ascii=False)}\n"
        f"CANDIDATES: {json.dumps(brief, ensure_ascii=False)[:3000]}\n"
        "idx MUST be one of the idx values in CANDIDATES, or -1.\n"
        "Match by meaning: prefer visible text/aria-label matching the keywords "
        "(the intent's language and the element's text may differ), then "
        "placeholder, then id.\n"
        '{"idx": <chosen idx or -1>}'
    )
    raw = ""
    try:
        raw = ai_complete(prompt, max_tokens=256, timeout=45, want_json=True,
                          usage_tag="automation_tiebreak",
                          on_retry=lambda m: cb(f"    {m}", "dim"))
        data = parse_json_robust(raw)
        if isinstance(data, list) and data:
            data = data[0]
        idx = int((data or {}).get("idx", -1))
    except CreditBalanceError:
        raise
    except Exception as e:
        # Regex fallback for a malformed-but-close-enough reply (e.g. stray
        # prose around the JSON, or idx returned as a quoted string) — makes
        # this call nearly unbreakable for the one thing that actually
        # matters here, before giving up entirely.
        m = re.search(r'"idx"\s*:\s*"?(-?\d+)', raw or "")
        if not m:
            cb(f"    tiebreak error: {str(e)[:60]}", "warn")
            return None
        idx = int(m.group(1))
    if idx < 0:
        return None
    # idx MUST reference a real candidate — a hallucinated idx not in the
    # shortlist would otherwise silently corrupt the match (the caller's
    # `next(...)` already guards this structurally, but making the intent
    # explicit here rather than relying on that fallthrough).
    return next((e for e in shortlist if e.get("idx") == idx), None)


def _to_locator(el):
    """Pick the most STABLE locator for the SAVED test (generated Java), so it
    survives re-renders. Order: data-testid > data-svgicon > stable id > name >
    aria-label > short visible text > css path > xpath. Framework-generated ids
    (PrimeNG pn_id_*, Angular/React, GUIDs) are skipped — they change every run."""
    if not el:
        return None
    tid = (el.get("testid") or "").strip()
    if tid:
        return {"by": "css", "value": '[data-testid="%s"]' % tid}
    svg = (el.get("svgicon") or "").strip()
    if svg:
        return {"by": "css", "value": '[data-svgicon="%s"]' % svg}
    eid = (el.get("id") or "").strip()
    if eid and not _VOLATILE_ID.search(eid):
        return {"by": "id", "value": eid}
    nm = (el.get("name") or "").strip()
    if nm:
        return {"by": "name", "value": nm}
    al = (el.get("aria") or el.get("aname") or "").strip()
    if al:
        return {"by": "xpath", "value": "//*[@aria-label=%s]" % _xq(al)}
    txt = (el.get("text") or "").strip()
    if txt and len(txt) <= 40:
        return {"by": "xpath",
                "value": "//%s[normalize-space()=%s]" % (el.get("tag", "*"), _xq(txt))}
    css = (el.get("css") or "")
    if css and not _VOLATILE_ID.search(css):
        return {"by": "css", "value": css}
    return {"by": "xpath", "value": el.get("xpath", "")}


_KIND_SYNONYMS = {"press": "click", "tap": "click", "open": "navigate",
                  "goto": "navigate", "input": "type", "enter": "type",
                  "fill": "type", "verify": "none", "check": "none",
                  "assert": "none"}
_VALID_MATCH_KINDS = ("click", "type", "select", "navigate", "none")


def _rank_elements_by_step(action, elements):
    """Cheap deterministic pre-filter for _match_step_to_element — token
    overlap between the raw step-action text and each element's haystack
    (_el_haystack). Exists because sending up to 120 elements/6000 chars to a
    weak model is a needle-in-haystack task it handles poorly: it tends to
    pick from whichever elements happen to be early in the list, and the
    6000-char truncation can silently cut the correct element out entirely
    on a busy page. Pre-ranking and sending only the strongest handful both
    fixes the truncation risk and cuts token cost substantially. No LLM
    involved — same spirit as _rank_candidates (STAGE 2 scoring), just
    against a raw action string instead of a structured intent."""
    toks = [t for t in re.split(r"[\s,.:؛،]+", _norm(action or "")) if len(t) > 2]
    scored = []
    for el in elements:
        if not el.get("visible", True):
            continue
        hay = _el_haystack(el)
        score = sum(1.0 for t in toks if t in hay)
        scored.append((score, el))
    scored.sort(key=lambda t: t[0], reverse=True)
    return [el for _s, el in scored]


def _match_step_to_element(action, elements, cb):
    """Ask the AI which real DOM element best matches a step's action.
    Returns (element_dict_or_None, kind) where kind in
    {'type','click','select','navigate','none'} and a value to type if relevant."""
    # Pre-rank by cheap token overlap and cap at ~30 — see
    # _rank_elements_by_step's docstring for why this matters more than any
    # wording change to the prompt itself. Falls back to the natural DOM
    # order for anything beyond the token-overlap signal (stable sort keeps
    # ties in original order), so this never does worse than the old
    # first-120 cutoff, only better.
    ranked = _rank_elements_by_step(action, elements)[:30]
    brief = [{"idx": e["idx"], "tag": e["tag"], "type": e["type"], "id": e["id"],
              "name": e["name"], "text": e["text"], "placeholder": e["placeholder"],
              "aria": e["aria"], "visible": e["visible"]}
             for e in ranked]
    prompt = (
        "You map a single UI test step to ONE real element on the current page.\n"
        "The step action may be in Arabic or English. Choose the best matching element.\n\n"
        f"STEP ACTION: {action}\n\n"
        f"ELEMENTS (JSON, use the 'idx'):\n{json.dumps(brief, ensure_ascii=False)[:6000]}\n\n"
        "idx MUST be an idx value from ELEMENTS, or -1.\n"
        "If the step only describes an expected result or verification (no user "
        "action performed), return {\"idx\":-1,\"kind\":\"none\",\"value\":\"\"}.\n"
        "The step's language may differ from the element's text — match by "
        "meaning, not exact words.\n"
        "Reply with ONLY a JSON object, no markdown:\n"
        '{\"idx\": <element idx or -1 if none fits>, '
        '\"kind\": \"click|type|select|navigate|none\", '
        '\"value\": \"<text to type if kind==type/select, else empty>\"}'
    )
    try:
        raw = ai_complete(prompt, max_tokens=1024, timeout=60,
                          usage_tag="automation_match_element",
                          on_retry=lambda m: cb(f"    {m}", "dim"))
        if not (raw or "").strip():
            # reasoning models (e.g. NVIDIA qwen) sometimes spend the whole
            # budget thinking and return empty — retry once before giving up.
            # (Largely superseded by ai_complete's own max_tokens-doubling
            # retry on empty responses, but kept as a final fallback for the
            # "returned literally whitespace with a normal stop reason" case
            # that doubling doesn't cover.)
            cb("    matcher returned empty — retrying once…", "dim")
            raw = ai_complete(prompt, max_tokens=2048, timeout=60,
                              usage_tag="automation_match_element",
                              on_retry=lambda m: cb(f"    {m}", "dim"))
        data = parse_json_robust(raw)
        if isinstance(data, list) and data:
            data = data[0]
        idx = int((data or {}).get("idx", -1))
        kind_raw = ((data or {}).get("kind") or "none").strip().lower()
        kind = _KIND_SYNONYMS.get(kind_raw, kind_raw)
        if kind not in _VALID_MATCH_KINDS:
            kind = "none"
        value = (data or {}).get("value", "") or ""
        if idx is None or idx < 0:
            return None, kind, value
        # idx must reference one of the elements actually offered (the
        # pre-ranked+capped list, not the full original set) — a
        # hallucinated idx would otherwise silently resolve to the wrong
        # element or None via a coincidental match elsewhere in `elements`.
        match = next((e for e in ranked if e.get("idx") == idx), None)
        if match is None:
            return None, "none", ""
        # Cross-check kind against the matched element's own tag — a "type"
        # verdict on a non-input element (or "select" on a non-select-ish
        # one) is a clear self-contradiction; downgrade rather than trust it
        # blindly, same philosophy as the dedup prompts' structured backstop.
        tag = (match.get("tag") or "").lower()
        if kind == "type" and tag not in ("input", "textarea"):
            kind = "click" if tag in ("button", "a") else "none"
        elif kind == "select" and not (tag == "select" or match.get("role") in ("combobox", "listbox")):
            kind = "click"
        return match, kind, value
    except CreditBalanceError:
        raise  # propagate so explore_and_map can auto-stop on repeated hits
    except Exception as e:
        cb(f"  match error: {str(e)[:80]}", "warn")
        return None, "none", ""


def _harvest_dom(driver):
    """Return the list of interactive elements on the current page."""
    try:
        return driver.execute_script("return (function(){" + _HARVEST_JS + "})();") or []
    except Exception:
        return []


def _harvest_errors(driver):
    """Return visible error/validation/notification message nodes."""
    try:
        return driver.execute_script(
            "return (function(){" + _ERROR_HARVEST_JS + "})();") or []
    except Exception:
        return []


def _verify_logged_in(driver, login_url, cb):
    """Best-effort check that login actually succeeded. Returns (ok, reason)."""
    import time as _t
    _t.sleep(1.0)
    cur = (driver.current_url or "").rstrip("/")
    base_login = (login_url or "").rstrip("/")
    moved = cur != base_login
    still_pw = False
    try:
        from selenium.webdriver.common.by import By
        pw = driver.find_elements(By.CSS_SELECTOR, "input[type=password]")
        still_pw = any(e.is_displayed() for e in pw)
    except Exception:
        pass
    err = False
    try:
        body = (driver.find_element("tag name", "body").text or "").lower()
        for kw in ("invalid", "incorrect", "failed", "غير صحيح", "خطأ",
                   "wrong password", "try again", "بيانات غير"):
            if kw in body:
                err = True; break
    except Exception:
        pass
    if still_pw and not moved:
        return False, "still on the login form (login likely failed)"
    if err and not moved:
        return False, "an error message is shown on the login page"
    return True, ("logged in — now at " + cur)


def explore_and_map(stories_payload, login, site_url, cb=None, should_stop=None,
                    headless=False, wait_secs=3):
    """LIVE-WALK explorer: logs in (verifying success), then for each test case
    walks its steps in order — matching each step's action to a real element on
    the live page, recording the EXACT locator, and performing the action so the
    page advances to the next step's state.

    Returns the stories_payload enriched so each step gains:
        step["locator"]      = {"by": "id|name|css|xpath", "value": "..."} or None
        step["locator_src"]  = "live" | "snapshot" | "guess"
        step["assert_locator"] (best-effort target for the expected result)
    Also attaches each story's accumulated DOM snapshots for fallback generation.

    NOTE: this performs REAL clicks/typing on the target site. Use a TEST env.
    """
    cb = cb or (lambda *a, **k: None)
    should_stop = should_stop or (lambda: False)
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    import time as _t

    opts = Options()
    if headless:
        opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--window-size=1440,900")
    opts.add_argument("--disable-gpu")

    cb("Launching Chrome…", "dim")
    driver = webdriver.Chrome(options=opts)
    driver.set_page_load_timeout(45)
    wait = WebDriverWait(driver, 20)
    all_snapshots = []   # union of every element seen (for fallback)

    def wait_dom_ready():
        """Wait until document.readyState is complete (page loaded)."""
        try:
            WebDriverWait(driver, 25).until(
                lambda d: d.execute_script("return document.readyState") == "complete")
        except Exception:
            pass

    def find_first(selectors, timeout=20):
        """Wait for and return the first element matching any CSS selector in the
        comma-or-list group. Tries each selector; returns None if none appear."""
        if isinstance(selectors, str):
            selectors = [s.strip() for s in selectors.split(",") if s.strip()]
        end = _t.time() + timeout
        while _t.time() < end:
            for sel in selectors:
                try:
                    els = driver.find_elements(By.CSS_SELECTOR, sel)
                    for e in els:
                        if e.is_displayed() and e.is_enabled():
                            return e
                except Exception:
                    pass
            _t.sleep(0.4)
        return None

    def snapshot(tag="", with_errors=False):
        els = _harvest_dom(driver)
        if with_errors:
            errs = _harvest_errors(driver)
            base = len(els)
            for j, e in enumerate(errs):
                e2 = dict(e); e2["idx"] = base + j
                els.append(e2)
            if errs and tag:
                cb(f"  + {len(errs)} message/error element(s)", "dim")
        all_snapshots.extend(els)
        if tag:
            cb(f"  captured {len(els)} elements ({tag})", "dim")
        return els

    def _el_key(e):
        return (e.get("id"), e.get("name"), e.get("css"), e.get("xpath"))

    def to_locator(el):
        """Stable locator strategy (module-level _to_locator, see there)."""
        return _to_locator(el)

    def _dedup_reindex(els):
        """De-dup the merged snapshot union and assign fresh unique idx values,
        so the matcher can reference elements unambiguously (each per-page harvest
        restarts idx at 0, which would otherwise collide across the union)."""
        seen = set(); out = []
        for e in els:
            key = (e.get("id"), e.get("name"), e.get("css"), e.get("xpath"))
            if key in seen:
                continue
            seen.add(key)
            e2 = dict(e); e2["idx"] = len(out)
            out.append(e2)
        return out

    def find_live(el):
        """Locate the actual Selenium element for an harvested element dict."""
        try:
            if el.get("id"):
                return driver.find_element(By.ID, el["id"])
            if el.get("name"):
                return driver.find_element(By.NAME, el["name"])
            if el.get("css"):
                return driver.find_element(By.CSS_SELECTOR, el["css"])
            return driver.find_element(By.XPATH, el["xpath"])
        except Exception:
            return None

    def _safe_click(el):
        """Click that survives overlays: scroll into view, try native click,
        then fall back to a JS click if something intercepts it."""
        try:
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
        except Exception:
            pass
        try:
            el.click(); return True
        except Exception:
            try:
                driver.execute_script("arguments[0].click();", el); return True
            except Exception as e:
                cb(f"    click fallback failed: {str(e)[:70]}", "warn"); return False

    def _describe(el):
        """Short, human-readable identity for a harvested element (for the log)."""
        if not el:
            return "?"
        if el.get("id"):
            return "#" + el["id"]
        if el.get("name"):
            return "[name=" + el["name"] + "]"
        t = (el.get("text") or el.get("aria") or el.get("placeholder") or "").strip()
        if t:
            return '"' + t[:24] + '"'
        return (el.get("css") or el.get("xpath") or "?")[:32]

    def _flash(live_el):
        """Briefly outline the element in the real browser so the human watching
        can SEE which element each step touched."""
        try:
            driver.execute_script(
                "var o=arguments[0];var p=o.style.outline;var q=o.style.outlineOffset;"
                "o.scrollIntoView({block:'center'});"
                "o.style.outline='3px solid #6A4DFF';o.style.outlineOffset='2px';"
                "setTimeout(function(){o.style.outline=p;o.style.outlineOffset=q;},900);",
                live_el)
        except Exception:
            pass

    def _settle(timeout=8):
        """Wait for the page to stop being busy: readyState complete, no
        aria-busy, and common spinner/loader overlays gone."""
        wait_dom_ready()
        end = _t.time() + timeout
        while _t.time() < end:
            try:
                busy = driver.execute_script(
                    "var b=document.querySelector('[aria-busy=true]');"
                    "var s=document.querySelector("
                    "  '.spinner,.loading,.loader,.MuiBackdrop-root,[class*=spinner i],[class*=loading i]');"
                    "function vis(e){return e&&(e.offsetWidth||e.offsetHeight||e.getClientRects().length);}"
                    "return !!(vis(b)||vis(s));")
            except Exception:
                busy = False
            if not busy:
                return
            _t.sleep(0.25)

    def _dismiss_overlays():
        """Best-effort dismissal of cookie/consent banners and stray modals that
        would intercept clicks. Only clicks clearly-dismissive controls."""
        labels = ["accept", "accept all", "agree", "i agree", "got it", "ok", "close",
                  "dismiss", "no thanks", "موافق", "قبول", "أوافق", "إغلاق", "تم", "حسنا"]
        try:
            btns = _harvest_dom(driver)
            for el in btns:
                txt = _norm(el.get("text") or el.get("aname") or el.get("aria"))
                if not txt:
                    continue
                if any(_norm(l) == txt or _norm(l) in txt for l in labels):
                    live = find_live(el)
                    if live is not None and live.is_displayed():
                        try:
                            driver.execute_script("arguments[0].click();", live)
                            _t.sleep(0.4)
                            return True
                        except Exception:
                            pass
        except Exception:
            pass
        return False

    def _topmost_ok(live_el):
        """True if the element is the topmost hit at its center (not covered by an
        overlay). Used to decide whether to dismiss an overlay before clicking."""
        try:
            return driver.execute_script(
                "var e=arguments[0];var r=e.getBoundingClientRect();"
                "if(!r.width||!r.height)return false;"
                "var x=r.left+r.width/2,y=r.top+r.height/2;"
                "var t=document.elementFromPoint(x,y);"
                "return !!t&&(t===e||e.contains(t)||t.contains(e));", live_el)
        except Exception:
            return True

    def _act(el_dict, verb, value, empty_ok=False):
        """STAGE 3 — perform an action so it survives overlays and timing.
        Re-finds the element, scrolls to center, waits to settle, clears overlays
        if it's covered, then runs a retry ladder (native → JS → ActionChains).
        Returns the live element acted on (or None)."""
        live = find_live(el_dict)
        if live is None:
            return None
        _flash(live)
        try:
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", live)
        except Exception:
            pass
        _settle(timeout=4)
        # type / select don't need topmost; clicks do
        if verb == "type":
            try:
                if empty_ok:
                    live.clear()
                else:
                    live.clear(); live.send_keys(value or "test")
                return live
            except Exception as e:
                cb(f"      type failed: {str(e)[:50]}", "warn"); return live
        if verb == "select":
            from selenium.webdriver.support.ui import Select
            # native <select>
            if (el_dict.get("tag") or "").lower() == "select":
                try:
                    Select(live).select_by_visible_text(value); return live
                except Exception:
                    try:
                        Select(live).select_by_value(value); return live
                    except Exception:
                        pass
            # custom dropdown (PrimeNG/Material/etc.): open the trigger, then click
            # the option whose visible text matches `value` (options render in an
            # overlay appended to <body> only after opening).
            self_open = _act(el_dict, "click", "", empty_ok)
            _settle(timeout=3); _t.sleep(0.4)
            want = _norm(value)
            want_toks = [t for t in want.split(" ") if len(t) > 1]
            best = None
            for o in _harvest_dom(driver):
                if not o.get("visible", True):
                    continue
                hay = _norm(" ".join(str(o.get(k, "")) for k in ("text", "aname", "aria")))
                if not hay:
                    continue
                if (want and want in hay) or any(tok in hay for tok in want_toks) or hay in want:
                    best = o; break
            if best is not None:
                lo = find_live(best)
                if lo is not None:
                    _flash(lo)
                    try:
                        lo.click()
                    except Exception:
                        try:
                            driver.execute_script("arguments[0].click();", lo)
                        except Exception:
                            pass
                    return lo
            cb(f"      select: option '{value[:24]}' not found after opening", "warn")
            return self_open or live
        # click / hover / navigate(default) — make it interception-proof
        if not _topmost_ok(live):
            if _dismiss_overlays():
                live = find_live(el_dict) or live
        try:
            WebDriverWait(driver, 6).until(EC.element_to_be_clickable(live))
        except Exception:
            pass
        # retry ladder
        for attempt in range(3):
            try:
                live.click(); return live
            except Exception:
                live = find_live(el_dict) or live      # handle staleness
                try:
                    from selenium.webdriver.common.action_chains import ActionChains
                    ActionChains(driver).move_to_element(live).pause(0.1).click().perform()
                    return live
                except Exception:
                    try:
                        driver.execute_script("arguments[0].click();", live); return live
                    except Exception:
                        _dismiss_overlays(); _t.sleep(0.3)
        cb("      could not click after retries", "warn")
        return live

    try:
        # ── login + verify ──
        login_url = (login or {}).get("url") or site_url
        # Keycloak (and similar) login URLs often carry one-time session params
        # (execution, tab_id, session_code, code, client_data). Hitting that exact
        # URL later yields "Cookie not found" because the session is gone. Strip
        # those so the IdP issues a FRESH login page + cookie.
        try:
            from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode
            parts = urlsplit(login_url)
            if parts.query:
                stale = {"execution", "tab_id", "session_code", "code",
                         "client_data", "auth_session_id", "kc_locale"}
                kept = [(k, v) for k, v in parse_qsl(parts.query)
                        if k.lower() not in stale]
                login_url = urlunsplit((parts.scheme, parts.netloc, parts.path,
                                        urlencode(kept), ""))
            # If the path points at the one-time login-actions endpoint, drop back
            # to the site URL so Keycloak starts a clean auth flow.
            if "login-actions/authenticate" in parts.path:
                cb("Login URL had one-time session params — using the site URL "
                   "to start a fresh login.", "warn")
                login_url = site_url
        except Exception:
            pass
        if login and login.get("user") and login.get("password"):
            def do_login(fresh=False):
                """Run the login flow on a clean login page. fresh=True clears the
                session first (used to re-establish auth after a negative-login
                case). Returns (ok, reason)."""
                if fresh:
                    try:
                        driver.delete_all_cookies()
                    except Exception:
                        pass
                cb("Opening login page\u2026", "dim")
                driver.get(login_url)
                wait_dom_ready()
                try:
                    user_sel = login.get("user_locator") or (
                        "#username,input[type=email],input[name=email],input[name=username],"
                        "input[name*=user i],input[id*=user i],input[type=text]")
                    cb("Waiting for the username field\u2026", "dim")
                    u = find_first(user_sel, timeout=25)
                    if u is None:
                        raise RuntimeError("username/email field did not appear")
                    u.clear(); u.send_keys(login["user"])
                    pass_sel = login.get("pass_locator") or "#password,input[type=password]"
                    p = find_first(pass_sel, timeout=10)
                    if p is None:
                        nxt = find_first("button[type=submit],#kc-login,button,input[type=submit]", timeout=5)
                        if nxt is not None:
                            nxt.click(); wait_dom_ready()
                        p = find_first(pass_sel, timeout=15)
                    if p is None:
                        raise RuntimeError("password field did not appear")
                    p.clear(); p.send_keys(login["password"])
                    submit_sel = login.get("submit_locator") or (
                        "#kc-login,button[type=submit],input[type=submit],button")
                    btn = find_first(submit_sel, timeout=10)
                    if btn is None:
                        p.submit()
                    else:
                        btn.click()
                    cb("Submitted login \u2014 verifying\u2026", "dim")
                    try:
                        WebDriverWait(driver, 20).until(
                            lambda d: (d.current_url or "").rstrip("/") != login_url.rstrip("/")
                                      or not d.find_elements(By.CSS_SELECTOR, "input[type=password]"))
                    except Exception:
                        pass
                    wait_dom_ready()
                except Exception as e:
                    raise RuntimeError(f"Login step failed: {str(e)[:160]}")
                return _verify_logged_in(driver, login_url, cb)

            have_creds = True
            ok, reason = do_login()
            if not ok:
                raise RuntimeError(f"Login could not be verified — {reason}. "
                                   f"Aborting so locators aren't captured from the wrong page.")
            cb(f"Login verified — {reason}", "ok")
        else:
            have_creds = False
            cb("No login provided — exploring as anonymous user.", "warn")

        # Navigate to the starting page
        cb(f"Opening target page…", "dim")
        driver.get(site_url)
        wait_dom_ready()
        _t.sleep(1.0)

        live_count = snap_count = guess_count = 0
        todos = []                 # steps that will carry a // TODO verify locator
        credit_hits = 0            # repeated AI credit-limit errors -> auto-stop
        CREDIT_STOP = 5
        abort_credit = False

        def _credit_guard(fn, *a, **k):
            """Run an AI-calling fn; count credit errors and trip abort_credit."""
            nonlocal credit_hits, abort_credit
            try:
                return fn(*a, **k)
            except CreditBalanceError:
                credit_hits += 1
                cb(f"AI credit limit hit ({credit_hits}/{CREDIT_STOP}).", "err")
                if credit_hits >= CREDIT_STOP:
                    abort_credit = True
                return None

        MIN_SCORE = 2.0   # below this, there is no real keyword/kind hit — don't
                          # auto-bind (that's how a click landed on #pn_id_*_header)

        def bind_target(intent, pool):
            """Deterministic-first binding of an intent to a live element.
            Returns (element_or_None, source) with source in
            {'live','snapshot','guess'}. The AI is used ONLY to break ties among a
            short candidate list — never to invent a locator. A weak best candidate
            (below MIN_SCORE) is NOT taken silently: the AI may rescue it, else it
            becomes a guess rather than a wrong click."""
            ranked = _rank_candidates(intent, pool)
            if ranked and ranked[0][0] >= MIN_SCORE:
                top = ranked[0][0]
                second = ranked[1][0] if len(ranked) > 1 else 0.0
                if len(ranked) == 1 or (top - second) >= 1:
                    return ranked[0][1], "live"          # confident — no AI call
                shortlist = [e for _, e in ranked[:5]]
                chosen = _credit_guard(_tiebreak_with_ai, intent, shortlist, cb)
                return (chosen or ranked[0][1]), "live"  # AI tie-break, else best deterministic
            if ranked:
                # weak candidates only — let the AI decide from the shortlist; if it
                # declines, fall through rather than clicking a low-confidence guess
                chosen = _credit_guard(_tiebreak_with_ai, intent, [e for _, e in ranked[:5]], cb)
                if chosen is not None:
                    return chosen, "live"
            # nothing solid on this page → try the union of everything seen so far
            union = _dedup_reindex(all_snapshots)
            if union:
                q = intent.get("target") or " ".join(intent.get("keywords") or [])
                r = _credit_guard(_match_step_to_element, q, union, cb)
                if r and r[0]:
                    return r[0], "snapshot"
            return None, "guess"

        def assign(step_idxs, locator, src, as_assert=False):
            """Write a captured locator back onto the ORIGINAL step(s) the intent
            came from, so the generated Java still mirrors the authored test case."""
            nonlocal live_count, snap_count, guess_count
            for n in step_idxs:
                if 1 <= n <= len(steps):
                    if as_assert:
                        steps[n - 1]["assert_locator"] = locator
                    else:
                        steps[n - 1]["locator"] = locator
                        steps[n - 1]["locator_src"] = src
            if not as_assert:
                if src == "live":      live_count += 1
                elif src == "snapshot": snap_count += 1
                elif src == "guess":    guess_count += 1

        def _todo(story, tc, idxs, target, kind):
            for n in idxs:
                todos.append({"s": story.get("id"), "tc": tc.get("title", ""),
                              "n": n, "a": (target or "")[:32], "kind": kind})

        # walk each test case (intent-driven)
        for sp in stories_payload:
            if should_stop() or abort_credit:
                break
            story = sp.get("story", {})
            cb(f"\u25b8 Story {story.get('id')} \u2014 {story.get('title','')}", "story")
            for tc in sp.get("test_cases", []):
                if should_stop() or abort_credit:
                    break
                steps = tc.get("steps", []) or []
                ctype = _classify_case(tc)
                is_neg = (ctype == "negative_login")
                pctx = _infer_page_context(tc, ctype)
                cb(f"  walking '{tc.get('title','')}'  [{ctype} \u00b7 {pctx}-page]  "
                   f"({len(steps)} steps)", "info")

                # STAGE 1 — compile messy steps into typed intents (collapses
                # restated/duplicate steps, routes preconditions away from clicks)
                intents = _credit_guard(compile_test_case, tc, story, cb, ctype) or []
                if not intents:
                    intents = _intents_from_raw_steps(tc)
                n_act = sum(1 for it in intents if it["role"] == "action")
                n_ass = sum(1 for it in intents if it["role"] == "assertion")
                cb(f"    compiled \u2192 {n_act} action(s), {n_ass} assertion(s), "
                   f"{len(intents) - n_act - n_ass} precondition(s)", "dim")

                # start page — login-page cases (incl. negative-login) walk on a
                # FRESH logged-out login page, where the language dropdown exists;
                # app cases walk on the authenticated page (single toggle).
                if pctx == "login":
                    cb(f"    \u21b3 {ctype} \u2014 walking on a fresh logged-out login "
                       f"page (where the language dropdown lives)", "info")
                    if have_creds:
                        try:
                            driver.delete_all_cookies()
                        except Exception:
                            pass
                    try:
                        driver.get(login_url); wait_dom_ready(); _t.sleep(wait_secs)
                    except Exception:
                        pass
                else:
                    try:
                        cb("    loading start page (authenticated app)\u2026", "dim")
                        driver.get(site_url); _t.sleep(wait_secs)
                    except Exception:
                        pass

                last_before = None   # snapshot keys just before the latest action
                for it in intents:
                    if should_stop() or abort_credit:
                        break
                    role = it["role"]; fs = it.get("from_steps") or []

                    if role == "precondition":
                        cb(f"    \u2022 precondition (no UI action): "
                           f"{(it.get('target') or '')[:40]}", "dim")
                        for n in fs:
                            if 1 <= n <= len(steps):
                                steps[n - 1].setdefault("locator", None)
                                steps[n - 1]["locator_src"] = "precondition"
                        continue

                    if role == "assertion":
                        # STAGE 2 (assert) — bind by DOM-diff: prefer elements that
                        # newly appeared/changed since the last action (the menu that
                        # opened, the error that showed). Same mechanism powers
                        # negative-login error capture.
                        _settle(timeout=4)
                        after = snapshot(with_errors=True)
                        new_pool = ([e for e in after if _el_key(e) not in last_before]
                                    if last_before is not None else after)
                        el, src = bind_target(it, new_pool or after)
                        assign(fs, to_locator(el) if el else None, src, as_assert=True)
                        if el:
                            tag = " (new)" if (last_before is not None and
                                               _el_key(el) not in last_before) else ""
                            cb(f"    \u2713 assertion \u2192 {_describe(el)}{tag}", "ok")
                        else:
                            cb("    ? assertion target not found on page", "warn")
                        continue

                    # role == 'action'
                    verb = it.get("verb") or "click"
                    cb(f"    \u2192 {verb}: {(it.get('target') or '')[:40]}", "dim")
                    cur = snapshot(with_errors=is_neg)
                    last_before = set(_el_key(e) for e in cur)
                    el, src = bind_target(it, cur)
                    if abort_credit:
                        break
                    if el is None:
                        assign(fs, None, "guess")
                        cb("      GUESS: no element matched \u2014 // TODO verify locator", "warn")
                        _todo(story, tc, fs, it.get("target"), "guess")
                        continue
                    assign(fs, to_locator(el), src)
                    if src == "snapshot":
                        cb(f"      SNAPSHOT: using {_describe(el)} from an earlier page "
                           f"\u2014 // TODO verify (from snapshot)", "warn")
                        _todo(story, tc, fs, it.get("target"), "snapshot")
                        continue
                    if verb == "navigate":
                        cb("      (navigate) \u2014 already on the target page", "dim")
                        continue
                    # STAGE 3 — interception-proof action
                    empty_ok = (verb == "type" and not (it.get("value") or "").strip()
                                and _wants_empty_field((it.get("target", "") + " " +
                                                        " ".join(it.get("keywords") or []))))
                    _act(el, verb, it.get("value", ""), empty_ok=empty_ok)
                    cb(f"      {verb} {_describe(el)}"
                       f"{' (left empty)' if empty_ok else ''}", "ok")
                    _settle(timeout=4); _t.sleep(0.6)

                # restore the authenticated session after a login-page case so
                # later app cases don't capture locators from the login page
                if pctx == "login" and have_creds and not (should_stop() or abort_credit):
                    cb("    \u21b3 re-establishing login after login-page case\u2026", "dim")
                    try:
                        ok2, reason2 = do_login(fresh=True)
                        cb(f"    \u21b3 re-login {'verified' if ok2 else 'NOT verified'} "
                           f"\u2014 {reason2}", "ok" if ok2 else "warn")
                    except Exception as e:
                        cb(f"    \u21b3 re-login failed: {str(e)[:80]}", "warn")

        if abort_credit:
            cb(f"Stopped automatically \u2014 the AI credit limit was hit {credit_hits} "
               f"times. Top up credits or switch provider, then run again.", "err")

        cb(f"Exploration done \u2014 {live_count} exact, {snap_count} from snapshots, "
           f"{guess_count} still need a guess.", "ok")
        # summary of every step that will carry a // TODO verify locator
        if todos:
            cb(f"{len(todos)} step(s) will need a // TODO verify locator:", "warn")
            for _t2 in todos:
                cb(f"   - story {_t2['s']} / {(_t2['tc'] or '')[:26]} / step {_t2['n']}  "
                   f"[{_t2['kind']}]  {_t2['a'][:32]}", "warn")
        else:
            cb("No TODOs \u2014 every walked step captured a real locator.", "ok")
        # de-dup snapshots
        seen = set(); uniq = []
        for e in all_snapshots:
            key = (e.get("id"), e.get("name"), e.get("css"))
            if key in seen:
                continue
            seen.add(key); uniq.append(e)
        return {"stories_payload": stories_payload, "dom_snapshot": uniq[:300],
                "stats": {"live": live_count, "snapshot": snap_count,
                          "guess": guess_count}}
    finally:
        try:
            driver.quit()
        except Exception:
            pass


def _driver_factory(pkg):
    return f"""package {pkg}.core;

import io.github.bonigarcia.wdm.WebDriverManager;
import org.openqa.selenium.WebDriver;
import org.openqa.selenium.chrome.ChromeDriver;
import org.openqa.selenium.chrome.ChromeOptions;

public final class DriverFactory {{
    private DriverFactory() {{}}

    public static WebDriver create() {{
        WebDriverManager.chromedriver().setup();
        ChromeOptions options = new ChromeOptions();
        // options.addArguments("--headless=new");
        options.addArguments("--start-maximized");
        return new ChromeDriver(options);
    }}
}}
"""

# ── Incremental generation: survive prior stories, append only new test cases ──
def _manifest_path(project_dir):
    return os.path.join(project_dir, ".qastudio", "manifest.json")

def load_manifest(project_dir):
    """Read the per-project manifest that records which stories/test-cases have
    already been generated. Missing/corrupt → empty manifest."""
    try:
        with open(_manifest_path(project_dir), "r", encoding="utf-8") as f:
            m = json.load(f)
    except Exception:
        m = {}
    if not isinstance(m, dict):
        m = {}
    m.setdefault("stories", {})
    return m

def save_manifest(project_dir, m):
    try:
        os.makedirs(os.path.dirname(_manifest_path(project_dir)), exist_ok=True)
        with open(_manifest_path(project_dir), "w", encoding="utf-8") as f:
            json.dump(m, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _tc_key(tc):
    """Stable identity for a test case: Azure id if present, else a title slug."""
    k = tc.get("id")
    if k is not None and str(k).strip():
        return str(k)
    return "t_" + _safe_class_name(tc.get("title", ""), "TC")

def _method_name(title):
    n = _safe_class_name(title, "test")
    return (n[:1].lower() + n[1:]) if n else "test"

def classify_selection(project_dir, stories_payload):
    """Compare the current selection against the on-disk manifest. Returns
    (new, grew, done, new_tcs):
      new  = story ids never generated here
      grew = story ids already generated that have NEW test cases to add
      done = story ids already generated with nothing new
      new_tcs = {story_id: [fresh test-case dicts]} for the grew set
    All ids are strings."""
    m = load_manifest(project_dir)
    new, grew, done, new_tcs = [], [], [], {}
    for sp in stories_payload:
        sid = str(sp.get("story", {}).get("id"))
        rec = m["stories"].get(sid)
        tcs = sp.get("test_cases", []) or []
        if not rec:
            new.append(sid); continue
        have = set(rec.get("test_cases", {}).keys())
        fresh = [tc for tc in tcs if _tc_key(tc) not in have]
        if fresh:
            grew.append(sid); new_tcs[sid] = fresh
        else:
            done.append(sid)
    return new, grew, done, new_tcs

def _validate_remote_url(url):
    """Validate a git remote URL before we hand it to git, so a wrong URL gives a
    clear message (and a suggested fix) instead of a cryptic git failure.
    Returns (ok: bool, message_or_suggestion: str)."""
    import urllib.parse
    u = (url or "").strip()
    if not u:
        return (False, "Repository URL is empty. Paste your repo URL, e.g. "
                       "https://github.com/owner/repo.git")
    # SSH remotes (git@host:owner/repo.git or ssh://…) — accept as-is.
    if u.startswith("git@") or u.startswith("ssh://"):
        return (True, u)
    if not (u.startswith("https://") or u.startswith("http://")):
        return (False, "Repository URL must start with https:// (or be an SSH "
                       f"git@ URL). Got: {u[:70]}")
    try:
        p = urllib.parse.urlparse(u)
    except Exception:
        return (False, f"Repository URL could not be parsed: {u[:70]}")
    if not p.netloc or "." not in p.netloc:
        return (False, f"Repository URL has no valid host: {u[:70]}")
    segs = [s for s in p.path.split("/") if s]
    if len(segs) < 2:
        return (False, "Repository URL is missing the owner/repo, e.g. "
                       f"https://{p.netloc}/<owner>/<repo>.git")
    owner, repo = segs[0], segs[1]
    clean_repo = repo.split(".git")[0] if ".git" in repo else repo
    suggestion = f"https://{p.netloc}/{owner}/{clean_repo}.git"
    host = p.netloc.lower()
    github_like = any(h in host for h in ("github.com", "gitlab.com", "bitbucket.org"))
    # ".git" appearing anywhere but the end (e.g. ".gitm", ".git/extra") is malformed
    path_no_slash = p.path.rstrip("/")
    bad_git = (".git" in path_no_slash and not path_no_slash.endswith(".git"))
    # GitHub/GitLab repos are exactly owner/repo; extra path segments are wrong
    extra_segs = github_like and len(segs) > 2
    if bad_git or extra_segs:
        return (False, "Repository URL looks malformed (check for a typo like "
                       "'.gitm' or extra text after the repo name). Try: " + suggestion)
    return (True, u)


def _parse_github_owner_repo(remote_url):
    """(owner, repo) for a github.com HTTPS remote, else (None, None) — used to
    scope auto-create to GitHub only (Azure DevOps/GitLab repos still have to be
    created by the user; their creation APIs differ and aren't wired up here)."""
    import urllib.parse
    try:
        u = urllib.parse.urlparse((remote_url or "").strip())
        if "github.com" not in (u.netloc or "").lower():
            return None, None
        parts = [x for x in (u.path or "").split("/") if x]
        if len(parts) < 2:
            return None, None
        owner, repo = parts[0], parts[1]
        if repo.endswith(".git"):
            repo = repo[:-4]
        return owner, repo
    except Exception:
        return None, None


def _github_ensure_repo(owner, repo, token, cb):
    """Create the GitHub repo if it doesn't exist yet (idempotent — a no-op, not
    an error, when it's already there). Returns (ok, msg)."""
    headers = {"Authorization": f"Bearer {token}", "Accept": "application/vnd.github+json"}
    try:
        r = requests.get(f"https://api.github.com/repos/{owner}/{repo}",
                         headers=headers, timeout=15)
        if r.status_code == 200:
            return True, "exists"
        if r.status_code != 404:
            return False, f"GitHub API error checking the repo ({r.status_code}): {r.text[:160]}"
    except Exception as ex:
        return False, f"Couldn't reach the GitHub API to check the repo: {ex}"

    cb(f"Repo {owner}/{repo} doesn't exist on GitHub yet — creating it…", "dim")
    my_login = None
    try:
        me = requests.get("https://api.github.com/user", headers=headers, timeout=15)
        if me.status_code == 200:
            my_login = (me.json() or {}).get("login")
    except Exception:
        pass
    create_url = ("https://api.github.com/user/repos"
                 if my_login and my_login.lower() == owner.lower()
                 else f"https://api.github.com/orgs/{owner}/repos")
    try:
        r = requests.post(create_url, headers=headers, timeout=20,
                          json={"name": repo, "private": True,
                                "description": "Automation tests generated by QA Studio"})
        if r.status_code == 201:
            cb(f"Created {owner}/{repo} on GitHub.", "ok")
            return True, "created"
        return False, f"Couldn't create the GitHub repo ({r.status_code}): {r.text[:200]}"
    except Exception as ex:
        return False, f"Couldn't create the GitHub repo: {ex}"


def _detect_project_framework(out_dir):
    """Best-effort detect which framework an on-disk generated project targets,
    by its marker file. Needed because 'push a forgotten run' (an existing folder
    picked via Browse) can be pushed WITHOUT ever calling
    generate_and_push_selfhealing in this session, so we can't just trust
    self._auto_target — the folder itself is the source of truth."""
    if os.path.exists(os.path.join(out_dir, "playwright.config.js")):
        return "playwright"
    if os.path.exists(os.path.join(out_dir, "cypress.config.js")):
        return "cypress"
    if os.path.exists(os.path.join(out_dir, "pom.xml")):
        return "selenium"
    return None


def sync_project_readme(out_dir, cb=None):
    """Rewrite README.md (+ its matching .gitignore) to the CURRENT template for
    whatever framework is actually on disk in `out_dir` — every time, unconditionally
    (write-if-changed). This is what makes the fix real: a folder generated by an
    older QA Studio build, or one where the README was hand-patched in an IDE and
    drifted again, gets corrected right before every push instead of only on a
    fresh 'Generate automation scripts' run. README/.gitignore are framework-fixed
    scaffolding (unlike test specs/locators/config), so always resyncing them can
    never clobber generated tests or hand-written config.
    Returns (changed: bool, framework: str | None). framework is None when the
    folder doesn't look like a QA Studio project yet (nothing to detect off of).
    """
    cb = cb or (lambda *a, **k: None)
    fw = _detect_project_framework(out_dir)
    if fw is None:
        return False, None
    if fw == "selenium":
        readme, gi = _sh_readme("", ""), _sh_gitignore()
    else:
        import automation_targets as _AT
        readme = _AT._PW_README if fw == "playwright" else _AT._CY_README
        gi = _AT._PW_GITIGNORE if fw == "playwright" else _AT._CY_GITIGNORE
    changed = False

    def _sync_one(name, content):
        nonlocal changed
        path = os.path.join(out_dir, name)
        try:
            cur = open(path, "r", encoding="utf-8").read() if os.path.exists(path) else None
        except Exception:
            cur = None
        if cur != content:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            changed = True

    _sync_one("README.md", readme)
    _sync_one(".gitignore", gi)
    if changed:
        cb(f"Synced README.md/.gitignore to the current {fw} template "
           f"(was out of date).", "warn")
    return changed, fw


def push_to_git(repo_dir, remote_url, token, branch="main", message="Add QA Studio automation tests", cb=None, force=False):
    """Init/commit/push the generated project to a Git remote using the git CLI.
    `token` is passed to the push subprocess via a one-shot GIT_CONFIG_* env
    override (http.extraheader) — see the SECURITY comment below. It is never
    embedded into the remote URL and never written to .git/config.
    Returns (ok, message).
    """
    cb = cb or (lambda *a, **k: None)
    import subprocess

    # --- pre-flight: the folder must exist and contain the generated project ---
    if not repo_dir or not os.path.isdir(repo_dir):
        return False, f"Project folder not found: {repo_dir or '(blank)'}. Generate scripts first."
    try:
        entries = [e for e in os.listdir(repo_dir) if e != ".git"]
    except Exception:
        entries = []
    if not entries:
        return False, ("Project folder is empty — nothing to push. Generate the "
                       "automation scripts to this folder first.")

    # --- pre-flight: validate the remote URL with a friendly message ---
    ok_url, url_msg = _validate_remote_url(remote_url)
    if not ok_url:
        return False, url_msg
    remote_url = url_msg if url_msg.startswith(("http", "git@", "ssh://")) else remote_url

    def run(args, **kw):
        return subprocess.run(args, cwd=repo_dir, capture_output=True, text=True, **kw)

    # check git is available
    try:
        v = subprocess.run(["git", "--version"], capture_output=True, text=True)
        if v.returncode != 0:
            return False, "Git is not installed on this machine."
    except Exception:
        return False, "Git is not installed or not on PATH."

    # SECURITY: auth is supplied per-invocation via git's native
    # GIT_CONFIG_COUNT/KEY/VALUE env-var mechanism on the push command itself
    # (see _push_env below) — the token is never embedded in the remote URL,
    # never written to .git/config, and never appears in the subprocess's
    # argv. Previously this built `https://<token>@host/...`, ran `git remote
    # add` with it (so the PAT sat in .git/config in plaintext), and relied on
    # a post-push `git remote set-url` to scrub it — a cleanup step that would
    # never run if the process crashed or was killed mid-push, leaving the
    # token on disk indefinitely. An interim fix passed it as a `-c
    # http.extraheader=...` argument instead, which fixed the disk-persistence
    # problem but still left the (base64-encoded, so trivially reversible)
    # token visible in the process's command line to any other process on the
    # same machine for the life of the call (e.g. /proc/<pid>/cmdline on
    # Linux, or Task Manager's "Command line" column on Windows). Env vars
    # aren't surfaced in either of those views, so this closes that too.
    def _auth_header(tok):
        import base64
        return "AUTHORIZATION: basic " + base64.b64encode(
            f"x-access-token:{tok}".encode()).decode()

    # Create the remote repo, but ONLY if this folder has never been pushed
    # before (no 'git_pushed' marker in its manifest) — an already-pushed folder
    # is assumed to already have its repo, so we don't hit the GitHub API on
    # every single push. GitHub only: Azure DevOps/GitLab repo creation isn't
    # wired up, so those remotes fall through unchanged (existing behaviour —
    # the repo must already exist, same as before this fix).
    _m = load_manifest(repo_dir)
    if not _m.get("git_pushed"):
        _owner, _gh_repo = _parse_github_owner_repo(remote_url)
        if _owner and _gh_repo:
            _ok_repo, _repo_msg = _github_ensure_repo(_owner, _gh_repo, token, cb)
            if not _ok_repo:
                return False, f"Couldn't prepare the GitHub repo: {_repo_msg}"

    # Resync README.md/.gitignore to the current framework template BEFORE staging,
    # so every push — including 'push a forgotten run' folders picked via Browse
    # that this session never (re)generated — always ships a correct, up-to-date
    # README instead of whatever was last written (possibly by an older QA Studio
    # build, or hand-patched and since drifted).
    try:
        sync_project_readme(repo_dir, cb=cb)
    except Exception:
        pass

    cb("Initializing git repository…", "dim")
    if not os.path.isdir(os.path.join(repo_dir, ".git")):
        run(["git", "init"])
    run(["git", "checkout", "-B", branch])
    # NOTE: .gitignore is now written by sync_project_readme() above, using the
    # correct per-framework template (Maven/.idea vs node_modules/ for the JS
    # targets). A hardcoded Maven-only .gitignore used to be force-written here
    # unconditionally, silently clobbering the Playwright/Cypress .gitignore
    # (and undoing the sync above) on every single push.
    run(["git", "add", "-A"])
    # ensure identity exists (use a neutral default if unset)
    if run(["git", "config", "user.email"]).stdout.strip() == "":
        run(["git", "config", "user.email", "qastudio@local"])
        run(["git", "config", "user.name", "QA Studio"])
    c = run(["git", "commit", "-m", message])
    if c.returncode != 0 and "nothing to commit" not in (c.stdout + c.stderr).lower():
        cb(c.stdout + c.stderr, "warn")

    # set remote — always the plain, token-less URL; see _auth_header above
    run(["git", "remote", "remove", "origin"])
    run(["git", "remote", "add", "origin", remote_url])

    cb(f"Pushing to {branch}…", "dim")
    _push = ["git", "push", "-u", "origin", branch]
    if force:
        _push.append("--force")
    _push_env = None
    if remote_url.startswith("https://") and token:
        _push_env = {**os.environ,
                    "GIT_CONFIG_COUNT": "1",
                    "GIT_CONFIG_KEY_0": "http.extraheader",
                    "GIT_CONFIG_VALUE_0": _auth_header(token)}
    p = run(_push, env=_push_env)
    out = (p.stdout + p.stderr)
    # scrub token from any echoed output (defense in depth — the header is
    # base64, not the raw token, but keep this in case git ever echoes argv)
    if token:
        out = out.replace(token, "***")
    if p.returncode == 0:
        cb("Push complete.", "ok")
        try:
            _m2 = load_manifest(repo_dir)
            _m2["git_pushed"] = {"remote": remote_url, "branch": branch,
                                 "at": time.strftime("%Y-%m-%dT%H:%M:%S")}
            save_manifest(repo_dir, _m2)
        except Exception:
            pass
        return True, "Pushed successfully."
    # Recognize a plain non-fast-forward rejection and translate the raw git
    # hint dump into one actionable sentence instead of surfacing plumbing
    # text (git's "[rejected] ... fetch first" + multi-line hints) straight
    # into the Activity log for a non-technical QA user.
    low = out.lower()
    if not force and ("[rejected]" in out or "non-fast-forward" in low) and \
            ("fetch first" in low or "non-fast-forward" in low):
        return False, ("rejected: the remote already has commits this folder doesn't "
                       "(someone/something else pushed since). Pull those changes in "
                       "first, or push with Force to overwrite the remote with what's "
                       "in this folder.")
    return False, out.strip()[:400] or "git push failed."


# ═══════════════════════════════════════════════════════════════════════════════
#  AUTO-UPDATE — compare a VERSION file in the GitHub repo to the local one
# ═══════════════════════════════════════════════════════════════════════════════
GITHUB_OWNER  = "AhmedSayedRepo"
GITHUB_REPO   = "qa-studio"
GITHUB_BRANCH = "main"

def _github_token():
    """Optional token for private-repo update checks. Read from the env var
    QA_STUDIO_GH_TOKEN or GITHUB_TOKEN, or a 'gh_token.txt' file next to the app.
    Without it, update checks only work for a PUBLIC repo."""
    for var in ("QA_STUDIO_GH_TOKEN", "GITHUB_TOKEN"):
        t = (os.environ.get(var) or "").strip()
        if t:
            return t
    for base in (_exe_dir(), _app_dir()):
        try:
            with open(os.path.join(base, "gh_token.txt"), "r", encoding="utf-8") as f:
                t = f.read().strip()
                if t:
                    return t
        except Exception:
            pass
    return ""

def _app_dir():
    return os.path.dirname(os.path.abspath(__file__))

def _exe_dir():
    """Folder of the running program: the .exe's folder when frozen, else source."""
    import sys
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return _app_dir()

def _resource_dir():
    """Where bundled read-only files (e.g. VERSION) live. For a PyInstaller/flet
    onefile build that's sys._MEIPASS; otherwise the source folder."""
    import sys
    if getattr(sys, "frozen", False):
        return getattr(sys, "_MEIPASS", _exe_dir())
    return _app_dir()

def _clean_ver(s):
    """Tolerate a VERSION file written as UTF-16 (e.g. PowerShell `echo x > VERSION`),
    which adds a BOM and a null byte between characters. Strip BOM, nulls, and any
    non [digit/dot] noise so '\\ufeff1\\x00.\\x008\\x00.\\x004' still reads as '1.8.4'."""
    s = (s or "").replace("\x00", "").lstrip("\ufeff\ufffe").strip()
    m = re.search(r"\d+(?:\.\d+){0,3}", s)
    return m.group(0) if m else s

def local_version():
    """Read the local VERSION file (next to this module). Returns str or '0.0.0'."""
    try:
        with open(os.path.join(_resource_dir(), "VERSION"), "rb") as f:
            raw = f.read().decode("utf-8-sig", "ignore")
        return _clean_ver(raw) or "0.0.0"
    except Exception:
        return "0.0.0"


# A source/exe update can finish after the user closes the window.  Keep this
# tiny, non-sensitive hand-off outside the install folder, otherwise the only
# completion state lives in the old process and the next launch has no way to
# show the release notes or confirm success.
def _pending_update_notice_path():
    try:
        import platform_caps as _pc_update
        d = _pc_update.app_data_dir()
        os.makedirs(d, exist_ok=True)
        return os.path.join(d, "pending_update_notice.json")
    except Exception:
        return ""


def save_pending_update_notice(version):
    """Persist one post-update notice for the next application launch."""
    path = _pending_update_notice_path()
    if not path:
        return False
    try:
        data = {"version": _clean_ver(str(version or "")), "created_at": time.time()}
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
        os.replace(tmp, path)
        return True
    except Exception:
        try:
            if os.path.exists(path + ".tmp"):
                os.remove(path + ".tmp")
        except Exception:
            pass
        return False


def get_pending_update_notice():
    """Return the pending update notice without consuming it."""
    path = _pending_update_notice_path()
    if not path:
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        version = _clean_ver(str((data or {}).get("version") or ""))
        return {"version": version} if version else {}
    except Exception:
        return {}


def clear_pending_update_notice():
    """Mark the one-time post-update confirmation as seen."""
    path = _pending_update_notice_path()
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception:
        pass

def _parse_ver(v):
    """'1.2.0' -> (1,2,0); tolerant of extra/missing parts."""
    parts = re.findall(r"\d+", str(v))
    nums = [int(p) for p in parts[:4]]
    while len(nums) < 3:
        nums.append(0)
    return tuple(nums)

def _ver_newer(remote, local):
    return _parse_ver(remote) > _parse_ver(local)

# The rolling Android release tag + the version-marker asset that build-apk.yml
# writes AFTER a successful build. See check_mobile_update.
MOBILE_RELEASE_TAG = "android-apk"
MOBILE_VERSION_ASSET = "mobile-version.txt"

def check_mobile_update(timeout=6):
    """MOBILE update check — reads the version of the APK that is ACTUALLY
    PUBLISHED, not the repo's VERSION on main.

    THE RACE THIS FIXES: release.bat bumps VERSION + cuts the vX.Y.Z release
    IMMEDIATELY, but the APK takes ~10 min to build and attach. check_for_update
    (which reads main's VERSION) would tell mobile users an update exists during
    that whole window — they'd tap Download and get a 404 or the OLD apk, since
    the new one isn't attached yet. So mobile must NOT key off the repo VERSION.

    build-apk.yml publishes the APK to the rolling `android-apk` release and, in
    the SAME step, uploads a `mobile-version.txt` asset containing the version
    it just built. Because the marker and the .apk are written together at the
    END of a successful build, the advertised version can never run ahead of an
    available APK. This reads that marker. Same dict shape as check_for_update;
    fail-soft (update=False) on any network/parse error.
    """
    local = local_version()
    import time as _t
    url = (f"https://github.com/{GITHUB_OWNER}/{GITHUB_REPO}/releases/download/"
           f"{MOBILE_RELEASE_TAG}/{MOBILE_VERSION_ASSET}?_={int(_t.time())}")
    hdr = {"Cache-Control": "no-cache"}
    tok = _github_token()
    if tok:
        hdr["Authorization"] = f"Bearer {tok}"
    try:
        r = requests.get(url, headers=hdr, timeout=timeout, allow_redirects=True)
        if r.status_code != 200:
            return {"update": False, "local": local, "remote": None, "error": None}
        remote = _clean_ver(r.text or "")
        if not remote:
            return {"update": False, "local": local, "remote": None, "error": None}
        return {"update": _ver_newer(remote, local), "local": local,
                "remote": remote, "error": None}
    except Exception:
        return {"update": False, "local": local, "remote": None, "error": None}

def check_for_update(timeout=6):
    """Fetch the repo's VERSION file and compare to local.
    Returns dict: {"update": bool, "local": str, "remote": str|None, "error": str|None}.
    Network failures are swallowed (update=False) so startup is never blocked.

    Sources tried in order:
      1) GitHub API contents endpoint (works for private repos when a token is set)
      2) Cache-busted raw.githubusercontent.com URL (public repos)
    A token (see _github_token) is sent when available so PRIVATE repos work.
    """
    import time as _t, base64 as _b64
    local = local_version()
    bust = int(_t.time())
    token = _github_token()

    def _auth_headers(extra):
        h = dict(extra)
        if token:
            h["Authorization"] = f"Bearer {token}"
            h["X-GitHub-Api-Version"] = "2022-11-28"
        return h

    def _via_api():
        url = (f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/"
               f"contents/VERSION?ref={GITHUB_BRANCH}")
        r = requests.get(url, timeout=timeout, headers=_auth_headers({
            "Accept": "application/vnd.github.raw+json",
            "Cache-Control": "no-cache",
        }))
        if r.status_code == 404:
            raise RuntimeError("API 404 — repo or VERSION file not found "
                               "(private repo needs a token; or commit a VERSION file).")
        if r.status_code in (401, 403):
            raise RuntimeError(f"API {r.status_code} — auth/rate-limit; set a valid token.")
        if r.status_code != 200:
            raise RuntimeError(f"API HTTP {r.status_code}")
        txt = r.text or ""
        if txt.lstrip().startswith("{"):
            import json as _json
            data = _json.loads(txt)
            txt = _b64.b64decode(data.get("content", "")).decode("utf-8", "ignore")
        return _clean_ver(txt)

    def _via_raw():
        url = (f"https://raw.githubusercontent.com/{GITHUB_OWNER}/{GITHUB_REPO}/"
               f"{GITHUB_BRANCH}/VERSION?cb={bust}")
        r = requests.get(url, timeout=timeout, headers=_auth_headers(
            {"Cache-Control": "no-cache", "Pragma": "no-cache"}))
        if r.status_code == 404:
            raise RuntimeError("raw 404 — VERSION file missing at repo root (or private repo).")
        if r.status_code != 200:
            raise RuntimeError(f"raw HTTP {r.status_code}")
        return _clean_ver(r.text or "")

    remote = None
    err = None
    for fetch in (_via_api, _via_raw):
        try:
            remote = fetch()
            if remote:
                break
        except Exception as e:
            err = str(e)[:160]
            continue
    if not remote:
        return {"update": False, "local": local, "remote": None, "error": err}
    return {"update": _ver_newer(remote, local), "local": local,
            "remote": remote, "error": None}

_RESTART_BAT = r'''@echo off
set "PID=__PID__"
:wait
tasklist /FI "PID eq %PID%" 2>nul | find "%PID%" >nul && (ping -n 2 127.0.0.1 >nul & goto wait)
ping -n 2 127.0.0.1 >nul
__LAUNCH__
del "%~f0" >nul 2>&1
'''

def schedule_restart():
    """Arm a detached helper that waits for THIS process to exit, then relaunches
    QA Studio (the exe when frozen, else `pythonw main.py`). Returns True if armed."""
    import sys, tempfile, subprocess
    try:
        if getattr(sys, "frozen", False):
            launch = f'start "" "{os.path.abspath(sys.executable)}"'
        else:
            pyw = os.path.abspath(sys.executable)
            mainpy = os.path.join(_app_dir(), "main.py")
            launch = f'start "" /d "{_app_dir()}" "{pyw}" "{mainpy}"'
        bat = os.path.join(tempfile.gettempdir(), "qastudio_restart.bat")
        script = _RESTART_BAT.replace("__PID__", str(os.getpid())).replace("__LAUNCH__", launch)
        with open(bat, "w", encoding="ascii", errors="ignore", newline="\r\n") as f:
            f.write(script)
        DETACHED, NEW_GROUP, NO_WINDOW = 0x00000008, 0x00000200, 0x08000000
        subprocess.Popen(["cmd", "/c", bat],
                         creationflags=DETACHED | NEW_GROUP | NO_WINDOW, close_fds=True)
        return True
    except Exception:
        return False

_SWAP_BAT = r'''@echo off
set "PID=__PID__"
set "NEW=__NEW__"
set "CUR=__CUR__"
:wait
tasklist /FI "PID eq %PID%" 2>nul | find "%PID%" >nul && (ping -n 2 127.0.0.1 >nul & goto wait)
set "n=0"
:try
ping -n 2 127.0.0.1 >nul
copy /y "%NEW%" "%CUR%" >nul && goto done
set /a n+=1
if %n% lss 20 goto try
:done
del "%NEW%" >nul 2>&1
start "" "%CUR%"
del "%~f0" >nul 2>&1
'''

def _latest_release(timeout=6):
    """Return (tag, (asset_name, asset_url), sums) for the newest GitHub release.
    `asset` is None when the release has no .exe attached; `sums` is the
    (name, url) of a published SHA-256 checksum file (SHA256SUMS / *.sha256), or
    None when the release doesn't publish one."""
    headers = {"Accept": "application/vnd.github+json"}
    token = _github_token()
    if token:
        headers["Authorization"] = f"Bearer {token}"
        headers["X-GitHub-Api-Version"] = "2022-11-28"
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/releases/latest"
    r = requests.get(url, timeout=timeout, headers=headers)
    r.raise_for_status()
    data = r.json()
    tag = (data.get("tag_name") or "").lstrip("vV")
    asset = None
    sums = None
    for a in data.get("assets", []):
        nm = str(a.get("name", "")).lower()
        if asset is None and nm.endswith(".exe"):
            asset = (a["name"], a["browser_download_url"])
        if sums is None and (nm in ("sha256sums", "sha256sums.txt", "checksums.txt")
                             or nm.endswith(".sha256")):
            sums = (a["name"], a["browser_download_url"])
    return tag, asset, sums


def _sha256_file(path):
    import hashlib
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _verify_download(path, name, sums, headers, cb):
    """Verify a downloaded artifact against the release's published SHA-256.
    SECURITY: fails CLOSED — a release that doesn't publish a checksum file is
    rejected rather than installed with a warning. This is the self-update
    path: it downloads and re-launches an executable that replaces the running
    app, so silently proceeding without any integrity check would let anyone
    who can write to the GitHub release (compromised token, supply-chain
    compromise) achieve code execution on every auto-updating install just by
    omitting the checksum file. If you start attaching a built .exe to
    releases again, publish a SHA256SUMS (or *.sha256) file alongside it —
    see _latest_release() above for the accepted file names."""
    if not sums:
        return False, ("This release doesn't publish a checksum file, so the "
                       "update can't be verified. Update aborted for your "
                       "safety — see the release page.")
    try:
        sr = requests.get(sums[1], timeout=30, headers=headers)
        sr.raise_for_status()
        text = sr.text
    except Exception as e:
        return False, f"Couldn't fetch the release checksum file: {str(e)[:140]}"
    want = None
    for line in text.splitlines():
        parts = line.replace("*", " ").split()
        if len(parts) >= 2 and parts[1].lower().endswith(name.lower()):
            want = parts[0].lower()
            break
        if len(parts) == 1 and len(parts[0]) == 64:   # bare single-asset hash
            want = parts[0].lower()
    if not want:
        return False, "The release checksum file has no entry for this download."
    if _sha256_file(path).lower() != want:
        return False, ("Checksum mismatch — the download may be corrupted or "
                       "tampered with. Update aborted.")
    cb("Checksum verified.", "ok")
    return True, ""

def _apply_update_exe(cb):
    """Frozen-build updater: download the new .exe, then hand off to a detached
    .bat that waits for THIS process to exit, swaps the file, and relaunches."""
    import sys, time, tempfile, subprocess
    cb = cb or (lambda *a, **k: None)
    try:
        tag, asset, sums = _latest_release()
    except Exception as e:
        return (False, f"Couldn't reach GitHub releases: {str(e)[:160]}")
    if not asset:
        return (False, "The latest release has no .exe attached. Upload the built "
                       "exe as a release asset so the app can self-update.")
    _name, dl_url = asset
    cur = os.path.abspath(sys.executable)
    new = os.path.join(tempfile.gettempdir(), f"qastudio_new_{int(time.time())}.exe")
    cb("Downloading the new version…", "dim")
    try:
        headers = {}
        token = _github_token()
        if token:
            headers["Authorization"] = f"Bearer {token}"
        with requests.get(dl_url, stream=True, timeout=120, headers=headers) as r:
            r.raise_for_status()
            with open(new, "wb") as f:
                for chunk in r.iter_content(65536):
                    if chunk:
                        f.write(chunk)
    except Exception as e:
        return (False, f"Download failed: {str(e)[:160]}")
    # SECURITY: verify the downloaded binary against the release SHA-256 before we
    # ever swap/execute it. Fails CLOSED — aborts the update on a checksum
    # mismatch AND when the release publishes no checksum at all (see
    # _verify_download's docstring for why: this replaces an executable that
    # gets re-launched, so an unverified download is a code-execution risk).
    ok_v, vmsg = _verify_download(new, _name, sums, headers, cb)
    if not ok_v:
        try: os.remove(new)
        except Exception: pass
        return (False, vmsg)
    bat = os.path.join(tempfile.gettempdir(), "qastudio_update.bat")
    script = (_SWAP_BAT.replace("__PID__", str(os.getpid()))
                       .replace("__NEW__", new).replace("__CUR__", cur))
    try:
        with open(bat, "w", encoding="ascii", errors="ignore", newline="\r\n") as f:
            f.write(script)
    except Exception as e:
        return (False, f"Couldn't write the updater helper: {str(e)[:160]}")
    DETACHED, NEW_GROUP, NO_WINDOW = 0x00000008, 0x00000200, 0x08000000
    try:
        subprocess.Popen(["cmd", "/c", bat],
                         creationflags=DETACHED | NEW_GROUP | NO_WINDOW,
                         close_fds=True)
    except Exception as e:
        return (False, f"Couldn't start the updater: {str(e)[:160]}")
    cb("Update ready.", "ok")
    return (True, "Update downloaded. Close QA Studio and it will reopen on the "
                  "new version automatically.")

def _latest_zipball(timeout=6):
    """Return (tag_name, zipball_url) for the newest release (source archive)."""
    headers = {"Accept": "application/vnd.github+json"}
    token = _github_token()
    if token:
        headers["Authorization"] = f"Bearer {token}"
        headers["X-GitHub-Api-Version"] = "2022-11-28"
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/releases/latest"
    r = requests.get(url, timeout=timeout, headers=headers)
    r.raise_for_status()
    data = r.json()
    return (data.get("tag_name") or ""), (data.get("zipball_url") or "")


def _resolve_branch_sha(headers, timeout=15):
    """Resolve GITHUB_BRANCH to its current commit SHA via the (read-only,
    cheap) Git API, so the update download is pinned to one immutable commit
    instead of the mutable branch ref (a ref that could resolve to different
    content between two requests, or be force-pushed mid-update)."""
    url = (f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
           f"/commits/{GITHUB_BRANCH}")
    r = requests.get(url, timeout=timeout, headers=headers)
    r.raise_for_status()
    sha = r.json().get("sha")
    if not sha:
        raise RuntimeError("GitHub API returned no commit sha for the branch.")
    return sha


def _validate_update_tree(src_root):
    """Sanity-check that an extracted update actually looks like a QA Studio
    source tree before we copy it over the live, running install. This is not
    cryptographic integrity (see _apply_update_zip's docstring for why a
    branch zipball can't get the same checksum guarantee the .exe path has)
    — it's a fails-closed guard against a truncated/empty/wrong-repo archive
    silently clobbering the app with garbage."""
    required = ("main.py", "engine.py", "VERSION")
    missing = [f for f in required if not os.path.isfile(os.path.join(src_root, f))]
    return (not missing), missing


def _apply_update_zip(cb):
    """Source (non-git) updater: download the latest release's source zip and copy
    it over the app folder in place, then reinstall deps. Used by ZIP/.bat
    installs that aren't git clones and aren't frozen exes.

    SECURITY: the .exe update path (_apply_update_exe/_verify_download) checks
    the download against a maintainer-published SHA-256 checksum file, which
    doesn't exist for an on-the-fly branch zipball — GitHub doesn't publish
    one for those, so there's nothing to compare against, and that's a real,
    structural gap vs. the exe path that a client-side check alone can't
    close. What this DOES add, to raise the bar as far as it reasonably can
    without a published checksum: (1) pin the download to one immutable
    commit SHA resolved just before downloading, instead of trusting a
    mutable branch ref that could change between requests; (2) fail closed
    if that resolution fails, rather than silently falling back to the old
    unpinned-branch behavior; (3) validate the extracted tree actually looks
    like this app's source before overwriting the live install; (4) log the
    exact commit SHA + downloaded file's own SHA-256 locally (diag_log) for
    every applied update, so there's always a durable, after-the-fact record
    of precisely what was installed and from where."""
    import sys, tempfile, zipfile, shutil, subprocess
    cb = cb or (lambda *a, **k: None)
    headers = {"Accept": "application/vnd.github+json"}
    token = _github_token()
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        sha = _resolve_branch_sha(headers)
    except Exception as e:
        return (False, f"Couldn't resolve the update branch to a commit: "
                       f"{str(e)[:140]}. Update aborted for your safety.")
    # Pull the latest code from that PINNED COMMIT (not a mutable branch name) —
    # still the same source check_for_update reads its VERSION from, and the
    # same archive install.bat uses; pinning just removes the "which content
    # did we actually get" ambiguity a plain branch-name URL has.
    zb = (f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
          f"/zipball/{sha}")
    tmp = tempfile.mkdtemp(prefix="qastudio_up_")
    zpath = os.path.join(tmp, "src.zip")
    cb("Downloading the latest version…", "dim")
    try:
        with requests.get(zb, headers=headers, stream=True, timeout=180,
                          allow_redirects=True) as r:
            r.raise_for_status()
            with open(zpath, "wb") as f:
                for chunk in r.iter_content(65536):
                    if chunk:
                        f.write(chunk)
    except Exception as e:
        shutil.rmtree(tmp, ignore_errors=True)
        return (False, f"Download failed: {str(e)[:160]}")
    try:
        zip_sha256 = _sha256_file(zpath)
    except Exception:
        zip_sha256 = "?"
    cb("Installing…", "dim")
    try:
        with zipfile.ZipFile(zpath) as z:
            # SECURITY (zip-slip): don't trust member names. Even though the zip
            # is GitHub's zipball pinned to an immutable commit SHA over TLS, a
            # single malicious/crafted entry ("../…" or an absolute path) would
            # let extractall() write OUTSIDE tmp. Verify every member resolves
            # inside tmp before extracting, and fail closed if any doesn't.
            _tmp_root = os.path.realpath(tmp)
            for _m in z.namelist():
                _dest = os.path.realpath(os.path.join(tmp, _m))
                if _dest != _tmp_root and not _dest.startswith(_tmp_root + os.sep):
                    raise ValueError(f"unsafe path in update archive: {_m!r}")
            z.extractall(tmp)
    except Exception as e:
        shutil.rmtree(tmp, ignore_errors=True)
        return (False, f"Couldn't unpack the update: {str(e)[:160]}")
    roots = [os.path.join(tmp, d) for d in os.listdir(tmp)
             if os.path.isdir(os.path.join(tmp, d))]
    if not roots:
        shutil.rmtree(tmp, ignore_errors=True)
        return (False, "Update archive was empty.")
    src_root, dst = roots[0], _app_dir()
    ok_tree, missing = _validate_update_tree(src_root)
    if not ok_tree:
        shutil.rmtree(tmp, ignore_errors=True)
        return (False, "The downloaded update doesn't look like a QA Studio "
                       f"source tree (missing {', '.join(missing)}). Update "
                       "aborted for your safety.")
    try:
        import diag_log
        diag_log.log_warn("engine._apply_update_zip",
                           f"applying commit {sha} — zip sha256={zip_sha256}")
    except Exception:
        pass
    try:
        for name in os.listdir(src_root):
            s = os.path.join(src_root, name)
            d = os.path.join(dst, name)
            if os.path.isdir(s):
                shutil.copytree(s, d, dirs_exist_ok=True)
            else:
                shutil.copy2(s, d)
    except Exception as e:
        shutil.rmtree(tmp, ignore_errors=True)
        return (False, f"Couldn't write update files: {str(e)[:160]}. "
                       f"Close the app and try again.")
    shutil.rmtree(tmp, ignore_errors=True)
    # Drop stale bytecode so Python recompiles from the freshly-copied sources. Zip
    # file mtimes can be OLDER than the cached .pyc, in which case Python would keep
    # running the old code even though the .py files were replaced.
    try:
        # RECURSIVELY — subpackages (e.g. tracker/) have their OWN __pycache__,
        # and clearing only the top-level one left their stale .pyc in place, so
        # a self-update could keep running old package code after the .py sources
        # were replaced (the same mtime-vs-.pyc trap described above, one level
        # down). Any package directory added later is covered automatically.
        for _root, _dirs, _files in os.walk(dst):
            if os.path.basename(_root) == "__pycache__":
                shutil.rmtree(_root, ignore_errors=True)
    except Exception:
        pass
    # best-effort: install any new dependencies the update introduced
    try:
        req = os.path.join(dst, "requirements.txt")
        if os.path.exists(req):
            cb("Updating dependencies…", "dim")
            subprocess.run([sys.executable, "-m", "pip", "install", "-r", req,
                            "--disable-pip-version-check"],
                           creationflags=0x08000000, timeout=300)
    except Exception:
        pass
    cb("Update installed.", "ok")
    return (True, "Updated to the latest version.")

def apply_update(cb=None):
    """Self-update. For a frozen .exe build, download + swap the binary; for a
    source/git clone, `git pull`. Returns (ok, message).
    """
    import sys
    cb = cb or (lambda *a, **k: None)
    if getattr(sys, "frozen", False):
        return _apply_update_exe(cb)
    import subprocess
    d = _app_dir()
    if not os.path.isdir(os.path.join(d, ".git")):
        return _apply_update_zip(cb)
    try:
        v = subprocess.run(["git", "--version"], capture_output=True, text=True)
        if v.returncode != 0:
            return (False, "Git isn't installed, so the app can't self-update. "
                           "Download the latest version from GitHub.")
    except Exception:
        return (False, "Git isn't installed. Download the latest version from GitHub.")

    def run(args):
        return subprocess.run(args, cwd=d, capture_output=True, text=True)

    cb("Fetching the latest version…", "dim")
    # discard local edits to tracked files so pull can't conflict, then pull
    run(["git", "stash", "--include-untracked"])
    p = run(["git", "pull", "--ff-only", "origin", GITHUB_BRANCH])
    out = (p.stdout + p.stderr).strip()
    if p.returncode != 0:
        # try a non-ff pull as a fallback
        p2 = run(["git", "pull", "origin", GITHUB_BRANCH])
        out = (p2.stdout + p2.stderr).strip()
        if p2.returncode != 0:
            return (False, out[:300] or "git pull failed.")
    cb("Update downloaded.", "ok")
    return (True, "Updated. Please restart QA Studio to use the new version.")

# ═══════════════════════════════════════════════════════════════════════════════
#  SELF-HEALING AUTOMATION (no live browser in QA Studio)
#  The Automation screen compiles each story's cases into intents, orders them
#  into a logical sequence (logged-out negatives/validation/login-page cases →
#  successful login → app cases), and GENERATES a Maven/TestNG/Selenium project
#  whose runtime heals locators by calling the configured AI provider when a seed
#  locator fails. QA Studio never drives the browser; IntelliJ runs `mvn test` and the
#  generated framework self-heals + caches locators.
# ═══════════════════════════════════════════════════════════════════════════════

def validate_and_sequence_suite(stories_payload, log=None, want_ai=True,
                                should_stop=lambda: False, on_error=None, gate=None):
    """Validate/repair each test case for automatability and order the suite.

    Pause/stop hooks:
      should_stop()  -> True to abort.
      gate()         -> called per case; blocks while the user paused; returns
                        False if stopping (so we abort cleanly).
      on_error(msg)  -> called on a recoverable AI error (e.g. low credit). It
                        blocks until the user switches provider + resumes
                        ('retry') or stops ('stop'); we retry the compile on
                        'retry' so the new provider takes effect.

    Ordering buckets (so we never log out to re-test invalids):
      0  login-page negative/validation  (logged OUT)
      1  login-page presence/interaction (logged OUT, e.g. language dropdown)
      2  the successful-login case        (transition; synthesized if absent)
      3  app cases                        (logged IN; e.g. language toggle)"""
    log = log or (lambda *a, **k: None)
    # `log` is now called concurrently from the compile-worker pool below as
    # well as this function's own main-thread code — shadow it with a lock so
    # every individual call stays atomic (same cb-locking pattern already
    # used for run_steps'/run_titles' worker pools).
    _log_lock = _threading.Lock()
    _real_log = log
    def log(msg, tone="dim"):
        with _log_lock:
            _real_log(msg, tone)

    out = []
    _total = sum(len(sp.get("test_cases", []) or []) for sp in stories_payload)
    _done = 0
    _todo_running = 0
    _COMPILE_WORKERS = 2
    import concurrent.futures as _cf_vs
    log("Sequencing %d test case(s) — this is the slow part (one AI pass each)…"
        % _total, "info")

    def _compile_one(tc, story, case_no, seq_no=None, seq_total=None):
        """Runs on a worker thread: classify/bucket + compile_test_case (with
        its own retry-on-recoverable-error loop) for ONE case. Returns a
        result dict; does NOT touch has_app/has_positive_login/_todo_running/
        cases directly — the main thread applies those from the result,
        mirroring run_steps' _gen_and_write/_apply_result split. has_app and
        has_positive_login are computed and returned even when the case is
        later skipped as non-automatable — the ORIGINAL sequential code set
        both accumulators BEFORE its skip check ran, so a skipped bucket-3
        case could still flip has_app; that quirk is preserved here
        deliberately, not "fixed", to avoid changing sequencing behavior.
        seq_no/seq_total are the same "k/N" already announced by the
        submission loop's "Sequencing case k/N" line — stamped onto this
        title line too because the two can log seconds apart and completely
        out of order (compile runs concurrently across worker threads), with
        nothing else tying a given title back to which case it belongs to."""
        _ctitle = (tc.get("title", "") or "").strip()
        # Always have a description ready, even when the generated case has no
        # title — otherwise the entry renders blank with no way to tell which
        # case it was (previously silently skipped when _ctitle was empty).
        _body = _ctitle[:90] if _ctitle else "(untitled test case)"
        # Logging the title immediately (when a worker STARTS a case) used to
        # put it on its own line, with the matching "⏱ elapsed" line landing
        # much later once that case's AI call actually finished — with
        # _COMPILE_WORKERS running concurrently, those two lines end up pages
        # apart, interleaved with OTHER cases' title/heartbeat lines, and even
        # with each self-labeled "#k/N" (previous fix) the overall log reads
        # as numbers jumping around at random. Merged into ONE line, logged
        # once, when the case is actually done — see the single log() call at
        # the bottom of this function.
        _case_start = time.time()   # ⏱ per-case sequencing time, same convention
                                    # as the Run log's create/update timing
        # Announce the case when a worker ACTUALLY starts it — this line
        # used to be logged by the submission loop, which submits every case
        # upfront, so all N "Sequencing case k/N" lines printed instantly in
        # one burst (confirmed live: 27 lines at once) and told the user
        # nothing about progress. Logged here, at most _COMPILE_WORKERS of
        # these appear between completion lines — real liveness.
        if seq_no:
            log("Sequencing case %d/%d" % (seq_no, seq_total), "case")
        _case_label = ("#%d/%d · " % (seq_no, seq_total)) if seq_no else ""
        ctype = _classify_case(tc)
        pctx = _infer_page_context(tc, ctype)
        # bucket + priority
        if pctx == "login" and ctype == "negative_login":
            bucket = 0
        elif pctx == "login":
            bucket = 1
        else:
            bucket = 3
        low = _norm(tc.get("title", ""))
        positive_login = (pctx == "login" and ctype not in ("negative_login", "presence")
                          and any(k in low for k in ("نجاح", "صحيح", "الصحيحة", "valid",
                                                     "success", "successful")))
        if positive_login:
            bucket = 2
        # compile to intents, pausing (not aborting) on a recoverable AI error
        intents = []
        _page_meta = {}   # filled by compile_test_case with the AI's own page judgment
        if want_ai:
            while True:
                if should_stop():
                    return {"abort": True}
                try:
                    # Run the (blocking, up-to-90s) AI call stop-aware: without
                    # this, clicking Stop mid-case had to wait out the current
                    # compile_test_case call before should_stop() was checked
                    # again — up to a 90s delay. _run_stopaware polls should_stop
                    # every 0.3s in a separate thread and raises StopRequested at
                    # once when Stop is clicked; the abandoned request finishes
                    # on its own timeout in the background and is harmless since
                    # compiling a case has no side effects.
                    # on_slow: the LAST silent window in this path. Retries were
                    # already visible (on_retry), but a long attempt that hasn't
                    # errored yet — or a backoff wait whose one collapsed line
                    # sits far above the log's bottom — showed nothing for
                    # minutes (reported live: "stuck" 5+ min at "Sequencing
                    # case 49/66"). The "— Ns so far…" shape is folded by
                    # _auto_logmsg's keyed collapse into the SAME "#k/N" line
                    # as this case's retry/cooldown notes: one status line per
                    # case, always current.
                    intents = _run_stopaware(
                        lambda: compile_test_case(tc, story, log, ctype,
                                                  case_label=_case_label,
                                                  meta_out=_page_meta),
                        should_stop=should_stop,
                        on_slow=lambda s: log("  %sstill compiling with the AI — %ds so far…"
                                              % (_case_label, s), "dim")) or []
                    break
                except StopRequested:
                    return {"abort": True}   # Stop clicked mid-call — abort now,
                                              # don't wait for this case's AI call.
                except Exception as e:
                    # compile_test_case only raises for RECOVERABLE provider
                    # errors (credit, expired/invalid key, rate limit, outage).
                    # Pause and let the user fix it / switch provider + Resume.
                    # _auto_on_ai_error's threading.Condition-based pause gate
                    # is already safe for multiple concurrent callers (each
                    # worker independently waits/wakes), so no extra
                    # coordination is needed for two workers to hit this at once.
                    decision = on_error(friendly_ai_error(e)) if on_error else "stop"
                    if decision == "retry":
                        log(f"  {_case_label}Retrying compile with the current provider…", "dim")
                        continue
                    return {"abort": True}   # user chose Stop (should_stop is now True)
        if not intents:
            intents = _intents_from_raw_steps(tc)
        # AI page-context override. The title-keyword heuristic
        # (_infer_page_context/_is_negative_login_tc) misreads a title that
        # merely ENUMERATES menu items ("… تغيير كلمة المرور - تسجيل الخروج")
        # as a login-page case — confirmed live: the dropdown-content case
        # landed in the generated Playwright project's logged-out group with
        # a "user logged in" precondition it could never satisfy. The compile
        # model reads the WHOLE case (preconditions included) and now returns
        # its own "page" judgment (prompt rule 5); when it disagrees with the
        # heuristic, the AI wins — the heuristic remains the fallback whenever
        # the AI was unavailable (want_ai=False / needs_review). Same
        # prompt-contract + programmatic-integration philosophy as the dedup
        # structured-output fix.
        _ai_page = _page_meta.get("page", "")
        if _ai_page and _ai_page != pctx:
            log("  %spage context: AI judged '%s' (title heuristic said '%s') — using the AI's"
                % (_case_label, _ai_page, pctx), "dim")
            pctx = _ai_page
            if pctx == "login":
                bucket = 0 if ctype == "negative_login" else 1
            else:
                bucket = 3
            positive_login = (pctx == "login" and ctype not in ("negative_login", "presence")
                              and any(k in low for k in ("نجاح", "صحيح", "الصحيحة", "valid",
                                                         "success", "successful")))
            if positive_login:
                bucket = 2
        # One line per case: elapsed time + "#k/N" as the META half, the
        # case's own title as the BODY half, joined with "\x1f" (an invisible
        # control character, never actually rendered) instead of being
        # concatenated into one plain string. main.py's _auto_log_line splits
        # on it and renders the two halves as SEPARATE Text controls — mixing
        # an LTR tag directly into the same string as an RTL Arabic title hit
        # Unicode's bidi reordering (confirmed live: "#3/36" rendered back as
        # "3/36#", scrambled) because Flutter determines paragraph direction
        # from the first strong-directional character in the WHOLE string,
        # not a fixed left-to-right rule — two independent Text widgets each
        # get their own direction resolved from their own content, so neither
        # can reorder relative to the other.
        _meta = "  ⏱ %s" % _fmt_mmss(time.time() - _case_start)
        if seq_no:
            _meta += " · #%d/%d" % (seq_no, seq_total)
        log("%s\x1f%s" % (_meta, _body), "case")
        # A case that drives username/password fields belongs on the LOGIN page,
        # not the app page — otherwise it runs logged-in against BASE_URL where
        # those fields don't exist (guaranteed failure). Pull it back to a
        # logged-out login bucket.
        if bucket == 3 and _ai_page != "app":
            # Only pull an app case back to the login page when its TITLE says it's
            # about login. A login PRECONDITION in the steps/intents (present in
            # almost every app case) must NOT drag a real app case onto the login
            # page — that's what made every case land "logged-out".
            # Skipped entirely when the AI explicitly judged "app" (override
            # above) — otherwise a title that merely MENTIONS password/logout
            # words (menu items, "change password" app screens) would drag the
            # case right back to the login bucket the AI just corrected.
            _blob = _norm(tc.get("title", ""))
            if any(s in _blob for s in ("password", "username", "كلمة المرور",
                                        "كلمه المرور", "تسجيل الدخول",
                                        "login button", "login field", "login submit")):
                bucket = 0
                pctx = "login"
        has_app = (bucket == 3)
        n_act = sum(1 for i in intents if i["role"] == "action")
        has_assert = any(i["role"] == "assertion" for i in intents)
        needs_review = False
        if n_act == 0 and bucket != 2 and not has_assert:
            # A presence case ("verify X exists") is still automatable as a
            # single visibility check — synthesize one rather than drop it.
            # A non-presence case landing here means BOTH the AI compiler
            # AND the raw-steps fallback parser came up empty — usually
            # because the AI call kept failing (e.g. a persistent provider
            # rate limit) rather than the case being genuinely
            # non-automatable. Dropping it silently made real, automatable
            # cases (confirmed live: ID/email-uniqueness validation cases)
            # vanish from the output with no trace. Generate the same
            # best-effort visibility check instead, and flag it distinctly
            # so it's visible rather than lost.
            title = tc.get("title", "")
            intents.append({
                "role": "assertion", "verb": "", "target": title,
                "keywords": [w for w in re.split(r"\s+", title) if len(w) > 1][:6],
                "kind": "any", "value": "", "check": "visible",
                "expected": "", "from_steps": []})
            has_assert = True
            if ctype != "presence":
                needs_review = True
                # Same meta/body split as the combined case line above — the
                # tag mixed into the same string as an Arabic title scrambles
                # under Unicode bidi reordering.
                _rev_meta = ("needs review — AI compile unavailable, placeholder check only"
                             + (" · #%d/%d" % (seq_no, seq_total) if seq_no else ""))
                log("    %s\x1f%s" % (_rev_meta, title[:40]), "review")
        case_dict = {"tc": tc, "title": tc.get("title", ""), "ctype": ctype,
                    "page_context": pctx, "bucket": bucket, "intents": intents,
                    "needs_review": needs_review}
        return {"case": case_dict, "case_no": case_no,
                "has_app": has_app, "has_positive_login": positive_login,
                "todo_delta": _count_null_seeds(bucket, intents)}

    for sp in stories_payload:
        if should_stop():
            return out
        story = sp.get("story", {})
        # Announce the story BEFORE sequencing its cases — with multiple
        # stories queued up, "Sequencing case k/N" alone doesn't say which
        # story k belongs to; this makes the run trackable story-by-story,
        # same "story" tone/symbol as the Run log's per-story lines.
        _n_cases = len(sp.get("test_cases", []) or [])
        log(f"Story {story.get('id', '')} · {story.get('title', '')} "
            f"({_n_cases} test case" + ("s" if _n_cases != 1 else "") + ")", "story")
        cases_by_no = {}
        has_app = False
        has_positive_login = False
        aborted = False
        # Compile this story's cases with up to _COMPILE_WORKERS concurrent
        # AI passes. Only this story's OWN cases are pooled together (stories
        # still run one after another) — keeps abort/pause semantics simple:
        # a stop/pause mid-story still discards just that story's progress,
        # same as the original "return out" behavior.
        with _cf_vs.ThreadPoolExecutor(max_workers=_COMPILE_WORKERS) as _ex:
            _futs = {}
            for i, tc in enumerate(sp.get("test_cases", []) or []):
                if should_stop():
                    aborted = True
                    break
                if gate and not gate():   # manual pause point (returns False on stop)
                    aborted = True
                    break
                _done += 1
                # The "Sequencing case k/N" announce line moved INTO
                # _compile_one (logged when a worker actually starts the
                # case) — logging it here, at submission time, printed all N
                # lines in one instant burst because this loop submits every
                # case upfront and lets the 2-worker pool chew through them.
                # _done is captured NOW (submission order) and threaded through to
                # _compile_one so its own title log line can stamp the same "k/N"
                # onto itself. Cases compile concurrently across _COMPILE_WORKERS
                # threads, so a case's title can log seconds after (and completely
                # out of order relative to) its own "Sequencing case k/N" line —
                # confirmed live: cases 3-18 of a 19-case story showed no title at
                # all, only 1, 2 and the last few, because those finished fast
                # enough to log right after their own progress line while the
                # others' titles landed much later, buried under retry noise with
                # nothing tying them back to a case number. Self-labeling each
                # title line makes it identifiable regardless of arrival order.
                fut = _ex.submit(_compile_one, tc, story, i, _done, _total)
                _futs[fut] = i
            for fut in _cf_vs.as_completed(_futs):
                res = fut.result()
                if res.get("abort"):
                    aborted = True
                    continue
                if res.get("has_app"):
                    has_app = True
                if res.get("has_positive_login"):
                    has_positive_login = True
                if res.get("skip"):
                    continue
                cases_by_no[res["case_no"]] = res["case"]
                # live TODO tally — grows as each case is sequenced (hidden control line
                # the UI reads to update the TODO counter without cluttering the log).
                _todo_running += res.get("todo_delta", 0)
                log("TODO_LIVE: %d" % _todo_running, "meta")

        if aborted or should_stop():
            return out

        # Original authored order (NOT completion order) — the stable sort
        # below relies on cases within a bucket keeping their original
        # authored order (the order they appear in the story's test_cases).
        cases = [cases_by_no[i] for i in sorted(cases_by_no)]

        # synthesize a successful-login transition if app cases exist but no
        # explicit positive-login case was authored
        if has_app and not has_positive_login:
            cases.append({"tc": {"title": "Successful login (auto-inserted)", "steps": []},
                          "title": "Successful login (auto-inserted)",
                          "ctype": "login", "page_context": "login", "bucket": 2,
                          "intents": [{"role": "action", "verb": "login", "target": "login",
                                       "keywords": [], "kind": "any", "value": "",
                                       "check": "", "expected": "", "from_steps": []}],
                          "synthetic_login": True})
            log("    + inserted a successful-login step before app cases", "dim")
        # Sort by BUCKET ONLY — Python's sort is stable, so cases within a bucket
        # keep their original authored order (the order they appear in the story's
        # test_cases). The previous secondary sort key, c["title"], reordered every
        # bucket ALPHABETICALLY by title text instead — scrambling any intended
        # step sequence (e.g. "create X" needs to run before "verify X exists"),
        # and for mixed Arabic/English titles the Unicode ranges don't sort
        # anywhere near authored/logical order either. This was the real
        # 'wrong sequencing' bug — nothing to do with RTL/LTR text rendering.
        cases.sort(key=lambda c: c["bucket"])
        for i, c in enumerate(cases):
            c["priority"] = c["bucket"] * 100 + i
        out.append({"story": story, "cases": cases})
        log(f"  sequenced story {story.get('id')}: "
            f"{sum(1 for c in cases if c['bucket']<2)} logged-out, "
            f"{sum(1 for c in cases if c['bucket']==2)} login, "
            f"{sum(1 for c in cases if c['bucket']>2)} app case(s)", "info")
    return out


def _seed_locator_for_intent(intent, app_page=False):
    """Best-effort seed Selenium By for an intent, BEFORE any runtime healing.
    'Stable where known, // TODO only where unknown.' Returns (by, value, known).

    `app_page=True` (a bucket-3 case, running on the app AFTER login) disables
    the LOGIN-FORM seed rules below: an app-page field whose keywords mention
    email/username (e.g. the profile-edit email field) used to get the login
    form's own '#username,input[name=username],input[type=email]' seed — on the
    wrong page that selector still MATCHES something (the login username box),
    so the healer trusts it, types into the wrong field, and the test fails
    with no healing ever attempted. Confirmed live in a generated Playwright
    project (story 101049's profile email cases)."""
    kws = [k for k in (intent.get("keywords") or []) if str(k).strip()]
    kind = intent.get("kind", "any")
    low = _norm(" ".join(kws) + " " + (intent.get("target", "") or ""))
    # known stable patterns we validated against the real DOM
    if any(k in low for k in ("languageswitch", "language toggle", "زر اللغة", "زر اللغه")):
        return ("cssSelector", '[data-svgicon="languageSwitch"]', True)
    if "kc-current-locale" in low or ("locale" in low and "dropdown" in low):
        return ("id", "kc-current-locale-link", True)
    if kind == "input" and any(k in low for k in ("username", "user", "email", "اسم المستخدم")):
        if app_page:
            if "email" in low or "بريد" in low:
                return ("cssSelector", "input[type=email]", False)
            return ("cssSelector", "TODO_RESOLVE_AT_RUNTIME", False)
        return ("cssSelector", "#username,input[name=username],input[type=email]", True)
    if kind == "input" and any(k in low for k in ("password", "كلمة المرور", "كلمة مرور")):
        return ("cssSelector", "#password,input[type=password]", True)
    if any(k in low for k in ("kc-login", "submit", "تسجيل الدخول", "login button", "sign in")):
        return ("cssSelector", "#kc-login,button[type=submit],input[type=submit]", True)
    # text-based xpath for an option/button with a clear label
    label = next((k for k in kws if len(str(k)) >= 2 and not str(k).isascii() or
                  (str(k).isascii() and len(str(k)) >= 3)), None)
    if label and kind in ("menuitem", "link", "button"):
        tag = {"menuitem": "a", "link": "a", "button": "button"}.get(kind, "*")
        return ("xpath", '//%s[normalize-space()="%s"]' % (tag, label), False)
    return ("cssSelector", "TODO_RESOLVE_AT_RUNTIME", False)


def _count_null_seeds(bucket, intents):
    """How many of a case's intents become RUNTIME-resolved locators (no stable
    seed) — i.e. the TODO tally. Mirrors _emit_intent's filtering so the live count
    matches what the specs actually emit (precondition = no locator; a null-seed
    text/message assertion becomes assertText = no locator; bucket-2 action intents
    are handled by performLogin)."""
    n = 0
    for intent in intents:
        role = intent.get("role")
        if role == "precondition":
            continue
        if bucket == 2 and role == "action":
            continue
        _, val, _ = _seed_locator_for_intent(intent, app_page=(bucket == 3))
        if val != "TODO_RESOLVE_AT_RUNTIME":
            continue
        if role == "assertion":
            kind = (intent.get("kind") or "").lower()
            if kind in ("text", "message", "menu", "validation", "error") or not kind:
                continue
        n += 1
    return n


def _sh_pom(group_id, artifact_id):
    return ("""<?xml version="1.0" encoding="UTF-8"?>
<project xmlns="http://maven.apache.org/POM/4.0.0"
         xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
         xsi:schemaLocation="http://maven.apache.org/POM/4.0.0 http://maven.apache.org/xsd/maven-4.0.0.xsd">
  <modelVersion>4.0.0</modelVersion>
  <groupId>__GID__</groupId>
  <artifactId>__AID__</artifactId>
  <version>1.0.0</version>
  <properties>
    <maven.compiler.source>17</maven.compiler.source>
    <maven.compiler.target>17</maven.compiler.target>
    <project.build.sourceEncoding>UTF-8</project.build.sourceEncoding>
  </properties>
  <dependencies>
    <dependency><groupId>org.seleniumhq.selenium</groupId><artifactId>selenium-java</artifactId><version>4.21.0</version></dependency>
    <dependency><groupId>org.testng</groupId><artifactId>testng</artifactId><version>7.10.2</version></dependency>
    <dependency><groupId>io.github.bonigarcia</groupId><artifactId>webdrivermanager</artifactId><version>5.9.2</version></dependency>
    <dependency><groupId>com.fasterxml.jackson.core</groupId><artifactId>jackson-databind</artifactId><version>2.17.1</version></dependency>
  </dependencies>
  <build><plugins>
    <plugin><groupId>org.apache.maven.plugins</groupId><artifactId>maven-surefire-plugin</artifactId><version>3.2.5</version>
      <configuration><suiteXmlFiles><suiteXmlFile>testng.xml</suiteXmlFile></suiteXmlFiles></configuration>
    </plugin>
  </plugins></build>
</project>
""".replace("__GID__", group_id).replace("__AID__", artifact_id))


def _healer_ai_meta():
    """(provider, ai_base_url, model) for the generated runtime self-healer, taken
    from the SAME AI connection used to generate the tests. Non-secret — baked into
    Config; the API key is read from the environment at run time."""
    prov = AI_PROVIDER
    cfg = AI_CONFIG.get(prov, {})
    model = current_model() or ""
    if prov == "anthropic":
        base = "https://api.anthropic.com"
    elif prov == "gemini":
        base = "https://generativelanguage.googleapis.com/v1beta"
    else:  # openai-compatible (openai/nvidia/groq/cerebras/openrouter/deepseek/qwen/mistral/ollama)
        base = (cfg.get("base_url") or "https://api.openai.com/v1")
    return prov, base.rstrip("/"), model


def _sh_config(pkg, base_url, login_url, ai_provider="anthropic",
               ai_base_url="https://api.anthropic.com", ai_model="claude-sonnet-4-6"):
    # NOTE: base_url / login_url are intentionally NOT baked into this committed
    # source — they live in config.properties (git-ignored). ai_* stay as a baked
    # FALLBACK so a fresh clone can still heal even without config.properties.
    return ("""package __PKG__.core;

import java.io.FileInputStream;
import java.util.Properties;

/** Runtime config. Each value resolves in order: environment variable, then
 *  config.properties (git-ignored, environment-specific), then a safe fallback.
 *  Environment URLs live only in config.properties, so this committed source is
 *  environment-agnostic — the same generated project runs against any environment
 *  with no regeneration. Secrets (USER / PASS / API_KEY) are ENV-ONLY and are never
 *  read from the file. */
public final class Config {
    private Config() {}
    private static final Properties P = new Properties();
    static {
        try (FileInputStream in = new FileInputStream("config.properties")) { P.load(in); }
        catch (Exception ignored) {}
    }
    public static final String BASE_URL  = val("APP_BASE_URL",  "app.base.url",  "");
    public static final String LOGIN_URL = val("APP_LOGIN_URL", "app.login.url", "");
    public static final String USER      = env("APP_USER",  "");   // secret: env only
    public static final String PASS      = env("APP_PASS",  "");   // secret: env only
    // ── AI self-healing (used only when a seed locator fails at run time) ──
    public static final String AI_PROVIDER = val("QA_AI_PROVIDER", "ai.provider", "__PROVIDER__");
    public static final String AI_BASE_URL = val("QA_AI_BASE_URL", "ai.base.url", "__AI_BASE__");
    public static final String MODEL       = val("QA_AI_MODEL",    "ai.model",    "__MODEL__");
    // Key: prefer the neutral QA_AI_API_KEY, else fall back to common provider vars.
    public static final String API_KEY   = firstEnv("QA_AI_API_KEY", "ANTHROPIC_API_KEY",
                                                     "OPENAI_API_KEY", "NVIDIA_API_KEY",
                                                     "GROQ_API_KEY", "GEMINI_API_KEY");
    public static final boolean HEAL      = !API_KEY.isEmpty();
    /** env var → config.properties → default. */
    private static String val(String envKey, String propKey, String d) {
        String v = System.getenv(envKey);
        if (v != null && !v.isEmpty()) return v;
        v = P.getProperty(propKey);
        return (v == null || v.isEmpty()) ? d : v;
    }
    private static String env(String k, String d) {
        String v = System.getenv(k); return (v == null || v.isEmpty()) ? d : v;
    }
    private static String firstEnv(String... keys) {
        for (String k : keys) { String v = System.getenv(k);
            if (v != null && !v.isEmpty()) return v; }
        return "";
    }
}
""".replace("__PKG__", pkg)
   .replace("__PROVIDER__", ai_provider).replace("__AI_BASE__", ai_base_url)
   .replace("__MODEL__", ai_model))


def _sh_config_properties(base_url, login_url, ai_provider, ai_base_url, ai_model):
    """Real, environment-specific values — written to config.properties (GIT-IGNORED)
    once, and never clobbered on a re-run so your edits survive."""
    return (
        "# QA Studio — environment-specific config (GIT-IGNORED).\n"
        "# Edit these for your environment. Environment variables of the same name\n"
        "# (APP_BASE_URL, APP_LOGIN_URL, QA_AI_PROVIDER, QA_AI_BASE_URL, QA_AI_MODEL)\n"
        "# override anything here at runtime.\n"
        "app.base.url=%s\n"
        "app.login.url=%s\n"
        "ai.provider=%s\n"
        "ai.base.url=%s\n"
        "ai.model=%s\n"
        "# Secrets are ENV-ONLY — never put them here: APP_USER, APP_PASS, QA_AI_API_KEY\n"
        % (base_url, login_url, ai_provider, ai_base_url, ai_model))


def _sh_config_properties_example():
    """Committed template so teammates know which keys to set. No real URLs/secrets."""
    return (
        "# Copy this file to config.properties (git-ignored) and fill in your\n"
        "# environment. Environment variables of the same name override these.\n"
        "app.base.url=https://your-app.example.com\n"
        "app.login.url=https://your-login.example.com/realms/.../protocol/openid-connect/auth\n"
        "ai.provider=anthropic\n"
        "ai.base.url=https://api.anthropic.com\n"
        "ai.model=claude-sonnet-4-6\n"
        "# Secrets are ENV-ONLY — do NOT put them here: APP_USER, APP_PASS, QA_AI_API_KEY\n")


def _sh_locator_store(pkg):
    return ("""package __PKG__.core;

import com.fasterxml.jackson.databind.ObjectMapper;
import org.openqa.selenium.By;
import java.io.File;
import java.time.Instant;
import java.util.HashMap;
import java.util.Map;

/** Single source of truth for locators, persisted to locators.json (COMMITTED).
 *  The generator seeds it with the locators captured at generation time; when a
 *  step is resolved at RUNTIME by the AI, the verified locator is written back
 *  here immediately. Because the UI is identical across environments, this file
 *  is shared — every later run, on any machine or environment, reuses the saved
 *  locator so the AI is asked at most once per step. */
public final class LocatorStore {
    private static final File FILE = new File("locators.json");
    private static final ObjectMapper M = new ObjectMapper();
    private final Map<String, Map<String, Object>> cache;

    @SuppressWarnings("unchecked")
    public LocatorStore() {
        Map<String, Map<String, Object>> c = new HashMap<>();
        try { if (FILE.exists()) c = M.readValue(FILE, Map.class); } catch (Exception ignored) {}
        this.cache = c;
    }
    public By get(String key) {
        Map<String, Object> e = cache.get(key);
        if (e == null || e.get("value") == null) return null;
        return Healer.toBy(str(e.get("by")), str(e.get("value")));
    }
    /** Save a resolved locator and flush to disk at once (so a crash mid-suite
     *  still keeps everything resolved so far). */
    public void put(String key, String by, String value) {
        Map<String, Object> e = new HashMap<>();
        e.put("by", by);
        e.put("value", value);
        e.put("resolvedAt", Instant.now().toString());
        e.put("provider", Config.AI_PROVIDER);
        cache.put(key, e);
        try { M.writerWithDefaultPrettyPrinter().writeValue(FILE, cache); } catch (Exception ignored) {}
    }
    private static String str(Object o) { return o == null ? null : o.toString(); }
}
""".replace("__PKG__", pkg))


def _sh_ai_client(pkg):
    return ("""package __PKG__.core;

import com.fasterxml.jackson.databind.JsonNode;
import com.fasterxml.jackson.databind.ObjectMapper;
import com.fasterxml.jackson.databind.node.ArrayNode;
import com.fasterxml.jackson.databind.node.ObjectNode;
import java.net.URI;
import java.net.http.HttpClient;
import java.net.http.HttpRequest;
import java.net.http.HttpResponse;
import java.util.Map;

/** Provider-agnostic AI client used ONLY to pick a locator when a seed fails.
 *  Speaks Anthropic, Gemini, or any OpenAI-compatible endpoint (OpenAI, NVIDIA,
 *  Groq, Cerebras, OpenRouter, DeepSeek, Qwen, Mistral, Ollama) — chosen by
 *  Config.AI_PROVIDER, which is baked from the connection that generated these
 *  tests. Returns {"by","value"} or null. */
public final class AiClient {
    private static final ObjectMapper M = new ObjectMapper();
    private static final HttpClient HTTP = HttpClient.newHttpClient();

    public static Map<String, String> pickLocator(String intentJson, String candidatesJson) {
        if (Config.API_KEY.isEmpty()) return null;
        String prompt = "You resolve a Selenium locator for a UI test step that failed to "
            + "find its element. Choose the ONE element that matches the intent and return a "
            + "STABLE locator. Reply ONLY JSON: {\\"by\\":\\"id|name|cssSelector|xpath\\",\\"value\\":\\"...\\"}. "
            + "Prefer id (non-generated) > data-testid/[data-svgicon] css > name > aria/text xpath. "
            + "Never use framework ids like pn_id_*, cdk-, mat-, GUIDs.\\n\\nINTENT: " + intentJson
            + "\\n\\nCANDIDATES: " + candidatesJson;
        try {
            String p = Config.AI_PROVIDER == null ? "" : Config.AI_PROVIDER.toLowerCase();
            String text;
            if (p.equals("anthropic"))      text = callAnthropic(prompt);
            else if (p.equals("gemini"))    text = callGemini(prompt);
            else                            text = callOpenAICompatible(prompt);
            if (text == null || text.isEmpty()) return null;
            int a = text.indexOf('{'), b = text.lastIndexOf('}');
            if (a < 0 || b < a) return null;
            JsonNode loc = M.readTree(text.substring(a, b + 1));
            String by = loc.path("by").asText(""), value = loc.path("value").asText("");
            if (by.isEmpty() || value.isEmpty()) return null;
            return Map.of("by", by, "value", value);
        } catch (Exception e) {
            System.out.println("[heal] pickLocator failed: " + e.getMessage());
            return null;
        }
    }

    private static String base() {
        String b = Config.AI_BASE_URL == null ? "" : Config.AI_BASE_URL;
        while (b.endsWith("/")) b = b.substring(0, b.length() - 1);
        return b;
    }

    private static String send(HttpRequest req) throws Exception {
        HttpResponse<String> resp = HTTP.send(req, HttpResponse.BodyHandlers.ofString());
        if (resp.statusCode() / 100 != 2) {
            System.out.println("[heal] API error " + resp.statusCode() + ": " + resp.body());
            return null;
        }
        return resp.body();
    }

    private static ObjectNode chatBody(String prompt) {
        ObjectNode body = M.createObjectNode();
        body.put("model", Config.MODEL);
        body.put("max_tokens", 256);
        ArrayNode msgs = body.putArray("messages");
        ObjectNode msg = msgs.addObject(); msg.put("role", "user"); msg.put("content", prompt);
        return body;
    }

    /** OpenAI-compatible /chat/completions (OpenAI, NVIDIA, Groq, Cerebras, …). */
    private static String callOpenAICompatible(String prompt) throws Exception {
        HttpRequest req = HttpRequest.newBuilder()
            .uri(URI.create(base() + "/chat/completions"))
            .header("Authorization", "Bearer " + Config.API_KEY)
            .header("content-type", "application/json")
            .POST(HttpRequest.BodyPublishers.ofString(M.writeValueAsString(chatBody(prompt))))
            .build();
        String resp = send(req);
        if (resp == null) return null;
        return M.readTree(resp).path("choices").path(0).path("message").path("content").asText("");
    }

    private static String callAnthropic(String prompt) throws Exception {
        HttpRequest req = HttpRequest.newBuilder()
            .uri(URI.create(base() + "/v1/messages"))
            .header("x-api-key", Config.API_KEY)
            .header("anthropic-version", "2023-06-01")
            .header("content-type", "application/json")
            .POST(HttpRequest.BodyPublishers.ofString(M.writeValueAsString(chatBody(prompt))))
            .build();
        String resp = send(req);
        if (resp == null) return null;
        return M.readTree(resp).path("content").path(0).path("text").asText("");
    }

    private static String callGemini(String prompt) throws Exception {
        ObjectNode body = M.createObjectNode();
        ArrayNode contents = body.putArray("contents");
        ArrayNode parts = contents.addObject().putArray("parts");
        parts.addObject().put("text", prompt);
        String url = base() + "/models/" + Config.MODEL + ":generateContent?key=" + Config.API_KEY;
        HttpRequest req = HttpRequest.newBuilder()
            .uri(URI.create(url))
            .header("content-type", "application/json")
            .POST(HttpRequest.BodyPublishers.ofString(M.writeValueAsString(body)))
            .build();
        String resp = send(req);
        if (resp == null) return null;
        return M.readTree(resp).path("candidates").path(0).path("content")
                 .path("parts").path(0).path("text").asText("");
    }
}
""".replace("__PKG__", pkg))


def _sh_healer(pkg):
    return ("""package __PKG__.core;

import org.openqa.selenium.*;
import org.openqa.selenium.support.ui.ExpectedConditions;
import org.openqa.selenium.support.ui.WebDriverWait;
import java.nio.charset.StandardCharsets;
import java.time.Duration;
import java.util.Map;

/** Finds elements with runtime self-healing. Order per step key:
 *  1) locator saved in locators.json  2) the generated seed locator
 *  3) ask the configured AI provider to pick one from the live DOM, then save it
 *     back into locators.json so it is reused on every later run. */
public final class Healer {
    private final WebDriver driver;
    private final WebDriverWait wait;
    private final LocatorStore store = new LocatorStore();
    private static String HARVEST_JS;

    public Healer(WebDriver driver) {
        this.driver = driver;
        this.wait = new WebDriverWait(driver, Duration.ofSeconds(12));
    }

    public WebElement find(String key, By seed, String intentJson) {
        By cached = store.get(key);
        if (cached != null) {
            WebElement e = tryFind(cached);
            if (e != null) return e;
        }
        if (seed != null) {
            WebElement e = tryFind(seed);
            if (e != null) return e;
        }
        if (Config.HEAL) {
            System.out.println("[heal] resolving '" + key + "' via AI (" + Config.AI_PROVIDER + ")");
            String dom = harvest();
            Map<String, String> picked = AiClient.pickLocator(intentJson, dom);
            if (picked != null) {
                By by = toBy(picked.get("by"), picked.get("value"));
                WebElement e = tryFind(by);
                if (e != null) {
                    store.put(key, picked.get("by"), picked.get("value"));
                    System.out.println("[heal] '" + key + "' -> " + picked.get("by")
                        + "=" + picked.get("value"));
                    return e;
                }
            }
        }
        throw new NoSuchElementException("Could not resolve step '" + key
            + "'. Seed=" + seed + ". Set QA_AI_API_KEY (or your provider's key) to enable AI healing.");
    }

    public void act(String key, String verb, By seed, String intentJson, String value) {
        if ("navigate".equals(verb) || "wait".equals(verb)) return;
        WebElement el = find(key, seed, intentJson);
        ((JavascriptExecutor) driver).executeScript(
            "arguments[0].scrollIntoView({block:'center'});", el);
        switch (verb == null ? "click" : verb) {
            case "type":
                el.clear(); if (value != null && !value.isEmpty()) el.sendKeys(value); break;
            case "select":
                el.click(); break;   // custom dropdowns: open; the next intent picks the option
            case "hover":
                el.click(); break;
            default:
                try { el.click(); }
                catch (Exception e) {
                    ((JavascriptExecutor) driver).executeScript("arguments[0].click();", el);
                }
        }
    }

    public boolean assertVisible(String key, By seed, String intentJson) {
        try { return find(key, seed, intentJson).isDisplayed(); }
        catch (Exception e) { return false; }
    }

    /** Verify by PAGE TEXT — true if the rendered page contains ANY of the given
     *  keywords. Used for error/validation/message checks instead of locating an
     *  element (no AI call; tolerant of where the message renders). Polls up to 8s
     *  so async messages have time to appear. Case-insensitive. */
    /** Arabic-aware normalization for text assertions: strips tashkeel and
     *  tatweel and folds alef/teh-marbuta/yaa variants, so the app's copy
     *  and the AI-authored keywords match despite orthography differences
     *  (e.g. \\u0639\\u0630\\u0631\\u064B\\u0627 vs its undiacritized form). */
    private static String normText(String s) {
        if (s == null) return "";
        return s.toLowerCase()
                .replaceAll("[\\u064B-\\u0652\\u0670\\u0640]", "")
                .replaceAll("[\\u0623\\u0625\\u0622]", "\\u0627")
                .replace("\\u0629", "\\u0647")
                .replace("\\u0649", "\\u064A");
    }

    public boolean assertTextPresent(String[] keywords) {
        long end = System.currentTimeMillis() + 8000;
        do {
            String txt = "";
            try {
                Object r = ((JavascriptExecutor) driver).executeScript(
                    "return document.body ? document.body.innerText : '';");
                txt = normText(r == null ? "" : r.toString());
            } catch (Exception ignored) {}
            for (String k : keywords) {
                if (k != null && !k.isEmpty() && txt.contains(normText(k))) return true;
            }
            try { Thread.sleep(300); } catch (InterruptedException e) { break; }
        } while (System.currentTimeMillis() < end);
        return false;
    }

    private WebElement tryFind(By by) {
        try { return wait.until(ExpectedConditions.visibilityOfElementLocated(by)); }
        catch (Exception e) { return null; }
    }

    private String harvest() {
        try {
            if (HARVEST_JS == null) {
                var in = Healer.class.getResourceAsStream("/harvest.js");
                HARVEST_JS = new String(in.readAllBytes(), StandardCharsets.UTF_8);
            }
            Object r = ((JavascriptExecutor) driver).executeScript(
                "return JSON.stringify((function(){" + HARVEST_JS + "})());");
            return r == null ? "[]" : r.toString();
        } catch (Exception e) { return "[]"; }
    }

    public static By toBy(String by, String value) {
        if (by == null) return By.cssSelector(value);
        switch (by) {
            case "id":          return By.id(value);
            case "name":        return By.name(value);
            case "xpath":       return By.xpath(value);
            case "linkText":    return By.linkText(value);
            case "className":   return By.className(value);
            default:            return By.cssSelector(value);
        }
    }
}
""".replace("__PKG__", pkg))


def _sh_base_test(pkg):
    return ("""package __PKG__.tests;

import __PKG__.core.*;
import org.openqa.selenium.*;
import org.testng.annotations.AfterClass;
import org.testng.annotations.BeforeClass;
import java.time.Duration;

/** One browser per test class. Logged-out cases run first (priority order),
 *  then the successful-login step, then app cases — no logout-to-retest. */
public abstract class BaseTest {
    protected WebDriver driver;
    protected Healer heal;

    @BeforeClass
    public void setUp() {
        driver = DriverFactory.create();
        driver.manage().timeouts().implicitlyWait(Duration.ofSeconds(6));
        heal = new Healer(driver);
        openLoginPage();
    }
    @AfterClass
    public void tearDown() { if (driver != null) driver.quit(); }

    /** Fresh, logged-out login page (clears any session).
     *  Navigates to the app ROOT and lets it redirect to a freshly-issued login
     *  page. OAuth/OIDC params (code_challenge, nonce, state) are generated per
     *  session, so a frozen LOGIN_URL can go stale — only fall back to it if the
     *  app doesn't redirect to a login form on its own. */
    /** Force the login page into the generated tests' language: the frozen
     *  LOGIN_URL can carry a stale ui_locales=… (seen live on the Playwright
     *  output: ui_locales=en rendered the page in English while every
     *  expected message was Arabic — behavior passed, language failed). */
    private static String withUiLocale(String url) {
        String loc = "__UI_LOCALE__";
        if (loc.isEmpty() || url == null || !url.startsWith("http")) return url;
        if (url.matches(".*[?&]ui_locales=.*"))
            return url.replaceAll("([?&])ui_locales=[^&]*", "$1ui_locales=" + loc);
        return url + (url.contains("?") ? "&" : "?") + "ui_locales=" + loc;
    }

    protected void openLoginPage() {
        try { driver.manage().deleteAllCookies(); } catch (Exception ignored) {}
        driver.get(Config.BASE_URL);
        if (!loginFieldPresent(4000)) {
            driver.get(withUiLocale(Config.LOGIN_URL));
            loginFieldPresent(4000);
        }
    }

    /** True once a username/password field is on the page (polls up to ms). */
    protected boolean loginFieldPresent(long ms) {
        long end = System.currentTimeMillis() + ms;
        do {
            try {
                Object r = ((JavascriptExecutor) driver).executeScript(
                    "return !!document.querySelector("
                    + "'input[type=password],#username,input[name=username]');");
                if (Boolean.TRUE.equals(r)) return true;
            } catch (Exception ignored) {}
            try { Thread.sleep(250); } catch (InterruptedException e) { return false; }
        } while (System.currentTimeMillis() < end);
        return false;
    }

    /** Perform a real successful login using the seed login locators + healing. */
    protected void performLogin() {
        heal.act("login.username", "type",
            Healer.toBy("cssSelector", "#username,input[name=username],input[type=email]"),
            "{\\"target\\":\\"username field\\",\\"kind\\":\\"input\\"}", Config.USER);
        heal.act("login.password", "type",
            Healer.toBy("cssSelector", "#password,input[type=password]"),
            "{\\"target\\":\\"password field\\",\\"kind\\":\\"input\\"}", Config.PASS);
        String beforeUrl = driver.getCurrentUrl();
        heal.act("login.submit", "click",
            Healer.toBy("cssSelector", "#kc-login,button[type=submit],input[type=submit]"),
            "{\\"target\\":\\"login submit button\\",\\"kind\\":\\"button\\"}", "");
        // Wait for the post-login navigation (URL change) instead of a blind sleep.
        long end = System.currentTimeMillis() + 12000;
        while (System.currentTimeMillis() < end) {
            try { if (!driver.getCurrentUrl().equals(beforeUrl)) break; } catch (Exception ignored) {}
            try { Thread.sleep(250); } catch (InterruptedException e) { break; }
        }
    }
}
""".replace("__PKG__", pkg)
       .replace("__UI_LOCALE__", "ar" if _is_arabic_out() else "en"))


def _sh_gitignore():
    # locators.json is COMMITTED on purpose: the UI is identical across
    # environments, so resolved locators are shared. config.properties is
    # IGNORED: it holds environment-specific URLs (and never secrets — those are
    # env-only). config.properties.example is committed as a template.
    return ("target/\n*.iml\n.idea/\n"
            "# environment-specific config (URLs stay out of git):\n"
            "config.properties\n")


def _sh_readme(base_url, login_url):
    return ("""# QA Studio — self-healing automation

Generated by QA Studio. Locators live in `locators.json` (committed): it is seeded
at generation time, and when a step's locator fails at RUNTIME the framework asks
your AI provider to pick the right element from the live DOM, then writes the
verified locator back into `locators.json` — so the AI is asked at most once per
step, and every later run reuses it.

Because the UI and code are identical across your environments, `locators.json` is
shared across all of them (it is committed, not git-ignored). What differs between
environments — the app URLs — lives in `config.properties` (git-ignored), so the same
committed project + locators runs against any environment with no regeneration. Copy
`config.properties.example` to `config.properties` and fill it in (or set the matching
env vars, which override the file).

The AI provider, endpoint and model resolve in order: environment variable →
`config.properties` → a baked fallback (the QA Studio connection you generated with,
e.g. Anthropic, NVIDIA, Groq, OpenAI, Gemini). Only the API KEY is read from the
environment — healing runs as a SEPARATE process here, so it needs its own key
(QA Studio isn't running when `mvn test` executes).

## Run
1. Configure the environment: copy `config.properties.example` → `config.properties`
   and set `app.base.url`, `app.login.url` (optionally `ai.provider`, `ai.base.url`,
   `ai.model`). Env vars `APP_BASE_URL`, `APP_LOGIN_URL`, `QA_AI_PROVIDER`,
   `QA_AI_BASE_URL`, `QA_AI_MODEL` override the file.
2. Set secrets as environment variables (never put these in any file):
   - `QA_AI_API_KEY`  — key for the provider (enables healing; without it, only seed
     locators are used). Provider vars also work as fallbacks: `ANTHROPIC_API_KEY`,
     `OPENAI_API_KEY`, `NVIDIA_API_KEY`, `GROQ_API_KEY`, `GEMINI_API_KEY`.
   - `APP_USER`, `APP_PASS`  (login credentials)
3. `mvn test`

## Sequence
Tests run by TestNG priority: logged-out cases (invalid login, validation, the
login-page language dropdown) first, then the successful-login step, then the
authenticated app cases — all in one browser, no logout-to-retest.

## Healing log
Watch stdout for `[heal] resolving '<step>' ...` and `[heal] '<step>' -> by=value`.
Each resolved locator is written back into `locators.json` and should be committed
so your team (and every environment) reuses it — the AI is then only called for
steps that are genuinely new or whose element changed.
""")  # base_url / login_url now live in config.properties, not the README


def _java_str(s):
    """Escape a Python string for a Java double-quoted literal."""
    return (str(s or "").replace("\\", "\\\\").replace('"', '\\"')
            .replace("\n", " ").replace("\r", " ").replace("\t", " "))


def _java_ident(s, fallback):
    out = re.sub(r"[^A-Za-z0-9_]", "", (s or "").title().replace(" ", ""))
    if not out or not (out[0].isalpha() or out[0] == "_"):
        out = fallback
    return out[:60]


def _emit_intent(lines, key, intent, seed_sink=None, app_page=False):
    """Append the Java for one intent to `lines`. When `seed_sink` is provided,
    record {key: {by, value}} for every KEYED step that has a known seed locator,
    so the project's committed locators.json can be pre-seeded at generation.
    `app_page` — see _seed_locator_for_intent: disables the login-form seed
    rules for bucket-3 (post-login app page) cases."""
    role = intent.get("role")
    target = intent.get("target", "")
    ij = json.dumps({"target": target, "keywords": intent.get("keywords", []),
                     "kind": intent.get("kind", "any"), "verb": intent.get("verb", "")},
                    ensure_ascii=False)
    by, val, _known = _seed_locator_for_intent(intent, app_page=app_page)
    seed = ("null" if val == "TODO_RESOLVE_AT_RUNTIME"
            else 'Healer.toBy("%s", "%s")' % (by, _java_str(val)))
    todo = "  // TODO verify locator (resolved at runtime)" if val == "TODO_RESOLVE_AT_RUNTIME" else ""

    def _seed(k):
        if seed_sink is not None and val != "TODO_RESOLVE_AT_RUNTIME":
            seed_sink[k] = {"by": by, "value": val}
    if role == "precondition":
        lines.append('        // precondition (no UI action): %s' % _java_str(target)[:70])
        return
    if role == "assertion":
        kind = (intent.get("kind") or "").lower()
        kws = [k for k in (intent.get("keywords") or []) if k] or ([target] if target else [])
        # Text/message/menu checks with no locatable element → verify by PAGE TEXT
        # (no AI heal): faster, cheaper, and robust to where the message renders.
        if seed == "null" and (kind in ("text", "message", "menu", "validation", "error")
                               or not kind):
            arr = ", ".join('"%s"' % _java_str(k) for k in kws)
            lines.append('        org.testng.Assert.assertTrue('
                         'heal.assertTextPresent(new String[]{%s}),' % arr)
            lines.append('            "expected text: %s");' % _java_str(target)[:60])
            return
        lines.append('        org.testng.Assert.assertTrue(heal.assertVisible("%s", %s, "%s"),'
                     % (_java_str(key), seed, _java_str(ij)))
        lines.append('            "expected: %s");%s' % (_java_str(target)[:60], todo))
        _seed(key)
        return
    verb = intent.get("verb") or "click"
    value = _java_str(intent.get("value", ""))
    lines.append('        heal.act("%s", "%s", %s, "%s", "%s");%s'
                 % (_java_str(key), verb, seed, _java_str(ij), value, todo))
    _seed(key)


def generate_selfhealing_test_class(story, cases, pkg, seed_sink=None):
    """Emit a TestNG class for one story DETERMINISTICALLY from compiled intents
    (no LLM writing Java). Cases run by priority: logged-out → login → app.
    `seed_sink`, if given, collects {key: {by, value}} for known seed locators."""
    sid = str(story.get("id", "0"))
    cls = "Story%sTests" % re.sub(r"[^A-Za-z0-9]", "", sid)
    L = []
    L.append("package %s.tests;" % pkg)
    L.append("")
    L.append("import %s.core.*;" % pkg)
    L.append("import org.testng.annotations.Test;")
    L.append("")
    L.append("/** Story %s — %s */" % (sid, _java_str(story.get("title", ""))))
    L.append("public class %s extends BaseTest {" % cls)
    for ci, c in enumerate(cases):
        bucket = c["bucket"]; pr = c.get("priority", bucket * 100 + ci)
        mname = "tc_%d_%s" % (pr, _java_ident(c.get("title", ""), "case%d" % ci))
        L.append("")
        L.append('    @Test(priority = %d, description = "%s")'
                 % (pr, _java_str(c.get("title", ""))))
        L.append("    public void %s() {" % mname)
        L.append('        // [%s · %s-page]' % (c["ctype"], c["page_context"]))
        if c.get("needs_review"):
            L.append('        // NEEDS REVIEW: AI compile was unavailable for this case '
                     '(e.g. provider rate limit) — this is a placeholder visibility '
                     'check only, verify the real steps manually.')
        if bucket < 2:
            L.append("        openLoginPage();")
        elif bucket == 2:
            L.append("        openLoginPage();")
            L.append("        performLogin();")
        else:
            L.append("        driver.get(Config.BASE_URL);")
        # emit intents (for the login-transition case, skip its action intents —
        # performLogin() already did the real login — but keep its assertions)
        for ii, intent in enumerate(c.get("intents", [])):
            if bucket == 2 and intent.get("role") == "action":
                continue
            _emit_intent(L, "%s.%d.%d" % (sid, ci, ii), intent, seed_sink=seed_sink,
                         app_page=(bucket == 3))
        L.append("    }")
    L.append("}")
    return "\n".join(L) + "\n", cls


def _sh_write_testng(out_dir, pkg, m):
    """(Re)write testng.xml from EVERY story recorded in the manifest so a push
    after a partial/resumed run carries all generated classes, not just this run's."""
    classes = sorted({r.get("test_class") for r in (m.get("stories") or {}).values()
                      if r.get("test_class")})
    items = "\n".join('      <class name="%s.tests.%s"/>' % (pkg, c) for c in classes)
    with open(os.path.join(out_dir, "testng.xml"), "w", encoding="utf-8") as f:
        f.write('<?xml version="1.0" encoding="UTF-8"?>\n'
                '<!DOCTYPE suite SYSTEM "https://testng.org/testng-1.0.dtd">\n'
                '<suite name="QA Studio Self-Healing Suite" verbose="1">\n'
                '  <test name="Sequenced Tests"><classes>\n%s\n  </classes></test>\n'
                '</suite>\n' % items)


def _write_seed_locators(out_dir, seeds, cb=None):
    """Merge generation-time seed locators into the COMMITTED locators.json
    without clobbering any entry already resolved/healed at runtime (or hand-
    edited). New keys are added with source='seed'; existing keys are left as-is."""
    cb = cb or (lambda *a, **k: None)
    path = os.path.join(out_dir, "locators.json")
    data = {}
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
    except Exception:
        data = {}
    added = 0
    for key, seed in (seeds or {}).items():
        if key not in data:            # never overwrite a healed / edited locator
            data[key] = {"by": seed["by"], "value": seed["value"], "source": "seed"}
            added += 1
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        return 0
    if added:
        cb("Seeded %d new locator(s) into locators.json (%d total)." % (added, len(data)), "dim")
    return added


def _prune_generated_orphans(out_dir, pkg, m, cb=None):
    """Delete generated .java files QA Studio no longer owns — e.g. a renamed
    framework class (AnthropicClient -> AiClient) or a test class whose story was
    dropped from the manifest. STRICTLY scoped to the managed com/qastudio core &
    tests packages; never touches any other package or user-authored code."""
    cb = cb or (lambda *a, **k: None)
    pkg_path = pkg.replace(".", "/")
    core_dir = os.path.join(out_dir, "src", "main", "java", pkg_path, "core")
    tests_dir = os.path.join(out_dir, "src", "test", "java", pkg_path, "tests")
    core_owned = {"DriverFactory", "Config", "LocatorStore", "AiClient", "Healer"}
    tests_owned = {"BaseTest"} | {r.get("test_class")
                                  for r in (m.get("stories") or {}).values()
                                  if r.get("test_class")}
    removed = []
    for d, owned in ((core_dir, core_owned), (tests_dir, tests_owned)):
        if not os.path.isdir(d):
            continue
        for fn in sorted(os.listdir(d)):
            if fn.endswith(".java") and fn[:-5] not in owned:
                try:
                    os.remove(os.path.join(d, fn))
                    removed.append(os.path.relpath(os.path.join(d, fn), out_dir))
                except Exception:
                    pass
    if removed:
        cb("Pruned %d stale file(s): %s" % (len(removed), ", ".join(removed)), "dim")
    return removed


def _sha1_text(s):
    import hashlib
    return hashlib.sha1((s or "").encode("utf-8")).hexdigest()


def _guarded_write_test_class(path, content, prev_hash, cb=None):
    """Write a generated test class UNLESS the on-disk file was hand-edited since
    we last generated it. We know our own last output by the hash recorded in the
    manifest; if the file on disk no longer matches that hash (and isn't already
    identical to the fresh generation), a human changed it — so we KEEP their file,
    drop a `<name>.java.new` sibling carrying the fresh generation for them to diff
    and merge, and warn. Returns (wrote_bool, hash_to_record).
    Note: the `.new` sibling ends in `.new`, not `.java`, so orphan-pruning leaves
    it alone."""
    cb = cb or (lambda *a, **k: None)
    new_hash = _sha1_text(content)
    if prev_hash and os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                cur = f.read()
        except Exception:
            cur = None
        cur_hash = _sha1_text(cur) if cur is not None else None
        if cur_hash and cur_hash != prev_hash and cur_hash != new_hash:
            side = path + ".new"
            try:
                with open(side, "w", encoding="utf-8") as f:
                    f.write(content)
            except Exception:
                pass
            cb("Kept your manual edits to %s — fresh generation saved as %s "
               "(diff & merge by hand)." % (os.path.basename(path),
                                            os.path.basename(side)), "ok")
            return False, prev_hash
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return True, new_hash


def build_selfhealing_project(out_dir, sequenced, base_url, login=None,
                              group_id="com.qastudio", artifact_id="automation-tests",
                              cb=None, should_stop=lambda: False, orig_tcs=None):
    """Write a full self-healing Maven/TestNG project (no browser was driven).
    `sequenced` = output of validate_and_sequence_suite(). Returns written paths.
    `orig_tcs` = {story_id: [original test-case dicts]} so each generated story is
    recorded in the manifest — that's what lets a stopped/closed run RESUME instead
    of regenerating from scratch (classify_selection reads the manifest)."""
    cb = cb or (lambda *a, **k: None)
    pkg = group_id
    pkg_path = pkg.replace(".", "/")
    src_main = os.path.join(out_dir, "src", "main", "java", pkg_path, "core")
    src_test = os.path.join(out_dir, "src", "test", "java", pkg_path, "tests")
    res_dir = os.path.join(out_dir, "src", "test", "resources")
    for d in (src_main, src_test, res_dir):
        os.makedirs(d, exist_ok=True)
    login_url = (login or {}).get("url") or base_url
    written = []
    orig_tcs = orig_tcs or {}
    m = load_manifest(out_dir)          # resume: record of already-generated stories
    m.setdefault("stories", {})
    m["manifest_version"] = 1

    def _w(path, content):
        # write-if-changed: skip files whose content is identical so re-runs don't
        # churn git blame / diffs / mtimes on the deterministic framework core.
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    if f.read() == content:
                        return
        except Exception:
            pass
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        written.append(os.path.relpath(path, out_dir))

    _ai_prov, _ai_base, _ai_model = _healer_ai_meta()
    cb(f"Writing self-healing framework (Config, Healer, AI client \u00b7 {_ai_prov})\u2026", "dim")
    _w(os.path.join(out_dir, "pom.xml"), _sh_pom(group_id, artifact_id))
    _w(os.path.join(out_dir, ".gitignore"), _sh_gitignore())
    _w(os.path.join(out_dir, "README.md"), _sh_readme(base_url, login_url))
    _w(os.path.join(src_main, "DriverFactory.java"), _driver_factory(pkg))
    _w(os.path.join(src_main, "Config.java"),
       _sh_config(pkg, base_url, login_url, _ai_prov, _ai_base, _ai_model))
    _w(os.path.join(src_main, "LocatorStore.java"), _sh_locator_store(pkg))
    _w(os.path.join(src_main, "AiClient.java"), _sh_ai_client(pkg))
    _w(os.path.join(src_main, "Healer.java"), _sh_healer(pkg))
    _w(os.path.join(src_test, "BaseTest.java"), _sh_base_test(pkg))
    _w(os.path.join(res_dir, "harvest.js"), _HARVEST_JS)
    # Environment-specific config: the committed template is always refreshed;
    # the real (git-ignored) config.properties is written ONCE so a re-run never
    # clobbers a user's environment edits.
    _w(os.path.join(out_dir, "config.properties.example"), _sh_config_properties_example())
    _cfg_props = os.path.join(out_dir, "config.properties")
    if not os.path.exists(_cfg_props):
        with open(_cfg_props, "w", encoding="utf-8") as f:
            f.write(_sh_config_properties(base_url, login_url, _ai_prov, _ai_base, _ai_model))
        written.append("config.properties")

    test_classes = []
    seed_sink = {}                      # generation-time locators → committed locators.json
    _todo_total = 0                     # locators marked // TODO (resolved at runtime)
    for entry in sequenced:
        if should_stop():
            break
        story = entry["story"]
        if not entry.get("cases"):
            continue
        sid = str(story.get("id"))
        cb(f"  generating tests for story {story.get('id')} "
           f"({len(entry['cases'])} case(s))", "dim")
        java, cls = generate_selfhealing_test_class(story, entry["cases"], pkg,
                                                    seed_sink=seed_sink)
        _todo_total += java.count(", null,")   # null seed = resolved at runtime (TODO)
        tpath = os.path.join(src_test, "%s.java" % cls)
        prior = m.get("stories", {}).get(sid) or {}
        wrote, chash = _guarded_write_test_class(tpath, java, prior.get("hash"), cb)
        if wrote:
            written.append(os.path.relpath(tpath, out_dir))
        test_classes.append(cls)
        # Record + persist this story immediately so a stop / pause / app-close keeps
        # it; the keys mirror what classify_selection() computes from the original
        # test cases, so a re-run recognises it as done and skips it. `hash` records
        # OUR last generated content so a later run can detect manual edits; the
        # provenance fields carry over unchanged when we KEEP a hand-edited file.
        _stamp = time.strftime("%Y-%m-%dT%H:%M:%S")
        m["stories"][sid] = {
            "test_class": cls,
            "hash": chash,
            "provider": _ai_prov if wrote else prior.get("provider", _ai_prov),
            "model": _ai_model if wrote else prior.get("model", _ai_model),
            "generatedAt": _stamp if wrote else prior.get("generatedAt", _stamp),
            "test_cases": {_tc_key(tc): _method_name(tc.get("title", ""))
                           for tc in orig_tcs.get(sid, [])},
        }
        try:
            save_manifest(out_dir, m)
            _sh_write_testng(out_dir, pkg, m)
        except Exception:
            pass

    # Final testng + manifest across EVERY recorded class (this run + prior runs).
    _sh_write_testng(out_dir, pkg, m)
    save_manifest(out_dir, m)
    # Seed the committed locators.json (merge, never clobber healed entries), then
    # remove any generated file we no longer own (renamed core class, dropped story).
    _write_seed_locators(out_dir, seed_sink, cb)
    _prune_generated_orphans(out_dir, pkg, m, cb)
    cb(f"TODO_LIVE: {_todo_total}", "meta")   # reconcile the live counter to the emitted total
    cb(f"TODO: {_todo_total} locator(s) to resolve at runtime "
       f"(the rest are seeded).", "warn")
    cb(f"Wrote {len(written)} files, {len(test_classes)} test class(es) this run.", "ok")
    return written


def generate_and_push_selfhealing(out_dir, stories_payload, base_url, login=None,
                                  group_id="com.qastudio", artifact_id="automation-tests",
                                  cb=None, should_stop=lambda: False, want_ai=True,
                                  on_error=None, gate=None, target="selenium"):
    """End-to-end no-browser path: validate+sequence → generate self-healing
    project. (Push is done separately via push_to_git, as today.)
    on_error/gate enable pause-on-error and manual pause (see
    validate_and_sequence_suite)."""
    cb = cb or (lambda *a, **k: None)
    cb("Validating and sequencing test cases (no browser)\u2026", "info")
    sequenced = validate_and_sequence_suite(stories_payload, log=cb, want_ai=want_ai,
                                            should_stop=should_stop, on_error=on_error,
                                            gate=gate)
    if should_stop():
        return []
    # Original test cases per story → recorded in the manifest for resume support.
    _orig_tcs = {str(sp.get("story", {}).get("id")): (sp.get("test_cases", []) or [])
                 for sp in stories_payload}
    if target and target != "selenium":
        # JS targets (playwright | cypress) reuse the framework-agnostic IR and the
        # same manifest/seed/prune/hash-guard behaviour, via a separate emitter.
        import automation_targets
        _ai_prov, _ai_base, _ai_model = _healer_ai_meta()
        cfg = {"base_url": base_url,
               "login_url": (login or {}).get("url") or base_url,
               # Login-page language for the generated auth helpers — must
               # match the generated tests' expected-message language (the
               # frozen login URL can carry a stale ui_locales=en; see the
               # auth.js template comment for the live failure this caused).
               "ui_locale": ("ar" if _is_arabic_out() else "en"),
               "ai_provider": _ai_prov, "ai_base_url": _ai_base, "ai_model": _ai_model}
        return automation_targets.build(target, out_dir, sequenced, cfg,
                                        _seed_locator_for_intent, _HARVEST_JS,
                                        cb=cb, should_stop=should_stop, orig_tcs=_orig_tcs)
    return build_selfhealing_project(out_dir, sequenced, base_url, login,
                                     group_id, artifact_id, cb, should_stop,
                                     orig_tcs=_orig_tcs)
