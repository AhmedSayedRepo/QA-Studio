"""regression.py — Regression Plan tab for QA Studio.

Builds a regression plan from an existing sprint + selected user stories,
estimates effort from the test cases that ALREADY exist in the chosen test plan
(reference-only — nothing is generated), balances work across named resources,
and exports to Word / Excel / JSON / PDF.

Integration hooks live in main.py (import, T.NAV registration, rail clickable,
render routing). Shared UI helpers and the theme are reused from main.py via a
deferred import so there is no circular import at load time.
"""
import os, re, json, threading
import time as _time
import contextlib as _ctxlib

# ── Lightweight perf logging ────────────────────────────────────────────────
# The app runs under pythonw.exe (no console), so prints are invisible. Instead
# we APPEND timings to a log file next to the app: "qa_perf.log". On by default;
# set QASTUDIO_PERF=0 to silence. Use this to find the real hotspot behind UI
# slowness (build_rows, table refresh, full render). One perf_counter() pair per
# block when enabled.
_PERF_ON = os.environ.get("QASTUDIO_PERF", "1") != "0"
try:
    _PERF_LOG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qa_perf.log")
except Exception:
    _PERF_LOG = os.path.join(os.path.expanduser("~"), "qa_perf.log")

def _perf_log(msg):
    if not _PERF_ON:
        return
    try:
        with open(_PERF_LOG, "a", encoding="utf-8") as _fh:
            _fh.write(f"{_time.strftime('%H:%M:%S')}  {msg}\n")
    except Exception:
        pass

@_ctxlib.contextmanager
def _perf(label):
    if not _PERF_ON:
        yield
        return
    _t0 = _time.perf_counter()
    try:
        yield
    finally:
        _perf_log(f"{label}: {(_time.perf_counter() - _t0) * 1000:.0f} ms")
from datetime import datetime

import flet as ft
import theme as T
import engine as E

# ── Effort model (HARDCODED — change here if your team's numbers differ) ───────
AVG_MINUTES_PER_CASE = 8          # manual execution time per existing test case
DEFAULT_PRIORITY     = 3          # used when a story has no ADO priority set
PRIORITY_BOOST = {1: 1.30, 2: 1.15, 3: 1.00, 4: 0.90}
_PRI_FULL = {1: "P1 (highest)", 2: "P2", 3: "P3", 4: "P4 (lowest)"}


def _digits_only():
    """Numeric-only input filter so count fields reject letters/symbols."""
    try:
        return ft.NumbersOnlyInputFilter()
    except Exception:
        try:
            return ft.InputFilter(regex_string=r"[0-9]", allow=True)
        except Exception:
            return None


def _keep_scroll(app, off):
    """Re-assert a scroll offset after an in-place table rebuild so collapse /
    pagination doesn't jump when the visible content height changes."""
    if not off:
        return
    col = getattr(app, "_left_scroll", None)
    if col is None:
        return

    def _do():
        try:
            col.scroll_to(offset=off, duration=0)
        except Exception:
            pass
    try:
        app.ui_safe(_do)
    except Exception:
        pass
    # a couple of delayed shots in case the layout settles a frame later
    for _d in (0.05, 0.16):
        try:
            threading.Timer(_d, lambda: app.ui_safe(_do)).start()
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════════════════════
#  DISK CACHE  — persist the per-project suite/count/meta/feature caches so the
#  FIRST generate after an app restart is fast (no ~90 s recount) when nothing in
#  Azure changed. The "Regenerate" button force-refreshes, so stale data always
#  has an escape hatch. Keyed by project; JSON (int keys stored as strings).
# ═══════════════════════════════════════════════════════════════════════════════
_CACHE_DIR = os.path.join(os.path.expanduser("~"), ".qa_tool")
_CACHE_FILE = os.path.join(_CACHE_DIR, "reg_cache.json")


def _cache_load(app):
    """Load the saved caches for the current project once (per project)."""
    proj = getattr(app, "project", None)
    if not proj:
        return
    if getattr(app, "_reg_cache_loaded_project", None) == proj:
        return
    app._reg_cache_loaded_project = proj
    # Reset first so switching projects never carries another project's caches.
    app._reg_case_count_cache = {}
    app._reg_meta_cache = {}
    app._reg_suite_cache = {}
    app._reg_feature_name_cache = {}
    app._reg_story_features = {}
    try:
        with open(_CACHE_FILE, "r", encoding="utf-8") as f:
            d = (json.load(f) or {}).get(proj) or {}
        app._reg_case_count_cache = {int(k): v for k, v in d.get("counts", {}).items()}
        app._reg_meta_cache = {int(k): v for k, v in d.get("meta", {}).items()}
        app._reg_suite_cache = {int(k): {int(sk): sv for sk, sv in v.items()}
                                for k, v in d.get("suites", {}).items()}
        app._reg_feature_name_cache = {int(k): v
                                       for k, v in d.get("fnames", {}).items()}
        app._reg_story_features = {int(k): tuple(v)
                                   for k, v in d.get("features", {}).items()}
    except Exception:
        pass


def _cache_save(app):
    """Persist the current project's caches to disk (best-effort)."""
    proj = getattr(app, "project", None)
    if not proj:
        return
    try:
        os.makedirs(_CACHE_DIR, exist_ok=True)
        try:
            with open(_CACHE_FILE, "r", encoding="utf-8") as f:
                allp = json.load(f) or {}
        except Exception:
            allp = {}
        allp[proj] = {
            "counts": {str(k): v for k, v
                       in (getattr(app, "_reg_case_count_cache", {}) or {}).items()},
            "meta": {str(k): v for k, v
                     in (getattr(app, "_reg_meta_cache", {}) or {}).items()},
            "suites": {str(k): {str(sk): sv for sk, sv in v.items()} for k, v
                       in (getattr(app, "_reg_suite_cache", {}) or {}).items()},
            "fnames": {str(k): v for k, v
                       in (getattr(app, "_reg_feature_name_cache", {}) or {}).items()},
            "features": {str(k): list(v) for k, v
                         in (getattr(app, "_reg_story_features", {}) or {}).items()},
        }
        with open(_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(allp, f)
    except Exception:
        pass


def clear_caches(app):
    """Clear the in-memory AND on-disk regression/sprint caches (case counts,
    work-item meta, suites, feature names). Safe to call any time — the caches
    are just a speed-up and are rebuilt from Azure on the next generate."""
    for attr in ("_reg_case_count_cache", "_reg_meta_cache", "_reg_suite_cache",
                 "_reg_feature_name_cache", "_reg_story_features"):
        try:
            setattr(app, attr, {})
        except Exception:
            pass
    app._reg_cache_loaded_project = None
    try:
        if os.path.exists(_CACHE_FILE):
            with open(_CACHE_FILE, "w", encoding="utf-8") as f:
                f.write("{}")
    except Exception:
        pass


def set_perf(on):
    """Enable/disable appending timings to qa_perf.log at runtime."""
    global _PERF_ON
    _PERF_ON = bool(on)
    return _PERF_ON


def perf_on():
    return bool(_PERF_ON)


# ═══════════════════════════════════════════════════════════════════════════════
#  DATA GATHERING  (Azure DevOps — reference existing only)
# ═══════════════════════════════════════════════════════════════════════════════
def _fetch_meta(app, ids):
    """Fetch work item metadata with per-story caching (fast on re-generate)."""
    cache = getattr(app, "_reg_meta_cache", {})
    missing = [i for i in ids if i not in cache]
    if missing:
        org, proj = E.AZURE_ORG, app.project
        for i in range(0, len(missing), 200):
            batch = missing[i:i + 200]
            joined = ",".join(map(str, batch))
            url = (f"https://dev.azure.com/{org}/{proj}/_apis/wit/workitems"
                   f"?ids={joined}"
                   f"&fields=System.Id,System.Title,System.State,"
                   f"Microsoft.VSTS.Common.Priority,System.Parent&api-version=7.0")
            try:
                data = E._azure_get(url)
            except Exception:
                data = {}
            for w in data.get("value", []):
                f = w.get("fields", {})
                try:
                    pri = int(f.get("Microsoft.VSTS.Common.Priority", DEFAULT_PRIORITY)
                              or DEFAULT_PRIORITY)
                except Exception:
                    pri = DEFAULT_PRIORITY
                cache[int(w["id"])] = {"title": f.get("System.Title", ""),
                                       "state": f.get("System.State", "Unknown"),
                                       "priority": pri,
                                       "parent_id": f.get("System.Parent")}
        app._reg_meta_cache = cache
    return {i: cache.get(i, {}) for i in ids}


def _fetch_feature_names(app, parent_ids):
    """Feature (parent work-item) titles, cached per id on app._reg_feature_name_cache
    so re-selecting plans / re-generating never re-requests a name we already have."""
    ids = [i for i in parent_ids if i]
    if not ids:
        return {}
    cache = getattr(app, "_reg_feature_name_cache", None)
    if cache is None:
        cache = {}
        app._reg_feature_name_cache = cache
    names = {i: cache[i] for i in ids if i in cache}
    missing = [i for i in ids if i not in cache]
    if missing:
        org, proj = E.AZURE_ORG, app.project
        for i in range(0, len(missing), 200):
            batch = missing[i:i + 200]
            joined = ",".join(map(str, batch))
            url = (f"https://dev.azure.com/{org}/{proj}/_apis/wit/workitems"
                   f"?ids={joined}&fields=System.Id,System.Title&api-version=7.0")
            try:
                data = E._azure_get(url)
            except Exception:
                data = {}
            for w in data.get("value", []):
                f2 = w.get("fields", {})
                nm = f2.get("System.Title", str(w["id"]))
                names[int(w["id"])] = nm
                cache[int(w["id"])] = nm
            # Every caller runs off the UI thread. Yield the GIL each batch so a
            # large background resolve can't monopolize it and freeze the plan
            # table's collapse / pagination for the minute-or-two it runs.
            _time.sleep(0.04)
    return names


def _resolve_reg_features(app, stories, gen=None):
    """Resolve each story's parent → Feature (id + title), cache it, and tag the
    stories with feature_id/feature_name. Called from _reload_plan_stories DURING
    story load (so it's covered by the 'Loading stories…' wait) instead of as a
    separate background pass — by Generate time the caches are warm, so build_rows
    does no feature network and nothing lingers to freeze the plan table later.
    Aborts early if the selection changed (gen token)."""
    try:
        feat_cache = dict(getattr(app, "_reg_story_features", {}) or {})
        sids = [int(s["id"]) for s in stories if int(s["id"]) not in feat_cache]
        if not sids:
            return
        org, proj = E.AZURE_ORG, app.project
        parent_map = {}
        for i in range(0, len(sids), 200):
            if gen is not None and gen != getattr(app, "_reg_stories_gen", gen):
                return                  # selection changed -> stop
            joined = ",".join(map(str, sids[i:i + 200]))
            url = (f"https://dev.azure.com/{org}/{proj}/_apis/wit/workitems"
                   f"?ids={joined}&fields=System.Id,System.Parent&api-version=7.0")
            try:
                data = E._azure_get(url)
            except Exception:
                data = {}
            for w in data.get("value", []):
                pid2 = w.get("fields", {}).get("System.Parent")
                if pid2:
                    parent_map[int(w["id"])] = pid2
            _time.sleep(0.04)           # yield GIL during the load
        names = _fetch_feature_names(app, list(set(parent_map.values())))
        for sid, pid2 in parent_map.items():
            feat_cache[sid] = (pid2, names.get(pid2, ""))
        app._reg_story_features = feat_cache
        for s in stories:
            fdata = feat_cache.get(int(s["id"]), (None, ""))
            s["feature_id"] = fdata[0]
            s["feature_name"] = fdata[1]
    except Exception:
        pass


def _fetch_cp_complexity(app, ids):
    """Rough 'amount of work' units per story from its CONTENT — acceptance
    criteria (weighted most), description, and title word counts. Lets the sprint
    estimate reflect each story's complexity/size instead of a random number.
    Returns {story_id: units(float)}; falls back to {} on error."""
    ids = [int(i) for i in ids if i]
    if not ids:
        return {}
    org, proj = E.AZURE_ORG, app.project
    flds = ("System.Id,System.Title,System.Description,"
            "Microsoft.VSTS.Common.AcceptanceCriteria")
    out = {}

    def _plain(html):
        s = re.sub(r"<[^>]+>", " ", html or "")
        return re.sub(r"\s+", " ", s).strip()

    for i in range(0, len(ids), 200):
        batch = ids[i:i + 200]
        url = (f"https://dev.azure.com/{org}/{proj}/_apis/wit/workitems"
               f"?ids={','.join(map(str, batch))}&fields={flds}&api-version=7.0")
        try:
            data = E._azure_get(url)
        except Exception:
            data = {}
        for w in data.get("value", []):
            f = w.get("fields", {})
            crit = _plain(f.get("Microsoft.VSTS.Common.AcceptanceCriteria", ""))
            desc = _plain(f.get("System.Description", ""))
            title = f.get("System.Title", "") or ""
            units = (len(crit.split()) * 1.0
                     + len(desc.split()) * 0.4
                     + len(title.split()) * 0.2)
            out[int(w["id"])] = max(1.0, units)
        _time.sleep(0.05)   # background; yield GIL so the Sprint table stays responsive
    return out


def _loading_field(text):
    """A bordered field that shows a small spinner + message — used while plans /
    stories are being fetched in the background."""
    return ft.Container(
        ft.Row([ft.ProgressRing(width=16, height=16, stroke_width=2, color=T.VIOLET),
                ft.Text(text, size=13, color=T.INK_3, expand=True)],
               spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
        padding=ft.Padding.symmetric(vertical=12, horizontal=12),
        bgcolor=T.CARD, border=ft.Border.all(1, T.BORDER), border_radius=T.R)


def _discover_suite_map(app, plan_id, story_ids):
    """One HTTP call per plan → {story_id: suite_id}. Result cached."""
    story_ids = set(story_ids)
    cache = getattr(app, "_reg_suite_cache", {})
    if plan_id in cache:
        return cache[plan_id]
    url = (f"https://dev.azure.com/{E.AZURE_ORG}/{app.project}"
           f"/_apis/testplan/Plans/{plan_id}/Suites?api-version=7.0&$expand=true")
    try:
        data = E._azure_get(url)
    except Exception:
        return {}
    smap = {}
    for suite in data.get("value", []):
        sid = suite.get("id")
        req_id = suite.get("requirementId")
        if req_id and int(req_id) in story_ids:
            smap[int(req_id)] = sid
            continue
        name = suite.get("name", "")
        try:
            cand = int(name.split(":")[0].strip())
            if cand in story_ids:
                smap[cand] = sid
        except (ValueError, IndexError):
            pass
    try:
        app._reg_suite_cache[plan_id] = smap
    except Exception:
        pass
    return smap


def _count_cases(app, selected, progress=None):
    """Phase 1: suite map per plan (1 call/plan, parallel, cached).
    Phase 2: test case count per suite (parallel, up to 16 workers).

    progress(done, total) — optional callback invoked as suites are counted, so the
    UI can show live progress instead of a blank 'Generating…'."""
    import concurrent.futures as _cf
    counts = {int(s["id"]): 0 for s in selected}
    by_plan = {}
    for s in selected:
        by_plan.setdefault(s.get("plan_id"), []).append(int(s["id"]))

    tasks = []
    plan_items = [(pid, sids) for pid, sids in by_plan.items() if pid]
    with _cf.ThreadPoolExecutor(max_workers=min(8, len(plan_items) or 1)) as ex:
        futs = {ex.submit(_discover_suite_map, app, pid, sids): (pid, sids)
                for pid, sids in plan_items}
        for fut, (pid, sids) in futs.items():
            smap = fut.result()
            for sid in sids:
                suite_id = smap.get(sid)
                if suite_id:
                    tasks.append((sid, pid, suite_id))

    # Per-suite count cache: a suite's case count doesn't change between
    # generates in a session, so re-generating (or tweaking resources and
    # re-running) costs ZERO network for suites we've already counted. This is
    # the single biggest lever — the first count of 433 suites took ~92 s.
    cache = getattr(app, "_reg_case_count_cache", None)
    if cache is None:
        cache = {}
        app._reg_case_count_cache = cache

    todo = []
    for sid, pid, suite_id in tasks:
        if suite_id in cache:
            counts[sid] = cache[suite_id]
        else:
            todo.append((sid, pid, suite_id))

    def _one(t):
        sid, pid, suite_id = t
        try:
            n = len(E.fetch_test_cases_for_suite(app.project, pid, suite_id))
        except Exception:
            n = 0
        return sid, suite_id, n

    if todo:
        # 16 workers: 24 measured SLOWER (Azure rate-limits) and starved the UI
        # thread's GIL share, which is what made renders freeze during a parallel
        # generate. Fewer concurrent JSON parses = a more responsive UI.
        _total, _done = len(todo), 0
        if progress:
            try: progress(0, _total)
            except Exception: pass
        with _cf.ThreadPoolExecutor(max_workers=min(16, _total)) as ex:
            _futs = [ex.submit(_one, t) for t in todo]
            for _f in _cf.as_completed(_futs):    # as-completed -> accurate progress
                sid, suite_id, n = _f.result()
                counts[sid] = n
                cache[suite_id] = n
                _done += 1
                if progress and (_done % 5 == 0 or _done == _total):
                    try: progress(_done, _total)
                    except Exception: pass
                if _done % 8 == 0:
                    _time.sleep(0.01)             # yield GIL so the UI stays responsive
    return counts


def build_rows(app, selected, progress=None):
    _cache_load(app)   # warm caches from disk on the first generate of a session
    ids = [int(s["id"]) for s in selected]
    with _perf(f"build_rows.fetch_meta ({len(ids)} stories)"):
        meta = _fetch_meta(app, ids)
    with _perf(f"build_rows.count_cases ({len(selected)} stories)"):
        counts = _count_cases(app, selected, progress=progress)
    _cache_save(app)   # persist freshly-fetched counts/meta/features for next launch
    story_features = getattr(app, "_reg_story_features", {})
    feat_cache = getattr(app, "_reg_feature_name_cache", {}) or {}
    # Each story's parent work item IS its Feature. _fetch_meta already pulled
    # System.Parent (free), so the feature id is known with no extra request. The
    # feature NAME is read from cache only here — resolving uncached names is done
    # OFF the generate path (background, below) so a cold cache can never block the
    # UI / freeze the nav while the heavy test-case counting is running.
    parent_of = {}
    for s in selected:
        sid = int(s["id"])
        p = meta.get(sid, {}).get("parent_id")
        if p:
            parent_of[sid] = int(p)
    rows = []
    for s in selected:
        sid = int(s["id"])
        m = meta.get(sid, {})
        pri = m.get("priority", DEFAULT_PRIORITY)
        cases = counts.get(sid, 0)
        boost = PRIORITY_BOOST.get(pri, 1.0)
        hours = round(cases * (AVG_MINUTES_PER_CASE / 60.0) * boost, 2)
        _pfeat = parent_of.get(sid)
        fid = (s.get("feature_id") or story_features.get(sid, (None, ""))[0] or _pfeat)
        fname = (s.get("feature_name") or story_features.get(sid, (None, ""))[1]
                 or (feat_cache.get(_pfeat) if _pfeat else "") or "")
        rows.append({"id": sid, "title": m.get("title", "") or s.get("title", ""),
                     "state": m.get("state", "Unknown"), "priority": pri,
                     "cases": cases, "boost": boost, "hours": hours,
                     "plan_id": s.get("plan_id"), "assignee": "",
                     "feature_id": fid, "feature_name": fname})
    # Feature ids come from meta (free); names normally come from the cache that
    # _reload_plan_stories warmed during story load (fast path -> the loop below is
    # a no-op). Fallback: if that pre-warm didn't apply to THIS selection — after
    # Regenerate / Clear caches (which drop the feature caches), or a selection that
    # didn't reload stories — resolve the still-missing names here so the plan still
    # groups by feature. No network happens when the cache is already warm.
    _missing = sorted({r["feature_id"] for r in rows
                       if r.get("feature_id") and not r.get("feature_name")})
    if _missing:
        try:
            _names = _fetch_feature_names(app, _missing)
            for _r in rows:
                if not _r.get("feature_name"):
                    _r["feature_name"] = _names.get(_r.get("feature_id"), "")
        except Exception:
            pass
    return rows


def assign_resources(rows, names):
    """Greedy balance: largest story → least-loaded resource. Sets r['assignee'].
    Returns {name: total_hours}.

    Stories with no test cases have 0 estimated hours; balancing on hours alone
    dumps all of them onto whichever resource is least-loaded (it stays least-
    loaded after a 0-hour add). So zero-effort stories are instead spread by head-
    count, and hour-ties are broken by head-count too, keeping the split even."""
    if not names:
        for r in rows:
            r["assignee"] = ""
        return {}
    load = {n: 0.0 for n in names}
    cnt = {n: 0 for n in names}
    for r in sorted(rows, key=lambda x: -x["hours"]):
        if r["hours"] > 0:
            n = min(names, key=lambda k: (load[k], cnt[k]))   # by effort, then count
        else:
            n = min(names, key=lambda k: cnt[k])              # no cases -> even by count
        r["assignee"] = n
        load[n] = round(load[n] + r["hours"], 2)
        cnt[n] += 1
    return load


def plan_payload(app):
    if getattr(app, "_reg_mode", "existing") == "create":
        return _cp_payload(app)
    rows = [dict(r) for r in (app._reg_selected_rows or [])]
    names = list(app._reg_res_names or [])
    load = assign_resources(rows, names)
    count = app._reg_res_count or len(names) or 1
    total_cases = sum(r["cases"] for r in rows)
    total_hours = round(sum(r["hours"] for r in rows), 2)
    per_person = round(total_hours / count, 2) if count else total_hours
    workload = []
    if names:
        cnt = {n: 0 for n in names}
        ccases = {n: 0 for n in names}
        for r in rows:
            a = r.get("assignee")
            if a in cnt:
                cnt[a] += 1
                ccases[a] += r["cases"]
        workload = [{"name": n, "stories": cnt[n], "cases": ccases[n],
                     "hours": round(load.get(n, 0.0), 2)} for n in names]
    plans = list(app._reg_plans_selected or [])
    # Show only the sprint number ("Sprint 22") in exports/email instead of the
    # full ADO test-plan name ("<Project>_Sprint 22"). Fall back to the full name
    # if a plan has no recognizable sprint number.
    plan_names = ", ".join(_sprint_num(p.get("name", "")) or p.get("name", "")
                           for p in plans) \
        or (_sprint_num(getattr(app, "plan_name", "") or "")
            or (getattr(app, "plan_name", "") or ""))
    plan_ids = ", ".join(str(p["id"]) for p in plans)
    return {"generated": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "project": app.project, "plan_id": plan_ids,
            "plan_name": plan_names, "plans": plans,
            "report_title": "Regression Test Plan", "mode": "existing",
            "avg_minutes_per_case": AVG_MINUTES_PER_CASE,
            "priority_boost": PRIORITY_BOOST, "resources_count": count,
            "resource_names": names, "stories": rows, "workload": workload,
            "total_stories": len(rows), "total_cases": total_cases,
            "total_hours": total_hours, "hours_per_person": per_person}


# ═══════════════════════════════════════════════════════════════════════════════
#  EXPORTERS
# ═══════════════════════════════════════════════════════════════════════════════
# ── email visual helpers ───────────────────────────────────────────────────────
_AV_COLORS = ["#3A57D6", "#1F9D57", "#6A52F0", "#1C80E0", "#C2860C", "#6A33A8"]

def _av(name):
    """Return (initial, color) for an assignee avatar."""
    nm = (name or "").strip()
    if not nm:
        return "—", "#9FA2B2"
    return nm[0].upper(), _AV_COLORS[sum(ord(c) for c in nm) % len(_AV_COLORS)]

def _email_pri(p):
    """Return (bg, fg, label) for a priority chip in the email."""
    return {1: ("#FCEBEC", "#E0474D"), 2: ("#FAF1DD", "#C2860C"),
            4: ("#E5F6EC", "#1F9D57")}.get(p, ("#F4F6FB", "#6E7180")) + (f"P{p}",)


def _wi_url(project, sid):
    """Azure DevOps work-item URL for a story id — used to hyperlink story IDs in
    every export (xlsx/docx/pdf/json) and the email so a reader can jump straight
    to the work item. Org comes from the configured Azure org; project is encoded
    so names with spaces still produce a valid link."""
    from urllib.parse import quote
    try:
        org = E.AZURE_ORG or ""
    except Exception:
        org = ""
    return (f"https://dev.azure.com/{quote(str(org), safe='')}/"
            f"{quote(str(project or ''), safe='')}/_workitems/edit/{sid}")


def _plan_html(d):
    """Polished, Outlook-safe HTML summary of the plan for the email body.

    Table-based with inline styles so it renders across Outlook / Gmail / Apple
    Mail. Colours track theme.py. The external-sender warning some inboxes show
    is injected by the mail server, not here.
    """
    # KPI tiles
    def _kpi(label, val, unit, fg, bg, bd, width="25%"):
        u = (f"<span style='font-size:12px;color:#9aa4b8;font-weight:600'> {unit}</span>"
             if unit else "")
        return (
            f"<td width='{width}' valign='top' style='padding:0 5px'>"
            f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
            f"style='background:{bg};border:1px solid {bd};border-radius:12px'><tr>"
            f"<td style='padding:13px 14px'>"
            f"<div style='font-size:10px;letter-spacing:.6px;text-transform:uppercase;"
            f"color:{fg};font-weight:700'>{label}</div>"
            f"<div style='font-family:Consolas,monospace;font-size:23px;font-weight:700;"
            f"color:{fg};margin-top:4px'>{val}{u}</div></td></tr></table></td>")

    # Sprint Plan estimates from an hours range (no existing test cases), so the
    # test-case KPI + column are hidden whenever the plan has zero cases.
    show_cases = (d.get("total_cases") or 0) > 0
    kw = "25%" if show_cases else "33.33%"
    kpi_cells = [_kpi("Stories", d["total_stories"], "", "#181A24", "#F6F8FC", "#EBEFF7", kw)]
    if show_cases:
        kpi_cells.append(_kpi("Test cases", d["total_cases"], "", "#181A24", "#F6F8FC", "#EBEFF7", kw))
    kpi_cells.append(_kpi("Total effort", d["total_hours"], "h", "#2940C2", "#E7ECFF", "#D6DEFF", kw))
    kpi_cells.append(_kpi("Per person", d["hours_per_person"], "h", "#1F9D57", "#E5F6EC", "#D2EEDF", kw))
    kpis = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        + "".join(kpi_cells) + "</tr></table>")

    # story rows
    def _story_html_row(r):
        bg, fg, lab = _email_pri(r["priority"])
        init, acol = _av(r.get("assignee", ""))
        who = (r.get("assignee") or "").strip()
        asg = (
            f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
            f"<td width='24' style='vertical-align:middle'>"
            f"<div style='width:24px;height:24px;line-height:24px;border-radius:50%;"
            f"background:{acol};color:#fff;text-align:center;font-size:10px;font-weight:700;"
            f"font-family:Segoe UI,Arial,sans-serif'>{init}</div></td>"
            f"<td style='padding-left:9px;font-size:13px;font-weight:600;color:#39435c;"
            f"white-space:nowrap;vertical-align:middle'>{who or '—'}</td></tr></table>")
        cases_cell = (
            f"<td style='padding:12px 8px;text-align:right;font-family:Consolas,monospace;"
            f"font-size:13.5px;color:#46506a'>{r['cases']}</td>") if show_cases else ""
        return (
            f"<tr style='border-top:1px solid #f0f3f9'>"
            f"<td style='padding:12px 14px;font-family:Consolas,monospace;font-size:13px;"
            f"font-weight:600;white-space:nowrap'>"
            f"<a href='{_wi_url(d['project'], r['id'])}' "
            f"style='color:#3A57D6;text-decoration:none'>{r['id']}</a></td>"
            f"<td style='padding:12px 8px;font-size:13.5px;font-weight:600;color:#1f2940'>"
            f"{(r['title'] or '—')}</td>"
            f"<td style='padding:12px 8px;text-align:center'>"
            f"<span style='font-family:Consolas,monospace;font-size:11px;font-weight:700;"
            f"padding:3px 8px;border-radius:6px;background:{bg};color:{fg}'>{lab}</span></td>"
            + cases_cell +
            f"<td style='padding:12px 8px;text-align:right;font-family:Consolas,monospace;"
            f"font-size:13.5px;font-weight:700;color:#1f2940'>{r['hours']}</td>"
            f"<td style='padding:12px 14px'>{asg}</td></tr>")

    from collections import OrderedDict as _ODh
    _feat_grps_h = _ODh()
    for r in d["stories"]:
        _feat_grps_h.setdefault(
            (r.get("feature_id"), r.get("feature_name") or "No Feature"), []).append(r)
    _html_ncols = 4 + (1 if show_cases else 0)  # id + title + pri + [cases] + hours + asg
    srows = []
    for (_fid_h, _fname_h), _fstories_h in _feat_grps_h.items():
        _fid_label_h = f"[{_fid_h}]  " if _fid_h else ""
        srows.append(
            f"<tr style='background:#EEE8FF;border-top:2px solid #C8BAFF'>"
            f"<td colspan='{_html_ncols + 2}' style='padding:7px 14px;font-size:10.5px;"
            f"font-weight:700;color:#6A4DFF;letter-spacing:.4px;"
            f"text-transform:uppercase'>"
            f"Feature: {_fid_label_h}{_fname_h}</td></tr>")
        for r in _fstories_h:
            srows.append(_story_html_row(r))
    _cols = [("Story", "14px", ""), ("Title", "8px", ""),
             ("Pri", "8px", "text-align:center")]
    if show_cases:
        _cols.append(("Cases", "8px", "text-align:right"))
    _cols += [("Hours", "8px", "text-align:right"), ("Assignee", "14px", "")]
    story_tbl = (
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
        f"style='border:1px solid #EBEFF7;border-radius:12px'>"
        f"<tr style='background:#F6F8FC'>"
        + "".join(f"<td style='padding:10px {p};font-size:10.5px;letter-spacing:.5px;"
                  f"text-transform:uppercase;color:#98a1b5;font-weight:700;{a}'>{h}</td>"
                  for h, p, a in _cols)
        + "</tr>" + "".join(srows) + "</table>")

    # workload bars
    wl_block = ""
    wl = d.get("workload", [])
    if wl:
        maxw = max((w["hours"] for w in wl), default=0) or 1
        wrows = []
        for w in wl:
            init, acol = _av(w["name"])
            pct = max(4, int(round(w["hours"] / maxw * 100)))
            wrows.append(
                f"<tr><td width='118' style='padding:7px 0'>"
                f"<table role='presentation' cellpadding='0' cellspacing='0'><tr>"
                f"<td width='22' style='vertical-align:middle'>"
                f"<div style='width:22px;height:22px;line-height:22px;border-radius:6px;"
                f"background:{acol};color:#fff;text-align:center;font-size:10px;"
                f"font-weight:700;font-family:Segoe UI,Arial,sans-serif'>{init}</div></td>"
                f"<td style='padding-left:9px;font-size:13px;font-weight:700;color:#1f2940;"
                f"vertical-align:middle'>"
                f"{w['name']}</td></tr></table></td>"
                f"<td style='padding:7px 14px'>"
                f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
                f"style='background:#eef1f7;border-radius:99px'><tr>"
                f"<td height='8' style='background:{acol};border-radius:99px;width:{pct}%;"
                f"font-size:0;line-height:0'>&nbsp;</td>"
                f"<td style='font-size:0;line-height:0'>&nbsp;</td></tr></table></td>"
                f"<td width='118' align='right' style='padding:7px 0;white-space:nowrap'>"
                f"<span style='font-size:11.5px;color:#8a93a8'>{w['stories']} stories</span>"
                f"<span style='font-family:Consolas,monospace;font-size:14px;font-weight:700;"
                f"color:#1f2940;padding-left:8px'>{w['hours']} h</span></td></tr>")
        wl_block = (
            f"<tr><td style='padding:22px 32px 4px'>"
            f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
            f"<td style='font-size:11px;letter-spacing:.7px;text-transform:uppercase;"
            f"color:#8a93a8;font-weight:700'>Resource workload</td>"
            f"<td align='right'><span style='font-size:11.5px;font-weight:600;color:#1F9D57;"
            f"background:#E5F6EC;padding:5px 11px;border-radius:999px'>"
            f"&#8776; {d['hours_per_person']} h / person</span></td></tr></table>"
            f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0' "
            f"style='margin-top:12px'>" + "".join(wrows) + "</table></td></tr>")

    scope = d["plan_name"] or d["project"]
    return (
        f"<div style='background:#e9edf4;padding:28px 12px;"
        f"font-family:Segoe UI,Arial,sans-serif'>"
        f"<table role='presentation' align='center' width='680' cellpadding='0' "
        f"cellspacing='0' style='max-width:680px;width:100%;margin:0 auto;background:#fff;"
        f"border-radius:16px;overflow:hidden;border:1px solid #e4e9f2'>"
        # header band
        f"<tr><td style='padding:26px 32px 22px;background:#3A57D6;"
        f"background-image:linear-gradient(125deg,#1C80E0 0%,#3A57D6 55%,#6A33A8 100%)'>"
        f"<table role='presentation' width='100%' cellpadding='0' cellspacing='0'><tr>"
        f"<td style='color:#fff;font-weight:800;font-size:15px;letter-spacing:.2px'>"
        f"QA&nbsp;Studio</td>"
        f"<td align='right'><span style='font-family:Consolas,monospace;font-size:11px;"
        f"color:#d6ddf6;background:rgba(255,255,255,.14);padding:6px 11px;"
        f"border-radius:8px'>generated {d['generated']}</span></td></tr></table>"
        f"<div style='margin-top:20px;color:#fff;font-size:25px;font-weight:800;"
        f"letter-spacing:-.5px'>{d.get('report_title', 'Regression Test Plan')}</div>"
        f"<div style='margin-top:6px;color:#cdd5f0;font-size:13.5px;font-weight:500'>"
        f"{scope}</div></td></tr>"
        # KPI strip
        f"<tr><td style='padding:22px 27px 6px'>{kpis}</td></tr>"
        # story table
        f"<tr><td style='padding:16px 32px 6px'>"
        f"<div style='font-size:11px;letter-spacing:.7px;text-transform:uppercase;"
        f"color:#8a93a8;font-weight:700;margin-bottom:12px'>Stories in scope</div>"
        f"{story_tbl}</td></tr>"
        + wl_block +
        # footer
        f"<tr><td style='padding:22px 32px 26px'>"
        f"<div style='border-top:1px solid #eef1f7;padding-top:18px;font-size:12px;"
        f"color:#9aa4b8;line-height:1.6'>Sent automatically from "
        f"<b style='color:#6b7790'>QA Studio</b>. The full plan is attached as a Word "
        f"document." + (f"<br>Estimates use {d['avg_minutes_per_case']}&nbsp;min / test "
        f"case weighted by Azure DevOps priority." if show_cases else
        f"<br>Effort is estimated per story and balanced across the team.")
        + f"</div></td></tr>"
        f"</table></div>")



def _out_dir():
    d = os.path.join(os.path.expanduser("~"), "QA Studio", "Regression Plans")
    os.makedirs(d, exist_ok=True)
    return d


def _stamp(app):
    if getattr(app, "_reg_mode", "existing") == "create":
        base = getattr(app, "_cp_sprint_name", "") or "sprint"
    else:
        base = ((getattr(app, "plan_name", "") or "")
                or (", ".join(p["name"] for p in (app._reg_plans_selected or [])) or "plan"))
    base = re.sub(r"[^A-Za-z0-9_-]+", "_", base).strip("_") or "plan"
    prefix = "SprintPlan" if getattr(app, "_reg_mode", "existing") == "create" else "RegressionPlan"
    return f"{prefix}_{base}_{datetime.now():%Y%m%d-%H%M}"


def _ask_save_path(fmt, default_name):
    """Open a native OS 'Save As' dialog (tkinter) and return the chosen path.
        str   -> the path the user picked
        None  -> the user cancelled
        False -> no native dialog available (caller should fall back)
    Must be called OFF the UI thread — it spins up its own hidden Tk root."""
    try:
        import tkinter as tk
        from tkinter import filedialog
    except Exception:
        return False
    try:
        root = tk.Tk()
        root.withdraw()
        try:
            root.attributes("-topmost", True)
        except Exception:
            pass
        path = filedialog.asksaveasfilename(
            parent=root, title="Save regression plan",
            initialfile=default_name, defaultextension="." + fmt,
            filetypes=[(f"{fmt.upper()} file", f"*.{fmt}"), ("All files", "*.*")])
        try:
            root.update(); root.destroy()
        except Exception:
            pass
        return path or None
    except Exception:
        return False


def export_json(app):
    d = plan_payload(app)
    for s in d.get("stories", []):          # link each story id to its Azure work item
        s["url"] = _wi_url(d["project"], s["id"])
    p = os.path.join(_out_dir(), _stamp(app) + ".json")
    with open(p, "w", encoding="utf-8") as f:
        json.dump(d, f, ensure_ascii=False, indent=2)
    return p


def export_xlsx(app):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    d = plan_payload(app)
    wb = Workbook()
    ws = wb.active
    ws.title = d.get("report_title", "Regression Plan")[:31]
    head = Font(bold=True, color="FFFFFF", name="Segoe UI")
    fill = PatternFill("solid", fgColor="3A57D6")          # brand indigo
    thin = Border(*[Side(style="thin", color="E6E8F1")] * 4)
    _AR = Alignment(horizontal="right", vertical="center")
    _AC = Alignment(horizontal="center", vertical="center")
    _AL = Alignment(horizontal="left", vertical="center", wrap_text=True)
    r = 1
    for k, v in (("Project", d["project"]), ("Test plan", d["plan_name"]),
                 ("Plan ID", d["plan_id"]),
                 ("Generated", d["generated"]),
                 ("Avg min / case", d["avg_minutes_per_case"]),
                 ("Resources", d["resources_count"]),
                 ("Resource names", ", ".join(d["resource_names"]) or "—")):
        ws.cell(r, 1, k).font = Font(bold=True)
        ws.cell(r, 2, v)
        r += 1
    r += 1
    # Sprint Plan has no test-case counts (estimates are complexity-based), so the
    # cases column/total is dropped when total_cases is 0.
    show_cases = (d.get("total_cases") or 0) > 0
    cols = (["Feature", "Story ID", "Title", "State", "Priority"]
            + (["Test cases"] if show_cases else [])
            + ["Est. hours", "Assignee"])
    ncol = len(cols)
    _hours_col = ncol - 1
    _cases_col = 6 if show_cases else None
    for c, name in enumerate(cols, 1):
        cell = ws.cell(r, c, name)
        cell.font = head
        cell.fill = fill
        cell.alignment = _AC
    _header_row = r
    r += 1
    from collections import OrderedDict as _ODx
    feat_grps = _ODx()
    for s in d["stories"]:
        feat_grps.setdefault(s.get("feature_name") or "No Feature", []).append(s)
    feat_fill = PatternFill("solid", fgColor="EEE8FF")
    for feat_name, fstories in feat_grps.items():
        ws.cell(r, 1, feat_name).font = Font(bold=True, color="2940C2", name="Segoe UI")
        ws.cell(r, 1).fill = feat_fill
        for c2 in range(2, ncol + 1): ws.cell(r, c2).fill = feat_fill
        r += 1
        for s in fstories:
            vals = ([feat_name, s["id"], s["title"], s["state"],
                     _PRI_FULL.get(s["priority"], s["priority"])]
                    + ([s["cases"]] if show_cases else [])
                    + [s["hours"], s.get("assignee") or "—"])
            for c, v in enumerate(vals, 1):
                cell = ws.cell(r, c, v)
                cell.border = thin
                if c == 2:                                       # story id -> Azure link
                    cell.hyperlink = _wi_url(d["project"], s["id"])
                    cell.font = Font(color="3A57D6", underline="single")
                if c in (2, _hours_col) or c == _cases_col:      # ids / hours / cases
                    cell.alignment = _AR
                elif c in (4, 5):                                # state / priority
                    cell.alignment = _AC
                elif c == 3:                                      # title
                    cell.alignment = _AL
            r += 1
    ws.freeze_panes = ws.cell(_header_row + 1, 1).coordinate
    ws.cell(r, 5, "TOTAL").font = Font(bold=True)
    if show_cases:
        ws.cell(r, 6, d["total_cases"]).font = Font(bold=True)
    ws.cell(r, ncol - 1, d["total_hours"]).font = Font(bold=True)
    r += 2
    if d["workload"]:
        ws.cell(r, 1, "Resource workload").font = Font(bold=True, size=12)
        r += 1
        wl_cols = (["Resource", "Stories"] + (["Test cases"] if show_cases else [])
                   + ["Hours"])
        for c, name in enumerate(wl_cols, 1):
            cell = ws.cell(r, c, name)
            cell.font = head
            cell.fill = fill
        r += 1
        for w in d["workload"]:
            wvals = ([w["name"], w["stories"]]
                     + ([w.get("cases", 0)] if show_cases else []) + [w["hours"]])
            for c, v in enumerate(wvals, 1):
                ws.cell(r, c, v)
            r += 1
    _widths = [28, 12, 52, 14, 16] + ([12] if show_cases else []) + [12, 18]
    for c, wdt in zip("ABCDEFGH", _widths):
        ws.column_dimensions[c].width = wdt
    p = os.path.join(_out_dir(), _stamp(app) + ".xlsx")
    wb.save(p)
    return p


def export_docx(app):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    d = plan_payload(app)

    def _add_hyperlink(paragraph, url, text, color="3A57D6"):
        """python-docx has no native hyperlink API — build the w:hyperlink run by
        hand (external relationship + blue underlined run) so the story id links to
        its Azure work item."""
        r_id = paragraph.part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/"
                 "relationships/hyperlink", is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), color); rpr.append(col)
        und = OxmlElement("w:u"); und.set(qn("w:val"), "single"); rpr.append(und)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = str(text); run.append(t)
        link.append(run)
        paragraph.append(link)
        return link
    doc = Document()
    try:
        _ns = doc.styles["Normal"]
        _ns.font.name = "Segoe UI"
        _ns.font.size = Pt(10)
    except Exception:
        pass
    h = doc.add_heading(d.get("report_title", "Regression Test Plan"), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = sub.add_run(f"{d['plan_name'] or d['project']}")
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0x6A, 0x4D, 0xFF)
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Light List Accent 1"
    for k, v in (("Project", d["project"]), ("Test plan", d["plan_name"]),
                 ("Plan ID", str(d["plan_id"])),
                 ("Generated", d["generated"]),
                 ("Resources", str(d["resources_count"])),
                 ("Resource names", ", ".join(d["resource_names"]) or "—"),
                 ("Effort model", f"{d['avg_minutes_per_case']} min/case "
                                   f"+ priority weighting")):
        cells = meta.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)
        for p in cells[0].paragraphs:
            for x in p.runs:
                x.font.bold = True
    # Sprint Plan has no test-case counts → drop the cases column/totals.
    show_cases = (d.get("total_cases") or 0) > 0
    doc.add_heading("Stories", level=1)
    from collections import OrderedDict as _ODd
    feat_grps_d = _ODd()
    for s in d["stories"]:
        feat_grps_d.setdefault(s.get("feature_name") or "No Feature", []).append(s)
    _heads = (["Story", "Title", "State", "Priority"]
              + (["Test cases"] if show_cases else []) + ["Est. hours", "Assignee"])
    for feat_name, fstories in feat_grps_d.items():
        fh = doc.add_heading(f"Feature: {feat_name}", level=2)
        fh.runs[0].font.color.rgb = RGBColor(0x6A, 0x4D, 0xFF)
        tbl = doc.add_table(rows=1, cols=len(_heads))
        tbl.style = "Medium Shading 1 Accent 1"
        for i, hd in enumerate(_heads):
            tbl.rows[0].cells[i].text = hd
        for s in fstories:
            c = tbl.add_row().cells
            _vals = ([str(s["id"]), s["title"], s["state"],
                      _PRI_FULL.get(s["priority"], str(s["priority"]))]
                     + ([str(s["cases"])] if show_cases else [])
                     + [str(s["hours"]), s.get("assignee") or "—"])
            _last = len(_vals) - 1
            for i, v in enumerate(_vals):
                if i == 0:                                  # story id -> Azure link
                    _add_hyperlink(c[i].paragraphs[0], _wi_url(d["project"], s["id"]), v)
                else:
                    c[i].text = v
                if i in (0, _last - 1) or (show_cases and i == 4):     # id / hours / cases
                    _al = WD_ALIGN_PARAGRAPH.RIGHT
                elif i in (2, 3):                                       # state / priority
                    _al = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    _al = WD_ALIGN_PARAGRAPH.LEFT
                for _p in c[i].paragraphs:
                    _p.alignment = _al
        doc.add_paragraph()
    doc.add_paragraph()
    tot = doc.add_table(rows=0, cols=2)
    tot.style = "Light List Accent 1"
    _tot_rows = ([("Total stories", d["total_stories"])]
                 + ([("Total test cases", d["total_cases"])] if show_cases else [])
                 + [("Total estimated hours", d["total_hours"]),
                    ("Hours per resource (target)", d["hours_per_person"])])
    for k, v in _tot_rows:
        cells = tot.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)
        for p in cells[0].paragraphs:
            for x in p.runs:
                x.font.bold = True
    if d["workload"]:
        doc.add_heading("Resource workload", level=1)
        _wh = (["Resource", "Stories"] + (["Test cases"] if show_cases else [])
               + ["Hours"])
        wt = doc.add_table(rows=1, cols=len(_wh))
        wt.style = "Medium Shading 1 Accent 1"
        for i, hd in enumerate(_wh):
            wt.rows[0].cells[i].text = hd
        for w in d["workload"]:
            c = wt.add_row().cells
            _wv = ([w["name"], str(w["stories"])]
                   + ([str(w.get("cases", 0))] if show_cases else []) + [str(w["hours"])])
            for i, v in enumerate(_wv):
                c[i].text = v
    foot = doc.add_paragraph()
    fr = foot.add_run("Generated by QA Studio · effort references existing "
                      "test cases only.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p = os.path.join(_out_dir(), _stamp(app) + ".docx")
    doc.save(p)
    return p


_PDF_FONT_NAME = None


def _pdf_font():
    """Register an Arabic-capable TTF for reportlab ONCE and return its name.
    reportlab's built-in fonts have no Arabic glyphs, so Arabic story titles
    rendered as blank — making the PDF look empty. Falls back to Helvetica
    (Latin-only) if no suitable font is found."""
    global _PDF_FONT_NAME
    if _PDF_FONT_NAME is not None:
        return _PDF_FONT_NAME
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        for fp in (r"C:\Windows\Fonts\segoeui.ttf", r"C:\Windows\Fonts\tahoma.ttf",
                   r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\ARIALUNI.TTF",
                   "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"):
            if os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("PlanFont", fp))
                _PDF_FONT_NAME = "PlanFont"
                return _PDF_FONT_NAME
    except Exception:
        pass
    _PDF_FONT_NAME = "Helvetica"
    return _PDF_FONT_NAME


def _ar(s):
    """Shape Arabic text (connected glyphs + RTL order) for the PDF when the
    reshaper libs are present; otherwise return it unchanged."""
    s = "" if s is None else str(s)
    if not any('؀' <= c <= 'ۿ' for c in s):
        return s
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(s))
    except Exception:
        return s


def export_pdf(app):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer)
    from reportlab.lib.styles import getSampleStyleSheet
    d = plan_payload(app)
    p = os.path.join(_out_dir(), _stamp(app) + ".pdf")
    doc = SimpleDocTemplate(p, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    _fn = _pdf_font()
    for _sn in ("Title", "Normal", "Heading2"):
        try: styles[_sn].fontName = _fn
        except Exception: pass
    from reportlab.lib.styles import ParagraphStyle
    _id_style = ParagraphStyle("wid", parent=styles["Normal"], fontName=_fn,
                               fontSize=7.5, alignment=1,        # 1 = centered
                               textColor=colors.HexColor("#3A57D6"))

    def _id_cell(sid):                       # story id -> clickable Azure link
        return Paragraph(
            f'<link href="{_wi_url(d["project"], sid)}"><u>{sid}</u></link>', _id_style)
    elems = [Paragraph(_ar(d.get("report_title") or "Regression Test Plan"), styles["Title"]),
             Paragraph(_ar(d['plan_name'] or d['project']),
                       styles["Normal"]),
             Spacer(1, 8 * mm)]
    # Sprint Plan has no test-case counts → drop the cases column/total.
    show_cases = (d.get("total_cases") or 0) > 0
    data = [["Story", "Title", "State", "Pri"]
            + (["Cases"] if show_cases else []) + ["Hours", "Assignee"]]
    for s in d["stories"]:
        data.append([_id_cell(s["id"]), _ar((s["title"] or "")[:38]), _ar(s["state"]),
                     str(s["priority"])]
                    + ([str(s["cases"])] if show_cases else [])
                    + [str(s["hours"]), _ar(s.get("assignee") or "—")])
    data.append(["", "", "", "TOT"]
                + ([str(d["total_cases"])] if show_cases else [])
                + [str(d["total_hours"]), ""])
    _cw = [18*mm, 54*mm, 22*mm, 10*mm] + ([14*mm] if show_cases else []) + [14*mm, 28*mm]
    tbl = Table(data, colWidths=_cw, repeatRows=1)
    _pri_col = 3
    _cases_col = 4 if show_cases else None
    tbl.setStyle(TableStyle([
        # header: brand indigo band, white, centered
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3A57D6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("FONTNAME", (0, 0), (-1, -1), _fn),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, -1), 7.5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E6E8F1")),
        ("LINEBELOW", (0, 0), (-1, 0), 0.8, colors.HexColor("#2C44BE")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2),
         [colors.white, colors.HexColor("#F6F8FC")]),
        # alignment by meaning: story id centred, numerics right, pri centred
        ("ALIGN", (0, 1), (0, -1), "CENTER"),          # story id
        ("ALIGN", (_pri_col, 1), (_pri_col, -1), "CENTER"),  # priority
        ("ALIGN", (-2, 1), (-2, -1), "RIGHT"),         # hours
        ("TEXTCOLOR", (-2, 1), (-2, -1), colors.HexColor("#181A24")),
        # totals row
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#EEF1FF")),
        ("LINEABOVE", (0, -1), (-1, -1), 0.8, colors.HexColor("#3A57D6")),
        ("TEXTCOLOR", (0, -1), (-1, -1), colors.HexColor("#2940C2")),
    ] + ([("ALIGN", (_cases_col, 1), (_cases_col, -1), "RIGHT")] if _cases_col else [])))
    elems += [tbl, Spacer(1, 6 * mm)]
    if d["workload"]:
        wd = [["Resource", "Stories"] + (["Test cases"] if show_cases else []) + ["Hours"]] + \
             [[_ar(w["name"]), str(w["stories"])] + ([str(w.get("cases", 0))] if show_cases else [])
              + [str(w["hours"])] for w in d["workload"]]
        _wcw = [55*mm, 25*mm] + ([30*mm] if show_cases else []) + [25*mm]
        wt = Table(wd, colWidths=_wcw, repeatRows=1)
        wt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3A57D6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 0), (-1, -1), _fn),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E6E8F1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#F6F8FC")]),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),   # numeric columns
        ]))
        elems += [Paragraph(_ar("Resource workload"), styles["Heading2"]), wt]
    doc.build(elems)
    return p


EXPORTERS = {"docx": export_docx, "xlsx": export_xlsx,
             "json": export_json, "pdf": export_pdf}
_MISSING_DEP = {"xlsx": "openpyxl  ·  pip install openpyxl",
                "docx": "python-docx  ·  pip install python-docx",
                "pdf":  "reportlab  ·  pip install reportlab"}


def _export_row(app, set_status=None):
    """Export buttons for the sprint report. Each saves via the shared exporter
    (which now reads the mode-aware payload) into ~/QA Studio/Regression Plans.

    set_status(kind, text) -- if provided, called in-place instead of render().
    """
    from main import green_btn, ghost_btn

    def _notify(kind, text):
        app._cp_msg = (kind, text)
        try:
            app.ui_safe(lambda k=kind, t=text: (app._toast(t) if k == "ok"
                                                else app._err(t)))
            return
        except Exception:
            pass
        if set_status is not None:
            try:
                app.ui_safe(lambda k=kind, t=text: set_status(k, t))
                return
            except Exception:
                pass
        app.ui_safe(app.render)

    def _go(fmt):
        def _do(e):
            if not app._cp_rows:
                _notify("err", "Pick a sprint first.")
                return

            def work():
                try:
                    dest = _ask_save_path(fmt, _stamp(app) + "." + fmt)
                    if dest is None:
                        return                       # user cancelled the dialog
                    path = EXPORTERS[fmt](app)
                    if dest and dest is not False:   # a path was chosen → move there
                        if not dest.lower().endswith("." + fmt):
                            dest += "." + fmt
                        import shutil
                        if os.path.abspath(dest) != os.path.abspath(path):
                            shutil.move(path, dest)
                        path = dest
                    try:
                        os.startfile(os.path.dirname(path))
                    except Exception:
                        pass
                    _notify("ok", f"Saved: {path}")
                except ModuleNotFoundError:
                    _notify("err", f"Missing dependency: {_MISSING_DEP.get(fmt, fmt)}")
                except Exception as ex:
                    _notify("err", f"Export failed: {str(ex)[:160]}")
            threading.Thread(target=work, daemon=True).start()
        return _do

    def _exp_btn(label, icon, color, fmt):
        return ft.OutlinedButton(
            content=ft.Row([ft.Icon(icon, size=17, color=color),
                            ft.Text(label, size=13.5, weight=ft.FontWeight.W_600,
                                    color=T.INK)],
                           spacing=8, tight=True),
            on_click=_go(fmt), height=44,
            style=ft.ButtonStyle(
                bgcolor={"": T.CARD},
                side=ft.BorderSide(1, T.BORDER),
                shape=ft.RoundedRectangleBorder(radius=T.R),
                padding=ft.Padding.symmetric(horizontal=15, vertical=0)))

    btns = ft.Row([
        _exp_btn("Word", ft.Icons.DESCRIPTION, T.BRAND_GRAD_1, "docx"),
        _exp_btn("Excel", ft.Icons.TABLE_CHART, T.GREEN, "xlsx"),
        _exp_btn("PDF", ft.Icons.PICTURE_AS_PDF, T.RED, "pdf"),
        _exp_btn("JSON", ft.Icons.DATA_OBJECT, T.STORY, "json"),
    ], spacing=8, wrap=True)

    _m_kind = app._cp_msg[0] if app._cp_msg else "ok"
    _m_text = app._cp_msg[1] if app._cp_msg else ""
    _m_col = T.GREEN if _m_kind == "ok" else T.RED
    _status_icon = ft.Icon(
        ft.Icons.CHECK_CIRCLE if _m_kind == "ok" else ft.Icons.ERROR_OUTLINE,
        size=15, color=_m_col)
    _status_text = _txt(_m_text, color=_m_col, size=12, expand=True)
    status = ft.Container(
        ft.Row([_status_icon, _status_text], spacing=8),
        padding=10, border_radius=T.R, margin=ft.Margin.only(top=10),
        bgcolor=T.CARD, border=ft.Border.all(1, T.BORDER_2),
        visible=bool(app._cp_msg))

    if set_status is None:
        # legacy: caller didn't provide a callback, nothing to wire
        pass
    # Note: set_status wiring happens externally after this returns

    return ft.Column([btns, status], spacing=0), status


# ═══════════════════════════════════════════════════════════════════════════════
#  UI helpers
# ═══════════════════════════════════════════════════════════════════════════════
def _pill(text, fg, bg):
    return ft.Container(ft.Text(text, size=11, weight=ft.FontWeight.BOLD, color=fg),
                        padding=ft.Padding.symmetric(vertical=2, horizontal=8),
                        bgcolor=bg, border_radius=20)


def _state_pill(state):
    s = (state or "").lower()
    if s in ("closed", "done", "resolved", "completed"):
        return _pill(state, T.GREEN, T.GREEN_SOFT)
    if s in ("active", "in progress", "committed"):
        return _pill(state, T.VIOLET_INK, T.VIOLET_SOFT)
    return _pill(state or "—", T.INK_2, T.CARD_2)


def _pri_pill(pri):
    if pri == 1:
        return _pill("P1", T.RED, T.RED_SOFT)
    if pri == 2:
        return _pill("P2", T.AMBER, T.AMBER_SOFT)
    if pri == 4:
        return _pill("P4", T.GREEN, T.GREEN_SOFT)
    return _pill(f"P{pri}", T.INK_2, T.CARD_2)


def _txt(s, **kw):
    kw.setdefault("size", 12)
    return ft.Text(s, **kw)


def _id_link(app, sid, **kw):
    """Story id rendered as a clickable link to its Azure DevOps work item — the
    in-app plan tables, mirroring the hyperlinks already in the exports. Opens via
    app._open_url (OS browser). Falls back to plain text on any error."""
    kw.setdefault("tooltip", f"Open story {sid} in Azure DevOps")
    txt = _txt(str(sid), **kw)
    try:
        url = _wi_url(app.project, sid)
    except Exception:
        return txt
    return ft.GestureDetector(content=txt,
                              on_tap=lambda e: app._open_url(url),
                              mouse_cursor=ft.MouseCursor.CLICK)


def _chip_hover(e):
    """Shared hover feedback for chips across the app — a small scale pop. Pair with
    `on_hover=_chip_hover, animate_scale=120` on a chip Container."""
    try:
        e.control.scale = 1.06 if (e.data in (True, "true", "True")) else 1.0
        e.control.update()
    except Exception:
        pass


def _clear_chip(on_click, label="Clear all"):
    """Small trailing 'clear all' pill placed after a row of selection chips."""
    return ft.Container(
        ft.Row([ft.Icon(ft.Icons.CLOSE_ROUNDED, size=12, color=T.INK_3),
                ft.Text(label, size=11, weight=ft.FontWeight.BOLD, color=T.INK_3)],
               spacing=4, tight=True),
        padding=ft.Padding.symmetric(vertical=4, horizontal=10),
        border_radius=999, border=ft.Border.all(1, T.BORDER_2),
        on_click=on_click, tooltip="Remove all",
        on_hover=_chip_hover, animate_scale=120)


def _feature_header(app, label, count, collapsed, on_toggle):
    """Polished, whole-row-clickable feature-group header for the plan tables: a
    violet gradient icon chip, the feature name, a count badge, and a chevron on the
    right, on a soft violet-tinted card. Clicking anywhere toggles collapse (an ink
    ripple gives click feedback)."""
    chev = (ft.Icons.KEYBOARD_ARROW_RIGHT if collapsed
            else ft.Icons.KEYBOARD_ARROW_DOWN)
    icon_chip = ft.Container(
        ft.Icon(ft.Icons.FOLDER_ROUNDED, size=15, color="#FFFFFF"),
        width=28, height=28, border_radius=9, alignment=ft.Alignment.CENTER,
        gradient=ft.LinearGradient(begin=ft.Alignment.TOP_LEFT,
                                   end=ft.Alignment.BOTTOM_RIGHT,
                                   colors=[T.VIOLET, T.VIOLET_INK]),
        shadow=ft.BoxShadow(blur_radius=8, spread_radius=-3, offset=ft.Offset(0, 3),
                            color=ft.Colors.with_opacity(0.40, T.VIOLET)))
    count_badge = ft.Container(
        ft.Text(f"{count} {'story' if count == 1 else 'stories'}",
                size=11, weight=ft.FontWeight.W_800, color=T.VIOLET_INK),
        padding=ft.Padding.symmetric(vertical=3, horizontal=10), border_radius=999,
        bgcolor=ft.Colors.with_opacity(0.14, T.VIOLET))
    return ft.Container(
        ft.Row([
            icon_chip,
            ft.Row([_txt(label, size=13, weight=ft.FontWeight.W_900, color=T.INK,
                         no_wrap=True),
                    count_badge],
                   spacing=10, expand=True,
                   vertical_alignment=ft.CrossAxisAlignment.CENTER),
            ft.Icon(chev, size=22, color=T.VIOLET_INK),
        ], spacing=12, vertical_alignment=ft.CrossAxisAlignment.CENTER),
        on_click=on_toggle, ink=True,
        tooltip="Collapse / expand this feature",
        padding=ft.Padding.only(left=10, right=12, top=9, bottom=9),
        margin=ft.Margin.only(top=10, bottom=4, left=2, right=2),
        gradient=ft.LinearGradient(
            begin=ft.Alignment.CENTER_LEFT, end=ft.Alignment.CENTER_RIGHT,
            colors=[ft.Colors.with_opacity(0.13, T.VIOLET),
                    ft.Colors.with_opacity(0.04, T.VIOLET)]),
        border=ft.Border.all(1, ft.Colors.with_opacity(0.40, T.VIOLET)),
        border_radius=12,
        shadow=ft.BoxShadow(blur_radius=12, spread_radius=-5, offset=ft.Offset(0, 5),
                            color=ft.Colors.with_opacity(0.16, T.VIOLET)))


def _avatar(name, size=26):
    """Round initial-avatar; colour is stable per name."""
    init, col = _av(name)
    return ft.Container(
        ft.Text(init, size=int(size * 0.42), weight=ft.FontWeight.BOLD, color="#FFFFFF"),
        width=size, height=size, bgcolor=col, border_radius=size,
        alignment=ft.Alignment.CENTER)


def _bar(frac, color=T.VIOLET, h=7):
    """Proportional fill bar (animates smoothly when its weight changes)."""
    f = max(1, int(round((frac or 0) * 100)))
    e = max(0, 100 - f)
    fill = ft.Container(height=h, bgcolor=color, border_radius=4, expand=f,
                        animate=700)
    inner = [fill] + ([ft.Container(expand=e)] if e > 0 else [])
    return ft.Container(ft.Row(inner, spacing=0), bgcolor=T.BORDER_2,
                        border_radius=4, height=h, clip_behavior=ft.ClipBehavior.HARD_EDGE)


def _kpi_tile(label, value, accent=None):
    accent = accent if accent is not None else T.INK   # call-time (theme-aware)
    try:
        from main import grad_text
        _stops = T.GRAD_GREEN if accent == T.GREEN else T.GRAD_LOGO
        _num = grad_text(value, size=23, weight=ft.FontWeight.BOLD,
                         stops=_stops, font_family=T.F_MONO)
    except Exception:
        _num = ft.Text(value, size=23, weight=ft.FontWeight.BOLD, color=accent,
                       font_family=T.F_MONO)
    return ft.Container(
        ft.Column([
            ft.Text(label, size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
            _num,
        ], spacing=4),
        expand=True, padding=14, bgcolor=T.CARD,
        border_radius=T.R, border=ft.Border.all(1, T.BORDER_2),
        shadow=ft.BoxShadow(blur_radius=16, spread_radius=-10, offset=ft.Offset(0, 6),
                            color=ft.Colors.with_opacity(0.08, "#1B1F3A")))


def locked_state(app, title, sub, msg, icon=None, steps=None):
    """Shared centered 'connect / select first' screen (also used by Automation).

    `steps` is an optional list of (icon, label) tuples shown as a 3-step path.
    """
    from main import primary_btn
    icon = icon or ft.Icons.LINK_OFF
    steps = steps or [(ft.Icons.TUNE, "Connect"),
                      (ft.Icons.CHECKLIST, "Select"),
                      (ft.Icons.AUTO_AWESOME, "Generate")]

    # "scanning for a connection" card — ft.ProgressBar(value=None) animates natively in Flet
    scan_card = ft.Container(
        ft.Column([
            ft.Container(ft.Icon(icon, size=28, color=T.VIOLET),
                         width=60, height=60, bgcolor=T.VIOLET_SOFT, border_radius=18,
                         alignment=ft.Alignment.CENTER),
            ft.Container(height=16),
            ft.ProgressBar(value=None, color=T.VIOLET, bgcolor=T.BORDER_2,
                           bar_height=6, border_radius=99, width=224),
            ft.Container(height=10),
            ft.Text("Scanning for a connection…", size=11, color=T.INK_3,
                    weight=ft.FontWeight.W_500, font_family=T.F_MONO),
        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=0),
        width=296, padding=ft.Padding.symmetric(vertical=20, horizontal=26),
        bgcolor=T.CARD, border=ft.Border.all(1, T.BORDER_2), border_radius=22)

    pill = ft.Container(
        ft.Row([ft.Container(width=7, height=7, border_radius=7, bgcolor=T.STORY),
                ft.Text("Awaiting connection", size=11, weight=ft.FontWeight.BOLD,
                        color=T.STORY)], spacing=7, tight=True),
        padding=ft.Padding.symmetric(vertical=6, horizontal=12),
        bgcolor=T.VIOLET_SOFT, border_radius=999,
        border=ft.Border.all(1, "#E0E5FF"))

    def _step(ic, label):
        return ft.Container(ft.Column([
            ft.Container(ft.Icon(ic, size=18, color=T.VIOLET), width=44, height=44,
                         bgcolor=T.VIOLET_SOFT, border_radius=13,
                         alignment=ft.Alignment.CENTER,
                         border=ft.Border.all(1, "#E0E5FF")),
            ft.Text(label, size=12, weight=ft.FontWeight.BOLD, color=T.INK_2),
        ], spacing=8, horizontal_alignment=ft.CrossAxisAlignment.CENTER), width=110)

    path = ft.Row([_step(ic, lab) for ic, lab in steps],
                  alignment=ft.MainAxisAlignment.CENTER, spacing=0)

    body = ft.Container(
        ft.Column([
            scan_card,
            ft.Container(height=12), pill,
            ft.Container(height=10),
            ft.Text("A few things first", size=20, weight=ft.FontWeight.BOLD,
                    color=T.INK),
            ft.Container(height=8),
            ft.Container(ft.Text(msg, size=13.5, color=T.INK_2,
                                 text_align=ft.TextAlign.CENTER), width=470),
            ft.Container(height=16), path,
            ft.Container(height=18),
            primary_btn("Go to Setup", icon=ft.Icons.ARROW_FORWARD,
                        on_click=lambda e: app.goto("setup")),
        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER,
           alignment=ft.MainAxisAlignment.CENTER, tight=True),
        alignment=ft.Alignment.CENTER, expand=True,
        padding=ft.Padding.symmetric(vertical=18, horizontal=20))
    return app.shell(title, sub, body)


# ═══════════════════════════════════════════════════════════════════════════════
#  SCREEN
# ═══════════════════════════════════════════════════════════════════════════════
def _init(app):
    for k, v in (("_reg_plans_selected", []), ("_reg_plan_stories", []),
                 ("_reg_stories_loading", False), ("_reg_selected", []),
                 ("_reg_plan_open", False), ("_reg_story_open", False),
                 ("_cp_sprint_open", False),
                 ("_reg_selected_rows", []), ("_reg_res_names", []),
                 ("_reg_res_count", None), ("_reg_busy", False),
                 ("_reg_export_msg", None), ("_reg_calc_msg", None),
                 ("_reg_email_to", ""), ("_reg_emailing", False),
                 ("_reg_plans_loading", False),
                 ("_reg_mode", "existing"),
                 ("_cp_iterations", []), ("_cp_iter_loading", False),
                 ("_cp_sprint_paths", []), ("_cp_sprint_name", ""),
                 ("_cp_stories_loading", False),
                 ("_cp_rows", []), ("_cp_res_names", []),
                 ("_cp_est_min", 1.0), ("_cp_est_max", 8.0),
                 ("_cp_res_count", None), ("_cp_calculated", False),
                 ("_cp_calc_msg", None),
                 ("_cp_msg", None), ("_cp_assigning", False),
                 ("_cp_busy", False),
                 ("_cp_email_to", ""), ("_cp_emailing", False),
                 ("_cp_collapsed_features", set()),
                 ("_cp_table_page", 0),
                 ("_reg_suite_cache", {}),
                 ("_reg_table_page", 0),
                 ("_reg_story_features", {}),
                 ("_reg_feature_name_cache", {}),
                 ("_reg_case_count_cache", {}),
                 ("_reg_collapsed_features", set()),
                 ("_reg_meta_cache", {})):
        if not hasattr(app, k):
            setattr(app, k, v)


def _sprint_num(text):
    """Pull a clean 'Sprint N' out of an iteration path or plan name; '' if none."""
    m = re.search(r"[Ss]print\s*\d+", text or "")
    return re.sub(r"\s+", " ", m.group(0)).strip() if m else ""


def _reload_plan_stories(app):
    """Aggregate the user stories that live in the currently-selected test plans
    (their requirement suites). Runs off the UI thread."""
    plans = list(app._reg_plans_selected or [])
    # Cancellation token: every (re)load bumps this; an in-flight fetch whose token
    # is stale must NOT apply its results. So select-all then deselect-all stops the
    # old fetch from repopulating the list instead of waiting for it to finish.
    app._reg_stories_gen = getattr(app, "_reg_stories_gen", 0) + 1
    _gen = app._reg_stories_gen
    if not plans:
        app._reg_plan_stories = []
        app._reg_stories_loading = False
        app.ui_safe(app.render)
        return
    app._reg_stories_loading = True
    app.ui_safe(app.render)

    def _work():
        import concurrent.futures as _cf_reload

        def _fetch_one_plan(p):
            # Fetch sprint label + stories for each plan concurrently.
            try:
                pj = E._azure_get(f"https://dev.azure.com/{E.AZURE_ORG}/{app.project}"
                                  f"/_apis/testplan/plans/{p['id']}?api-version=7.0")
                itr = pj.get("iteration") or ""
                sprint = _sprint_num(itr) or _sprint_num(p.get("name", "")) \
                    or itr.split("\\")[-1]
            except Exception:
                sprint = p.get("sprint", "")
            try:
                stories = E.fetch_stories_in_plan(app.project, p["id"])
            except Exception:
                stories = []
            return p["id"], sprint, stories

        agg, seen = [], set()
        with _cf_reload.ThreadPoolExecutor(max_workers=min(8, len(plans))) as ex:
            plan_results = list(ex.map(_fetch_one_plan, plans))

        plan_sprint = {pid: sprint for pid, sprint, _ in plan_results}
        for p in plans:
            p["sprint"] = plan_sprint.get(p["id"], p.get("sprint", ""))

        for pid, sprint, stories in plan_results:
            for s in stories:
                key = s["id"]   # dedupe by story id (a story in 2 plans = one entry)
                if key in seen:
                    continue
                seen.add(key)
                # group under the PLAN's sprint (what the user picked), not the
                # story's own iteration — so only selected sprints appear
                agg.append({"id": s["id"], "title": s.get("title", ""),
                            "sprint": sprint or _sprint_num(s.get("sprint", "")),
                            "plan_id": pid})
        # group by sprint then id so the picker lists them clustered by sprint
        agg.sort(key=lambda s: (s.get("sprint", "") or "~", s["id"]))
        if _gen != getattr(app, "_reg_stories_gen", _gen):
            return                      # selection changed mid-fetch -> drop results
        # Resolve each story's parent → Feature WHILE we're still loading (covered by
        # the "Loading stories…" wait) and tag the stories. This warms the feature
        # caches so a later Generate is fast and the plan groups by feature
        # immediately — with NO separate background pass churning afterwards (that
        # lingering pass is what used to freeze the table's collapse/paging).
        _resolve_reg_features(app, agg, _gen)
        if _gen != getattr(app, "_reg_stories_gen", _gen):
            return
        app._reg_plan_stories = agg
        app._reg_stories_loading = False
        app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


# ═══════════════════════════════════════════════════════════════════════════════
#  SPRINT REPORT (mode = "create") — a regression-style report built from a sprint
#  and its user stories, with a random (editable) per-story estimate and a random
#  (editable) assignment. Nothing is written to Azure here.
# ═══════════════════════════════════════════════════════════════════════════════
def _cp_is_sprint(it):
    """Keep only iteration nodes that look like a sprint (have 'Sprint N')."""
    return bool(_sprint_num(it.get("name", "")) or _sprint_num(it.get("path", "")))


def _cp_load_iterations(app):
    if app._cp_iterations or app._cp_iter_loading:
        return
    app._cp_iter_loading = True

    def _work():
        try:
            its = E.fetch_iterations(app.project) or []
        except Exception:
            its = []
        sprints = [it for it in its if _cp_is_sprint(it)] or its
        # Ascending by sprint number (Sprint 0, 1, 2, …) so Sprint 0 reads first
        # instead of being stranded at the end; un-numbered iterations sort last.
        sprints.sort(key=lambda it: (_sprint_sort_key(it) < 0, _sprint_sort_key(it)))
        app._cp_iterations = sprints
        app._cp_iter_loading = False
        app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


def _sprint_sort_key(it):
    m = re.search(r"\d+", _sprint_num(it.get("name", "")) or _sprint_num(it.get("path", "")))
    return int(m.group(0)) if m else -1


def _checkbox_multiselect(options, selected, on_toggle, on_all, *, is_open, on_open,
                          placeholder="Select…", height=240, empty="No options.",
                          page=None, app=None, invalid=False, sync_key=None):
    """Collapsible checkbox multiselect.

    When page= is supplied every interaction (open/close AND checkbox tick) is
    handled fully IN-PLACE via control mutation + page.update().  render() is
    never called from inside this component so the scroll position never jumps.

    Callers must NOT call render() inside on_toggle / on_all / on_open when
    page= is provided -- they should only mutate app state.  The component keeps
    its own mutable sel set and syncs all visible refs itself.

    When app= is also supplied the component registers a close() callable in
    app._dd_closers so _close_dropdowns can close it in-place on click-away.
    """
    # mutable state owned by this widget instance
    sel = set(selected or [])
    keys = [k for k, _ in options]

    def _all_on():
        return bool(keys) and all(k in sel for k in keys)

    # refs we need to mutate on tick / open
    field_label_ref = ft.Text(
        (f"{len(sel)} selected" if sel else placeholder),
        size=13, color=(T.INK if sel else T.INK_3), expand=True)
    arrow_icon = ft.Icon(
        ft.Icons.KEYBOARD_ARROW_UP if is_open else ft.Icons.KEYBOARD_ARROW_DOWN,
        size=20, color=T.INK_3)
    select_all_cb = ft.Checkbox(value=_all_on())
    count_text = ft.Text(f"{len(sel)} selected", size=11.5, color=T.INK_3,
                         weight=ft.FontWeight.BOLD)

    # per-row checkbox refs keyed by option key
    row_cbs = {}

    def _refresh_header():
        n = len(sel)
        field_label_ref.value = f"{n} selected" if sel else placeholder
        field_label_ref.color = T.INK if sel else T.INK_3
        count_text.value = f"{n} selected"
        select_all_cb.value = _all_on()

    def _do_toggle(kk, checked):
        if checked:
            sel.add(kk)
        else:
            sel.discard(kk)
        for k2, cb in row_cbs.items():
            if k2 != kk:
                cb.value = (k2 in sel)
        _refresh_header()
        try:
            on_toggle(kk, checked)
        except Exception:
            pass
        if page is not None:
            try:
                page.update()
            except Exception:
                pass

    def _do_all(checked):
        if checked:
            sel.update(keys)
        else:
            sel.clear()
        for k2, cb in row_cbs.items():
            cb.value = (k2 in sel)
        _refresh_header()
        try:
            on_all(checked)
        except Exception:
            pass
        if page is not None:
            try:
                page.update()
            except Exception:
                pass

    select_all_cb.on_change = lambda e: _do_all(e.control.value)

    # ── Panel rows are built LAZILY ──────────────────────────────────────────
    # Building one live control-set (Container+Row+Checkbox+Text) per option for
    # every story is the single most expensive thing a full render() does: a plan
    # with hundreds of stories yields thousands of controls that Flet must diff
    # and ship to the client on EVERY render — even while the dropdown is
    # collapsed, because visible=is_open only HIDES the wrapper, it does not skip
    # building it. That full rebuild is what freezes the UI right after "Generate"
    # (generation completes -> app.render() -> the whole hidden picker is rebuilt
    # and Flutter re-lays-out thousands of off-screen controls, so collapse / next
    # / prev / export / email / nav all stop responding until it settles).
    # Fix: when the panel is closed we build nothing; the rows are materialised
    # the first time the panel is actually opened.
    body = ft.Column([], spacing=0, scroll=ft.ScrollMode.AUTO)
    body_holder = ft.Container(body, height=64,
                               padding=ft.Padding.symmetric(vertical=4),
                               border=ft.Border.all(1, T.BORDER),
                               border_radius=ft.BorderRadius.only(
                                   bottom_left=T.R, bottom_right=T.R))
    _built = {"done": False}

    def _ensure_built():
        if _built["done"]:
            return
        _built["done"] = True
        row_cbs.clear()
        rows = []
        for k, label in options:
            cb = ft.Checkbox(value=(k in sel),
                             on_change=(lambda e, kk=k: _do_toggle(kk, e.control.value)))
            row_cbs[k] = cb
            rows.append(ft.Container(
                ft.Row([cb, ft.Text(label, size=12.5, color=T.INK, expand=True,
                                    no_wrap=False)],
                       spacing=6, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.Padding.only(left=10, right=10, top=2, bottom=2)))
        if rows:
            body.controls = rows
            body_holder.height = min(height, max(40, len(rows) * 34 + 8))
        else:
            body.controls = [ft.Container(ft.Text(empty, size=12, color=T.INK_3),
                                          padding=14, alignment=ft.Alignment.CENTER)]
            body_holder.height = 64

    head = ft.Container(
        ft.Row([select_all_cb,
                ft.Text("Select all", size=12.5, weight=ft.FontWeight.BOLD, color=T.INK),
                ft.Container(expand=True),
                count_text],
               spacing=6, vertical_alignment=ft.CrossAxisAlignment.CENTER),
        padding=ft.Padding.symmetric(vertical=8, horizontal=10), bgcolor=T.CARD_2,
        border_radius=ft.BorderRadius.only(top_left=T.R, top_right=T.R))

    panel_body = ft.Column([head, body_holder], spacing=0)
    panel_wrap = ft.Container(panel_body, padding=ft.Padding.only(top=6),
                              visible=is_open)

    # If the panel is persisted-open across this render, build now so the rows
    # are present on first paint.
    if is_open:
        _ensure_built()

    def _do_open(e):
        new_open = not panel_wrap.visible
        if new_open:
            _ensure_built()   # materialise rows lazily on first open
        panel_wrap.visible = new_open
        arrow_icon.name = (ft.Icons.KEYBOARD_ARROW_UP if new_open
                           else ft.Icons.KEYBOARD_ARROW_DOWN)
        field_container.border = ft.Border.all(
            1, T.VIOLET if new_open else (T.RED if invalid else T.BORDER))
        try:
            on_open()
        except Exception:
            pass
        if page is not None:
            try:
                page.update()
            except Exception:
                pass

    # Register a close() callable so _close_dropdowns can close this in-place.
    def _close_this():
        if panel_wrap.visible:
            panel_wrap.visible = False
            arrow_icon.name = ft.Icons.KEYBOARD_ARROW_DOWN
            field_container.border = ft.Border.all(1, T.RED if invalid else T.BORDER)
            try:
                on_open()   # sync the flag
            except Exception:
                pass
            return True
        return False

    if app is not None:
        try:
            app._dd_closers.append(_close_this)
        except Exception:
            pass

    # External in-place sync: when a chip is removed elsewhere, the caller calls
    # this with the authoritative selection so the dropdown's checkboxes + header
    # untick in place (no full render, and correct even while the panel is open).
    if app is not None and sync_key is not None:
        def _sync_selected(new_keys):
            new = {str(k) for k in (new_keys or [])}
            sel.clear()
            sel.update(new)
            for _k2, _cb in row_cbs.items():
                _cb.value = (_k2 in sel)
            _refresh_header()
            if page is not None:
                try:
                    page.update()
                except Exception:
                    pass
        try:
            app._dd_syncers[sync_key] = _sync_selected
        except Exception:
            pass

    field_container = ft.Container(
        ft.Row([field_label_ref, arrow_icon],
               vertical_alignment=ft.CrossAxisAlignment.CENTER),
        on_click=_do_open if page is not None else (lambda e: on_open()),
        padding=ft.Padding.symmetric(vertical=12, horizontal=12),
        bgcolor=T.CARD,
        border=ft.Border.all(
            1, T.VIOLET if is_open else (T.RED if invalid else T.BORDER)),
        border_radius=T.R)

    # legacy path (no page=)
    if page is None:
        if not is_open:
            return field_container
        return ft.Column([field_container,
                          ft.Container(panel_body, padding=ft.Padding.only(top=6))],
                         spacing=0)

    # in-place path
    return ft.Column([field_container, panel_wrap], spacing=0)

def _cp_load_stories(app):
    paths = list(app._cp_sprint_paths or [])
    # Cancellation token (see _reload_plan_stories): deselecting sprints aborts an
    # in-flight fetch instead of waiting for it to complete and repopulate.
    app._cp_stories_gen = getattr(app, "_cp_stories_gen", 0) + 1
    _gen = app._cp_stories_gen
    if not paths:
        app._cp_rows = []
        app._cp_stories_loading = False
        app.ui_safe(app.render)
        return
    app._cp_stories_loading = True
    app._cp_rows = []
    app._cp_table_page = 0
    app.ui_safe(app.render)

    def _work():
        import concurrent.futures as _cf

        # Fetch every selected sprint's stories CONCURRENTLY. Sequentially this was
        # one blocking call per sprint, so 23 sprints kept "Loading stories…" up
        # (and the Generate button disabled) for the sum of all of them.
        def _one_sprint(path):
            try:
                return E.fetch_stories_in_iteration(app.project, path) or []
            except Exception:
                return []

        agg, seen = [], set()
        with _perf(f"cp.fetch_stories ({len(paths)} sprints)"):
            with _cf.ThreadPoolExecutor(max_workers=min(8, len(paths) or 1)) as ex:
                for stories in ex.map(_one_sprint, paths):
                    for s in stories:
                        if s["id"] in seen:
                            continue
                        seen.add(s["id"])
                        agg.append({"id": s["id"], "title": s.get("title", ""),
                                    "hours": 0.0, "assignee": ""})
        if _gen != getattr(app, "_cp_stories_gen", _gen):
            return                      # sprint selection changed mid-fetch -> drop
        app._cp_rows = agg
        # Pull real Azure DevOps priority (+ state) for these stories so the plan
        # table and email show P1–P4 like the Regression Plan report (not a bare "P").
        # _fetch_meta also returns each story's parent (feature) id, which we use
        # to group the plan table by feature (collapsible, like the Regression Plan).
        try:
            with _perf(f"cp.fetch_meta+features ({len(agg)} stories)"):
                meta = _fetch_meta(app, [int(r["id"]) for r in agg])
                parent_ids = []
                for r in agg:
                    m = meta.get(int(r["id"]), {})
                    r["priority"] = m.get("priority", DEFAULT_PRIORITY)
                    if m.get("state"):
                        r["state"] = m["state"]
                    fid = m.get("parent_id")
                    r["feature_id"] = fid
                    if fid:
                        parent_ids.append(fid)
                fnames = _fetch_feature_names(app, list(set(parent_ids)))
                for r in agg:
                    r["feature_name"] = (fnames.get(r.get("feature_id"), "")
                                         if r.get("feature_id") else "")
                # Content-complexity units per story → drives a non-random estimate.
                comp = _fetch_cp_complexity(app, [int(r["id"]) for r in agg])
                for r in agg:
                    r["work_units"] = comp.get(int(r["id"]), 1.0)
        except Exception:
            pass
        _cp_estimate_and_assign(app)
        if _gen != getattr(app, "_cp_stories_gen", _gen):
            return                      # selection changed mid-fetch -> drop results
        app._cp_stories_loading = False
        # Only repaint if the user is still on the Sprint Plan screen — otherwise
        # this background completion would force a full (heavy) render of whatever
        # screen they navigated to, which reads as a freeze.
        if getattr(app, "active", None) == "testplan":
            app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


def _cp_estimate_and_assign(app):
    """Estimate each story's hours from its CONTENT COMPLEXITY (acceptance-criteria
    /description size × priority weight), mapped into the user's [min, max] range —
    so bigger / higher-priority stories get more hours instead of a random number —
    then balance the resulting HOURS across resources. Both stay editable."""
    names = list(app._cp_res_names or [])
    lo, hi = float(app._cp_est_min or 1.0), float(app._cp_est_max or 8.0)
    if hi < lo:
        lo, hi = hi, lo
    rows = app._cp_rows or []

    # complexity score = content work-units × priority weight
    def _score(r):
        units = float(r.get("work_units", 1.0) or 1.0)
        boost = PRIORITY_BOOST.get(r.get("priority", DEFAULT_PRIORITY), 1.0)
        return units * boost
    scores = [_score(r) for r in rows]
    smin = min(scores) if scores else 0.0
    smax = max(scores) if scores else 0.0
    span = (smax - smin) or 1.0
    # Map each score linearly into [lo, hi] (equal scores → midpoint), 0.5 h steps.
    for r, sc in zip(rows, scores):
        h = (lo + hi) / 2.0 if smax == smin else lo + (hi - lo) * (sc - smin) / span
        r["hours"] = round(h * 2) / 2.0

    # assignment: balance by HOURS, not story count. Greedy longest-processing-time
    # — assign each story (largest estimate first) to the least-loaded resource —
    # which is the same algorithm the Regression plan uses, so total hours per
    # person stay as even as possible. Deterministic (tie-break by id / name).
    if names:
        load = {n: 0.0 for n in names}
        for r in sorted(rows, key=lambda x: (-float(x.get("hours", 0) or 0), x["id"])):
            n = min(names, key=lambda nm: (load[nm], nm))
            r["assignee"] = n
            load[n] = round(load[n] + float(r.get("hours", 0) or 0), 2)
    else:
        for r in rows:
            r["assignee"] = ""


def _cp_payload(app):
    """Build the same payload shape plan_payload() returns, from the sprint rows —
    so every exporter and the email path work unchanged."""
    names = list(app._cp_res_names or [])
    rows = [{"id": r["id"], "title": r.get("title", ""), "state": r.get("state", ""),
             "priority": r.get("priority", DEFAULT_PRIORITY), "cases": 0, "boost": 1.0,
             "hours": round(float(r.get("hours", 0) or 0), 2),
             "assignee": r.get("assignee", "")} for r in (app._cp_rows or [])]
    total_hours = round(sum(r["hours"] for r in rows), 2)
    count = len(names) or 1
    per_person = round(total_hours / count, 2)
    workload = []
    if names:
        st = {n: 0 for n in names}
        hr = {n: 0.0 for n in names}
        for r in rows:
            a = r.get("assignee")
            if a in st:
                st[a] += 1
                hr[a] += r["hours"]
        workload = [{"name": n, "stories": st[n], "cases": 0,
                     "hours": round(hr[n], 2)} for n in names]
    return {"generated": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "project": app.project, "plan_id": "",
            "plan_name": app._cp_sprint_name or ", ".join(app._cp_sprint_paths or []),
            "report_title": "Sprint Plan", "mode": "create",
            "plans": [], "avg_minutes_per_case": 0,
            "priority_boost": {}, "resources_count": len(names),
            "resource_names": names, "stories": rows, "workload": workload,
            "total_stories": len(rows), "total_cases": 0,
            "total_hours": total_hours, "hours_per_person": per_person}


def _mode_toggle(app):
    """Segmented switch between the existing 'from test plans' report and the new
    'from a sprint' report."""
    def _seg(label, mode_id):
        on = (getattr(app, "_reg_mode", "existing") == mode_id)
        def _go(e, m=mode_id):
            app._reg_mode = m
            app.ui_safe(app.render)
        return ft.Container(
            ft.Text(label, size=12.5, weight=ft.FontWeight.BOLD,
                    color=("#FFFFFF" if on else T.INK_3)),
            on_click=_go, padding=ft.Padding.symmetric(vertical=8, horizontal=16),
            bgcolor=(T.VIOLET if on else T.CARD), border_radius=T.R,
            border=ft.Border.all(1, T.VIOLET if on else T.BORDER_2))
    return ft.Row([_seg("From test plans", "existing"),
                   _seg("From a sprint", "create")], spacing=8)


def test_plan_screen(app):
    """Entry point for the 'Test Plan' nav tab — a sprint-based effort report."""
    _init(app)
    from main import card, sec_head, field_label, green_btn, ghost_btn  # noqa: F401
    if not (app.connected and app.project):
        return locked_state(
            app, "Sprint Plan",
            "Plan & estimate test effort across a sprint’s stories",
            "Connect your Azure DevOps account on the Setup screen, then pick a "
            "sprint here.")
    app._reg_mode = "create"
    return _create_screen(app)


def _create_screen(app):
    from main import (card, sec_head, field_label, green_btn, ghost_btn,
                      primary_btn, searchable_dropdown)
    _cp_load_iterations(app)
    _flush_toasts(app)

    names = list(app._cp_res_names or [])
    count = app._cp_res_count
    mismatch = bool(count) and bool(names) and count != len(names)

    # ── Card 1: sprint(s) — checkbox multiselect with Select all ──
    def _sprint_names():
        names = []
        for p in app._cp_sprint_paths:
            it = next((x for x in app._cp_iterations if x["path"] == p), None)
            if it:
                names.append(_sprint_num(it["name"]) or it["name"])
        return names

    def _after_sprint_change():
        app._cp_sprint_name = ", ".join(_sprint_names())
        app._cp_calculated = False
        app._cp_calc_msg = None
        app._cp_sprint_invalid = False
        _cp_load_stories(app)

    def _toggle_sprint(key, checked):
        s = set(app._cp_sprint_paths)
        s.add(key) if checked else s.discard(key)
        app._cp_sprint_paths = [p for p in (x["path"] for x in app._cp_iterations) if p in s]
        _after_sprint_change()

    def _all_sprints(checked):
        app._cp_sprint_paths = [x["path"] for x in app._cp_iterations] if checked else []
        _after_sprint_change()

    def _open_sprints():
        app._cp_sprint_open = not app._cp_sprint_open
        # flag synced; in-place toggle handled by the component itself

    sprint_picker = (
        ft.Container(_txt("Loading sprints…", color=T.INK_3, size=12), padding=10)
        if app._cp_iter_loading else
        _checkbox_multiselect(
            [(it["path"], (_sprint_num(it["name"]) or it["name"]) + f"   ·   {it['path']}")
             for it in app._cp_iterations],
            app._cp_sprint_paths, _toggle_sprint, _all_sprints,
            is_open=app._cp_sprint_open, on_open=_open_sprints,
            placeholder="Select sprint(s)",
            empty="No sprints found for this project.",
            page=app.page, app=app, sync_key="cp_sprints",
            invalid=getattr(app, "_cp_sprint_invalid", False)))

    picked = ft.Container()
    if app._cp_sprint_paths:
        def _sprint_chip(path, name):
            return ft.Container(
                ft.Row([
                    ft.Text(name, size=12, weight=ft.FontWeight.BOLD, color=T.VIOLET_INK,
                            font_family=T.F_MONO),
                    ft.GestureDetector(
                        content=ft.Icon(ft.Icons.CLOSE, size=12, color=T.VIOLET_INK),
                        on_tap=(lambda e, pp=path: _toggle_sprint(pp, False)),
                        mouse_cursor=ft.MouseCursor.CLICK),
                ], spacing=5, tight=True),
                padding=ft.Padding.only(left=10, right=7, top=5, bottom=5),
                bgcolor=T.VIOLET_SOFT, border_radius=T.R_SM,
                border=ft.Border.all(1, "#D9D2FF"),
                on_hover=_chip_hover, animate_scale=120)
        _sp_chips = []
        for _p in app._cp_sprint_paths:
            _it = next((x for x in app._cp_iterations if x["path"] == _p), None)
            _nm = (_sprint_num(_it["name"]) or _it["name"]) if _it else _p
            _sp_chips.append(_sprint_chip(_p, _nm))
        if len(_sp_chips) > 1:
            _sp_chips.append(_clear_chip(lambda e: _all_sprints(False)))
        sprint_chips = ft.Row(_sp_chips, wrap=True, spacing=6, run_spacing=6)
        # While loading we show nothing here — the bottom "Loading sprint stories…"
        # spinner already indicates progress. Once loaded, show the story count.
        _stories_status = (ft.Container()
                           if app._cp_stories_loading
                           else _txt(f"· {len(app._cp_rows)} stories", color=T.INK_3))
        picked = ft.Container(
            ft.Column([
                ft.Row([ft.Icon(ft.Icons.CHECK_CIRCLE, size=15, color=T.GREEN),
                        ft.Column([sprint_chips], expand=True),
                        _stories_status],
                       spacing=8, vertical_alignment=ft.CrossAxisAlignment.START),
            ], spacing=0),
            padding=ft.Padding.only(top=10))

    card1 = card(ft.Column([
        sec_head("1", "Sprint"),
        ft.Container(height=10),
        ft.Column([field_label("Sprints", req=True), sprint_picker], spacing=6),
        picked,
    ], spacing=0))

    # ── Card 2: resources (count + names + chips), like the Regression screen ──
    # mutable cell so _refresh_chips_inplace can enable calc_btn after it's built
    _cp_calc_btn_cell = [None]

    def _refresh_chips_inplace():
        name_chips.controls = [_res_chip(n) for n in app._cp_res_names]
        has = bool(app._cp_res_names)
        name_chips_wrap.visible = has
        res_count_text.visible = has
        res_count_text.value = f"{len(app._cp_res_names)} resource(s)"
        # enable/disable calc_btn in-place based on current state
        cb = _cp_calc_btn_cell[0]
        if cb is not None:
            should_enable = bool(app._cp_rows and app._cp_res_names) \
                and not getattr(app, "readonly", False)
            try:
                cb.opacity = 1.0 if should_enable else 0.45
                cb.on_click = _calculate if should_enable else None
                cb.update()
            except Exception:
                pass
        try:
            name_chips_wrap.update()
            res_count_text.update()
            name_field.update()
        except Exception:
            pass

    def _add_name(e):
        for piece in re.split(r"[,\n]+", name_field.value or ""):
            nm = piece.strip()
            if nm and nm not in app._cp_res_names:
                app._cp_res_names.append(nm)
        name_field.value = ""
        app._cp_calculated = False
        app._cp_msg = None
        if app._cp_res_names and getattr(app, "_cp_res_invalid", False):
            app._cp_res_invalid = False
            name_field.border_color = T.BORDER
            try: name_field.update()
            except Exception: pass
        _refresh_chips_inplace()

    def _remove_name(nm):
        app._cp_res_names = [n for n in app._cp_res_names if n != nm]
        app._cp_calculated = False
        _refresh_chips_inplace()

    def _on_count(e):
        v = (count_field.value or "").strip()
        app._cp_res_count = int(v) if v.isdigit() and int(v) > 0 else None
        if app._cp_res_count:
            app._cp_count_invalid = False
        # in-place: just update the count field border color; no layout change needed
        new_mismatch = bool(app._cp_res_count is not None and app._cp_res_names
                            and app._cp_res_count != len(app._cp_res_names))
        count_field.border_color = T.RED if new_mismatch else T.BORDER
        try:
            count_field.update()
        except Exception:
            pass

    def _on_min(e):
        try:
            app._cp_est_min = float(e.control.value or 1)
        except Exception:
            pass

    def _on_max(e):
        try:
            app._cp_est_max = float(e.control.value or 8)
        except Exception:
            pass

    count_field = ft.TextField(
        value=("" if count is None else str(count)), hint_text="e.g. 3",
        keyboard_type=ft.KeyboardType.NUMBER, on_blur=_on_count, on_submit=_on_count,
        input_filter=_digits_only(),
        width=92, text_size=13,
        border_color=(T.RED if (mismatch or getattr(app, "_cp_count_invalid", False))
                      else T.BORDER), focused_border_color=T.VIOLET,
        border_radius=T.R,
        content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))
    name_field = ft.TextField(
        hint_text="Type a tester's name, press Enter (or paste comma-separated)",
        on_submit=_add_name, on_blur=_add_name, expand=True, text_size=13,
        border_color=(T.RED if getattr(app, "_cp_res_invalid", False) else T.BORDER),
        focused_border_color=T.VIOLET, border_radius=T.R,
        content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))

    def _res_chip(nm):
        init, col = _av(nm)
        return ft.Container(
            ft.Row([
                ft.Container(ft.Text(init, size=10, weight=ft.FontWeight.BOLD,
                                     color="#FFFFFF"),
                             width=20, height=20, bgcolor=col, border_radius=20,
                             alignment=ft.Alignment.CENTER),
                ft.Text(nm, size=12.5, weight=ft.FontWeight.BOLD, color=T.INK),
                ft.GestureDetector(
                    content=ft.Icon(ft.Icons.CLOSE, size=12, color=T.INK_3),
                    on_tap=(lambda e, x=nm: _remove_name(x)),
                    mouse_cursor=ft.MouseCursor.CLICK)],
               spacing=7, tight=True),
            padding=ft.Padding.only(left=5, right=9, top=4, bottom=4),
            bgcolor=T.CARD_2, border_radius=999, border=ft.Border.all(1, T.BORDER_2),
            on_hover=_chip_hover, animate_scale=120)

    name_chips = ft.Row([_res_chip(n) for n in app._cp_res_names],
                        wrap=True, spacing=8, run_spacing=8)
    name_chips_wrap = ft.Container(name_chips, padding=ft.Padding.only(top=10),
                                   visible=bool(app._cp_res_names))
    res_count_text = ft.Text(f"{len(app._cp_res_names)} resource(s)", size=11, color=T.INK_3,
                             weight=ft.FontWeight.BOLD, visible=bool(app._cp_res_names))

    warn = ft.Container()
    if mismatch:
        more = "more names than the number" if len(names) > count else \
               "fewer names than the number"
        warn = ft.Container(
            ft.Row([ft.Icon(ft.Icons.WARNING_AMBER, size=15, color=T.AMBER),
                    ft.Text(f"You set {count} resource(s) but added {len(names)} "
                            f"name(s) — {more}.", size=12, color=T.AMBER,
                            weight=ft.FontWeight.W_500, expand=True)], spacing=8),
            padding=10, bgcolor=T.AMBER_SOFT, border_radius=T.R,
            border=ft.Border.all(1, "#EAD9A8"), margin=ft.Margin.only(top=10))

    def _num(v, on_change):
        return ft.TextField(value=str(v), on_change=on_change, width=92, text_size=13,
                            border_color=T.BORDER, focused_border_color=T.VIOLET,
                            border_radius=T.R, keyboard_type=ft.KeyboardType.NUMBER,
                            content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))

    card2 = card(ft.Column([
        sec_head("2", "Resources & estimate"),
        ft.Container(height=10),
        ft.Row([
            ft.Column([field_label("Count", req=True), count_field], spacing=6, tight=True),
            ft.Column([field_label("Add a name", req=True),
                       ft.Row([name_field,
                               green_btn("Add", icon=ft.Icons.ADD, on_click=_add_name)],
                              spacing=8)], spacing=6, expand=True),
        ], spacing=14, vertical_alignment=ft.CrossAxisAlignment.START),
        name_chips_wrap,
        res_count_text,
        warn,
        ft.Container(height=14),
        ft.Row([
            ft.Column([field_label("Min h / story"), _num(app._cp_est_min, _on_min)],
                      spacing=6),
            ft.Column([field_label("Max h / story"), _num(app._cp_est_max, _on_max)],
                      spacing=6),
            ft.Container(
                ft.Row([
                    ft.Icon(ft.Icons.AUTO_GRAPH, size=16, color=T.VIOLET_INK),
                    ft.Column([
                        ft.Text("Estimates are complexity-based",
                                size=11.5, weight=ft.FontWeight.BOLD, color=T.VIOLET_INK),
                        _txt("Each story's hours come from its content size "
                             "(acceptance criteria + description) weighted by priority, "
                             "scaled into this Min–Max range — bigger / higher-priority "
                             "stories get more. Workload is then balanced by hours across "
                             "resources. Hours & assignees stay editable below.",
                             color=T.INK_2, size=11.5, no_wrap=False),
                    ], spacing=2, tight=True, expand=True),
                ], spacing=8, vertical_alignment=ft.CrossAxisAlignment.START),
                expand=True, padding=12,
                bgcolor=getattr(T, "VIOLET_SOFT", T.CARD_2), border_radius=T.R,
                border=ft.Border.all(1, "#D9D2FF")),
        ], spacing=14, vertical_alignment=ft.CrossAxisAlignment.START),
    ], spacing=0))

    # ── Assign & Estimate button ──
    # Mutable refs so _calculate can update these in-place.
    cp_calc_note_text = ft.Text("", size=12, color=T.AMBER, weight=ft.FontWeight.W_500,
                                expand=True)
    cp_calc_note_wrap = ft.Container(
        ft.Row([ft.Icon(ft.Icons.INFO_OUTLINE, size=15, color=T.AMBER), cp_calc_note_text],
               spacing=8),
        padding=10, bgcolor=T.AMBER_SOFT, border_radius=T.R,
        border=ft.Border.all(1, "#EAD9A8"), margin=ft.Margin.only(top=10),
        visible=bool(app._cp_calc_msg))
    if app._cp_calc_msg:
        cp_calc_note_text.value = app._cp_calc_msg

    def _calculate(e):
        if getattr(app, "readonly", False):
            return app._toast("Read-only — your role can’t generate plans.")
        # Don't let two plans generate at once — running both starves Python's GIL
        # and makes every render freeze for tens of seconds (see qa_perf.log).
        if app._reg_busy:
            app._err("The Regression plan is still generating — let it finish, then "
                     "generate the Sprint plan.")
            return
        if getattr(app, "_auto_running", False):
            app._err("Automation is running — let it finish before generating a plan.")
            return
        if not app._cp_rows:
            app._cp_calc_msg = "Pick a sprint with stories first."
            app._cp_sprint_invalid = True
            app.render()   # repaint so the sprint picker turns red
            return
        if not app._cp_res_count:
            app._cp_calc_msg = "Enter the resource count."
            app._cp_count_invalid = True
            app.render()   # repaint so the count field turns red
            return
        if not app._cp_res_names:
            app._cp_calc_msg = "Add at least one resource name first."
            app._cp_res_invalid = True
            app.render()   # repaint so the resource field turns red
            return
        if app._cp_res_count != len(app._cp_res_names):
            app._cp_calc_msg = (f"Resource count ({app._cp_res_count}) must match the "
                                f"number of names added ({len(app._cp_res_names)}).")
            app._cp_count_invalid = True
            app._cp_res_invalid = True
            app.render()   # repaint so both fields turn red
            return
        app._cp_calc_msg = None
        app._cp_sprint_invalid = False
        app._cp_count_invalid = False
        app._cp_res_invalid = False
        cp_calc_note_wrap.visible = False
        try: cp_calc_note_wrap.update()
        except Exception: pass

        # Flip to busy and re-render once so the previous sprint table is replaced
        # immediately by the spinner + skeleton. results is gated off while busy, so
        # this repaint is light; _do() then runs the estimate/assign and the final
        # render (deferred via ui_safe) swaps in the fresh table.
        app._cp_busy = True
        app.render()

        def _do():
            with _perf(f"cp.generate_render ({len(app._cp_rows)} stories)"):
                _cp_estimate_and_assign(app)
                app._cp_calculated = True
                app._cp_busy = False
                if getattr(app, "active", None) == "testplan":
                    app.render()   # result table must appear — full render here
        app.ui_safe(_do)

    calc_btn = primary_btn("Generating…" if app._cp_busy else "Generate Sprint Plan",
                           icon=ft.Icons.CALCULATE,
                           on_click=(None if app._cp_busy else _calculate),
                           disabled=app._cp_busy or not (app._cp_rows and app._cp_res_names))
    _cp_calc_btn_cell[0] = calc_btn   # wire so _refresh_chips_inplace can enable it

    # ── results / plan (after Assign & Estimate) ──
    results = None
    # While generating, don't build the old table — the body shows the spinner +
    # skeleton instead, then the fresh table replaces it on completion.
    if app._cp_calculated and app._cp_rows and app._cp_res_names and not app._cp_busy:

        # --- live, in-place builders (no full re-render → scroll is preserved) ---
        def _kpis():
            d2 = plan_payload(app)
            return [
                _kpi_tile("STORIES", str(d2["total_stories"])),
                _kpi_tile("TOTAL EFFORT", f"{d2['total_hours']} h"),
                _kpi_tile("PER PERSON", f"{d2['hours_per_person']} h", T.GREEN),
            ]

        def _workload():
            d2 = plan_payload(app)
            if not d2["workload"]:
                return ft.Container()
            maxw = max((w["hours"] for w in d2["workload"]), default=0) or 1
            cards_wl = [ft.Container(ft.Column([
                ft.Row([_avatar(w["name"], 32),
                        ft.Column([_txt(w["name"], color=T.INK, weight=ft.FontWeight.BOLD, size=14),
                                   _txt(f"{w['stories']} stories", color=T.INK_3, size=11)],
                                  spacing=1, tight=True, expand=True),
                        _txt(f"{w['hours']} h", color=T.INK, weight=ft.FontWeight.BOLD,
                             size=16, no_wrap=True)],
                       spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                ft.Container(height=12), _bar(w["hours"] / maxw, T.VIOLET, 8),
            ], spacing=0), width=300, padding=14, bgcolor=T.CARD,
                border=ft.Border.all(1, T.BORDER_2), border_radius=T.R)
                for w in d2["workload"]]
            return ft.Column([
                ft.Container(height=16),
                ft.Text("RESOURCE WORKLOAD", size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
                ft.Container(height=10),
                ft.Row(cards_wl, spacing=14, wrap=True, run_spacing=14,
                       vertical_alignment=ft.CrossAxisAlignment.START)], spacing=0)

        def _refresh_totals():
            kpi_strip.controls = _kpis()
            workload_holder.content = _workload()
            kpi_strip.update(); workload_holder.update()

        def _refresh_all():
            _cp_row_cache.clear()          # row set changed (delete) -> drop memo
            plan_col.controls = [hdr] + _trows()
            plan_col.update(); _refresh_totals()

        def _edit_hours(sid):
            def _h(e):
                try:
                    v = float(e.control.value or 0)
                except Exception:
                    return
                for r in app._cp_rows:
                    if r["id"] == sid:
                        r["hours"] = round(v, 2); break
                _refresh_totals()          # update KPIs + workload, keep field focused
            return _h

        def _edit_assignee(sid):
            def _a(e):
                for r in app._cp_rows:
                    if r["id"] == sid:
                        r["assignee"] = e.control.value or ""; break
                _refresh_totals()
            return _a

        def _delete_story(sid):
            def _do():
                app._cp_rows = [r for r in app._cp_rows if r["id"] != sid]
                _refresh_all()             # rebuild table + recalc, no scroll jump
                try:
                    app._toast(f"Removed story {sid} from the sprint plan.")
                except Exception:
                    pass
            def _d(e):
                if getattr(app, "readonly", False):
                    return app._toast("Read-only — your role can’t modify the plan.")
                app._confirm(
                    "Remove story?",
                    f"Remove story {sid} from the sprint plan and recalculate the "
                    "workload? This doesn't change anything in Azure DevOps.",
                    _do, yes_label="Remove")
            return _d

        hdr = ft.Container(
            ft.Row([ft.Container(width=34),
                    _txt("STORY", color=T.INK_2, size=10.5, weight=ft.FontWeight.BOLD, width=84),
                    _txt("TITLE", color=T.INK_2, size=10.5, weight=ft.FontWeight.BOLD, expand=True),
                    _txt("P", color=T.INK_2, size=10.5, weight=ft.FontWeight.BOLD, width=44),
                    _txt("HOURS", color=T.INK_2, size=10.5, weight=ft.FontWeight.BOLD, width=110),
                    _txt("ASSIGNEE", color=T.INK_2, size=10.5, weight=ft.FontWeight.BOLD, width=180)],
                   spacing=10),
            padding=ft.Padding.symmetric(vertical=11, horizontal=12),
            bgcolor=T.CARD_2,
            border=ft.Border.only(bottom=ft.BorderSide(1, T.BORDER)))

        # Memoize editable rows by story id so a collapse / page-flip REUSES the
        # already-built Dropdown+TextField controls instead of reconstructing them
        # (rebuilding ~25 dropdowns per interaction was the Sprint-table lag).
        _cp_row_cache = {}

        def _story_row(r, i):
            hours_f = ft.TextField(
                value=str(r["hours"]), on_change=_edit_hours(r["id"]),
                width=92, text_size=13, border_color=T.BORDER,
                focused_border_color=T.VIOLET, border_radius=T.R,
                keyboard_type=ft.KeyboardType.NUMBER,
                content_padding=ft.Padding.symmetric(vertical=8, horizontal=8))
            assignee_dd = ft.Dropdown(
                value=r["assignee"] or None, width=142, text_size=13,
                options=[ft.DropdownOption(key=n, text=n) for n in app._cp_res_names],
                on_select=_edit_assignee(r["id"]), border_color=T.BORDER,
                border_radius=T.R, content_padding=ft.Padding.symmetric(vertical=6, horizontal=8))
            assignee_cell = ft.Row(
                [_avatar(r.get("assignee", ""), 26), assignee_dd],
                spacing=8, vertical_alignment=ft.CrossAxisAlignment.CENTER)
            del_btn = ft.IconButton(
                icon=ft.Icons.DELETE_OUTLINE, icon_size=18, icon_color=T.RED,
                tooltip="Remove story & recalculate",
                on_click=_delete_story(r["id"]),
                width=34, height=34,
                style=ft.ButtonStyle(padding=ft.Padding.all(0),
                                     shape=ft.RoundedRectangleBorder(radius=8)))
            return ft.Container(
                ft.Row([ft.Container(del_btn, width=34),
                        _id_link(app, r["id"], color=T.VIOLET_INK,
                                 weight=ft.FontWeight.BOLD, width=84,
                                 font_family=T.F_MONO),
                        _txt(r["title"] or "—", color=T.INK, expand=True),
                        ft.Container(_pri_pill(r.get("priority", DEFAULT_PRIORITY)), width=44),
                        ft.Container(hours_f, width=110),
                        ft.Container(assignee_cell, width=180)],
                       spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.Padding.symmetric(vertical=6, horizontal=12),
                bgcolor=(T.CARD if i % 2 == 0 else T.CARD_2),
                border=ft.Border.only(bottom=ft.BorderSide(1, T.BORDER_2)))

        def _row_for(r, i):
            c = _cp_row_cache.get(r["id"])
            if c is None:
                c = _story_row(r, i)
                _cp_row_cache[r["id"]] = c
            else:
                c.bgcolor = (T.CARD if i % 2 == 0 else T.CARD_2)   # keep zebra correct
            return c

        def _toggle_cp_feature(fname):
            def _h(e):
                _off = getattr(app, "_scroll_offset", 0) or 0
                cf = set(getattr(app, "_cp_collapsed_features", set()))
                cf.discard(fname) if fname in cf else cf.add(fname)
                app._cp_collapsed_features = cf
                try:
                    plan_col.controls = [hdr] + _trows()
                    plan_col.update()
                    _keep_scroll(app, _off)
                except Exception:
                    app.render()
            return _h

        _CP_PAGE = 25  # paginate so a 500-story sprint doesn't build 4000+ controls

        def _trows():
            # Paginate FIRST (one small page of stories at a time), then group that
            # page by feature with collapsible headers — same pattern as the
            # Regression Plan table. Building every story's editable row at once was
            # the cause of the multi-second / multi-minute Sprint Plan renders.
            from collections import OrderedDict as _OD
            allrows = app._cp_rows or []
            # If the resource roster changed, cached dropdown options are stale —
            # drop the cache so rows rebuild with the new assignee list.
            _sig = tuple(app._cp_res_names or [])
            if _cp_row_cache.get("__sig") != _sig:
                _cp_row_cache.clear()
                _cp_row_cache["__sig"] = _sig
            total = max(1, -(-len(allrows) // _CP_PAGE))
            pg = max(0, min(getattr(app, "_cp_table_page", 0), total - 1))
            page_rows = allrows[pg * _CP_PAGE:(pg + 1) * _CP_PAGE]

            groups = _OD()
            for r in page_rows:
                groups.setdefault(r.get("feature_name") or "No Feature", []).append(r)
            has_features = any(g != "No Feature" for g in groups)

            out, i = [], 0
            if not has_features:
                out = [_row_for(r, k) for k, r in enumerate(page_rows)]
            else:
                coll = getattr(app, "_cp_collapsed_features", set())
                for fname, frows in groups.items():
                    is_c = fname in coll
                    fid = frows[0].get("feature_id") if frows else None
                    fl = f"[{fid}]  " if fid else ""
                    out.append(_feature_header(app, f"{fl}{fname}", len(frows),
                                               is_c, _toggle_cp_feature(fname)))
                    if not is_c:
                        for r in frows:
                            out.append(_row_for(r, i))
                            i += 1

            if total > 1:
                def _cp_go(delta, p=pg, t=total):
                    def _h(e):
                        _off = getattr(app, "_scroll_offset", 0) or 0
                        app._cp_table_page = max(0, min(p + delta, t - 1))
                        try:
                            plan_col.controls = [hdr] + _trows()
                            plan_col.update()
                            _keep_scroll(app, _off)
                        except Exception:
                            app.render()
                    return _h
                pager = ft.Row([
                    ghost_btn("← Prev", on_click=(None if pg == 0 else _cp_go(-1))),
                    _txt(f"Page {pg + 1} of {total}  ·  "
                         f"rows {pg * _CP_PAGE + 1}–{min((pg + 1) * _CP_PAGE, len(allrows))} "
                         f"of {len(allrows)}", size=12, color=T.INK_3),
                    ghost_btn("Next →", on_click=(None if pg >= total - 1 else _cp_go(1))),
                ], alignment=ft.MainAxisAlignment.CENTER, spacing=16)
                out = out + [ft.Container(pager,
                                          padding=ft.Padding.symmetric(vertical=10))]
            return out

        kpi_strip = ft.Row(_kpis(), spacing=14)
        plan_col = ft.Column([hdr] + _trows(), spacing=0)
        table = ft.Container(plan_col, border=ft.Border.all(1, T.BORDER),
                             border_radius=T.R, clip_behavior=ft.ClipBehavior.HARD_EDGE)
        workload_holder = ft.Container(content=_workload())
        workload_ui = workload_holder

        _cp_status_cell = [None]

        def _set_cp_status(kind, text):
            app._cp_msg = (kind, text)
            sc = _cp_status_cell[0]
            if sc is None:
                app.ui_safe(app.render); return
            ok = (kind == "ok")
            sc.bgcolor = T.CARD
            sc.visible = True
            try:
                row = sc.content
                row.controls[0].name = (ft.Icons.CHECK_CIRCLE if ok
                                        else ft.Icons.ERROR_OUTLINE)
                row.controls[0].color = T.GREEN if ok else T.RED
                row.controls[1].value = text
                row.controls[1].color = T.GREEN if ok else T.RED
                sc.update()
            except Exception:
                app.ui_safe(app.render)

        exports_col, _cp_export_status = _export_row(app, _set_cp_status)
        _cp_status_cell[0] = _cp_export_status

        def _email(e):
            to = [a.strip() for a in re.split(r"[,\s;]+", (email_field.value or ""))
                  if a.strip()]
            if not to:
                app._err("Enter at least one recipient email.")
                return
            if not getattr(E, "GMAIL_APP_PASS", ""):
                app._err("Set the Gmail App Password on the Setup screen first.")
                return
            app._cp_email_to = ", ".join(to)
            app._cp_emailing = True
            app._cp_msg = None
            app._toast("Sending the sprint plan…")

            def work():
                try:
                    d = plan_payload(app)
                    try:
                        attach = [export_docx(app)]
                    except Exception:
                        attach = []
                    subj = f"Sprint Plan — {d['plan_name'] or d['project']}"
                    ok, err = E.send_report(to, subj, _plan_html(d), attachments=attach)
                    kind = "ok" if ok else "err"
                    text = f"Emailed to {', '.join(to)}" if ok else (err or "Email failed.")
                except Exception as ex:
                    kind, text = "err", f"Email failed: {str(ex)[:160]}"
                app._cp_emailing = False
                app.ui_safe(lambda k=kind, t=text: (app._toast(t) if k == "ok"
                                                    else app._err(t)))
            threading.Thread(target=work, daemon=True).start()

        email_field = ft.TextField(
            value=app._cp_email_to, hint_text="name@company.com, another@company.com",
            on_change=lambda e: setattr(app, "_cp_email_to", e.control.value or ""),
            expand=True, text_size=13, border_color=T.BORDER,
            focused_border_color=T.VIOLET, border_radius=T.R,
            content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))
        email_row = ft.Column([
            ft.Text("EMAIL", size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
            ft.Container(height=8),
            ft.Row([email_field,
                    green_btn("Sending…" if app._cp_emailing else "Email plan",
                              icon=ft.Icons.SEND,
                              on_click=(None if app._cp_emailing else _email))],
                   spacing=8),
        ], spacing=0)

        def _assign_testers(e):
            rows = [{"id": r["id"], "name": r.get("assignee", "")}
                    for r in app._cp_rows if r.get("assignee")]
            if not rows:
                _set_cp_status("err", "Assign resources to stories first.")
                return
            app._cp_assigning = True
            app._cp_msg = None

            def work():
                try:
                    res = E.assign_testers(app.project, rows)
                except Exception as ex:
                    res = {"ok": 0, "errors": [str(ex)[:160]]}
                errs = res.get("errors", [])
                n = res.get("ok", 0)
                if n and not errs:
                    kind, text = "ok", (f"Assigned {n} stories to the "
                                        f"Assigned To Tester field in Azure.")
                elif n:
                    kind, text = "err", (f"Assigned {n}; {len(errs)} failed — "
                                         + "  ·  ".join(errs[:4])
                                         + ("  …" if len(errs) > 4 else ""))
                else:
                    kind, text = "err", ("  ·  ".join(errs[:5]) or "Nothing assigned.")
                app._cp_assigning = False
                app.ui_safe(lambda k=kind, t=text: _set_cp_status(k, t))
            threading.Thread(target=work, daemon=True).start()

        assign_note = ft.Container(
            ft.Row([ft.Icon(ft.Icons.INFO_OUTLINE, size=15, color=T.INK_3),
                    _txt("Writes each story's assignee into the Azure “Assigned To "
                         "Tester” field. Names are matched to that field's list — "
                         "you'll get a readable error for any that don't match.",
                         color=T.INK_3, size=11.5, expand=True)], spacing=8),
            padding=10, bgcolor=T.CARD, border_radius=T.R,
            border=ft.Border.all(1, T.BORDER_2), margin=ft.Margin.only(top=12))

        results = card(ft.Column([
            sec_head("3", "Plan"), ft.Container(height=12), kpi_strip,
            ft.Container(height=14), table, workload_ui,
            ft.Divider(height=22, color=T.BORDER),
            ft.Text("EXPORT", size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
            ft.Container(height=8), exports_col,
            ft.Divider(height=22, color=T.BORDER),
            email_row,
            ft.Container(height=14),
            ft.Row([green_btn("Assigning…" if app._cp_assigning
                              else "Assign to tester in Azure",
                              icon=ft.Icons.PERSON_ADD,
                              on_click=(None if app._cp_assigning else _assign_testers))]),
            assign_note,
        ], spacing=0))

    # Generating / loading indicator (same style as the Regression Plan spinner).
    _cp_spin_label = ("Generating sprint plan…" if app._cp_busy
                      else "Loading sprint stories from Azure DevOps…")
    cp_spinner = ft.Container(
        ft.Row([ft.ProgressRing(width=18, height=18, stroke_width=2.5, color=T.VIOLET),
                ft.Text(_cp_spin_label,
                        size=12.5, color=T.INK_3, weight=ft.FontWeight.W_500)],
               spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
        padding=ft.Padding.symmetric(vertical=10, horizontal=12),
        margin=ft.Margin.only(top=10),
        bgcolor=getattr(T, "VIOLET_SOFT", T.CARD_2), border_radius=T.R,
        visible=bool(app._cp_stories_loading or app._cp_busy))

    body_children = [card1, ft.Container(height=14), card2,
                     ft.Container(height=16), calc_btn, cp_spinner, cp_calc_note_wrap]
    if results is not None:
        body_children += [ft.Container(height=16), results]
    elif app._cp_busy:
        # Skeleton only while GENERATING (matches the Regression Plan). During the
        # pre-generation story fetch the small "Loading sprint stories…" spinner is
        # enough — the big skeleton there just made the screen jump/expand.
        from main import card as _card, skeleton_rows as _skel
        body_children += [ft.Container(height=16), _card(_skel(6))]
    body = ft.Column(body_children, spacing=0, scroll=ft.ScrollMode.AUTO, expand=True)
    return app.shell("Sprint Plan",
                     "Plan & estimate test effort across a sprint’s stories", body,
                     badge="STEP T")


def _flush_toasts(app):
    """Surface any newly-set status message (export / email / assign / calc) as a
    floating toast, in addition to the inline banner. Guards by remembered text so
    it fires once per new message, not on every render."""
    seen = getattr(app, "_toast_seen", None)
    if seen is None:
        seen = app._toast_seen = {}
    for slot in ("_reg_export_msg", "_reg_calc_msg", "_cp_msg", "_cp_calc_msg"):
        val = getattr(app, slot, None)
        if isinstance(val, (tuple, list)) and len(val) == 2:
            kind, text = val[0], val[1]
        elif isinstance(val, str) and val:
            kind, text = "err", val
        else:
            seen[slot] = None
            continue
        if not text or seen.get(slot) == text:
            continue
        seen[slot] = text
        try:
            if kind == "ok":
                app.ui_safe(lambda t=text: app._toast(t))
            else:
                app.ui_safe(lambda t=text: app._err(t))
        except Exception:
            pass


def screen(app):
    _init(app)
    _flush_toasts(app)
    from main import (card, sec_head, field_label, green_btn, ghost_btn,
                      primary_btn, searchable_dropdown)

    # ── gate: only needs a connection + a project ──
    if not (app.connected and app.project):
        return locked_state(
            app, "Regression Plan",
            "Build a regression plan from your test plans & their stories",
            "Connect your Azure DevOps account on the Setup screen. You can pick "
            "the test plans right here once connected.")

    app._reg_mode = "existing"

    # lazy-load test plans
    if not app._plans and not app._reg_plans_loading:
        app._reg_plans_loading = True

        def _lp():
            try:
                app._load_plans()
            except Exception:
                pass
            app._reg_plans_loading = False
            app.ui_safe(app.render)
        threading.Thread(target=_lp, daemon=True).start()

    selected_ids = [s["id"] for s in app._reg_selected]
    selected_plan_ids = [p["id"] for p in app._reg_plans_selected]

    # ── handlers ──
    def _add_plan(e):
        v = plan_dd.value
        if not v or not str(v).strip().isdigit():
            return
        pid = int(v)
        if pid not in selected_plan_ids:
            name = next((p["name"] for p in (app._plans or []) if p["id"] == pid), str(pid))
            app._reg_plans_selected.append({"id": pid, "name": name})
            # keep app.plan_id pointing at the first selected plan (Setup/back-compat)
            if not app.plan_id:
                app.plan_id = pid
                app.plan_name = name
            app._reg_selected_rows = []
            app._reg_export_msg = app._reg_calc_msg = None
            _reload_plan_stories(app)
        app.render()

    def _remove_plan(pid):
        app._reg_plans_selected = [p for p in app._reg_plans_selected if p["id"] != pid]
        # drop any selected stories that belonged to the removed plan
        app._reg_selected = [s for s in app._reg_selected if s.get("plan_id") != pid]
        # invalidate cached suite map for this plan
        try:
            app._reg_suite_cache.pop(pid, None)
        except Exception:
            pass
        if app.plan_id == pid:
            app.plan_id = (app._reg_plans_selected[0]["id"] if app._reg_plans_selected else None)
            app.plan_name = (app._reg_plans_selected[0]["name"] if app._reg_plans_selected else "")
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        _sync = (getattr(app, "_dd_syncers", {}) or {}).get("reg_plans")
        if _sync:                        # untick the plan in the open dropdown
            try:
                _sync([str(p["id"]) for p in app._reg_plans_selected])
            except Exception:
                pass
        _reload_plan_stories(app)
        app.render()

    def _add_story(e):
        v = story_dd.value
        if not v or not str(v).strip().isdigit():
            return
        sid = int(v)
        if sid not in selected_ids:
            src = next((s for s in app._reg_plan_stories if s["id"] == sid), None)
            if src:
                app._reg_selected.append({"id": sid, "title": src.get("title", ""),
                                          "plan_id": src.get("plan_id")})
                app._reg_selected_rows = []
                app._reg_export_msg = app._reg_calc_msg = None
        app.render()

    def _remove_story(sid):
        app._reg_selected = [s for s in app._reg_selected if s["id"] != sid]
        app._reg_export_msg = app._reg_calc_msg = None
        _had_results = bool(app._reg_selected_rows)
        app._reg_selected_rows = []
        _sync = (getattr(app, "_dd_syncers", {}) or {}).get("reg_stories")
        if _had_results or not _sync:
            app.render()                 # a generated table must clear -> full render
        else:
            try:
                _sync([str(s["id"]) for s in app._reg_selected])   # untick in dropdown
                _refresh_story_externals()                         # chips + button
            except Exception:
                app.render()

    def _delete_row(sid):
        # inline-delete from the calculated plan table + recalculate
        def _do():
            app._reg_selected_rows = [r for r in (app._reg_selected_rows or [])
                                      if r.get("id") != sid]
            app._reg_selected = [s for s in app._reg_selected if s["id"] != sid]
            app._reg_export_msg = app._reg_calc_msg = None
            app.render()
            try:
                app._toast(f"Removed story {sid} from the regression plan.")
            except Exception:
                pass
        def _d(e):
            if getattr(app, "readonly", False):
                return app._toast("Read-only — your role can’t modify the plan.")
            app._confirm(
                "Remove story?",
                f"Remove story {sid} from the regression plan and recalculate the "
                "effort? This doesn't change anything in Azure DevOps.",
                _do, yes_label="Remove")
        return _d

    def _use_setup_selection(e):
        # ensure the Setup plan is among the selected plans so its cases count
        if app.plan_id and app.plan_id not in selected_plan_ids:
            app._reg_plans_selected.append(
                {"id": app.plan_id, "name": getattr(app, "plan_name", "") or str(app.plan_id)})
            _reload_plan_stories(app)
        for sid in (app.story_ids or []):
            if int(sid) not in [s["id"] for s in app._reg_selected]:
                app._reg_selected.append({"id": int(sid), "title": "", "plan_id": app.plan_id})
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        app.render()

    def _add_name(e):
        raw = (name_field.value or "")
        # accept several names at once, separated by commas (or newlines)
        for piece in re.split(r"[,\n]+", raw):
            nm = piece.strip()
            if nm and nm not in app._reg_res_names:
                app._reg_res_names.append(nm)
        name_field.value = ""
        app._reg_export_msg = None
        if app._reg_res_names and getattr(app, "_reg_res_invalid", False):
            app._reg_res_invalid = False
            name_field.border_color = T.BORDER
        try:   # in-place update keeps the scroll position (full render snaps to top)
            name_chips.controls = [_res_chip(n, (lambda e, x=n: _remove_name(x)))
                                   for n in app._reg_res_names]
            name_chips_wrap.visible = bool(app._reg_res_names)
            name_chips_wrap.update(); name_field.update()
        except Exception:
            app.render()

    def _remove_name(nm):
        app._reg_res_names = [n for n in app._reg_res_names if n != nm]
        app._reg_export_msg = None
        try:
            name_chips.controls = [_res_chip(n, (lambda e, x=n: _remove_name(x)))
                                   for n in app._reg_res_names]
            name_chips_wrap.visible = bool(app._reg_res_names)
            name_chips_wrap.update()
        except Exception:
            app.render()

    def _on_count(e):
        v = (count_field.value or "").strip()
        app._reg_res_count = int(v) if v.isdigit() and int(v) > 0 else None
        app._reg_export_msg = None
        if app._reg_res_count:
            app._reg_count_invalid = False
        new_mismatch = bool(app._reg_res_count is not None and app._reg_res_names
                            and app._reg_res_count != len(app._reg_res_names))
        count_field.border_color = T.RED if new_mismatch else T.BORDER
        try:
            count_field.update()
        except Exception:
            pass

    # Mutable refs so _calculate can update these widgets in-place.
    # They are populated after the layout code below — Python closures capture
    # by reference so mutations made later are visible here.
    calc_note_text = ft.Text("", size=12, color=T.AMBER, weight=ft.FontWeight.W_500,
                             expand=True)
    calc_note_wrap = ft.Container(
        ft.Row([ft.Icon(ft.Icons.INFO_OUTLINE, size=15, color=T.AMBER), calc_note_text],
               spacing=8),
        padding=10, bgcolor=T.AMBER_SOFT, border_radius=T.R,
        border=ft.Border.all(1, "#EAD9A8"), margin=ft.Margin.only(top=10),
        visible=bool(app._reg_calc_msg))
    if app._reg_calc_msg:
        calc_note_text.value = app._reg_calc_msg
    # calc_btn_ref is set after primary_btn() is called below; use a list cell
    # as a mutable reference so the closure can reach the button.
    _calc_btn_cell = [None]
    _gen_spinner_cell = [None]

    def _get_calc_btn():
        return _calc_btn_cell[0]

    def _calculate(e):
        if getattr(app, "readonly", False):
            return app._toast("Read-only — your role can’t generate plans.")
        # Don't let two plans generate at once — running both starves Python's GIL
        # and makes every render freeze for tens of seconds (see qa_perf.log).
        if app._cp_busy or app._cp_stories_loading:
            app._err("The Sprint plan is still working — let it finish, then generate "
                     "the Regression plan.")
            return
        if getattr(app, "_auto_running", False):
            app._err("Automation is running — let it finish before generating a plan.")
            return
        if not app._reg_plans_selected:
            app._reg_calc_msg = "Add at least one test plan first so effort can be " \
                                "read from its existing test cases."
            app._reg_plan_invalid = True
            app.render()   # repaint so the test-plan picker turns red
            return
        if not app._reg_selected:
            app._reg_calc_msg = "Add at least one story."
            app._reg_story_invalid = True
            app.render()   # repaint so the stories picker turns red
            return
        if not app._reg_res_count:
            app._reg_calc_msg = "Enter the resource count."
            app._reg_count_invalid = True
            app.render()   # repaint so the count field turns red
            return
        if not app._reg_res_names:
            app._reg_calc_msg = "Add at least one resource name."
            app._reg_res_invalid = True
            app.render()   # repaint so the resource field turns red
            return
        if app._reg_res_count != len(app._reg_res_names):
            app._reg_calc_msg = (f"Resource count ({app._reg_res_count}) must match the "
                                 f"number of names added ({len(app._reg_res_names)}).")
            app._reg_count_invalid = True
            app._reg_res_invalid = True
            app.render()   # repaint so both fields turn red
            return
        app._reg_busy = True
        app._reg_calc_msg = app._reg_export_msg = None
        app._reg_plan_invalid = False
        app._reg_story_invalid = False
        app._reg_count_invalid = False
        app._reg_res_invalid = False
        calc_note_wrap.visible = False
        # Re-render once so the previous plan table is immediately replaced by the
        # progress spinner + loading skeleton (the button also flips to "Generating…"
        # and the table is rebuilt from scratch when _work completes below).
        # build_rows runs off the UI thread, so this single repaint stays snappy.
        app.render()

        def _work():
            def _progress(done, total):
                # Live count feedback in the spinner so a long cold count clearly
                # reads as "working", not stuck. Throttled by _count_cases.
                def _upd():
                    sp = getattr(app, "_reg_gen_spinner", None)   # the on-screen one
                    if sp is None:
                        return
                    try:
                        sp.content.controls[1].value = (
                            f"Counting test cases — {done} / {total} suites…")
                        sp.visible = True
                        sp.update()
                    except Exception:
                        pass
                app.ui_safe(_upd)
            try:
                rows = build_rows(app, app._reg_selected, progress=_progress)
            except Exception as ex:
                rows = []
                app._reg_calc_msg = f"Couldn't read plan: {ex}"
            app._reg_selected_rows = rows
            app._reg_busy = False
            # Only repaint if still on the Regression screen — otherwise this
            # background completion forces a heavy render of whatever screen the
            # user navigated to (reads as a freeze).
            if getattr(app, "active", None) == "regression":
                app.ui_safe(app.render)
        threading.Thread(target=_work, daemon=True).start()

    def _regenerate(e):
        if getattr(app, "readonly", False):
            return app._toast("Read-only — your role can’t regenerate plans.")
        # Force-refresh: drop the cached suite maps, counts and metadata so the plan
        # is rebuilt from a fresh Azure pull — this is how new test cases in existing
        # suites (and changed priorities/states) get reflected. We KEEP the feature
        # caches (_reg_story_features / _reg_feature_name_cache): feature names rarely
        # change, and re-resolving them on every regenerate is what made it slow and
        # briefly drop the grouping. (Use Settings → Clear caches for a full reset.)
        for _attr in ("_reg_case_count_cache", "_reg_meta_cache", "_reg_suite_cache"):
            try:
                setattr(app, _attr, {})
            except Exception:
                pass
        # Mark this project's cache as already "loaded" (now empty) so the
        # generate doesn't re-warm it from the stale disk copy. build_rows will
        # refetch and then _cache_save() overwrites disk with fresh data.
        app._reg_cache_loaded_project = getattr(app, "project", None)
        _calculate(e)

    def _do_export_to(fmt, dest=None):
        try:
            path = EXPORTERS[fmt](app)
        except ImportError:
            app._reg_export_msg = ("err", f"{fmt.upper()} needs {_MISSING_DEP.get(fmt, fmt)}")
            return None
        except Exception as ex:
            app._reg_export_msg = ("err", f"Export failed: {ex}")
            return None
        if dest:
            if not dest.lower().endswith("." + fmt):
                dest += "." + fmt
            try:
                import shutil
                if os.path.abspath(dest) != os.path.abspath(path):
                    shutil.move(path, dest)
                path = dest
            except Exception as ex:
                app._reg_export_msg = ("err", f"Couldn't save there: {ex}")
                return None
        app._reg_export_msg = ("ok", f"Saved {fmt.upper()}: {path}")
        try:
            os.startfile(os.path.dirname(path))
        except Exception:
            pass
        return path

    # Mutable status ref for in-place export/email banner updates.
    _status_cell = [None]   # [0] = ft.Container ref once results card is built

    def _set_status(kind, text):
        app._reg_export_msg = (kind, text)
        sc = _status_cell[0]
        if sc is None:
            app.ui_safe(app.render); return
        ok = (kind == "ok")
        sc.bgcolor = T.GREEN_SOFT if ok else T.RED_SOFT
        sc.visible = True
        try:
            row = sc.content
            row.controls[0].name = (ft.Icons.CHECK_CIRCLE if ok
                                    else ft.Icons.ERROR_OUTLINE)
            row.controls[0].color = T.GREEN if ok else T.RED
            row.controls[1].value = text
            row.controls[1].color = T.GREEN if ok else T.RED
            sc.update()
        except Exception:
            app.ui_safe(app.render)

    def _export(fmt):
        def _do(e):
            if not app._reg_selected_rows:
                app._err("Calculate the plan first.")
                return

            def work():
                dest = _ask_save_path(fmt, _stamp(app) + "." + fmt)
                if dest is False:
                    _do_export_to(fmt)
                elif dest:
                    _do_export_to(fmt, dest)
                else:
                    return
                msg = app._reg_export_msg
                if msg:
                    app.ui_safe(lambda m=msg: (app._toast(m[1]) if m[0] == "ok"
                                               else app._err(m[1])))
            threading.Thread(target=work, daemon=True).start()
        return _do

    def _on_email_to(e):
        app._reg_email_to = (email_field.value or "").strip()

    def _email(e):
        if not app._reg_selected_rows:
            app._err("Calculate the plan first.")
            return
        to = [a.strip() for a in re.split(r"[,\s;]+", (email_field.value or ""))
              if a.strip()]
        if not to:
            app._err("Enter at least one recipient email.")
            return
        if not E.GMAIL_APP_PASS:
            app._err("Set the Gmail App Password on the Setup screen first.")
            return
        app._reg_email_to = ", ".join(to)
        app._reg_emailing = True
        app._reg_export_msg = None
        app._toast("Sending the regression plan…")

        def work():
            try:
                d = plan_payload(app)
                # The mail report shows only the sprint number, not the long
                # "<Project>_Sprint N" test-plan name.
                trimmed = ", ".join(
                    (_sprint_num(p.get("name") or p.get("sprint") or "")
                     or (p.get("name") or "").strip())
                    for p in (app._reg_plans_selected or [])).strip(", ")
                if trimmed:
                    d["plan_name"] = trimmed
                try:
                    attach = [export_docx(app)]
                except Exception:
                    attach = []
                subj = f"Regression Test Plan — {d['plan_name'] or d['project']}"
                ok, err = E.send_report(to, subj, _plan_html(d), attachments=attach)
                kind = "ok" if ok else "err"
                text = f"Emailed to {', '.join(to)}" if ok else (err or "Email failed.")
            except Exception as ex:
                kind, text = "err", f"Email failed: {ex}"
            app._reg_emailing = False
            app.ui_safe(lambda: (app._toast(text) if kind == "ok" else app._err(text)))
        threading.Thread(target=work, daemon=True).start()

    # ── validation ──
    names = app._reg_res_names
    count = app._reg_res_count
    mismatch = bool(count is not None and names and count != len(names))

    def _txt(s, **kw):
        kw.setdefault("size", 12)
        return ft.Text(s, **kw)

    # ── Card 1: source + stories ──
    # ── Test plans: checkbox multiselect with Select all ──
    def _toggle_plan(key, checked):
        pid = int(key)
        ids = [p["id"] for p in app._reg_plans_selected]
        if checked and pid not in ids:
            name = next((p["name"] for p in (app._plans or []) if p["id"] == pid), str(pid))
            app._reg_plans_selected.append({"id": pid, "name": name})
        elif not checked:
            app._reg_plans_selected = [p for p in app._reg_plans_selected if p["id"] != pid]
            app._reg_selected = [s for s in app._reg_selected if s.get("plan_id") != pid]
        app.plan_id = (app._reg_plans_selected[0]["id"] if app._reg_plans_selected else None)
        app.plan_name = (app._reg_plans_selected[0]["name"] if app._reg_plans_selected else "")
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        if app._reg_plans_selected:
            app._reg_plan_invalid = False
        _reload_plan_stories(app)

    def _all_plans(checked):
        app._reg_plans_selected = ([{"id": p["id"], "name": p["name"]}
                                    for p in (app._plans or [])] if checked else [])
        if not checked:
            app._reg_selected = []
        app.plan_id = (app._reg_plans_selected[0]["id"] if app._reg_plans_selected else None)
        app.plan_name = (app._reg_plans_selected[0]["name"] if app._reg_plans_selected else "")
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        if app._reg_plans_selected:
            app._reg_plan_invalid = False
        _reload_plan_stories(app)

    def _open_plans():
        app._reg_plan_open = not app._reg_plan_open
        app._reg_story_open = False
        # flag synced; in-place toggle handled by the component itself

    plan_picker = (
        _loading_field("Loading test plans…")
        if app._reg_plans_loading else
        _checkbox_multiselect(
            [(str(p["id"]), f"[{p['id']}] {p['name']}") for p in (app._plans or [])],
            [str(p["id"]) for p in app._reg_plans_selected],
            _toggle_plan, _all_plans, is_open=app._reg_plan_open, on_open=_open_plans,
            placeholder="Select test plan(s)", height=200,
            empty="No test plans found for this project.",
            page=app.page, app=app, sync_key="reg_plans",
            invalid=getattr(app, "_reg_plan_invalid", False)))

    # ── Stories: checkbox multiselect with Select all ──
    def _toggle_story(key, checked):
        sid = int(key)
        ids = [s["id"] for s in app._reg_selected]
        if checked and sid not in ids:
            src = next((s for s in app._reg_plan_stories if s["id"] == sid), None)
            if src:
                app._reg_selected.append({"id": sid, "title": src.get("title", ""),
                                          "plan_id": src.get("plan_id")})
        elif not checked:
            app._reg_selected = [s for s in app._reg_selected if s["id"] != sid]
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        if app._reg_selected:
            app._reg_story_invalid = False
        _fn = getattr(app, "_reg_refresh_story_ext", None)
        if callable(_fn):
            _fn()                       # keep chips + "N stories selected" in sync

    def _all_stories(checked):
        if checked:
            have = {s["id"] for s in app._reg_selected}
            for s in app._reg_plan_stories:
                if s["id"] not in have:
                    app._reg_selected.append({"id": s["id"], "title": s.get("title", ""),
                                              "plan_id": s.get("plan_id")})
        else:
            app._reg_selected = []
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        if app._reg_selected:
            app._reg_story_invalid = False
        _fn = getattr(app, "_reg_refresh_story_ext", None)
        if callable(_fn):
            _fn()                       # keep chips + "N stories selected" in sync

    def _open_stories():
        app._reg_story_open = not app._reg_story_open
        app._reg_plan_open = False
        # flag synced; in-place toggle handled by the component itself

    _have_plans = bool(app._reg_plans_selected)

    def _disabled_field(text):
        # mirrors the closed dropdown field so the Stories control keeps its
        # border/placeholder even before a plan is chosen.
        return ft.Container(
            ft.Row([ft.Text(text, size=13, color=T.INK_3, expand=True),
                    ft.Icon(ft.Icons.KEYBOARD_ARROW_DOWN, size=20, color=T.INK_3)],
                   vertical_alignment=ft.CrossAxisAlignment.CENTER),
            padding=ft.Padding.symmetric(vertical=12, horizontal=12),
            bgcolor=T.CARD, border=ft.Border.all(1, T.BORDER), border_radius=T.R)

    if app._reg_stories_loading:
        story_picker = _loading_field("Loading stories…")
    elif not _have_plans:
        story_picker = _disabled_field("Select a test plan first")
    else:
        story_picker = _checkbox_multiselect(
            [(str(s["id"]),
              (f"[{s['sprint']}] " if s.get("sprint") else "")
              + f"[{s['id']}] {(s['title'] or '')[:60]}")
             for s in app._reg_plan_stories],
            [str(s["id"]) for s in app._reg_selected],
            _toggle_story, _all_stories, is_open=app._reg_story_open, on_open=_open_stories,
            placeholder="Select stories", height=260,
            empty="No stories in the selected plan(s).",
            page=app.page, app=app, sync_key="reg_stories",
            invalid=getattr(app, "_reg_story_invalid", False))

    def _chip(label, on_close):
        return ft.Container(
            ft.Row([ft.Text(label, size=12, weight=ft.FontWeight.BOLD,
                            color=T.VIOLET_INK, font_family=T.F_MONO),
                    ft.GestureDetector(
                        content=ft.Icon(ft.Icons.CLOSE, size=12, color=T.VIOLET_INK),
                        on_tap=on_close, mouse_cursor=ft.MouseCursor.CLICK)],
                   spacing=5, tight=True),
            padding=ft.Padding.only(left=10, right=7, top=5, bottom=5),
            bgcolor=T.VIOLET_SOFT, border_radius=T.R_SM,
            border=ft.Border.all(1, "#D9D2FF"),
            on_hover=_chip_hover, animate_scale=120)

    def _plan_chip_label(p):
        # Show ONLY the sprint number; fall back to iteration tail or id so a
        # chip never renders blank.
        return (_sprint_num(p.get("sprint") or "") or _sprint_num(p.get("name") or "")
                or (p.get("sprint") or "").strip() or f"[{p['id']}]")

    _plan_chip_list = [_chip(_plan_chip_label(p), (lambda e, x=p["id"]: _remove_plan(x)))
                       for p in app._reg_plans_selected]
    if len(_plan_chip_list) > 1:
        _plan_chip_list.append(_clear_chip(lambda e: _all_plans(False)))
    plan_chips = ft.Row(_plan_chip_list, wrap=True, spacing=6, run_spacing=6)

    _STORY_CHIP_CAP = 40

    def _clear_stories(e=None):
        app._reg_selected = []
        app._reg_selected_rows = []
        app._reg_export_msg = app._reg_calc_msg = None
        app.render()

    def _story_chip_controls():
        sel = app._reg_selected
        ctrls = [_chip(str(s["id"]), (lambda e, x=s["id"]: _remove_story(x)))
                 for s in sel[:_STORY_CHIP_CAP]]
        if len(sel) > _STORY_CHIP_CAP:
            ctrls.append(ft.Container(
                ft.Text(f"+{len(sel) - _STORY_CHIP_CAP} more", size=12,
                        weight=ft.FontWeight.BOLD, color=T.INK_3),
                padding=ft.Padding.only(left=10, right=10, top=5, bottom=5),
                bgcolor=T.CARD_2, border_radius=T.R_SM))
        if len(sel) > 1:
            ctrls.append(_clear_chip(_clear_stories))
        return ctrls

    story_chips = ft.Row(_story_chip_controls(), wrap=True, spacing=6, run_spacing=6)
    story_chips_wrap = ft.Container(story_chips, padding=ft.Padding.only(top=10),
                                    visible=bool(app._reg_selected))
    story_count_text = ft.Text(f"{len(app._reg_selected)} stories selected", size=11,
                               color=T.INK_3, weight=ft.FontWeight.BOLD)

    def _refresh_story_externals():
        # the multiselect updates its own header in place but never re-renders, so
        # these external controls must be patched here or they go stale (the
        # "53 vs 433" mismatch). We also re-evaluate the Generate button's enabled
        # state here, otherwise it stays disabled after an in-place story pick.
        try:
            story_chips.controls = _story_chip_controls()
            story_chips_wrap.visible = bool(app._reg_selected)
            story_count_text.value = f"{len(app._reg_selected)} stories selected"
            story_chips.update(); story_chips_wrap.update(); story_count_text.update()
        except Exception:
            try:
                app.render()
            except Exception:
                pass
        cb = _calc_btn_cell[0]
        if cb is not None:
            should = bool(app._reg_selected) and not app._reg_busy \
                and not getattr(app, "readonly", False)
            try:
                cb.opacity = 1.0 if should else 0.45
                cb.on_click = _calculate if should else None
                cb.update()
            except Exception:
                pass
    app._reg_refresh_story_ext = _refresh_story_externals

    card1 = card(ft.Column([
        sec_head("1", "Source & stories"),
        ft.Container(height=10),
        ft.Column([field_label("Test plans", req=True), plan_picker], spacing=6),
        ft.Container(plan_chips, padding=ft.Padding.only(top=10),
                     visible=bool(app._reg_plans_selected)),
        ft.Text(f"{len(app._reg_plans_selected)} plan(s) selected", size=11,
                color=T.INK_3, weight=ft.FontWeight.BOLD,
                visible=bool(app._reg_plans_selected)),
        ft.Container(height=14),
        ft.Column([field_label("Stories", req=True), story_picker], spacing=6),
        story_chips_wrap,
        story_count_text,
    ], spacing=0))

    # ── Card 2: resources ──
    count_field = ft.TextField(
        value=("" if count is None else str(count)), hint_text="e.g. 3",
        keyboard_type=ft.KeyboardType.NUMBER, on_blur=_on_count, on_submit=_on_count,
        input_filter=_digits_only(),
        width=92, text_size=13,
        border_color=(T.RED if (mismatch or getattr(app, "_reg_count_invalid", False))
                      else T.BORDER), focused_border_color=T.VIOLET,
        border_radius=T.R,
        content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))
    name_field = ft.TextField(
        hint_text="Type a name, press Enter", on_submit=_add_name, on_blur=_add_name,
        expand=True, text_size=13,
        border_color=(T.RED if getattr(app, "_reg_res_invalid", False) else T.BORDER),
        focused_border_color=T.VIOLET,
        border_radius=T.R,
        content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))
    def _res_chip(nm, on_close):
        init, col = _av(nm)
        return ft.Container(
            ft.Row([
                ft.Container(ft.Text(init, size=10, weight=ft.FontWeight.BOLD,
                                     color="#FFFFFF"),
                             width=20, height=20, bgcolor=col, border_radius=20,
                             alignment=ft.Alignment.CENTER),
                ft.Text(nm, size=12.5, weight=ft.FontWeight.BOLD, color=T.INK),
                ft.GestureDetector(
                    content=ft.Icon(ft.Icons.CLOSE, size=12, color=T.INK_3),
                    on_tap=on_close, mouse_cursor=ft.MouseCursor.CLICK)],
               spacing=7, tight=True),
            padding=ft.Padding.only(left=5, right=9, top=4, bottom=4),
            bgcolor=T.CARD_2, border_radius=999,
            border=ft.Border.all(1, T.BORDER_2),
            on_hover=_chip_hover, animate_scale=120)

    name_chips = ft.Row(
        [_res_chip(n, (lambda e, x=n: _remove_name(x))) for n in app._reg_res_names],
        wrap=True, spacing=8, run_spacing=8)
    name_chips_wrap = ft.Container(name_chips, padding=ft.Padding.only(top=10),
                                   visible=bool(app._reg_res_names))

    warn = ft.Container()
    if mismatch:
        more = "more names than the number" if len(names) > count else \
               "fewer names than the number"
        warn = ft.Container(
            ft.Row([ft.Icon(ft.Icons.WARNING_AMBER, size=15, color=T.AMBER),
                    ft.Text(f"You set {count} resource(s) but added {len(names)} "
                            f"name(s) — {more}. Match them to export.",
                            size=12, color=T.AMBER, weight=ft.FontWeight.W_500,
                            expand=True)], spacing=8),
            padding=10, bgcolor=T.AMBER_SOFT, border_radius=T.R,
            border=ft.Border.all(1, "#EAD9A8"), margin=ft.Margin.only(top=10))

    card2 = card(ft.Column([
        sec_head("2", "Resources"),
        ft.Container(height=10),
        ft.Row([
            ft.Column([field_label("Count", req=True), count_field],
                      spacing=6, tight=True),
            ft.Column([field_label("Add a name", req=True),
                       ft.Row([name_field,
                               green_btn("Add", icon=ft.Icons.ADD,
                                         on_click=_add_name)], spacing=8)],
                      spacing=6, expand=True),
        ], spacing=14, vertical_alignment=ft.CrossAxisAlignment.START),
        ft.Container(name_chips_wrap, padding=ft.Padding.only(top=0),
                     visible=True),
        warn,
    ], spacing=0))

    # ── Card 3: effort model ──
    card3 = card(ft.Column([
        sec_head("3", "How effort is estimated"),
        ft.Container(height=10),
        ft.Container(
            ft.Row([
                _pill("test cases", T.INK_2, T.CARD_2),
                ft.Text("×", size=14, color=T.INK_3, weight=ft.FontWeight.BOLD),
                _pill(f"{AVG_MINUTES_PER_CASE} min", T.INK_2, T.CARD_2),
                ft.Text("×", size=14, color=T.INK_3, weight=ft.FontWeight.BOLD),
                _pill("priority weight", T.VIOLET_INK, T.VIOLET_SOFT),
                ft.Text("=", size=14, color=T.INK_3, weight=ft.FontWeight.BOLD),
                _pill("estimated hours", T.GREEN, T.GREEN_SOFT),
            ], spacing=8, wrap=True), padding=ft.Padding.only(bottom=12)),
        ft.Text("Priority weight (from each story's Azure DevOps priority):",
                size=12, color=T.INK_2, weight=ft.FontWeight.W_500),
        ft.Container(height=8),
        ft.Row([_pill("P1 ×1.30", T.RED, T.RED_SOFT),
                _pill("P2 ×1.15", T.AMBER, T.AMBER_SOFT),
                _pill("P3 ×1.00", T.INK_2, T.CARD_2),
                _pill("P4 ×0.90", T.GREEN, T.GREEN_SOFT)], spacing=8, wrap=True),
        ft.Container(height=10),
        ft.Text(f"Example: a P1 story with 33 cases  →  33 × {AVG_MINUTES_PER_CASE} "
                f"min × 1.30 ≈ 5.7 h", size=11.5, color=T.INK_3,
                weight=ft.FontWeight.W_500),
    ], spacing=0))

    # ── results ──
    results = None
    # While (re)generating, don't build the old table — the body shows the
    # progress spinner + skeleton instead, then the fresh table on completion.
    if app._reg_selected_rows and not app._reg_busy:
        d = plan_payload(app)

        def _cell(w, content, expand=False):
            return ft.Container(content, width=(None if expand else w), expand=expand,
                                padding=ft.Padding.symmetric(vertical=0, horizontal=6),
                                alignment=ft.Alignment.CENTER_LEFT)

        def _hd(s, w, expand=False):
            return _cell(w, _txt(s, size=10.5, weight=ft.FontWeight.BOLD,
                                 color=T.INK_3), expand=expand)

        header = ft.Container(
            ft.Row([ft.Container(width=34),
                    _hd("STORY", 64), _hd("TITLE", 0, expand=True), _hd("STATE", 84),
                    _hd("PRI", 44), _hd("CASES", 52), _hd("HOURS", 128),
                    _hd("ASSIGNEE", 140)], spacing=4),
            padding=ft.Padding.symmetric(vertical=9, horizontal=8), bgcolor=T.CARD_2,
            border=ft.Border.only(bottom=ft.BorderSide(1, T.BORDER)))

        maxh_story = max((x["hours"] for x in d["stories"]), default=0) or 1

        # Memoize rows by story id so collapse / page-flip reuse built controls.
        _reg_row_cache = {}

        def _story_row(s, i):
            bg = T.CARD if i % 2 == 0 else ft.Colors.with_opacity(0.5, T.BG)
            asg = s.get("assignee")
            asg_ctl = (ft.Row([_avatar(asg, 24),
                               _txt(asg, color=T.INK, weight=ft.FontWeight.W_500)],
                              spacing=8, vertical_alignment=ft.CrossAxisAlignment.CENTER)
                       if asg else _txt("—", color=T.INK_3))
            hours_ctl = ft.Row([
                ft.Container(_bar(s["hours"] / maxh_story), width=70),
                _txt(str(s["hours"]), color=T.INK, weight=ft.FontWeight.BOLD),
            ], spacing=8, vertical_alignment=ft.CrossAxisAlignment.CENTER)
            return ft.Container(
                ft.Row([
                    _cell(34, ft.IconButton(
                        icon=ft.Icons.DELETE_OUTLINE, icon_size=18, icon_color=T.RED,
                        tooltip="Remove story & recalculate",
                        on_click=_delete_row(s["id"]), width=34, height=34,
                        style=ft.ButtonStyle(padding=ft.Padding.all(0),
                                             shape=ft.RoundedRectangleBorder(radius=8)))),
                    _cell(64, _id_link(app, s["id"], font_family=T.F_MONO,
                                       color=T.VIOLET_INK, weight=ft.FontWeight.BOLD)),
                    _cell(0, _txt(s["title"] or "—", color=T.INK, no_wrap=False),
                          expand=True),
                    _cell(84, _state_pill(s["state"])),
                    _cell(44, _pri_pill(s["priority"])),
                    _cell(52, _txt(str(s["cases"]), color=T.INK_2)),
                    _cell(128, hours_ctl),
                    _cell(140, asg_ctl),
                ], spacing=4, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.Padding.symmetric(vertical=9, horizontal=8), bgcolor=bg)

        def _reg_row_for(s, i):
            c = _reg_row_cache.get(s["id"])
            if c is None:
                c = _story_row(s, i)
                _reg_row_cache[s["id"]] = c
            else:
                c.bgcolor = (T.CARD if i % 2 == 0
                             else ft.Colors.with_opacity(0.5, T.BG))   # keep zebra
            return c

        _PAGE = 25  # small page -> fast page.update() -> no event-loop freeze

        # In-place table update (collapse / page-flip):
        # Instead of full app.render() which rebuilds 800+ controls and blocks
        # Flet's event loop while Flutter re-renders, we rebuild only the table
        # body column and call table_body_col.update().  Keeps page.update() tiny.
        def _build_table_body():
            _all = d["stories"]
            _total = max(1, -(-len(_all) // _PAGE))
            _pg = max(0, min(getattr(app, "_reg_table_page", 0), _total - 1))
            _ps = _all[_pg * _PAGE: (_pg + 1) * _PAGE]

            from collections import OrderedDict as _OD
            _fg = _OD()
            for s in _ps:
                _fg.setdefault(s.get("feature_name") or "No Feature", []).append(s)

            rows = []
            _ri = _pg * _PAGE
            _coll = getattr(app, "_reg_collapsed_features", set())
            for fname, fstories in _fg.items():
                _is_c = fname in _coll
                _fid = fstories[0].get("feature_id") if fstories else None
                _fl = f"[{_fid}]  " if _fid else ""

                def _make_tog(fn=fname):
                    def _h(e):
                        cf = set(getattr(app, "_reg_collapsed_features", set()))
                        if fn in cf:
                            cf.discard(fn)
                        else:
                            cf.add(fn)
                        app._reg_collapsed_features = cf
                        _refresh_table()
                    return _h

                rows.append(_feature_header(app, f"{_fl}{fname}", len(fstories),
                                            _is_c, _make_tog()))
                if not _is_c:
                    for s in fstories:
                        rows.append(_reg_row_for(s, _ri))
                        _ri += 1

            _pager = ft.Container()
            if _total > 1:
                def _go(delta, pg=_pg, tot=_total):
                    def _h2(e):
                        app._reg_table_page = max(0, min(pg + delta, tot - 1))
                        _refresh_table()
                    return _h2
                _pager = ft.Row([
                    ghost_btn("← Prev",
                              on_click=(None if _pg == 0 else _go(-1))),
                    _txt(f"Page {_pg + 1} of {_total}  ·  "
                         f"rows {_pg*_PAGE+1}–{min((_pg+1)*_PAGE, len(_all))} "
                         f"of {len(_all)}",
                         size=12, color=T.INK_3),
                    ghost_btn("Next →",
                              on_click=(None if _pg >= _total - 1 else _go(1))),
                ], alignment=ft.MainAxisAlignment.CENTER, spacing=16)

            return rows + ([_pager] if _total > 1 else [])

        table_body_col = ft.Column(_build_table_body(), spacing=0)

        def _refresh_table():
            """Swap table rows/pager in-place without a full page rebuild.
            Preserves the scroll position so collapse / page-flip never jumps."""
            _off = getattr(app, "_scroll_offset", 0) or 0
            try:
                with _perf("table.rebuild_rows"):
                    table_body_col.controls = _build_table_body()
                with _perf("table.flet_update"):
                    table_body_col.update()
            except Exception:
                _perf_log("table.refresh FELL BACK to full render()")
                with _perf("table.full_render_fallback"):
                    app.render()
                return
            _keep_scroll(app, _off)

        table = ft.Container(
            ft.Column([header, table_body_col], spacing=0),
            border=ft.Border.all(1, T.BORDER), border_radius=T.R,
            clip_behavior=ft.ClipBehavior.HARD_EDGE)

        kpi_strip = ft.Row([
            _kpi_tile("STORIES", str(d["total_stories"])),
            _kpi_tile("TEST CASES", str(d["total_cases"])),
            _kpi_tile("TOTAL EFFORT", f"{d['total_hours']} h", T.VIOLET),
            _kpi_tile("PER PERSON", f"{d['hours_per_person']} h", T.GREEN),
        ], spacing=10)

        workload_ui = ft.Container()
        if d["workload"]:
            maxw = max((w["hours"] for w in d["workload"]), default=0) or 1
            cards_wl = [ft.Container(ft.Column([
                ft.Row([_avatar(w["name"], 32),
                        ft.Column([_txt(w["name"], color=T.INK, weight=ft.FontWeight.BOLD,
                                        size=14),
                                   _txt(f"{w['stories']} stories · {w.get('cases', 0)} cases",
                                        color=T.INK_3, size=11)],
                                  spacing=1, tight=True, expand=True),
                        _txt(f"{w['hours']} h", color=T.INK, weight=ft.FontWeight.BOLD,
                             size=16, font_family=T.F_MONO, no_wrap=True)],
                       spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                ft.Container(height=12),
                _bar(w["hours"] / maxw, T.VIOLET, 8),
            ], spacing=0), width=300, padding=14, bgcolor=T.CARD,
                border=ft.Border.all(1, T.BORDER_2), border_radius=T.R)
                for w in d["workload"]]
            workload_ui = ft.Column([
                ft.Container(height=16),
                ft.Text("RESOURCE WORKLOAD", size=10.5, weight=ft.FontWeight.BOLD,
                        color=T.INK_3),
                ft.Container(height=10),
                ft.Row(cards_wl, spacing=12, wrap=True, run_spacing=12,
                       vertical_alignment=ft.CrossAxisAlignment.START)], spacing=0)

        if mismatch:
            exports = ft.Container(
                ft.Row([ft.Icon(ft.Icons.LOCK_OUTLINE, size=15, color=T.INK_3),
                        ft.Text("Resolve the resource mismatch above to export.",
                                size=12, color=T.INK_3, weight=ft.FontWeight.W_500)],
                       spacing=8),
                padding=10, bgcolor=T.CARD_2, border_radius=T.R)
        else:
            def _exp_btn(label, icon, color, fmt):
                return ft.OutlinedButton(
                    content=ft.Row([ft.Icon(icon, size=17, color=color),
                                    ft.Text(label, size=13.5, weight=ft.FontWeight.W_600,
                                            color=T.INK)], spacing=8, tight=True),
                    on_click=_export(fmt), height=44,
                    style=ft.ButtonStyle(
                        bgcolor={"": T.CARD},
                        side=ft.BorderSide(1, T.BORDER),
                        shape=ft.RoundedRectangleBorder(radius=T.R),
                        padding=ft.Padding.symmetric(horizontal=15, vertical=0)))
            exports = ft.Row([
                _exp_btn("Word", ft.Icons.DESCRIPTION, T.BRAND_GRAD_1, "docx"),
                _exp_btn("Excel", ft.Icons.TABLE_CHART, T.GREEN, "xlsx"),
                _exp_btn("PDF", ft.Icons.PICTURE_AS_PDF, T.RED, "pdf"),
                _exp_btn("JSON", ft.Icons.DATA_OBJECT, T.STORY, "json"),
            ], spacing=8, wrap=True)

        _s_kind = app._reg_export_msg[0] if app._reg_export_msg else "ok"
        _s_text = app._reg_export_msg[1] if app._reg_export_msg else ""
        _s_ok = (_s_kind == "ok")
        _status_icon = ft.Icon(
            ft.Icons.CHECK_CIRCLE if _s_ok else ft.Icons.ERROR_OUTLINE,
            size=16, color=(T.GREEN if _s_ok else T.RED))
        _status_txt = ft.Text(_s_text, size=12, color=(T.GREEN if _s_ok else T.RED),
                              weight=ft.FontWeight.W_500, selectable=True, expand=True)
        status = ft.Container(
            ft.Row([_status_icon, _status_txt], spacing=8),
            padding=10, bgcolor=(T.GREEN_SOFT if _s_ok else T.RED_SOFT),
            border_radius=T.R, margin=ft.Margin.only(top=10),
            visible=bool(app._reg_export_msg))
        _status_cell[0] = status   # wire mutable ref for in-place updates

        email_field = ft.TextField(
            value=app._reg_email_to or "", on_change=_on_email_to,
            hint_text="recipient@company.com (comma-separate for several)",
            expand=True, text_size=13, border_color=T.BORDER,
            focused_border_color=T.VIOLET, border_radius=T.R,
            content_padding=ft.Padding.symmetric(vertical=12, horizontal=10))
        email_row = ft.Column([
            ft.Divider(height=20, color=T.BORDER),
            ft.Text("EMAIL THE PLAN", size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
            ft.Container(height=8),
            ft.Row([email_field,
                    green_btn("Sending…" if app._reg_emailing else "Send",
                              icon=ft.Icons.SEND, on_click=_email)],
                   spacing=10),
            ft.Text("Attaches the Word plan and an inline summary. Uses the Gmail "
                    "sender configured on Setup.", size=11, color=T.INK_3,
                    weight=ft.FontWeight.W_500),
        ], spacing=6)

        # Regenerate: re-run the plan on the current selection without scrolling
        # back up to the Generate button. With the per-suite caches this is
        # near-instant on an unchanged selection.
        regen_btn = ghost_btn(
            "Regenerating…" if app._reg_busy else "Regenerate",
            icon=ft.Icons.REFRESH,
            on_click=(None if app._reg_busy else _regenerate))

        results = card(ft.Column([
            ft.Row([sec_head("4", "Plan"), ft.Container(expand=True), regen_btn],
                   vertical_alignment=ft.CrossAxisAlignment.CENTER),
            ft.Container(height=12), kpi_strip,
            ft.Container(height=14), table,
            workload_ui, ft.Divider(height=22, color=T.BORDER),
            ft.Text("EXPORT", size=10.5, weight=ft.FontWeight.BOLD, color=T.INK_3),
            ft.Container(height=8), exports, email_row, status,
        ], spacing=0))

    calc_btn = primary_btn("Generating…" if app._reg_busy else "Generate Regression Plan",
                           icon=ft.Icons.CALCULATE, on_click=_calculate,
                           disabled=app._reg_busy or not app._reg_selected)
    _calc_btn_cell[0] = calc_btn   # store ref so _calculate can mutate it

    gen_spinner = ft.Container(
        ft.Row([ft.ProgressRing(width=18, height=18, stroke_width=2.5, color=T.VIOLET),
                ft.Text("Generating plan — reading test cases from Azure DevOps…",
                        size=12.5, color=T.INK_3, weight=ft.FontWeight.W_500)],
               spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
        padding=ft.Padding.symmetric(vertical=10, horizontal=12),
        margin=ft.Margin.only(top=10),
        bgcolor=getattr(T, "VIOLET_SOFT", T.CARD_2), border_radius=T.R,
        visible=app._reg_busy)
    _gen_spinner_cell[0] = gen_spinner   # so _calculate can show it in place
    app._reg_gen_spinner = gen_spinner   # current on-screen spinner, for live count
                                         # progress (the _work closure may predate the
                                         # busy re-render, so target the latest one)

    body_children = [card1, ft.Container(height=14),
                     ft.Row([ft.Container(card2, expand=1),
                             ft.Container(card3, expand=1)],
                            spacing=14,
                            vertical_alignment=ft.CrossAxisAlignment.START),
                     ft.Container(height=16), calc_btn, gen_spinner, calc_note_wrap]
    if results is not None:
        body_children += [ft.Container(height=16), results]
    elif app._reg_busy:
        from main import card as _card, skeleton_rows as _skel
        body_children += [ft.Container(height=16), _card(_skel(6))]

    body = ft.Column(body_children, spacing=0, scroll=ft.ScrollMode.AUTO, expand=True)
    return app.shell("Regression Plan",
                     "Build a regression plan from your test plans & their stories", body,
                     right=ghost_btn("Use Setup selection", icon=ft.Icons.DOWNLOAD,
                                     on_click=_use_setup_selection),
                     badge="STEP R")
# perf: lazy-build dropdown rows to keep full renders cheap
