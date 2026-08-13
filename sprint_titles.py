"""sprint_titles.py — the "Sprint Report" screen.

Pick a sprint and build a Sprint (closure) Report: its User Stories grouped by
status (Completed vs In-progress / carried over), with the titles AI-translated to
the chosen language, plus a Bugs summary (total, regression vs sprint, and a
by-status breakdown). Shown in-app and downloadable as a Word document — RTL for
Arabic, LTR for English.

Self-contained: reuses regression.py's multiselect/helpers and engine.py's Azure
+ AI calls, but keeps its own `_st_*` state so it never collides with Sprint Plan.
main.py imports this and dispatches `screen(app)` for the "titles" nav tab.
"""
import os
import re
import json
import threading
from datetime import datetime

import flet as ft
import theme as T
import engine as E
import backend_setup

# Story states that count as "done" for the report.
_DONE = {"done", "closed", "completed", "resolved", "accepted"}

# States that count as "completed" in the SPRINT CLOSURE REPORT specifically.
# Per the team's Azure workflow a story is done for the report once it reaches
# "Product Owner Review"; every OTHER state is in-progress / carried. Kept
# separate from _DONE (which still colours the state-distribution chips).
_REPORT_DONE = {"product owner review"}

# Localized labels (headings are fixed strings; story titles are AI-translated).
_L = {
    "ar": {
        "title": "تقرير إغلاق السبرنت", "sprint": "السبرنت", "date": "التاريخ",
        "completed": "النقاط المكتملة", "carried": "نقاط قيد التنفيذ / مرحّلة",
        "bugs": "الأخطاء (Bugs)", "total_bugs": "إجمالي الأخطاء",
        "regression_bugs": "أخطاء الـ Regression", "sprint_bugs": "أخطاء السبرنت",
        "by_status": "حسب الحالة", "stories": "قصة", "none": "لا يوجد.",
        "other": "أهداف أخرى", "objectives": "أهداف السبرنت",
    },
    "en": {
        "title": "Sprint Closure Report", "sprint": "Sprint", "date": "Date",
        "completed": "Completed", "carried": "In progress / carried over",
        "bugs": "Bugs", "total_bugs": "Total bugs",
        "regression_bugs": "Regression bugs", "sprint_bugs": "Sprint bugs",
        "by_status": "By status", "stories": "stories", "none": "None.",
        "other": "Other objectives", "objectives": "Sprint objectives",
    },
    "fr": {
        "title": "Rapport de clôture du sprint", "sprint": "Sprint", "date": "Date",
        "completed": "Terminés", "carried": "En cours / reportés",
        "bugs": "Bugs", "total_bugs": "Total des bugs",
        "regression_bugs": "Bugs de régression", "sprint_bugs": "Bugs du sprint",
        "by_status": "Par statut", "stories": "récits", "none": "Aucun.",
        "other": "Autres objectifs", "objectives": "Objectifs du sprint",
    },
    "tr": {
        "title": "Sprint Kapanış Raporu", "sprint": "Sprint", "date": "Tarih",
        "completed": "Tamamlanan", "carried": "Devam eden / aktarılan",
        "bugs": "Hatalar", "total_bugs": "Toplam hata",
        "regression_bugs": "Regresyon hataları", "sprint_bugs": "Sprint hataları",
        "by_status": "Duruma göre", "stories": "hikaye", "none": "Yok.",
        "other": "Diğer hedefler", "objectives": "Sprint hedefleri",
    },
    "es": {
        "title": "Informe de cierre del sprint", "sprint": "Sprint", "date": "Fecha",
        "completed": "Completadas", "carried": "En progreso / trasladadas",
        "bugs": "Errores", "total_bugs": "Total de errores",
        "regression_bugs": "Errores de regresión", "sprint_bugs": "Errores del sprint",
        "by_status": "Por estado", "stories": "historias", "none": "Ninguna.",
        "other": "Otros objetivos", "objectives": "Objetivos del sprint",
    },
    "de": {
        "title": "Sprint-Abschlussbericht", "sprint": "Sprint", "date": "Datum",
        "completed": "Abgeschlossen", "carried": "In Bearbeitung / übertragen",
        "bugs": "Fehler", "total_bugs": "Fehler gesamt",
        "regression_bugs": "Regressionsfehler", "sprint_bugs": "Sprint-Fehler",
        "by_status": "Nach Status", "stories": "Storys", "none": "Keine.",
        "other": "Weitere Ziele", "objectives": "Sprint-Ziele",
    },
    "nl": {
        "title": "Sprint-afsluitrapport", "sprint": "Sprint", "date": "Datum",
        "completed": "Voltooid", "carried": "In uitvoering / doorgeschoven",
        "bugs": "Bugs", "total_bugs": "Totaal aantal bugs",
        "regression_bugs": "Regressiebugs", "sprint_bugs": "Sprintbugs",
        "by_status": "Op status", "stories": "stories", "none": "Geen.",
        "other": "Overige doelen", "objectives": "Sprintdoelen",
    },
}

# Ordinal prefixes for epic group headings (mirrors the reference report's
# أولاً / ثانياً / ثالثاً … grouping).
_ORD = {
    "ar": ["أولاً", "ثانياً", "ثالثاً", "رابعاً", "خامساً", "سادساً",
           "سابعاً", "ثامناً", "تاسعاً", "عاشراً"],
    "en": ["First", "Second", "Third", "Fourth", "Fifth", "Sixth",
           "Seventh", "Eighth", "Ninth", "Tenth"],
    "fr": ["Premièrement", "Deuxièmement", "Troisièmement", "Quatrièmement",
           "Cinquièmement", "Sixièmement", "Septièmement", "Huitièmement",
           "Neuvièmement", "Dixièmement"],
    "tr": ["Birincisi", "İkincisi", "Üçüncüsü", "Dördüncüsü", "Beşincisi",
           "Altıncısı", "Yedincisi", "Sekizincisi", "Dokuzuncusu", "Onuncusu"],
    "es": ["Primero", "Segundo", "Tercero", "Cuarto", "Quinto", "Sexto",
           "Séptimo", "Octavo", "Noveno", "Décimo"],
    "de": ["Erstens", "Zweitens", "Drittens", "Viertens", "Fünftens",
           "Sechstens", "Siebtens", "Achtens", "Neuntens", "Zehntens"],
    "nl": ["Ten eerste", "Ten tweede", "Ten derde", "Ten vierde", "Ten vijfde",
           "Ten zesde", "Ten zevende", "Ten achtste", "Ten negende", "Ten tiende"],
}


def _group_by_epic(rows):
    """Group stories by their epic, preserving first-seen order; stories with no
    epic ("") are collected into a final group. Returns [(epic_name, [stories])]."""
    groups, order = {}, []
    for s in rows:
        e = (s.get("epic") or "").strip()
        if e not in groups:
            groups[e] = []
            order.append(e)
        groups[e].append(s)
    named = [e for e in order if e]
    rest = [e for e in order if not e]
    return [(e, groups[e]) for e in named + rest]


def _init(app):
    for k, v in (("_st_iterations", []), ("_st_iter_loading", False),
                 ("_st_sprint_paths", []), ("_st_open", False),
                 ("_st_lang", getattr(app, "lang", "ar")),
                 ("_st_busy", False), ("_st_report", None),
                 ("_st_done", False), ("_st_msg", None)):
        if not hasattr(app, k):
            setattr(app, k, v)


def _sprint_num(text):
    m = re.search(r"[Ss]print\s*\d+", text or "")
    return re.sub(r"\s+", " ", m.group(0)).strip() if m else ""


def _sprint_range_label(labels):
    """Compact multi-sprint header, mirroring regression.py's _sprint_range_label:
    "Sprint 1 to Sprint 4" instead of listing every one out. Only labels that
    parse as "Sprint N" get compressed; anything else is left alone and appended
    comma-separated."""
    uniq = list(dict.fromkeys(l for l in labels if l))
    if not uniq:
        return ""
    if len(uniq) == 1:
        return uniq[0]
    numbered, other = [], []
    for lbl in uniq:
        m = re.search(r"\d+", lbl)
        if m and re.search(r"[Ss]print", lbl):
            numbered.append((int(m.group(0)), lbl))
        else:
            other.append(lbl)
    parts = []
    if numbered:
        numbered.sort(key=lambda t: t[0])
        parts.append(numbered[0][1] if len(numbered) == 1
                     else f"{numbered[0][1]} to {numbered[-1][1]}")
    parts.extend(other)
    return ", ".join(parts)


def _sort_key(it):
    m = re.search(r"\d+", _sprint_num(it.get("name", "")) or _sprint_num(it.get("path", "")))
    return int(m.group(0)) if m else -1


def _load_iterations(app):
    if not (app.connected and app.project):
        return                      # no connection → don't fetch (read-only viewers)
    # Key on project so an empty result counts as "loaded" (falsy `if app._st_iterations`
    # otherwise reloads + full-renders every pass → the flashing loop).
    if app._st_iter_loading or getattr(app, "_st_iter_for", None) == app.project:
        return
    app._st_iter_loading = True
    _proj = app.project
    app._st_iter_for = _proj

    def _work():
        try:
            # Backend-aware: Azure returns its iterations; Jira/Xray return their
            # board sprints. Was E.fetch_iterations (Azure-only), so the sprint
            # dropdown was empty on every non-Azure backend.
            its = backend_setup.fetch_sprints(app, _proj) or []
        except Exception:
            its = []
        sprints = [it for it in its
                   if (_sprint_num(it.get("name", "")) or _sprint_num(it.get("path", "")))] or its
        sprints.sort(key=lambda it: (_sort_key(it) < 0, _sort_key(it)))
        app._st_iter_loading = False
        if app.project != _proj:      # project changed mid-load → drop stale result
            app._st_iter_for = None
            return
        app._st_iterations = sprints
        app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


# Arabic-Indic (٠-٩) and Persian (۰-۹) digits → ASCII, so a model that numbers
# its Arabic output with native digits still parses.
_AR_DIGITS = {ord(c): str(i) for i, c in enumerate("٠١٢٣٤٥٦٧٨٩")}
_AR_DIGITS.update({ord(c): str(i) for i, c in enumerate("۰۱۲۳۴۵۶۷۸۹")})
# RTL/LTR marks and the bidi isolate chars a model may prepend to each line.
_BIDI_MARKS = "".join(("‎", "‏", "‪", "‫", "‬",
                       "⁦", "⁧", "⁨", "⁩", "﻿"))


def _clean_line(s):
    return s.translate(_AR_DIGITS).strip().strip(_BIDI_MARKS).strip()


# Report text is PLAIN TEXT: no double quotes, no colons. Azure story titles
# routinely carry both ("Login: verify the \"Remember me\" box"), and the AI
# translation can reintroduce them, so this is applied once at the point the
# report's title is set — which covers the on-screen report, the Word export
# and the emailed HTML in one place, so they can never drift apart.
# Colons become a SPACE rather than being deleted, so "Login:verify" reads as
# "Login verify" instead of running the words together; quotes are dropped
# outright. Curly/typographic quotes and the full-width colon are included
# because both the AI and Azure produce them. Whitespace is collapsed
# afterwards so nothing is left with double spaces.
_PLAIN_DROP = '"“”„«»'      # " “ ” „ « »
_PLAIN_SPACE = ':：ː'                        # : ： ː


def _plain_text(s):
    """Strip double quotes and colons from report text (see above)."""
    if not s:
        return s
    out = str(s)
    for ch in _PLAIN_DROP:
        out = out.replace(ch, "")
    for ch in _PLAIN_SPACE:
        out = out.replace(ch, " ")
    return re.sub(r"\s{2,}", " ", out).strip()


def _parse_json_array(out, n):
    """Pull a JSON array of `n` strings out of a model response; None if it can't."""
    try:
        s = out[out.index("["):out.rindex("]") + 1]
        arr = json.loads(s)
    except Exception:
        return None
    if isinstance(arr, list) and len(arr) == n:
        return [(_clean_line(str(a)) or None) for a in arr]
    return None


def _parse_numbered(out, n):
    """Fallback line parser: numbered (digit-normalized) or positional."""
    by_num, ordered = {}, []
    for raw in out.splitlines():
        line = _clean_line(raw)
        if not line:
            continue
        m = re.match(r"^(\d+)[.)\-:]\s*(.+)$", line)
        if m:
            val = m.group(2).strip()
            by_num[int(m.group(1))] = val
            ordered.append(val)
        else:
            ordered.append(line)
    mapped = [by_num.get(i + 1) for i in range(n)]
    if sum(1 for v in mapped if v) < n and len(ordered) == n:
        mapped = ordered
    return mapped


def _parse_json_items(out, n):
    """Pull a JSON OBJECT wrapping the array of `n` strings — the shape strict
    json_object providers (Groq) require. Reads {"items":[...]} or any single
    list value / invented key. None if it cannot."""
    try:
        s = out[out.index("{"):out.rindex("}") + 1]
        obj = json.loads(s)
    except Exception:
        return None
    if not isinstance(obj, dict):
        return None
    arr = None
    for k in ("items", "translations", "result", "results", "strings", "values"):
        v = obj.get(k)
        if isinstance(v, list):
            arr = v
            break
    if arr is None:
        lists = [v for v in obj.values() if isinstance(v, list)]
        arr = lists[0] if lists else None
    if isinstance(arr, list) and len(arr) == n:
        return [(_clean_line(str(a)) or None) for a in arr]
    return None


_ARABIC_RE = re.compile(r"[\u0600-\u06FF]")


def _has_arabic(s):
    return bool(_ARABIC_RE.search(s or ""))


def _strip_bilingual(s, lang):
    """Clean one translated title. Groq's model sometimes (a) echoes
    'original -> translation', (b) self-corrects a word as 'typo -> fix', and
    (c) appends an editorial note like '(typo in thought, fixed)'. Strip all of
    that: drop model-commentary parentheticals, then collapse any arrow artifact
    to the fullest segment in the target language. Legit titles (no arrow, no
    commentary) pass through unchanged; real product parentheticals survive."""
    if not s:
        return s
    s = s.strip().strip("`").strip()
    # Cut model "thinking out loud" — deliberation that leaks into the string,
    # e.g. a title followed by "or I'll stick with ...?". Keep the text BEFORE the
    # first English deliberation marker (the first, usually-correct rendering); the
    # markers are specific enough not to hit real titles ("Add or Edit" is safe).
    _cut = re.sub(r"\s*(?:\bor\s+)?(?:i['’]?ll\b|let me\b|on second thought\b|"
                  r"i think\b|alternatively\b|\bhmm\b|wait,).*$",
                  "", s, flags=re.I | re.S).strip()
    if _cut:
        s = _cut
    s = re.sub(r"\s*[\(\[][^)\]]*(?:typo|corrected|correction|mistranslat|"
               r"in thought|thinking)[^)\]]*[\)\]]", "", s, flags=re.I)
    for sep in (" -> ", " → "):
        if sep in s:
            parts = [x.strip() for x in s.split(sep) if x.strip()]
            if len(parts) >= 2:
                if lang == "ar":
                    cand = [x for x in parts if _has_arabic(x)] or parts
                else:
                    cand = [x for x in parts if not _has_arabic(x)] or parts
                s = max(cand, key=len)
    # Collapse an immediately repeated word ("VALID VALID" -> "VALID"), a
    # common small-model garble. Safe: real titles rarely repeat a word
    # back-to-back.
    s = re.sub(r"(\b\w+\b)(?:\s+\1\b)+", r"\1", s, flags=re.I)
    return re.sub(r"\s{2,}", " ", s).strip()


def _translate_chunk(texts, target):
    """Translate one batch; returns a list aligned 1:1 (per-item fallback to the
    original on any miss). Asks for a JSON object wrapping the array (safe for
    strict json_object providers like Groq); falls back to bare-array then a
    numbered/positional line parse if the model ignores the JSON instruction."""
    payload = json.dumps(texts, ensure_ascii=False)
    prompt = (
        f"Translate each string in this JSON array into {target}. "
        f"Output ONLY the {target} translation of each item. Do NOT include the "
        "original text, do NOT return both languages, and do NOT join the original "
        "to the translation with arrows (->), slashes or dashes. Do NOT correct, "
        "comment on, or annotate typos, and add NO notes, parentheses, brackets, "
        "or explanations about your translation — return the final translated text "
        "only. Give exactly ONE final translation per item; never deliberate or "
        "offer alternatives, and never add a question mark. Translate the meaning "
        "naturally and concisely; keep IDs, version "
        "numbers and well-known product names intelligible. Return ONLY a JSON "
        "array of the translated strings, same length and same order as the input, "
        f"no keys, no commentary, no code fences.\n\n{payload}")
    # Let credit/error bubble up so the caller can tell the user (don't silently
    # fall back to English on an out-of-credit / failed provider).
    # want_json is OFF on purpose. Groq's json_object mode 400s with
    # "Failed to validate JSON" when the model returns a top-level array
    # (or anything not a bare object), which is exactly what a translation
    # list is. We ask for JSON in the prompt and parse it defensively below,
    # so we don't need — and don't want — the provider's strict JSON grammar.
    out = E.ai_complete(prompt, max_tokens=4096) or ""
    mapped = (_parse_json_items(out, len(texts))
              or _parse_json_array(out, len(texts))
              or _parse_numbered(out, len(texts)))
    tlang = "ar" if target == "Arabic" else "en"
    return [_strip_bilingual(mapped[i] or texts[i], tlang)
            for i in range(len(texts))]


def _translate(texts, lang):
    """AI-translate `texts` into `lang` ('ar'/'en'). Returns (results, err) where
    `results` is aligned 1:1 with the input (originals kept on any miss) and `err`
    is None on success, "credit" when the AI account is out of credit/quota, or
    "error:<msg>" for any other failure. Normalizes native-digit numbering, falls
    back to positional line alignment, and chunks long lists so nothing truncates."""
    texts = [t or "" for t in texts]
    if not any(t.strip() for t in texts):
        return list(texts), None
    target = E.LANGUAGES.get(lang, {}).get("name") or "English"
    out, CHUNK, err = [], 20, None
    for i in range(0, len(texts), CHUNK):
        chunk = texts[i:i + CHUNK]
        try:
            out.extend(_translate_chunk(chunk, target))
        except E.CreditBalanceError:
            err = "credit"
            out.extend(chunk)            # keep originals for this + remaining
            out.extend(texts[i + CHUNK:])
            break
        except Exception as ex:
            err = err or ("error:" + str(ex)[:160])
            out.extend(chunk)            # keep originals for this chunk, keep going
    # pad just in case a break left it short
    while len(out) < len(texts):
        out.append(texts[len(out)])
    return out[:len(texts)], err


def _generate(app):
    # Permission gate: read-only users can't generate.
    if getattr(app, "readonly", False):
        app._st_msg = ("err", "Your role doesn’t allow generating sprint reports.")
        app.ui_safe(app.render)
        return
    paths = list(app._st_sprint_paths or [])
    if not paths:
        app._st_msg = ("err", "Pick at least one sprint first.")
        app.ui_safe(app.render)
        return
    # The report body now routes through backend_setup.sprint_report_data, which
    # works on EVERY backend (Azure via the engine; others pull the sprint's
    # stories from the Jira/read source) — so the earlier Azure-only gate is
    # removed. Only the execution-rollup Sprint SUMMARY stays Azure-only.
    lang = app._st_lang
    app._st_busy = True
    app._st_done = False
    app._st_report = None
    app._st_msg = None
    app.ui_safe(app.render)

    def _work():
        try:
            names, seen_s, seen_b = [], set(), set()
            stories, bugs = [], []
            for p in paths:
                it = next((x for x in app._st_iterations if x["path"] == p), None)
                if it:
                    names.append(_sprint_num(it["name"]) or it["name"])
                try:
                    d = backend_setup.sprint_report_data(app, p)
                except Exception:
                    d = {"stories": [], "bugs": []}
                for s in d.get("stories", []):
                    if s["id"] in seen_s:
                        continue
                    seen_s.add(s["id"])
                    stories.append(s)
                for b in d.get("bugs", []):
                    if b["id"] in seen_b:
                        continue
                    seen_b.add(b["id"])
                    bugs.append(b)

            # Sort into Azure sprint-board order (StackRank/BacklogPriority/Id)
            # so the report lists stories like the other screens; _group_by_epic
            # below preserves this order within each epic group. Non-Azure /
            # unrankable backends fall back to a stable (sprint, id) order.
            stories = E.sort_stories_by_board(getattr(app, "project", ""), stories)
            # translate every title once, then split into sections by state
            originals = [s["title"] for s in stories]
            titles, terr = _translate(originals, lang)
            for s, tr in zip(stories, titles):
                # Plain text only — no quotes/colons (see _plain_text). Set here,
                # at the single point the report title is assigned, so the
                # on-screen report, the .docx export and the email all agree.
                s["t"] = _plain_text(tr)
            n_changed = sum(1 for o, t in zip(originals, titles)
                            if (t or "").strip() != (o or "").strip())
            completed = [s for s in stories if (s.get("state", "").lower() in _REPORT_DONE)]
            carried = [s for s in stories if (s.get("state", "").lower() not in _REPORT_DONE)]

            from collections import Counter
            reg = sum(1 for b in bugs if "regression" in (b.get("tags", "") or "").lower())
            app._st_report = {
                "sprint_name": _sprint_range_label(names),
                "date": datetime.now().strftime("%d-%m-%Y"),
                "lang": lang,
                "completed": completed, "carried": carried,
                "total_stories": len(stories),
                "bug_by_state": dict(Counter(b.get("state", "Unknown") for b in bugs)),
                "total_bugs": len(bugs), "regression_bugs": reg,
                "sprint_bugs": len(bugs) - reg,
            }
            app._st_done = True
            if terr == "credit":
                app._st_msg = ("err", "Report built, but the AI account is OUT OF CREDIT "
                               "— titles kept in their original language. Top up or switch "
                               "provider in Setup, then Generate again.")
            elif terr and terr.startswith("error:"):
                app._st_msg = ("err", f"Report built, but translation failed: "
                               f"{terr.split(':', 1)[1].strip()} — titles kept as-is.")
            elif stories and n_changed == 0:
                app._st_msg = ("err", "Stories & bugs loaded, but the AI returned no "
                               "translation (0 titles changed) — check the AI provider "
                               "in Setup. Showing the original titles.")
            else:
                app._st_msg = ("ok", f"Report ready — {len(stories)} stories "
                               f"({n_changed} translated), {len(bugs)} bugs.")
        except Exception as ex:
            app._st_msg = ("err", f"Couldn't build the report: {ex}")
        app._st_busy = False
        if getattr(app, "active", None) == "titles":
            app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


def _retranslate(app):
    """Re-translate an existing report's titles into the current language (used when
    the language toggle is flipped after a report was generated, so titles update
    live instead of only the static labels)."""
    r = app._st_report
    if not r:
        app.ui_safe(app.render)
        return
    lang = app._st_lang
    if r.get("lang") == lang:
        app.ui_safe(app.render)
        return
    app._st_busy = True
    app._st_msg = None
    app.ui_safe(app.render)

    def _work():
        try:
            rows = list(r.get("completed", [])) + list(r.get("carried", []))
            originals = [s["title"] for s in rows]
            titles, terr = _translate(originals, lang)
            for s, t in zip(rows, titles):
                # Same plain-text rule as _generate — otherwise flipping the
                # language toggle would reintroduce quotes/colons the report
                # was generated without.
                s["t"] = _plain_text(t)
            r["lang"] = lang
            n_changed = sum(1 for o, t in zip(originals, titles)
                            if (t or "").strip() != (o or "").strip())
            if terr == "credit":
                app._st_msg = ("err", "AI account is OUT OF CREDIT — titles kept as-is. "
                               "Top up or switch provider in Setup.")
            elif terr and terr.startswith("error:"):
                app._st_msg = ("err", f"Translation failed: "
                               f"{terr.split(':', 1)[1].strip()} — titles kept as-is.")
            elif rows and n_changed == 0:
                app._st_msg = ("err", "The AI returned no translation (0 titles "
                               "changed) — check the AI provider in Setup.")
            else:
                app._st_msg = ("ok", f"Translated to "
                               f"{'Arabic' if lang == 'ar' else 'English'} "
                               f"({n_changed}/{len(rows)} titles).")
        except Exception as ex:
            app._st_msg = ("err", f"Re-translate failed: {ex}")
        app._st_busy = False
        if getattr(app, "active", None) == "titles":
            app.ui_safe(app.render)
    threading.Thread(target=_work, daemon=True).start()


def _decode_assets():
    """Write the embedded brand images (logo + colour band) to temp PNGs and return
    (logo_path, band_path); ('' , '') if unavailable. Caller cleans up."""
    import base64, tempfile
    out = {}
    try:
        import report_assets as RA
        for key, b64 in (("logo", RA.LOGO_PNG_B64), ("band", RA.BAND_PNG_B64)):
            tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tf.write(base64.b64decode(b64))
            tf.close()
            out[key] = tf.name
    except Exception:
        return "", ""
    return out.get("logo", ""), out.get("band", "")


def _export_docx(app):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml

    r = app._st_report or {}
    lang = r.get("lang", "ar")
    # Direction follows the report LANGUAGE (requested): Arabic → RTL, English →
    # LTR. `rtl` drives every direction-dependent decision below — paragraph
    # <w:bidi> + RIGHT alignment (_set_bidi), run-level <w:rtl/> (_rtl_run), the
    # bullet side in _bullet, and the logo alignment. Content is unchanged either
    # way — Unicode's bidi algorithm shapes the Arabic glyphs correctly inside
    # each run regardless of paragraph direction.
    rtl = (lang == "ar")
    L = _L.get(lang, _L["en"])
    BLUE = RGBColor(0x4C, 0x94, 0xD8)         # brand blue from the reference report
    INK = RGBColor(0x1F, 0x1F, 0x1F)
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    logo_path, band_path = _decode_assets()

    doc = Document()
    try:
        ns = doc.styles["Normal"]
        ns.font.name = "Segoe UI"
        ns.font.size = Pt(11)
    except Exception:
        pass
    for sec in doc.sections:
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    def _set_bidi(p):
        # RTL paragraph. Two things matter:
        #  1) <w:bidi> sets the base direction (must precede <w:spacing>/<w:ind>/
        #     <w:jc>/<w:rPr> in <w:pPr> per the schema, or Word ignores it —
        #     LibreOffice is lenient).
        #  2) Alignment uses the LOGICAL value w:jc="start" (start-of-line = the
        #     RIGHT in RTL), NOT physical "right". Word SWAPS the meaning of
        #     jc=left/right inside a bidi paragraph, so "right" visually LEFT-aligns
        #     RTL text — the reported "Arabic still left aligned" bug — and Word
        #     versions disagree on that swap. "start"/"end" are unambiguous across
        #     every Word 2010+, so the export is right-aligned on all of them.
        try:
            p.alignment = RIGHT               # python-docx places <w:jc> correctly…
            pPr = p._p.get_or_add_pPr()
            _jc = pPr.find(qn("w:jc"))
            if _jc is not None:
                _jc.set(qn("w:val"), "start")  # …then make it logical-start = RIGHT in RTL
            if pPr.find(qn("w:bidi")) is None:
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                anchor = None
                for tag in ("w:spacing", "w:ind", "w:jc", "w:rPr"):
                    found = pPr.find(qn(tag))
                    if found is not None:
                        anchor = found
                        break
                if anchor is not None:
                    anchor.addprevious(bidi)
                else:
                    pPr.append(bidi)
        except Exception:
            pass
        return p

    def _rtl_run(run):
        # Mark the run itself RTL (<w:rtl/> in rPr) — required for Word to lay the
        # Arabic out right-to-left, not just right-aligned.
        try:
            rPr = run._element.get_or_add_rPr()
            el = OxmlElement("w:rtl")
            el.set(qn("w:val"), "1")
            rPr.append(el)
        except Exception:
            pass
        return run

    def _para(text="", size=11, bold=False, underline=False, color=None,
              align=None, before=3, after=3, rtl_p=False):
        p = doc.add_paragraph()
        run = None
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.underline = underline
            run.font.size = Pt(size)
            run.font.name = "Segoe UI"
            if color is not None:
                run.font.color.rgb = color
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if rtl_p:
            if run is not None:
                _rtl_run(run)
            _set_bidi(p)
        if align is not None:
            p.alignment = align
        return p

    def _bullet(text, size=11, bold=True, color=None):
        # RTL reports: bidi paragraph + RTL run so the bullet sits on the RIGHT
        # (like the reference). LTR reports: normal left bullet. Hanging indent
        # keeps wrapped lines aligned under the text, not the bullet.
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        if rtl:
            pf.right_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.25)
        else:
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.25)
        run = p.add_run("•  " + (text or ""))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Segoe UI"
        run.font.color.rgb = INK if color is None else color
        if rtl:
            _rtl_run(run)
            _set_bidi(p)
        return p

    def _img(path, width_in, align):
        if not path or not os.path.exists(path):
            return
        try:
            p = doc.add_paragraph()
            p.alignment = align
            p.add_run().add_picture(path, width=Inches(width_in))
        except Exception:
            pass

    def _setcell(cell, runs, align=LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        for txt, bold, col in runs:
            rr = p.add_run(str(txt))
            rr.bold = bold
            rr.font.size = Pt(10.5)
            rr.font.name = "Segoe UI"
            if col is not None:
                rr.font.color.rgb = col

    def _add_watermark():
        """Faint, centred logo behind the text on every page (via the header)."""
        if not logo_path or not os.path.exists(logo_path):
            return
        try:
            section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            rId, _image = header.part.get_or_add_image(logo_path)
            cx = int(Inches(5.2))
            cy = int(cx * 52 / 246)          # logo aspect 246×52
            xml = (
                f'<w:r {nsdecls("w", "wp", "a", "pic", "r")}><w:drawing>'
                '<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0" '
                'simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" '
                'relativeHeight="0">'
                '<wp:simplePos x="0" y="0"/>'
                '<wp:positionH relativeFrom="margin"><wp:align>center</wp:align></wp:positionH>'
                '<wp:positionV relativeFrom="margin"><wp:align>center</wp:align></wp:positionV>'
                f'<wp:extent cx="{cx}" cy="{cy}"/>'
                '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
                '<wp:docPr id="77" name="Watermark"/><wp:cNvGraphicFramePr/>'
                '<a:graphic><a:graphicData '
                'uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
                '<pic:pic><pic:nvPicPr><pic:cNvPr id="77" name="Watermark"/>'
                '<pic:cNvPicPr/></pic:nvPicPr>'
                f'<pic:blipFill><a:blip r:embed="{rId}"><a:alphaModFix amt="18000"/></a:blip>'
                '<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
                f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
                '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
                '</pic:pic></a:graphicData></a:graphic></wp:anchor>'
                '</w:drawing></w:r>')
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p._p.append(parse_xml(xml))
        except Exception:
            pass

    # ── Header band + logo ──────────────────────────────────────────────────
    _img(logo_path, 1.9, (RIGHT if rtl else LEFT))
    _img(band_path, 6.6, CENTER)

    # ── Title (always English, centred, black) + meta (English, left) ───────
    _para("Sprint Closure Report", size=16, bold=True, color=INK,
          align=CENTER, before=10, after=10)
    _para(f"Sprint: {r.get('sprint_name','')}", size=11, bold=True,
          align=LEFT, before=2, after=2)
    _para(f"Date: {r.get('date','')}", size=11, bold=True,
          align=LEFT, before=2, after=8)

    # ── Objectives + sections grouped by epic with ordinals ─────────────────
    _para(f"{L['objectives']}:", size=13, bold=True, color=BLUE,
          rtl_p=rtl, before=10, after=4)

    def _section(label, rows):
        _para(f"{label}:", size=13, bold=True, underline=True, color=BLUE,
              rtl_p=rtl, before=8, after=4)
        if not rows:
            _para(L["none"], size=11, color=INK, rtl_p=rtl)
            return
        groups = _group_by_epic(rows)
        show_groups = any(e for e, _ in groups)
        for gi, (epic, grp) in enumerate(groups):
            if show_groups:
                ordn = _ORD.get(lang, _ORD["en"])[gi] if gi < len(_ORD.get(lang, _ORD["en"])) else str(gi + 1)
                _para(f"{ordn}: {epic or L['other']}:", size=12, bold=True,
                      underline=True, color=INK, rtl_p=rtl, before=6, after=2)
            for s in grp:
                _bullet(s.get("t") or s.get("title") or "")

    _section(L["completed"], r.get("completed", []))
    _section(L["carried"], r.get("carried", []))

    # ── Bugs summary table (reference layout: merged Total cell, blue labels) ─
    _para("", before=6, after=0)
    statuses = list((r.get("bug_by_state") or {}).items())
    tbl = doc.add_table(rows=2 + len(statuses), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for row in tbl.rows:
        row.cells[0].width = Inches(2.3)
        row.cells[1].width = Inches(4.0)
    try:
        tbl.cell(0, 0).merge(tbl.cell(1, 0))
    except Exception:
        pass
    sprint_name = r.get("sprint_name", "")
    _setcell(tbl.cell(0, 0),
             [(f"Total No of Bugs: {r.get('total_bugs', 0)}", True, BLUE)])
    _setcell(tbl.cell(0, 1),
             [("Regression bugs:  ", True, BLUE), (r.get("regression_bugs", 0), True, INK)])
    _setcell(tbl.cell(1, 1),
             [(f"Sprint {sprint_name} bugs:  ", True, BLUE), (r.get("sprint_bugs", 0), True, INK)])
    for i, (stt, cnt) in enumerate(statuses, start=2):
        _setcell(tbl.cell(i, 0), [(str(stt), True, BLUE)])
        _setcell(tbl.cell(i, 1), [(cnt, True, INK)])

    _add_watermark()

    import platform_caps as _pc
    out_dir = os.path.join(_pc.export_base_dir(), "QA Studio", "Sprint Reports")
    os.makedirs(out_dir, exist_ok=True)
    base = re.sub(r"[^A-Za-z0-9_-]+", "_", r.get("sprint_name", "") or "sprint").strip("_") or "sprint"
    path = os.path.join(out_dir, f"SprintReport_{base}_{datetime.now():%Y%m%d-%H%M}.docx")
    doc.save(path)
    for tmp in (logo_path, band_path):
        try:
            if tmp:
                os.remove(tmp)
        except Exception:
            pass
    return path


def screen(app):
    _init(app)
    import regression as R
    from main import (card, sec_head, field_label, primary_btn, green_btn, ghost_btn)

    if not app.readonly and not (app.connected and app.project):
        return R.locked_state(
            app, "Sprint Report",
            "A sprint closure report — stories by status + bug summary, Arabic or English",
            "Connect your Azure DevOps account on the Setup screen, then pick a "
            "sprint here.")

    _load_iterations(app)
    lang = app._st_lang
    L = _L.get(lang, _L["en"])
    rtl = (lang == "ar")
    _ral = ft.TextAlign.RIGHT if rtl else ft.TextAlign.LEFT

    # in-place enable/disable of the Generate button when sprints change (the
    # picker is in-place, so without this the button stayed disabled until a render)
    _gen_cell = [None]

    def _sync_gen():
        b = _gen_cell[0]
        if b is None:
            return
        ok = bool(app._st_sprint_paths) and not app._st_busy \
            and not getattr(app, "readonly", False)
        try:
            b.opacity = 1.0 if ok else 0.45
            b.on_click = (lambda e: _generate(app)) if ok else None
            b.update()
        except Exception:
            pass

    def _set_lang(k):
        new = k if k in E.LANGUAGES else "ar"
        if new == app._st_lang:
            return
        app._st_lang = new
        # If a report is already on screen, re-translate its titles live so the
        # toggle changes the content, not just the static labels.
        if app._st_report and not app._st_busy:
            _retranslate(app)
        else:
            app.ui_safe(app.render)

    def _lang_seg():
        def seg(label, key):
            sel = (app._st_lang == key)
            return ft.Container(
                ft.Text(label, size=12, weight=ft.FontWeight.BOLD,
                        color=(T.VIOLET_INK if sel else T.INK_2)),
                height=32, alignment=ft.Alignment.CENTER,
                padding=ft.Padding.symmetric(horizontal=16),
                bgcolor=(T.VIOLET_SOFT if sel else None), border_radius=T.R_SM,
                border=ft.Border.all(1, T.VIOLET if sel else ft.Colors.TRANSPARENT),
                on_click=lambda e, k=key: _set_lang(k))
        _lang_opts = [ft.DropdownOption(key=_code, text=_info["native"])
                      for _code, _info in E.LANGUAGES.items()]
        _lang_kwargs = dict(
            value=(app._st_lang if app._st_lang in E.LANGUAGES else "ar"),
            options=_lang_opts,
            on_select=lambda e: _set_lang(e.control.value or app._st_lang),
            border_color=T.BORDER, focused_border_color=T.VIOLET, border_radius=T.R,
            content_padding=ft.Padding.symmetric(vertical=8, horizontal=10),
            text_size=12, filled=True, bgcolor=T.CARD, width=190)
        try:
            _lang_dd = ft.Dropdown(menu_height=280, **_lang_kwargs)
        except TypeError:
            _lang_dd = ft.Dropdown(**_lang_kwargs)
        return ft.Container(
            _lang_dd, width=200,
            padding=4, bgcolor=T.CARD_2, border_radius=T.R, border=ft.Border.all(1, T.BORDER))

    def _toggle(key, checked):
        s = set(app._st_sprint_paths)
        s.add(key) if checked else s.discard(key)
        app._st_sprint_paths = [p for p in (x["path"] for x in app._st_iterations) if p in s]
        _sync_gen()

    def _all(checked):
        app._st_sprint_paths = [x["path"] for x in app._st_iterations] if checked else []
        _sync_gen()

    def _open():
        app._st_open = not app._st_open

    picker = (ft.Container(R._txt("Loading sprints…", color=T.INK_3, size=12), padding=10)
              if app._st_iter_loading else
              R._checkbox_multiselect(
                  [(it["path"], (_sprint_num(it["name"]) or it["name"]) + f"   ·   {it['path']}")
                   for it in app._st_iterations],
                  app._st_sprint_paths, _toggle, _all, is_open=app._st_open, on_open=_open,
                  placeholder="Select sprint(s)", empty="No sprints found for this project.",
                  page=app.page, app=app, sync_key="st_sprints"))

    card1 = card(ft.Column([
        sec_head("1", "Sprint & language"),
        ft.Container(height=10),
        ft.Column([field_label("Sprint(s)", req=True), picker], spacing=6),
        ft.Container(height=12),
        ft.Row([field_label("Report language"), ft.Container(expand=True), _lang_seg()],
               vertical_alignment=ft.CrossAxisAlignment.CENTER),
    ], spacing=0))

    _ro = bool(getattr(app, "readonly", False))
    _can_gen = bool(app._st_sprint_paths) and not app._st_busy and not _ro
    gen_btn = primary_btn(
        "Generating…" if app._st_busy else "Generate sprint report",
        icon=ft.Icons.SUMMARIZE,
        on_click=((lambda e: _generate(app)) if _can_gen else None))
    try:
        # opacity must track the SAME condition as clickability (incl. read-only),
        # otherwise the button looks enabled but does nothing.
        gen_btn.opacity = 1.0 if _can_gen else 0.45
    except Exception:
        pass
    _gen_cell[0] = gen_btn

    body_children = [card1, ft.Container(height=16), gen_btn]

    if app._st_msg and not app._st_busy:
        kind, text = app._st_msg
        _ok = (kind == "ok")
        body_children += [ft.Container(
            ft.Row([ft.Icon(ft.Icons.CHECK_CIRCLE if _ok else ft.Icons.ERROR_OUTLINE,
                            color=(T.GREEN if _ok else T.RED), size=18),
                    R._txt(text, color=(T.GREEN if _ok else T.RED), size=12.5,
                           no_wrap=False, expand=True)],
                   spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
            padding=ft.Padding.symmetric(vertical=11, horizontal=14),
            margin=ft.Margin.only(top=14), border_radius=T.R,
            bgcolor=(T.GREEN_SOFT if _ok else T.RED_SOFT),
            border=ft.Border.all(1, ft.Colors.with_opacity(0.4, T.GREEN if _ok else T.RED)))]

    if app._st_busy:
        body_children += [ft.Container(
            ft.Row([ft.ProgressRing(width=18, height=18, stroke_width=2.5, color=T.VIOLET),
                    R._txt("Fetching stories & bugs, translating titles…",
                           color=T.INK_3, size=12.5)], spacing=10),
            padding=ft.Padding.symmetric(vertical=12, horizontal=14),
            margin=ft.Margin.only(top=14),
            bgcolor=getattr(T, "VIOLET_SOFT", T.CARD_2), border_radius=T.R)]
    elif app._st_done and app._st_report:
        r = app._st_report

        _cross = ft.CrossAxisAlignment.END if rtl else ft.CrossAxisAlignment.START
        _main = ft.MainAxisAlignment.END if rtl else ft.MainAxisAlignment.START

        def _drow(kids, **kw):
            return ft.Row(list(reversed(kids)) if rtl else kids, **kw)

        def _item(text_val, accent, stripe):
            bullet = ft.Container(width=7, height=7, border_radius=999, bgcolor=accent,
                                  margin=ft.Margin.only(top=6))
            txt = R._txt(text_val, color=T.INK, size=13.5, no_wrap=False,
                         expand=True, text_align=_ral)
            return ft.Container(
                _drow([bullet, txt], spacing=12,
                      vertical_alignment=ft.CrossAxisAlignment.START),
                padding=ft.Padding.symmetric(vertical=8, horizontal=12),
                border_radius=T.R_SM,
                bgcolor=(T.CARD_2 if stripe else ft.Colors.TRANSPARENT))

        def _epic_head(ordn, epic, accent):
            # Ordinal, an explicit ":" separator, and the (often English) epic
            # name as THREE separate spans in a direction-aware row — not one
            # mixed string, whose colon bidi-reorders unpredictably. For Arabic
            # the ordinal sits on the RIGHT, then ":", then the epic name on the
            # LEFT ("ثانياً : Member Portal"); for English it reads naturally.
            _mk = lambda v: R._txt(v, color=accent, weight=ft.FontWeight.W_800,
                                   size=12)
            return ft.Container(
                _drow([_mk(ordn), _mk(":"), _mk(epic)], spacing=4, alignment=_main),
                padding=ft.Padding.symmetric(vertical=6, horizontal=10),
                bgcolor=ft.Colors.with_opacity(0.10, accent),
                border_radius=999,
                margin=ft.Margin.only(top=14, bottom=6))

        _sec_icon = {T.GREEN: ft.Icons.CHECK_CIRCLE_ROUNDED,
                     T.AMBER: ft.Icons.SCHEDULE_ROUNDED}

        def _sec(label, rows, accent, soft):
            header = ft.Container(
                _drow([
                    ft.Icon(_sec_icon.get(accent, ft.Icons.LABEL_ROUNDED),
                            color=accent, size=17),
                    R._txt(label, color=T.INK, weight=ft.FontWeight.W_900, size=14),
                    ft.Container(expand=True),
                    ft.Container(R._txt(f"{len(rows)}", size=11, weight=ft.FontWeight.W_800,
                                        color=accent),
                                 padding=ft.Padding.symmetric(vertical=2, horizontal=9),
                                 bgcolor=ft.Colors.with_opacity(0.16, accent),
                                 border_radius=999),
                ], spacing=10, alignment=_main,
                   vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.Padding.symmetric(vertical=11, horizontal=13),
                bgcolor=soft, border_radius=T.R_SM,
                border=ft.Border.all(1, ft.Colors.with_opacity(0.35, accent)))
            if not rows:
                body = [ft.Container(R._txt(L["none"], color=T.INK_3, size=12.5),
                                     padding=ft.Padding.symmetric(vertical=9, horizontal=12))]
            else:
                groups = _group_by_epic(rows)
                show_groups = any(e for e, _ in groups)
                body, stripe = [], 0
                for gi, (epic, grp) in enumerate(groups):
                    if show_groups:
                        ordn = _ORD.get(lang, _ORD["en"])[gi] if gi < len(_ORD.get(lang, _ORD["en"])) else str(gi + 1)
                        body.append(_epic_head(ordn, epic or L['other'], accent))
                    for s in grp:
                        body.append(_item(s.get("t") or s.get("title") or "",
                                          accent, stripe % 2 == 1))
                        stripe += 1
            return ft.Column([header, ft.Container(height=8),
                              ft.Column(body, spacing=2)], spacing=0)

        # Same tile the Regression/Sprint Plan KPI strip uses (R._kpi_tile) —
        # gradient mono numbers on a shadowed card, instead of this screen's
        # own flatter, unshadowed `_stat` tile — so all three "plan" screens
        # share one visual language for a labeled headline number.
        import platform_caps as _pc_sr
        bug_stats = _drow(R.kpi_tiles_mobile([
            R._kpi_tile(L["total_bugs"].upper(), str(r["total_bugs"]), T.VIOLET_INK),
            R._kpi_tile(L["regression_bugs"].upper(), str(r["regression_bugs"]), T.RED),
            R._kpi_tile(L["sprint_bugs"].upper(), str(r["sprint_bugs"]), T.AMBER),
        ]), spacing=10, wrap=_pc_sr.is_mobile(), run_spacing=10)

        def _state_pair(st):
            s = (st or "").lower()
            if s in _DONE:
                return T.GREEN, T.GREEN_SOFT
            if s in ("active", "in progress", "committed", "doing"):
                return T.VIOLET_INK, T.VIOLET_SOFT
            if s in ("new", "to do", "proposed", "open"):
                return T.AMBER, T.AMBER_SOFT
            return T.INK_2, T.CARD_2

        def _status_chip(st, n):
            fg, bg = _state_pair(st)
            return ft.Container(
                _drow([ft.Container(width=7, height=7, border_radius=999, bgcolor=fg),
                       R._txt(str(st), color=T.INK_2, size=12),
                       ft.Container(R._txt(str(n), color=fg, size=12,
                                           weight=ft.FontWeight.W_800),
                                    padding=ft.Padding.symmetric(vertical=1, horizontal=7),
                                    bgcolor=ft.Colors.with_opacity(0.14, fg),
                                    border_radius=999)],
                      spacing=8, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.Padding.symmetric(vertical=6, horizontal=10),
                bgcolor=bg, border_radius=999, border=ft.Border.all(1, T.BORDER))

        status_items = [_status_chip(st, n)
                        for st, n in (r.get("bug_by_state") or {}).items()]
        bug_section = ft.Column([
            _drow([ft.Icon(ft.Icons.BUG_REPORT_ROUNDED, color=T.VIOLET, size=17),
                   R._txt(L["bugs"], color=T.INK, weight=ft.FontWeight.W_900, size=14)],
                  spacing=10, alignment=_main,
                  vertical_alignment=ft.CrossAxisAlignment.CENTER),
            ft.Container(height=12), bug_stats, ft.Container(height=16),
            R._txt(L["by_status"], color=T.INK_3, size=11, weight=ft.FontWeight.BOLD,
                   text_align=_ral),
            ft.Container(height=8),
            (ft.Row(list(reversed(status_items)) if rtl else status_items,
                    wrap=True, spacing=8, run_spacing=8, alignment=_main)
             if status_items else R._txt(L["none"], color=T.INK_3, size=12)),
        ], spacing=0, horizontal_alignment=_cross)

        def _copy(e):
            # Same shared clipboard-with-fallback mechanism the Run/Automation
            # logs use (main.py's _copy_text_to_clipboard): tries
            # page.set_clipboard() first, then falls back to a direct Windows
            # clipboard write if that silently fails (a real, documented Flet
            # IPC limitation), and only shows a red error toast if BOTH paths
            # fail. The previous bare try/except here swallowed set_clipboard's
            # failure with a plain `pass` — so on machines where the IPC path
            # fails, the "Report copied" toast never even fired and there was
            # no fallback, which is why Copy looked like it did nothing.
            # (Uses the raw-text variant, not _copy_log_text, since that one
            # strips blank lines — which here are the deliberate spacers
            # between sections, not log noise.)
            lines = [f"{L['title']} — {L['sprint']} {r['sprint_name']} ({r['date']})", ""]
            for label, rows in ((L["completed"], r["completed"]), (L["carried"], r["carried"])):
                lines.append(f"{label} ({len(rows)}):")
                lines += ["  • " + (s.get("t") or s.get("title") or "") for s in rows] or ["  -"]
                lines.append("")
            lines += [f"{L['bugs']}: {L['total_bugs']} {r['total_bugs']} · "
                      f"{L['regression_bugs']} {r['regression_bugs']} · "
                      f"{L['sprint_bugs']} {r['sprint_bugs']}"]
            app._copy_text_to_clipboard("\n".join(lines), "Report copied to clipboard.")

        def _download(e):
            def _w():
                try:
                    import platform_caps as _pc
                    # Ask WHERE to save (native 'Save As') — same flow as the Sprint
                    # Plan / Regression exporters. Mobile has no tkinter dialog, so
                    # _ask_save_path returns False and we fall back to the download
                    # popup. Runs off the UI thread (app._bg), as _ask_save_path needs.
                    _rr = getattr(app, "_st_report", None) or {}
                    _base = (re.sub(r"[^A-Za-z0-9_-]+", "_",
                                    (_rr.get("sprint_name") or "sprint")).strip("_")
                             or "sprint")
                    _default = f"SprintReport_{_base}_{datetime.now():%Y%m%d-%H%M}.docx"
                    dest = R._ask_save_path("docx", _default)
                    if dest is None:
                        return                        # user cancelled the dialog
                    p = _export_docx(app)
                    if dest and dest is not False:    # a location was chosen → move there
                        if not dest.lower().endswith(".docx"):
                            dest += ".docx"
                        import shutil
                        if os.path.abspath(dest) != os.path.abspath(p):
                            shutil.move(p, dest)
                        p = dest
                    if _pc.is_mobile():
                        app.ui_safe(lambda pp=p: app._mobile_download_popup(
                            pp, "Word document ready"))
                    else:
                        _pc.open_folder(os.path.dirname(p))
                        app.ui_safe(lambda pp=p: app._toast(f"Saved Word document: {pp}"))
                except ImportError:
                    app.ui_safe(lambda: app._err("Word export needs python-docx."))
                except Exception as ex:
                    app.ui_safe(lambda e=ex: app._err(f"Export failed: {ex}"))
            app._bg(_w)

        # "Hero" treatment matching the plan email's masthead (cont'd #29) —
        # a small violet-soft eyebrow pill above the title instead of the
        # title sitting flatly next to the icon with nothing to anchor it,
        # plus a shadowed gradient icon so this reads as the report's
        # headline rather than just another list item.
        title_band = ft.Container(
            _drow([
                ft.Container(ft.Icon(ft.Icons.SUMMARIZE_ROUNDED, color=ft.Colors.WHITE, size=22),
                             width=46, height=46, alignment=ft.Alignment.CENTER,
                             border_radius=T.R,
                             gradient=ft.LinearGradient(
                                 begin=ft.Alignment.TOP_LEFT, end=ft.Alignment.BOTTOM_RIGHT,
                                 colors=[T.VIOLET, T.VIOLET_H]),
                             shadow=ft.BoxShadow(blur_radius=14, spread_radius=-6,
                                                 offset=ft.Offset(0, 5),
                                                 color=ft.Colors.with_opacity(0.35, T.VIOLET))),
                ft.Column([
                    ft.Container(
                        R._txt(("سبرنت مغلق" if lang == "ar" else "SPRINT SNAPSHOT"),
                               color=T.VIOLET_INK, size=10, weight=ft.FontWeight.W_800,
                               text_align=_ral),
                        padding=ft.Padding.symmetric(vertical=2, horizontal=9),
                        bgcolor=T.VIOLET_SOFT, border_radius=999),
                    ft.Container(height=4),
                    R._txt(L["title"], color=T.INK, weight=ft.FontWeight.W_900, size=17,
                           text_align=_ral),
                    R._txt(f"{L['sprint']} {r['sprint_name']}  ·  {r['date']}  ·  "
                           f"{r['total_stories']} {L['stories']}",
                           color=T.INK_3, size=12, weight=ft.FontWeight.BOLD,
                           font_family=T.F_MONO, text_align=_ral),
                ], spacing=0, expand=True, horizontal_alignment=_cross),
            ], spacing=14, vertical_alignment=ft.CrossAxisAlignment.CENTER),
            padding=16, bgcolor=T.CARD_2, border_radius=T.R,
            border=ft.Border.all(1, T.BORDER))

        def _panel(child):
            return ft.Container(child, padding=14, bgcolor=T.CARD,
                                border=ft.Border.all(1, T.BORDER), border_radius=T.R)

        # Action row (Copy / Download). On mobile the single row overflowed —
        # "2 Report" + spacer + Copy + "Download Word" is wider than a phone,
        # so the green Download button ran off the right edge (reported live).
        # On mobile: drop the header onto its own line and give the two buttons
        # their own full-width row, each expanding to split the width evenly.
        # Desktop keeps the original one-line layout.
        import platform_caps as _pc_hdr
        _mob_hdr = _pc_hdr.is_mobile()
        _copy_btn = ghost_btn("Copy", icon=ft.Icons.CONTENT_COPY, on_click=_copy,
                              expand=_mob_hdr)
        _dl_btn = green_btn("Download Word", icon=ft.Icons.DESCRIPTION,
                            on_click=_download, expand=_mob_hdr)
        if _mob_hdr:
            _report_header = ft.Column([
                sec_head("2", "Report"),
                ft.Container(height=10),
                ft.Row([_copy_btn, _dl_btn], spacing=10),
            ], spacing=0)
        else:
            _report_header = ft.Row(
                [sec_head("2", "Report"), ft.Container(expand=True),
                 _copy_btn, _dl_btn],
                vertical_alignment=ft.CrossAxisAlignment.CENTER)
        results = card(ft.Column([
            _report_header,
            ft.Container(height=12),
            title_band,
            ft.Container(height=14),
            _panel(_sec(L["completed"], r["completed"], T.GREEN, T.GREEN_SOFT)),
            ft.Container(height=12),
            _panel(_sec(L["carried"], r["carried"], T.AMBER, T.AMBER_SOFT)),
            ft.Container(height=12),
            _panel(bug_section),
        ], spacing=0))
        body_children += [ft.Container(height=16), results]

    body = ft.Column(body_children, spacing=0, scroll=ft.ScrollMode.AUTO, expand=True)
    return app.shell("Sprint Report",
                     "Stories by status + bug summary from a sprint, in Arabic or English",
                     body, badge="SR")
